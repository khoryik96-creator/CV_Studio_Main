import copy
import io
import json
import os
from pathlib import Path
import re
import tempfile
import unittest
from unittest import mock
import zipfile

from owner_build_tools import build_protected


_TEMP = tempfile.TemporaryDirectory(prefix="cvstudio-long-cv-")
_ROOT = Path(__file__).resolve().parents[1]
os.environ["CVSTUDIO_DB_PATH"] = str(Path(_TEMP.name) / "state.sqlite3")
os.environ["LOCALAPPDATA"] = str(Path(_TEMP.name) / "local")
build_protected.write_test_receipt(_ROOT)

import app


class LongCvOutputCorrectiveTests(unittest.TestCase):
    def fixture(self):
        return {
            "candidate": {
                "name": "Saik Eng Joo",
                "current_position": "Insurance Agent/Financial Advisor (assumed from duties)",
                "current_company": "MANULIFE",
            },
            "work_experiences": [{
                "date_range": "2023 to Present",
                "company": "Example Sdn Bhd",
                "roles": [{
                    "title": "Insurance Agent/Financial Advisor (likely based on responsibilities)",
                    "date_range": "",
                    "reason_for_leaving": "",
                    "bullets": [
                        "Key responsibilities",
                        '{"heading":"Responsibilities","bullets":["Achievement"]}',
                        '{"bullets":["Successfully achieved over 100% of the annual sales target.","Honored with Best Performing Business Development Manager for year 2024."],"heading":"Business Set up"}',
                        '{"bullets":["Successfully achieved 120% of the annual sales target."],"heading":"Manpower"}',
                        '{"heading":"malformed","bullets":[}',
                    ],
                }],
            }],
            "education": [],
            "certifications": ["", "  "],
            "skills": [{"category": "Skills", "items": "Leadership"}, {"category": "", "items": ""}],
        }

    def test_long_cv_timeout_is_bounded_and_role_dense(self):
        self.assertEqual(app._cv_parse_backend_timeout_seconds("x" * 7999), 180)
        self.assertEqual(app._cv_parse_backend_timeout_seconds("x" * 8000), 300)
        # A real dense 8-page CV extracts to ~9-10k chars and must get the long
        # (300s) budget, not be cut off at 180s.
        self.assertEqual(app._cv_parse_backend_timeout_seconds("x" * 9808), 300)
        dense = "\n".join(["Key responsibilities"] * 7)
        self.assertEqual(app._cv_parse_backend_timeout_seconds(dense), 180)
        dense += "\nKey achievement"
        self.assertEqual(app._cv_parse_backend_timeout_seconds(dense), 300)

    def test_docx_extraction_preserves_real_list_markers_and_plain_headings(self):
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        document = Document()
        heading = document.add_paragraph("Implementation")
        heading.runs[0].bold = True
        document.add_paragraph("Configured the system", style="List Bullet")
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = "Rollout"
        cell.paragraphs[0].runs[0].bold = True
        cell.add_paragraph("Delivered the rollout", style="List Bullet")
        suppressed = document.add_paragraph("Intentionally plain", style="List Bullet")
        num_pr = suppressed._p.get_or_add_pPr().get_or_add_numPr()
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "0")
        num_pr.append(num_id)
        payload = io.BytesIO()
        document.save(payload)

        extracted = app._extract_docx_text_preserve_tables(payload.getvalue()).splitlines()
        self.assertIn("Implementation", extracted)
        self.assertIn("• Configured the system", extracted)
        self.assertIn("Rollout", extracted)
        self.assertIn("• Delivered the rollout", extracted)
        self.assertIn("Intentionally plain", extracted)
        self.assertNotIn("• Implementation", extracted)
        self.assertNotIn("• Rollout", extracted)
        self.assertNotIn("• Intentionally plain", extracted)

    def test_common_role_subheadings_render_as_sections_not_bullets(self):
        normalized = app._normalize_cv_bullet_items([
            "Implementation",
            "Configured the system",
            "Delivered the rollout",
            "Rollout:",
            "Deployed the release",
            "Activities Description",
            "Documented the process",
        ])
        self.assertEqual(normalized, [
            {
                "heading": "Implementation",
                "bullets": ["Configured the system", "Delivered the rollout"],
                "kind": "section",
            },
            {
                "heading": "Rollout",
                "bullets": ["Deployed the release"],
                "kind": "section",
            },
            {
                "heading": "Activities Description",
                "bullets": ["Documented the process"],
                "kind": "section",
            },
        ])

    def test_parse_route_passes_long_timeout_to_provider(self):
        provider_result = {
            "content": [{"type": "text", "text": json.dumps({
                "candidate": {}, "work_experiences": [], "education": [],
                "certifications": [], "skills": [],
            })}],
            "usage": {},
        }
        with (
            mock.patch.object(app, "call_llm", return_value=provider_result) as call,
            mock.patch.object(app, "_ai_spend_session_allowed", return_value=True),
        ):
            response = app.app.test_client().post(
                "/parse",
                json={"api_key": "fixture-key", "cv_text": "x" * 18000},
                headers={"Origin": "http://127.0.0.1:5000", "X-CV-Studio-Request": "1"},
            )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(call.call_args.args[2]["_timeout_seconds"], 300)

    def test_parse_retry_does_not_shorten_candidate_content(self):
        # When the first parse returns unrepairable JSON, the retry must NOT ask
        # the model to shorten the CV (a formatter must never silently drop
        # bullets/roles/skills). It re-requests full, complete JSON instead.
        valid = json.dumps({
            "candidate": {}, "work_experiences": [], "education": [],
            "certifications": [], "skills": [],
        })
        calls = []

        def fake_call_llm(provider, api_key, payload):
            calls.append(payload)
            text = "This is not valid JSON at all." if len(calls) == 1 else valid
            return {"content": [{"type": "text", "text": text}], "usage": {}}

        with (
            mock.patch.object(app, "call_llm", side_effect=fake_call_llm),
            mock.patch.object(app, "_ai_spend_session_allowed", return_value=True),
        ):
            response = app.app.test_client().post(
                "/parse",
                json={"api_key": "fixture-key", "cv_text": "Kelana Edy Zainudin\nWork Experience\n"},
                headers={"Origin": "http://127.0.0.1:5000", "X-CV-Studio-Request": "1"},
            )
        self.assertEqual(response.status_code, 200)
        self.assertGreaterEqual(len(calls), 2, "the retry strategy should have fired")
        retry_prompt = calls[1]["messages"][0]["content"]
        self.assertNotIn("max 15 words", retry_prompt)
        self.assertNotIn("max 8 bullet", retry_prompt)
        self.assertNotIn("keep output short", retry_prompt)
        self.assertIn("Preserve EVERY", retry_prompt)

    def test_every_parse_attempt_uses_deterministic_temperature(self):
        calls = []

        def fake_call_llm(provider, api_key, payload):
            calls.append(payload)
            return {
                "content": [{"type": "text", "text": "not valid json"}],
                "usage": {},
            }

        with (
            mock.patch.object(app, "call_llm", side_effect=fake_call_llm),
            mock.patch.object(app, "_ai_spend_session_allowed", return_value=True),
        ):
            response = app.app.test_client().post(
                "/parse",
                json={
                    "api_key": "fixture-key",
                    "cv_text": "Deterministic Candidate\nWork Experience\n",
                },
                headers={
                    "Origin": "http://127.0.0.1:5000",
                    "X-CV-Studio-Request": "1",
                },
            )

        self.assertEqual(response.status_code, 500)
        self.assertEqual(len(calls), 3)
        self.assertTrue(all(call.get("temperature") == 0 for call in calls))

    def test_truncated_parse_is_salvaged_and_flagged_degraded(self):
        # A truncated AI response that json.loads cannot read but the bracket
        # salvage (try_close_json) can. The CV must STILL be repaired and
        # returned, and the response must carry a degraded flag + warning so the
        # caller knows completeness is not guaranteed. Flagging never blocks the
        # repair.
        truncated = (
            '{"candidate": {"name": "Jane Doe", "current_company": "Acme"}, '
            '"work_experiences": [{"company": "Acme", "date_range": "2020 to Present", '
            '"roles": [{"title": "Engineer", "date_range": "", "reason_for_leaving": "", '
            '"bullets": ["Built systems", "Led a team'
        )
        calls = []

        def fake_call_llm(provider, api_key, payload):
            calls.append(payload)
            return {"content": [{"type": "text", "text": truncated}], "usage": {}}

        with (
            mock.patch.object(app, "call_llm", side_effect=fake_call_llm),
            mock.patch.object(app, "_ai_spend_session_allowed", return_value=True),
        ):
            response = app.app.test_client().post(
                "/parse",
                json={"api_key": "fixture-key", "cv_text": "Jane Doe\nWork Experience\n"},
                headers={"Origin": "http://127.0.0.1:5000", "X-CV-Studio-Request": "1"},
            )
        self.assertEqual(response.status_code, 200)
        body = response.get_json()
        # Salvaged in one shot -- no wasteful Strategy 2/3 re-send.
        self.assertEqual(len(calls), 1)
        # Repair still happened: the readable content is returned.
        self.assertEqual(body["data"]["candidate"]["name"], "Jane Doe")
        # And it is flagged so an incomplete parse is not reported as clean.
        self.assertTrue(body.get("degraded"))
        self.assertEqual(body.get("degraded_reason"), "truncated_response_bracket_salvage")
        self.assertTrue(str(body.get("warning") or "").strip())

    def test_clean_parse_is_not_flagged_degraded(self):
        # A first-try valid parse is lossless and must NOT carry the degraded
        # flag or a warning -- the happy-path response shape is unchanged.
        valid = json.dumps({
            "candidate": {"name": "John Smith"}, "work_experiences": [],
            "education": [], "certifications": [], "skills": [],
        })

        def fake_call_llm(provider, api_key, payload):
            return {"content": [{"type": "text", "text": valid}], "usage": {}}

        with (
            mock.patch.object(app, "call_llm", side_effect=fake_call_llm),
            mock.patch.object(app, "_ai_spend_session_allowed", return_value=True),
        ):
            response = app.app.test_client().post(
                "/parse",
                json={"api_key": "fixture-key", "cv_text": "John Smith\nWork Experience\n"},
                headers={"Origin": "http://127.0.0.1:5000", "X-CV-Studio-Request": "1"},
            )
        self.assertEqual(response.status_code, 200)
        body = response.get_json()
        self.assertEqual(body["data"]["candidate"]["name"], "John Smith")
        self.assertNotIn("degraded", body)
        self.assertNotIn("warning", body)

    def test_structured_normalization_is_idempotent_and_factual(self):
        normalized = app._normalize_cv_structured_content(self.fixture())
        role = normalized["work_experiences"][0]["roles"][0]
        self.assertEqual(normalized["candidate"]["current_position"], "")
        self.assertEqual(role["title"], "")
        # The orphan "Key responsibilities" label (no bullets of its own,
        # immediately followed by another section) is dropped; the real
        # "Key responsibilities" section with its bullet survives as the first.
        self.assertEqual(role["bullets"][0], {
            "heading": "Key responsibilities",
            "bullets": ["Achievement"],
            "kind": "section",
        })
        self.assertEqual(role["bullets"][1]["heading"], "Business Set up")
        self.assertEqual(role["bullets"][2]["heading"], "Manpower")
        self.assertEqual(role["bullets"][3], '{"heading":"malformed","bullets":[}')
        self.assertEqual(normalized["certifications"], [])
        self.assertEqual(normalized["skills"], [{"category": "Skills", "items": "Leadership"}])
        self.assertEqual(app._normalize_cv_structured_content(copy.deepcopy(normalized)), normalized)

    def test_standalone_sections_and_bounded_inference_variants_are_normalized(self):
        # A standalone label followed by loose bullets absorbs them (rather than
        # rendering as a bare label above orphaned siblings).
        self.assertEqual(app._normalize_cv_bullet_items(["Key achievement", "Delivered result"]), [
            {"heading": "Key achievements", "bullets": ["Delivered result"], "kind": "section"},
        ])
        for title in (
            "Advisor (inferred from duties)",
            "Advisor (implied from responsibilities)",
            "Advisor (assumed from duties)",
            "Advisor (guessed from context)",
            "Advisor (likely based on responsibilities)",
        ):
            self.assertEqual(app._strip_cv_inferred_title(title), "")
        self.assertEqual(app._strip_cv_inferred_title("Advisor"), "Advisor")
        self.assertEqual(app._strip_cv_inferred_title("Advisor (likely to succeed)"), "Advisor (likely to succeed)")

    def test_additional_information_strips_baked_markers_without_false_positives(self):
        data = {
            "candidate": {"name": "Joel Ezra Joe"},
            "work_experiences": [],
            "education": [],
            "certifications": ["• Professional Scrum Master", "-5% variance"],
            "skills": [
                {
                    "category": "Projects",
                    "items": [
                        "CryptoRadar",
                        "• Production TypeScript app",
                        "• Implemented response caching",
                    ],
                },
                {
                    "category": "Volunteer & Community",
                    "items": "VideoClaw Hackerhouse\n• Won Best Product Design",
                },
                {
                    "category": "Protected values",
                    "items": ["-5% variance", "3.5 years", "(Vendors: AWS, Azure)"],
                },
            ],
        }
        normalized = app._normalize_cv_structured_content(copy.deepcopy(data))
        self.assertEqual(
            normalized["certifications"],
            ["Professional Scrum Master", "-5% variance"],
        )
        self.assertEqual(
            normalized["skills"][0]["items"],
            ["CryptoRadar", "Production TypeScript app", "Implemented response caching"],
        )
        self.assertEqual(
            normalized["skills"][1]["items"],
            "VideoClaw Hackerhouse\nWon Best Product Design",
        )
        self.assertEqual(
            normalized["skills"][2]["items"],
            ["-5% variance", "3.5 years", "(Vendors: AWS, Azure)"],
        )
        self.assertEqual(
            app._normalize_cv_structured_content(copy.deepcopy(normalized)),
            normalized,
        )

        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": data},
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        import xml.etree.ElementTree as ET
        root = ET.fromstring(document_xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        numbered_text = []
        for paragraph in root.findall(".//w:body//w:p", ns):
            if paragraph.find("w:pPr/w:numPr/w:numId", ns) is None:
                continue
            numbered_text.append("".join(
                node.text or "" for node in paragraph.findall(".//w:t", ns)
            ))
        self.assertIn("Production TypeScript app", numbered_text)
        self.assertIn("Won Best Product Design", numbered_text)
        self.assertFalse(any(text.startswith("•") for text in numbered_text))
        self.assertIn("-5% variance", numbered_text)
        self.assertIn("3.5 years", numbered_text)
        self.assertIn("(Vendors: AWS, Azure)", numbered_text)

    def test_docx_xml_never_serializes_structured_bullets_or_inferred_title(self):
        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": self.fixture()},
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertNotIn('&quot;heading&quot;:&quot;Business Set up&quot;', document_xml)
        self.assertNotIn('&quot;heading&quot;:&quot;Responsibilities&quot;', document_xml)
        self.assertNotIn("assumed from duties", document_xml)
        self.assertNotIn("likely based on responsibilities", document_xml)
        self.assertIn("Key responsibilities", document_xml)
        heading_paragraphs = [
            paragraph for paragraph in re.findall(r"<w:p\b.*?</w:p>", document_xml, re.S)
            if "Key responsibilities" in paragraph
        ]
        self.assertTrue(heading_paragraphs)
        self.assertTrue(all("<w:numPr>" not in paragraph for paragraph in heading_paragraphs))
        self.assertIn("Business Set up", document_xml)
        self.assertIn("Successfully achieved over 100%", document_xml)
        self.assertEqual(document_xml.count("<w:t>Skills:</w:t>"), 1)
        self.assertNotIn("E D U C A T I O N S  &amp;  T R A I N I N G", document_xml)
        self.assertNotIn("<w:t>Example Sdn Bhd</w:t></w:r></w:p></w:tc>\n    <w:tc", document_xml)

    def test_docx_education_renders_single_year_and_omits_missing_date_separator(self):
        data = {
            "candidate": {"name": "Education Fixture"},
            "work_experiences": [],
            "education": [
                {
                    "institution": "University of Tenaga Nasional",
                    "degree": "Bachelor's Degree in Computer Science",
                    "date_range": "to 2001",
                },
                {
                    "institution": "Undated University",
                    "degree": "Certificate",
                    "date_range": "",
                },
            ],
            "certifications": [],
            "skills": [],
        }

        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": data},
            headers={"Origin": "http://127.0.0.1:5000"},
        )

        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("2001 | University of Tenaga Nasional", document_xml)
        self.assertNotIn("to 2001", document_xml)
        self.assertIn("Undated University", document_xml)
        self.assertNotIn("| Undated University", document_xml)

    def test_docx_fills_summary_textbox_in_both_word_compatibility_layers(self):
        import xml.etree.ElementTree as ET

        response = app.app.test_client().post(
            "/generate-docx",
            json={
                "data": {
                    "candidate": {"name": "Summary Fixture"},
                    "summary_bullets": [
                        "**Cloud platforms** leader across regional delivery.",
                        "Built engineering teams & delivery standards.",
                    ],
                    "work_experiences": [{
                        "company": "Example Sdn Bhd",
                        "date_range": "Jan 2020 to Present",
                        "roles": [{"title": "Director", "bullets": []}],
                    }],
                    "education": [],
                    "certifications": [],
                    "skills": [],
                }
            },
            headers={"Origin": "http://127.0.0.1:5000"},
        )

        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        ET.fromstring(document_xml)
        first_table = re.search(r"<w:tbl\b.*?</w:tbl>", document_xml, re.S)
        self.assertIsNotNone(first_table)
        self.assertNotIn("Cloud platforms", first_table.group(0))
        summary_boxes = [
            match.group(0)
            for match in re.finditer(
                r"<w:txbxContent\b[^>]*>.*?</w:txbxContent>",
                document_xml,
                re.S,
            )
            if "_CVStudioSummaryBox" in match.group(0)
        ]
        self.assertEqual(len(summary_boxes), 2)
        for summary_box in summary_boxes:
            self.assertIn("Cloud platforms", summary_box)
            self.assertIn("Built engineering teams &amp; delivery standards.", summary_box)
            self.assertIn('<w:numId w:val="2"/>', summary_box)
            self.assertIn('<w:textDirection w:val="btLr"/>', summary_box)
            self.assertEqual(summary_box.count('<w:spacing w:after="160"/>'), 1)
            self.assertNotIn("*Summary*", summary_box)
            self.assertNotIn("*Tech Stack*", summary_box)
            self.assertNotIn("*Key Things to Highlight*", summary_box)
            self.assertNotIn("*Certifications", summary_box)
        self.assertEqual(document_xml.count("_CVStudioSummaryBox"), 2)
        self.assertEqual(document_xml.count("ABOUT HIM / HER"), 2)
        self.assertNotIn("S U M M A R Y", document_xml)
        self.assertRegex(
            document_xml,
            r"<w:r><w:rPr><w:b/><w:bCs/>.*?<w:t>Cloud platforms</w:t></w:r>",
        )
        self.assertEqual(
            document_xml.count("Built engineering teams &amp; delivery standards."),
            2,
        )

    def test_docx_replaces_xml_invalid_controls_without_damaging_unicode(self):
        import xml.etree.ElementTree as ET

        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": {
                "candidate": {"name": "Bad\x00Name\x0b😀"},
                "summary_bullets": ["Built\x01safe\x0c output & kept Unicode 😀."],
                "work_experiences": [],
                "education": [],
                "certifications": [],
                "skills": [],
            }},
            headers={"Origin": "http://127.0.0.1:5000"},
        )

        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        ET.fromstring(document_xml)
        self.assertNotRegex(
            document_xml,
            r"[\x00-\x08\x0B\x0C\x0E-\x1F\uD800-\uDFFF\uFFFE\uFFFF]",
        )
        self.assertIn("Bad Name 😀", document_xml)
        self.assertIn("Built safe output &amp; kept Unicode 😀.", document_xml)

    def test_summary_box_autofit_resizes_safe_text_caps_long_text_on_page_one_and_can_be_disabled(self):
        def render(summary_bullets, setting_marker=None):
            body = {"data": {
                "candidate": {"name": "Auto-fit Fixture"},
                "summary_bullets": summary_bullets,
                "work_experiences": [],
                "education": [],
                "certifications": [],
                "skills": [],
            }}
            if setting_marker is not None:
                body["summary_box_autofit"] = setting_marker
            response = app.app.test_client().post(
                "/generate-docx",
                json=body,
                headers={"Origin": "http://127.0.0.1:5000"},
            )
            self.assertEqual(response.status_code, 200)
            with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
                return archive.read("word/document.xml").decode("utf-8")

        def modern_summary_shape(document_xml):
            shapes = [
                match.group(0)
                for match in re.finditer(
                    r"<wps:wsp>.*?</wps:wsp>", document_xml, re.S
                )
                if "_CVStudioSummaryBox" in match.group(0)
            ]
            self.assertEqual(len(shapes), 1)
            return shapes[0]

        def modern_summary_anchor(document_xml):
            anchors = [
                match.group(0)
                for match in re.finditer(
                    r"<wp:anchor\b[^>]*>.*?</wp:anchor>", document_xml, re.S
                )
                if "_CVStudioSummaryBox" in match.group(0)
            ]
            self.assertEqual(len(anchors), 1)
            return anchors[0]

        def vml_summary_opening(document_xml):
            for marker in re.finditer("_CVStudioSummaryBox", document_xml):
                rect_start = document_xml.rfind("<v:rect", 0, marker.start())
                rect_close = document_xml.rfind("</v:rect>", 0, marker.start())
                if rect_start > rect_close:
                    opening_end = document_xml.index(">", rect_start)
                    return document_xml[rect_start:opening_end + 1]
            self.fail("VML summary rectangle was not found")

        def vml_summary_group_opening(document_xml):
            groups = re.findall(
                r'<v:group\b(?=[^>]*\bid="Group 28")[^>]*>',
                document_xml,
            )
            self.assertEqual(len(groups), 1)
            return groups[0]

        resized = render(["A concise summary bullet.", "A second concise bullet."])
        self.assertLess(
            resized.index("_CVStudioSummaryBox"),
            resized.index("<w:tbl>"),
        )
        self.assertIn("<a:spAutoFit/>", modern_summary_shape(resized))
        self.assertIn("mso-fit-shape-to-text:t", vml_summary_opening(resized))
        self.assertIn("<wp:posOffset>2218690</wp:posOffset>", modern_summary_anchor(resized))
        self.assertIn('cy="3089275"', modern_summary_anchor(resized))

        fixed = render(["A concise summary bullet."], False)
        self.assertIn("<a:noAutofit/>", modern_summary_shape(fixed))
        self.assertNotIn("mso-fit-shape-to-text", vml_summary_opening(fixed))
        self.assertNotIn("mso-fit-text-to-shape", vml_summary_opening(fixed))
        self.assertIn("<wp:posOffset>2218690</wp:posOffset>", modern_summary_anchor(fixed))
        self.assertIn("margin-top:174.7pt", vml_summary_group_opening(fixed))
        self.assertIn("height:243.25pt", vml_summary_group_opening(fixed))

        very_long = render(["Long summary text " * 12 for _ in range(10)], True)
        self.assertLess(
            very_long.index("<w:tbl>"),
            very_long.index("_CVStudioSummaryBox"),
        )
        self.assertEqual(very_long.count("<w:tbl>"), 1)
        self.assertEqual(very_long.count("_CVStudioSummaryBox"), 2)
        long_shape = modern_summary_shape(very_long)
        self.assertIn("<a:noAutofit/>", long_shape)
        self.assertNotIn("<a:normAutofit/>", long_shape)
        self.assertIn('<w:sz w:val="22"/>', long_shape)
        self.assertIn("Long summary text", long_shape)
        self.assertNotIn("mso-fit-text-to-shape", vml_summary_opening(very_long))
        self.assertNotIn("mso-fit-shape-to-text:t", vml_summary_opening(very_long))
        long_anchor = modern_summary_anchor(very_long)
        self.assertIn("<wp:posOffset>194310</wp:posOffset>", long_anchor)
        self.assertIn('<wp:extent cx="5730875" cy="6229350"/>', long_anchor)
        self.assertRegex(
            long_anchor,
            r'<wpg:grpSpPr>\s*<a:xfrm>.*?<a:ext cx="5730875" cy="6229350"/>',
        )
        long_group = vml_summary_group_opening(very_long)
        self.assertIn("margin-top:15.3pt", long_group)
        self.assertIn("height:490.5pt", long_group)

    def test_summary_max_geometry_does_not_resize_an_unrelated_vml_group(self):
        unrelated = (
            '<v:group id="Group 28" style="margin-top:99pt;height:88pt">'
            '<v:rect id="Unrelated shape"/></v:group>'
        )
        summary = (
            '<v:group id="Group 28" style="margin-top:174.7pt;height:243.25pt">'
            '<v:group id="Nested group"><v:rect>'
            '<w:bookmarkStart w:id="1" w:name="_CVStudioSummaryBox1"/>'
            '</v:rect></v:group></v:group>'
        )

        patched = app._summary_docx_patch_max_geometry(unrelated + summary)

        self.assertIn(unrelated, patched)
        self.assertIn(
            'id="Group 28" style="margin-top:15.3pt;height:490.5pt"',
            patched,
        )
        self.assertEqual(patched.count("margin-top:15.3pt"), 1)

    def test_uploaded_docx_moves_only_a_max_summary_after_the_details_table(self):
        base = app.app.test_client().post(
            "/generate-docx",
            json={"data": {
                "candidate": {"name": "Summary Anchor Fixture"},
                "work_experiences": [],
                "education": [],
                "certifications": [],
                "skills": [],
            }},
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(base.status_code, 200)

        _, short_xml = self._apply_summary(
            base.data,
            ["A concise summary bullet.", "A second concise bullet."],
            autofit=True,
        )
        self.assertLess(
            short_xml.index("_CVStudioSummaryBox"),
            short_xml.index("<w:tbl>"),
        )

        _, long_xml = self._apply_summary(
            base.data,
            ["Long summary text " * 12 for _ in range(10)],
            autofit=True,
        )
        self.assertLess(
            long_xml.index("<w:tbl>"),
            long_xml.index("_CVStudioSummaryBox"),
        )
        self.assertIn("<wp:posOffset>194310</wp:posOffset>", long_xml)
        self.assertIn("margin-top:15.3pt", long_xml)
        self.assertEqual(long_xml.count("<w:tbl>"), 1)
        self.assertEqual(long_xml.count("_CVStudioSummaryBox"), 2)
        self.assertEqual(
            app._summary_docx_relocate_max_box_after_details_table(long_xml),
            long_xml,
        )

    def test_generated_docx_validator_rejects_invalid_word_xml(self):
        invalid = io.BytesIO()
        with zipfile.ZipFile(invalid, "w") as archive:
            archive.writestr("[Content_Types].xml", "<Types/>")
            archive.writestr(
                "word/document.xml",
                b'<w:document xmlns:w="urn:fixture"><w:body>bad\x0b</w:body></w:document>',
            )
        with self.assertRaisesRegex(RuntimeError, "invalid Word XML"):
            app._validate_generated_docx_bytes(invalid.getvalue())

    def test_generate_docx_multipart_updates_summary_textbox_idempotently(self):
        client = app.app.test_client()
        source_response = client.post(
            "/generate-docx",
            json={"data": {
                "candidate": {"name": "Summary Box Fixture"},
                "work_experiences": [], "education": [],
                "certifications": [], "skills": [],
            }},
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(source_response.status_code, 200)

        first_bytes, first_xml = self._apply_summary(
            source_response.data,
            ["**Cloud platforms** leader.", "Built regional teams."],
        )
        self.assertEqual(first_xml.count("_CVStudioSummaryBox"), 2)
        self.assertEqual(first_xml.count("Cloud platforms"), 2)
        self.assertNotIn("*Summary*", first_xml)
        self.assertNotIn("S U M M A R Y", first_xml)
        self.assertEqual(first_xml.count("<a:spAutoFit/>"), 1)
        self.assertEqual(first_xml.count("mso-fit-shape-to-text:t"), 1)
        self.assertEqual(first_xml.count('<w:spacing w:after="160"/>'), 2)

        _, max_xml = self._apply_summary(
            source_response.data,
            ["Long summary text " * 12 for _ in range(10)],
        )
        self.assertIn("<wp:posOffset>194310</wp:posOffset>", max_xml)
        self.assertIn('<wp:extent cx="5730875" cy="6229350"/>', max_xml)
        self.assertIn("margin-top:15.3pt", max_xml)
        self.assertIn("height:490.5pt", max_xml)

        _, replacement_xml = self._apply_summary(
            first_bytes,
            ["Replacement summary only."],
            name="CV - Summary.docx",
            autofit=False,
        )
        self.assertEqual(replacement_xml.count("_CVStudioSummaryBox"), 2)
        self.assertEqual(replacement_xml.count("Replacement summary only."), 2)
        self.assertNotIn("Cloud platforms", replacement_xml)
        self.assertNotIn("<a:spAutoFit/>", replacement_xml)
        self.assertNotIn("mso-fit-shape-to-text", replacement_xml)

    def test_generate_docx_multipart_replaces_summary_placeholder_and_is_idempotent(self):
        source_response = app.app.test_client().post(
            "/generate-docx",
            json={
                "data": {
                    "candidate": {"name": "Uploaded DOCX Fixture", "languages": "English"},
                    "work_experiences": [],
                    "education": [],
                    "certifications": [],
                    "skills": [],
                }
            },
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(source_response.status_code, 200)

        placeholder_source = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(source_response.data)) as source_archive:
            document_xml = source_archive.read("word/document.xml").decode("utf-8")
            document_xml = document_xml.replace("*Summary*", "*Overview*")
            table_end = document_xml.index("</w:tbl>") + len("</w:tbl>")
            placeholder = (
                "<w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Summary:</w:t></w:r></w:p>"
                "<w:p><w:r><w:t>PLACEHOLDER SUMMARY TEXT</w:t></w:r></w:p>"
            )
            document_xml = document_xml[:table_end] + placeholder + document_xml[table_end:]
            with zipfile.ZipFile(placeholder_source, "w") as patched_archive:
                for info in source_archive.infolist():
                    payload = document_xml.encode("utf-8") if info.filename == "word/document.xml" else source_archive.read(info.filename)
                    patched_archive.writestr(info, payload)
        source_bytes = placeholder_source.getvalue()

        patched_response = app.app.test_client().post(
            "/generate-docx",
            data={
                "source_docx": (io.BytesIO(source_bytes), "Hyppies CV - Fixture.docx"),
                "summary_bullets": json.dumps([
                    "**Cloud platforms** leader across regional delivery.",
                    "Built engineering teams & delivery standards.",
                ]),
            },
            content_type="multipart/form-data",
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(patched_response.status_code, 200)

        with (
            zipfile.ZipFile(io.BytesIO(source_bytes)) as source_archive,
            zipfile.ZipFile(io.BytesIO(patched_response.data)) as patched_archive,
        ):
            self.assertEqual(source_archive.namelist(), patched_archive.namelist())
            for name in source_archive.namelist():
                if name != "word/document.xml":
                    self.assertEqual(source_archive.read(name), patched_archive.read(name), name)
            patched_xml = patched_archive.read("word/document.xml").decode("utf-8")

        first_table = re.search(r"<w:tbl\b.*?</w:tbl>", patched_xml, re.S)
        self.assertIsNotNone(first_table)
        self.assertNotIn("Cloud platforms", first_table.group(0))
        self.assertEqual(patched_xml.count('_CVStudioSummary'), 1)
        self.assertIn("Summary:", patched_xml)
        self.assertNotIn("PLACEHOLDER SUMMARY TEXT", patched_xml)
        self.assertIn("Cloud platforms", patched_xml)

        replacement_response = app.app.test_client().post(
            "/generate-docx",
            data={
                "source_docx": (io.BytesIO(patched_response.data), "Hyppies CV - Fixture - Summary.docx"),
                "summary_bullets": json.dumps(["Replacement summary only."]),
            },
            content_type="multipart/form-data",
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(replacement_response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(replacement_response.data)) as replacement_archive:
            replacement_xml = replacement_archive.read("word/document.xml").decode("utf-8")
        self.assertEqual(replacement_xml.count('_CVStudioSummary'), 1)
        self.assertIn("Replacement summary only.", replacement_xml)
        self.assertNotIn("Cloud platforms", replacement_xml)

    def test_generate_docx_multipart_rejects_non_docx_and_invalid_packages(self):
        client = app.app.test_client()
        non_docx = client.post(
            "/generate-docx",
            data={
                "source_docx": (io.BytesIO(b"not a docx"), "candidate.pdf"),
                "summary_bullets": json.dumps(["Summary"]),
            },
            content_type="multipart/form-data",
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(non_docx.status_code, 400)
        self.assertIn("DOCX files only", non_docx.get_json()["error"])

        invalid_docx = client.post(
            "/generate-docx",
            data={
                "source_docx": (io.BytesIO(b"not a zip package"), "candidate.docx"),
                "summary_bullets": json.dumps(["Summary"]),
            },
            content_type="multipart/form-data",
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(invalid_docx.status_code, 400)
        self.assertIn("valid DOCX", invalid_docx.get_json()["error"])

    def _summary_source_docx(self, extra_body_xml):
        """A minimal generated CV DOCX with extra paragraphs after the table."""
        client = app.app.test_client()
        base = client.post(
            "/generate-docx",
            json={"data": {
                "candidate": {"name": "Summary Regression Fixture"},
                "work_experiences": [], "education": [],
                "certifications": [], "skills": [],
            }},
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(base.status_code, 200)
        out = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(base.data)) as src:
            document_xml = src.read("word/document.xml").decode("utf-8")
            # These cases exercise the legacy/body fallback, so neutralise the
            # modern template's textbox placeholder before adding body fixtures.
            document_xml = document_xml.replace("*Summary*", "*Overview*")
            table_end = document_xml.index("</w:tbl>") + len("</w:tbl>")
            document_xml = document_xml[:table_end] + extra_body_xml + document_xml[table_end:]
            with zipfile.ZipFile(out, "w") as patched:
                for info in src.infolist():
                    payload = (document_xml.encode("utf-8")
                               if info.filename == "word/document.xml"
                               else src.read(info.filename))
                    patched.writestr(info, payload)
        return out.getvalue()

    def _apply_summary(self, source_bytes, bullets, name="CV.docx", autofit=None):
        form = {
            "source_docx": (io.BytesIO(source_bytes), name),
            "summary_bullets": json.dumps(bullets),
        }
        if autofit is not None:
            form["summary_box_autofit"] = "true" if autofit else "false"
        resp = app.app.test_client().post(
            "/generate-docx",
            data=form,
            content_type="multipart/form-data",
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(resp.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(resp.data)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        return resp.data, xml

    def test_generate_docx_summary_replace_keeps_wellformed_word_xml(self):
        # Regression (HIGH): replacing an already-inserted Summary block must yield
        # well-formed OOXML. The block start was located with rfind("<w:p"), which
        # matches <w:pPr>; the replacement then left the paragraph's <w:p> open tag
        # dangling and Word flagged the 2nd-run file as corrupt. The existing
        # idempotency test only checked content/count, not XML validity.
        import xml.etree.ElementTree as ET
        source = self._summary_source_docx(
            "<w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Summary:</w:t></w:r></w:p>"
        )
        first_bytes, _ = self._apply_summary(source, ["First summary bullet."])
        replaced_bytes, replaced_xml = self._apply_summary(
            first_bytes, ["Replacement bullet one.", "Replacement bullet two."],
            name="CV - Summary.docx",
        )
        ET.fromstring(replaced_xml)  # raises ParseError on unbalanced <w:p> tags
        self.assertEqual(replaced_xml.count("_CVStudioSummary"), 1)
        self.assertIn("Replacement bullet one.", replaced_xml)
        self.assertNotIn("First summary bullet.", replaced_xml)

    def test_generate_docx_summary_preserves_following_employment_section(self):
        # Regression (HIGH): the boundary detector must recognise a real-world
        # "Professional Experience" heading; otherwise the Summary span runs past
        # it and the replacement deletes the employment section and job details.
        import xml.etree.ElementTree as ET

        def para(text, bold=False):
            rpr = "<w:rPr><w:b/></w:rPr>" if bold else ""
            return "<w:p><w:r>{}<w:t>{}</w:t></w:r></w:p>".format(rpr, text)

        source = self._summary_source_docx(
            para("Summary", bold=True)
            + para("Old summary line that should be replaced.")
            + para("Professional Experience", bold=True)
            + para("Acme Corp — Senior Engineer (2019-2024)")
            + para("Delivered regional platform migrations.")
            + para("Skills", bold=True)
            + para("Python, Go, Kubernetes")
        )
        _, out_xml = self._apply_summary(source, ["New summary bullet."])
        ET.fromstring(out_xml)
        self.assertIn("New summary bullet.", out_xml)
        self.assertNotIn("Old summary line", out_xml)         # old summary replaced
        self.assertIn("Professional Experience", out_xml)      # employment heading kept
        self.assertIn("Acme Corp", out_xml)                    # job details kept
        self.assertIn("Delivered regional platform migrations.", out_xml)
        self.assertIn("Python, Go, Kubernetes", out_xml)       # skills kept

    def test_generate_docx_summary_uses_complete_heading_not_sentence_prefix(self):
        # Regression (HIGH): prefix matching treated "Experience leading ..." as
        # the next section, leaving the old Summary in place. At the same time,
        # an unrecognised "Career Highlights" section was consumed as Summary.
        import xml.etree.ElementTree as ET

        def para(text, bold=False):
            rpr = "<w:rPr><w:b/></w:rPr>" if bold else ""
            return "<w:p><w:r>{}<w:t>{}</w:t></w:r></w:p>".format(rpr, text)

        source = self._summary_source_docx(
            para("Summary", bold=True)
            + para("Experience leading regional delivery teams.")
            + para("Old second summary sentence.")
            + para("Career Highlights")
            + para("Major achievement that must remain.")
            + para("Skills", bold=True)
            + para("Python, Go, Kubernetes")
        )
        _, out_xml = self._apply_summary(source, ["New summary bullet."])
        ET.fromstring(out_xml)
        self.assertIn("New summary bullet.", out_xml)
        self.assertNotIn("Experience leading regional delivery teams.", out_xml)
        self.assertNotIn("Old second summary sentence.", out_xml)
        self.assertIn("Career Highlights", out_xml)
        self.assertIn("Major achievement that must remain.", out_xml)
        self.assertIn("Python, Go, Kubernetes", out_xml)

    def test_summary_boundary_accepts_decorated_and_letter_spaced_headings(self):
        self.assertTrue(app._summary_docx_is_section_boundary(
            "W O R K   E X P E R I E N C E S __________________"
        ))
        self.assertTrue(app._summary_docx_is_section_boundary("2. Core Expertise:"))
        self.assertFalse(app._summary_docx_is_section_boundary(
            "Skills spanning Python, Go, and Kubernetes."
        ))

    def test_docx_restores_source_project_training_and_omits_untrusted_metadata(self):
        source = """Other Information
[PROJECT INVOLVEMENT HISTORY]:
* Firewalls Configuring & VPN Services
* Network Infrastructure Enhancement & Redesign
[PARTICIPATED TRAINING PROGRAMME]:
* Enhancing Performance Through Teamwork
* Professional Trainer Programme (Train The Trainer)
[SOFT SKILLS]:
* Vendor Management & Negotiation Skills
"""
        parsed = {
            "candidate": {"name": "Kwong Yew Leong"},
            "work_experiences": [],
            "education": [],
            "certifications": [],
            "skills": [
                {
                    "category": "Summary",
                    "items": (
                        "Position: Retrieved Resumes (SiVA folder: Prescreened); "
                        "Date Applied: 30 Mar 2022"
                    ),
                },
                {
                    "category": "Portfolio & Links",
                    "items": "GitHub: https://github.com/unknown",
                },
            ],
        }
        normalized = app._normalize_cv_data_for_output(parsed, source)

        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": normalized},
            headers={"Origin": "http://127.0.0.1:5000"},
        )

        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("Project Involvement History:", document_xml)
        self.assertIn("Network Infrastructure Enhancement &amp; Redesign", document_xml)
        self.assertIn("Participated Training Programme:", document_xml)
        self.assertIn("Professional Trainer Programme (Train The Trainer)", document_xml)
        self.assertNotIn("Retrieved Resumes", document_xml)
        self.assertNotIn("github.com/unknown", document_xml)

    def test_docx_route_omits_placeholder_github_without_source_context(self):
        response = app.app.test_client().post(
            "/generate-docx",
            json={
                "data": {
                    "candidate": {"name": "Source-Free Export"},
                    "work_experiences": [],
                    "education": [],
                    "certifications": [],
                    "skills": [
                        {
                            "category": "Portfolio & Links",
                            "items": "GitHub: https://github.com/unknown",
                        }
                    ],
                }
            },
            headers={"Origin": "http://127.0.0.1:5000"},
        )

        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertNotIn("github.com/unknown", document_xml)
        self.assertNotIn("Portfolio &amp; Links", document_xml)

    def test_prompt_forbids_inferred_titles_and_serialized_groups(self):
        self.assertIn("Never invent, infer, imply, annotate, or explain a title", app.SYSTEM_PROMPT)
        self.assertIn("never as JSON serialized inside a string", app.SYSTEM_PROMPT)

    def test_prompt_keeps_notice_verbatim_cert_dates_and_separate_stints(self):
        # #3 notice period must be copied verbatim, not coerced into a month count.
        self.assertIn("map to candidate.notice_period VERBATIM", app.SYSTEM_PROMPT)
        self.assertNotIn('"notice_period": "X month(s) or empty string"', app.SYSTEM_PROMPT)
        # #4 certification/training dates must be preserved.
        self.assertIn("CERTIFICATION/TRAINING DATES — KEEP THEM", app.SYSTEM_PROMPT)
        # #2 distinct stints at the same employer stay separate entries.
        self.assertIn("SEPARATE STINTS AT THE SAME EMPLOYER", app.SYSTEM_PROMPT)

    def test_prompt_omits_recruitment_metadata_and_forbids_invented_links(self):
        self.assertIn("RECRUITMENT-SYSTEM METADATA — OMIT ENTIRELY", app.SYSTEM_PROMPT)
        self.assertIn("https://github.com/unknown", app.SYSTEM_PROMPT)
        self.assertIn("Project Involvement History", app.SYSTEM_PROMPT)
        self.assertIn("Participated Training Programme", app.SYSTEM_PROMPT)

    def test_prompt_omits_salary_and_remuneration_details(self):
        # Candidates sometimes include current/expected salary or a remuneration
        # section in their source CV; the formatted CV must never carry it over.
        self.assertIn("SALARY / REMUNERATION — OMIT ENTIRELY", app.SYSTEM_PROMPT)
        self.assertIn("expected/asking/desired/target salary", app.SYSTEM_PROMPT)
        # The catch-all extra-section mapping must not re-capture salary content.
        self.assertIn(
            "never apply this catch-all to salary/remuneration/compensation content",
            app.SYSTEM_PROMPT,
        )
        # Notice period stays a first-class candidate field, not salary.
        self.assertIn("Notice period is NOT salary", app.SYSTEM_PROMPT)

    def test_ai_crawler_backend_no_longer_requires_password(self):
        with app.app.test_request_context("/jobadder/spider_options"):
            self.assertTrue(app._ai_crawler_lock_allowed({}))

    def test_company_header_span_recomputed_and_left_candidate_is_last_position(self):
        parsed = {
            "candidate": {"current_position": "Head of Modern Trade", "is_employed": True},
            "work_experiences": [
                {"company": "Mondelez", "date_range": "Jun 2024 to Jul 2026",
                 "roles": [{"title": "Head of Modern Trade", "date_range": ""}]},
                {"company": "Unilever", "date_range": "Jan 2020 to Dec 2019", "roles": [
                    {"title": "Head of Trade Marketing", "date_range": "Jan 2020 to Mar 2021"},
                    {"title": "Head of Sales Operation", "date_range": "Apr 2018 to Dec 2019"},
                    {"title": "National Sales Manager", "date_range": "Dec 2015 to Mar 2018"},
                ]},
            ],
        }
        out = app._order_same_company_roles_newest_first(parsed)
        # #1 backwards/incomplete company header recomputed from roles.
        self.assertEqual(out["work_experiences"][1]["date_range"], "Dec 2015 to Mar 2021")
        # Single-role employer with the date on the entry (not the role) is left as-is.
        self.assertEqual(out["work_experiences"][0]["date_range"], "Jun 2024 to Jul 2026")
        # #3 latest role ends on a concrete past date (no "Present") -> candidate has left.
        self.assertFalse(out["candidate"]["is_employed"])

    def test_present_latest_role_marks_candidate_employed(self):
        parsed = {
            "candidate": {},
            "work_experiences": [
                {"company": "Acme", "date_range": "",
                 "roles": [{"title": "Engineer", "date_range": "Jan 2024 to Present"}]},
            ],
        }
        out = app._order_same_company_roles_newest_first(parsed)
        self.assertTrue(out["candidate"]["is_employed"])

    def test_title_first_header_lines_are_authoritative_over_provider_drift(self):
        # "<Title> — <Company> <DateRange>" lines under an EXPERIENCE heading
        # (date trailing) must be extracted from the source and correct provider
        # hallucinations of employer name, title, and dates.
        cv_text = app._cv_pretranslate_iso_dates(
            "EXPERIENCE\n"
            "Senior Linux System Administrator — Zand Bank 2025-03 – Present\n"
            "AVP Non-Wintel System Administrator — UOB Indonesia 2020-06 – 2025-07\n"
            "Dispatcher Technical Support — PT. Supra Primatama Nusantara (Biznet) 2011-03 – 2013-02\n"
            "IT Support — PT. Elka Prakarsa Utama 2010-01 – 2011-12\n"
            "EDUCATION\n"
            "Bina Nusantara University 2019 – 2021\n"
        )
        rows = app._extract_authoritative_work_rows(cv_text, {})
        # Four work rows; the education line is not captured.
        self.assertEqual(len(rows), 4)
        self.assertEqual(rows[0]["title"], "Senior Linux System Administrator")
        self.assertEqual(rows[0]["company"], "Zand Bank")
        self.assertEqual(rows[1]["date_range"], "Jun 2020 to Jul 2025")
        self.assertEqual(rows[3]["company"], "PT. Elka Prakarsa Utama")

        # Provider hallucinated a different company/title; reconciliation corrects it.
        parsed = {"work_experiences": [
            {"company": "Nusa Network Prakarsa", "date_range": "Feb 2010 to Feb 2011",
             "roles": [{"title": "Customer Care Consultant", "date_range": "", "bullets": ["kept bullet"]}]},
        ]}
        out = app._reconcile_work_experience_with_authoritative_table(parsed, cv_text)
        companies = {e["company"] for e in out["work_experiences"]}
        titles = {r.get("title") for e in out["work_experiences"] for r in e["roles"]}
        self.assertIn("PT. Elka Prakarsa Utama", companies)
        self.assertNotIn("Nusa Network Prakarsa", companies)
        self.assertIn("IT Support", titles)
        self.assertNotIn("Customer Care Consultant", titles)

    def test_company_first_dash_headers_keep_fields_bullets_and_extra_roles(self):
        # "<Company> — <Title> | <DateRange>" headers (the reverse of the Zand
        # Bank layout). The authoritative-row extractor must NOT swap employer and
        # title, and reconciliation must not blank the bullets or drop a role the
        # single-line row regex could not see (multi-line / undated entries).
        cv_text = app._cv_pretranslate_iso_dates(
            "Work Experience\n"
            "TDCX – Product Support Engineer | Mar 2026 – Present\n"
            "Fujitsu Malaysia Sdn. Bhd. – System Engineer | Nov 2023 – Feb 2026\n"
            "Hong Leong Bank – IT Manager (Unix & Storage) | May 2022 – Nov 2023\n"
            "Hewlett-Packard – ITO Service Delivery | Aug 2009 – Dec 2018\n"
            "Education\n"
        )
        parsed = {
            "candidate": {"current_company": "TDCX", "current_position": "Product Support Engineer"},
            "work_experiences": [
                {"company": "TDCX", "date_range": "Mar 2026 to Present",
                 "roles": [{"title": "Product Support Engineer", "date_range": "Mar 2026 to Present", "bullets": ["b1", "b2", "b3"]}]},
                {"company": "Fujitsu Malaysia Sdn. Bhd.", "date_range": "Nov 2023 to Feb 2026",
                 "roles": [{"title": "System Engineer", "date_range": "Nov 2023 to Feb 2026", "bullets": ["b1"]}]},
                {"company": "Hong Leong Bank", "date_range": "May 2022 to Nov 2023",
                 "roles": [{"title": "IT Manager (Unix & Storage)", "date_range": "May 2022 to Nov 2023", "bullets": ["b1"]}]},
                # Present in the parse but NOT as a one-line table row (multi-line
                # in the real CV). Must survive reconciliation, not be dropped.
                {"company": "Toyota Malaysia", "date_range": "Dec 2018 to May 2022",
                 "roles": [{"title": "Resident Infrastructure Engineer", "date_range": "Dec 2018 to May 2022", "bullets": ["b1", "b2"]}]},
                {"company": "Hewlett-Packard", "date_range": "Aug 2009 to Dec 2018",
                 "roles": [{"title": "ITO Service Delivery", "date_range": "Aug 2009 to Dec 2018", "bullets": ["b1"]}]},
            ],
        }

        # Fix A: the extracted rows must not swap title/company.
        rows = app._extract_authoritative_work_rows(cv_text, parsed)
        tdcx_row = next(r for r in rows if app._cv_match_key(r["company"]) == app._cv_match_key("TDCX"))
        self.assertEqual(app._cv_match_key(tdcx_row["title"]), app._cv_match_key("Product Support Engineer"))

        # Fix B: reconciliation preserves every role (incl. the off-table one) with bullets.
        out = app._reconcile_work_experience_with_authoritative_table(
            __import__("json").loads(__import__("json").dumps(parsed)), cv_text)
        companies = [e["company"] for e in out["work_experiences"]]
        self.assertEqual(len(out["work_experiences"]), 5)
        self.assertIn("Toyota Malaysia", companies)
        self.assertEqual(out["work_experiences"][0]["company"], "TDCX")
        self.assertEqual(out["work_experiences"][0]["roles"][0]["title"], "Product Support Engineer")
        self.assertTrue(all(e["roles"][0]["bullets"] for e in out["work_experiences"]))
        self.assertEqual(out["candidate"]["current_company"], "TDCX")
        self.assertEqual(out["candidate"]["current_position"], "Product Support Engineer")

    def test_source_work_subsections_restore_lee_history_without_resorting(self):
        cv_text = (
            "PROFESSIONAL EXPERIENCE\n"
            "Senior Solution Consultant — Delivery & Presales | TrustDecision "
            "(Fintech Fraud & Risk Technology) Jan 2025 – Present\n"
            "● Owned two strategic Philippine banking and digital-wallet deployments end to end, "
            "from discovery through award, implementation and go-\n"
            "live. One was the company's first flagship banking client in that market.\n"
            "● Acted as project manager on a major e-wallet engagement and ran a structured handover.\n"
            "Senior IT Business Analyst / Solution Lead | Huawei Technologies — Asia Pacific "
            "Dec 2019 – Jan 2025\n"
            "● Directed a 2024 transformation portfolio of 21 initiatives.\n"
            "INDEPENDENT CONSULTING & DELIVERY\n"
            "Founder | Movaflow — SME Operations Consultancy, Kuala Lumpur 2026 – Present\n"
            "● Founded an early-stage consultancy focused on operations and process improvement.\n"
            "Consultant & Developer | De Hubs Logistics Sdn Bhd 2025 – 2026\n"
            "● Built a standalone Windows desktop billing application.\n"
            "Consultant & Developer | MySteel Construction Sdn Bhd 2025\n"
            "● Built a mobile project-tracking application for a steelwork design services firm.\n"
            "EARLIER EXPERIENCE\n"
            "Corporate Banking Specialist, Hong Leong Bank Berhad (Jul – Dec 2019) • "
            "Tax Operations Specialist, Royal Bank of Canada (Jan – Jun 2019)\n"
            "EDUCATION & CERTIFICATIONS\n"
        )
        parsed = {
            "candidate": {
                "current_company": "Delivery & Presales | TrustDecision (Fintech Fraud & Risk Technology)",
                "current_position": "Senior Solution Consultant",
            },
            "work_experiences": [
                {
                    "company": "Movaflow — SME Operations Consultancy, Kuala Lumpur",
                    "date_range": "2026 to Present",
                    "roles": [{"title": "Founder", "date_range": "", "bullets": ["Founded an early-stage consultancy focused on operations and process improvement."]}],
                },
                {
                    "company": "Delivery & Presales | TrustDecision (Fintech Fraud & Risk Technology)",
                    "date_range": "Jan 2025 to Present",
                    "roles": [{"title": "Senior Solution Consultant", "date_range": "", "bullets": []}],
                },
                {
                    "company": "De Hubs Logistics Sdn Bhd",
                    "date_range": "2025 to 2026",
                    "roles": [{"title": "Consultant & Developer", "date_range": "", "bullets": ["Built a standalone Windows desktop billing application."]}],
                },
                {
                    "company": "Huawei Technologies — Asia Pacific",
                    "date_range": "Dec 2019 to Jan 2025",
                    "roles": [{"title": "Senior IT Business Analyst / Solution Lead", "date_range": "", "bullets": ["Directed a 2024 transformation portfolio of 21 initiatives."]}],
                },
            ],
            "education": [{
                "date_range": "2018 to 2018",
                "institution": "Tunku Abdul Rahman University",
                "degree": "B.Sc. (Hons) Financial Mathematics",
            }],
        }

        reconciled = app._reconcile_work_experience_with_authoritative_table(parsed, cv_text)
        normalized = app._normalize_cv_data_for_output(reconciled, cv_text)

        work = normalized["work_experiences"]
        self.assertEqual(
            [entry["company"] for entry in work],
            [
                "Delivery & Presales | TrustDecision (Fintech Fraud & Risk Technology)",
                "Huawei Technologies — Asia Pacific",
                "Movaflow — SME Operations Consultancy, Kuala Lumpur",
                "De Hubs Logistics Sdn Bhd",
                "MySteel Construction Sdn Bhd",
                "Hong Leong Bank Berhad",
                "Royal Bank of Canada",
            ],
        )
        self.assertEqual(len(work[0]["roles"][0]["bullets"]), 2)
        self.assertIn("go-live", work[0]["roles"][0]["bullets"][0])
        self.assertEqual(
            work[2]["section_heading"],
            "Independent Consulting & Delivery",
        )
        self.assertEqual(work[4]["date_range"], "2025")
        self.assertEqual(
            work[5]["section_heading"],
            "Earlier Experience",
        )
        self.assertEqual(work[5]["date_range"], "Jul 2019 to Dec 2019")
        self.assertEqual(work[6]["date_range"], "Jan 2019 to Jun 2019")
        self.assertEqual(normalized["education"][0]["date_range"], "2018")
        self.assertNotIn("_work_experience_order_authoritative", normalized)
        normalized_again = app._normalize_cv_data_for_output(normalized, cv_text)
        self.assertEqual(
            [entry["company"] for entry in normalized_again["work_experiences"]],
            [entry["company"] for entry in work],
        )

        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": normalized_again},
            headers={"Origin": "http://127.0.0.1:5000"},
        )
        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        trust_position = document_xml.index("Jan 2025 to Present")
        huawei_position = document_xml.index("Dec 2019 to Jan 2025")
        independent_position = document_xml.index("Independent Consulting &amp; Delivery")
        movaflow_position = document_xml.index("2026 to Present")
        self.assertLess(trust_position, huawei_position)
        self.assertLess(huawei_position, independent_position)
        self.assertLess(independent_position, movaflow_position)
        self.assertIn("Owned two strategic Philippine banking", document_xml)
        self.assertIn("MySteel Construction Sdn Bhd", document_xml)
        self.assertIn("Earlier Experience", document_xml)
        self.assertIn("2018 | Tunku Abdul Rahman University", document_xml)
        self.assertNotIn("2018 to 2018", document_xml)

    def test_generate_docx_preserves_preview_work_order_without_subsection_heading(self):
        response = app.app.test_client().post(
            "/generate-docx",
            json={"data": {
                "candidate": {"name": "Concurrent Work Fixture"},
                "work_experiences": [
                    {
                        "company": "Primary Co",
                        "date_range": "Jan 2025 to Present",
                        "roles": [{"title": "Director", "date_range": "", "bullets": []}],
                    },
                    {
                        "company": "Side Co",
                        "date_range": "2026 to Present",
                        "roles": [{"title": "Founder", "date_range": "", "bullets": []}],
                    },
                ],
                "education": [],
                "certifications": [],
                "skills": [],
            }},
            headers={"Origin": "http://127.0.0.1:5000"},
        )

        self.assertEqual(response.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(response.data)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertLess(document_xml.index("Primary Co"), document_xml.index("Side Co"))

    def test_source_bullet_recovery_stops_at_awards_section(self):
        source = (
            "PROFESSIONAL EXPERIENCE\n"
            "Engineer | Acme Jan 2020 - Dec 2022\n"
            "● Maintained systems.\n"
            "AWARDS\n"
            "● Excellence Award 2019\n"
        )

        rows = app._extract_authoritative_work_rows(source)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].get("source_bullets"), ["Maintained systems."])

    def test_bare_year_programme_prose_does_not_become_an_employer(self):
        source = (
            "PROFESSIONAL EXPERIENCE\n"
            "Engineer | Acme Jan 2020 - Dec 2022\n"
            "● Maintained systems.\n"
            "Manager — Leadership Programme 2023\n"
        )

        rows = app._extract_authoritative_work_rows(source)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["company"], "Acme")
        self.assertEqual(
            rows[0].get("source_bullets"),
            ["Maintained systems."],
        )

    def test_bare_year_programme_prose_is_rejected_inside_work_subsection(self):
        source = (
            "PROFESSIONAL EXPERIENCE\n"
            "INDEPENDENT CONSULTING & DELIVERY\n"
            "Consultant | Acme 2024 - 2025\n"
            "● Delivered an operating model.\n"
            "Manager — Leadership Programme 2023\n"
        )

        rows = app._extract_authoritative_work_rows(source)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["company"], "Acme")
        self.assertEqual(
            rows[0].get("source_bullets"),
            ["Delivered an operating model."],
        )

    def test_letter_spaced_section_heading_stops_source_bullet_recovery(self):
        source = (
            "PROFESSIONAL EXPERIENCE\n"
            "Engineer | Acme Jan 2020 - Dec 2022\n"
            "● Maintained systems.\n"
            "A W A R D S ______\n"
            "● Excellence Award 2019\n"
            "L A N G U A G E S English\n"
        )

        rows = app._extract_authoritative_work_rows(source)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].get("source_bullets"), ["Maintained systems."])

    def test_role_level_achievements_heading_does_not_end_work_history(self):
        source = (
            "PROFESSIONAL EXPERIENCE\n"
            "Engineer | Acme Jan 2020 - Dec 2022\n"
            "● Maintained systems.\n"
            "ACHIEVEMENTS:\n"
            "● Reduced incidents by 20%.\n"
            "Analyst | Beta Jan 2018 - Dec 2019\n"
            "● Produced reports.\n"
            "EDUCATION\n"
        )

        rows = app._extract_authoritative_work_rows(source)

        self.assertEqual([row["company"] for row in rows], ["Acme", "Beta"])
        self.assertEqual(
            rows[0].get("source_bullets"),
            ["Maintained systems.", "Reduced incidents by 20%."],
        )
        self.assertEqual(rows[1].get("source_bullets"), ["Produced reports."])

    def test_top_level_achievements_heading_does_not_enter_last_job(self):
        source = (
            "PROFESSIONAL EXPERIENCE\n"
            "Engineer | Acme Jan 2020 - Dec 2022\n"
            "● Maintained systems.\n"
            "ACHIEVEMENTS:\n"
            "● Excellence Award 2019\n"
        )

        rows = app._extract_authoritative_work_rows(source)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].get("source_bullets"), ["Maintained systems."])

    def test_same_employer_subsection_heading_is_a_hard_merge_boundary(self):
        parsed = {
            "work_experiences": [
                {
                    "company": "Acme",
                    "date_range": "Jan 2024 to Dec 2024",
                    "roles": [{"title": "Employee", "date_range": "", "bullets": []}],
                },
                {
                    "company": "Acme",
                    "date_range": "Jan 2025 to Dec 2025",
                    "section_heading": "Independent Consulting",
                    "roles": [{"title": "Consultant", "date_range": "", "bullets": []}],
                },
            ]
        }

        normalized = app._normalize_cv_data_for_output(parsed)

        self.assertEqual(len(normalized["work_experiences"]), 2)
        self.assertEqual(
            normalized["work_experiences"][1]["section_heading"],
            "Independent Consulting",
        )

    def test_prompt_allows_exact_single_year_work_entries(self):
        self.assertIn(
            '"date_range": "Mon YYYY to Mon YYYY, exact source YYYY, or empty"',
            app.SYSTEM_PROMPT,
        )
        self.assertIn('preserve it as exactly "YYYY"', app.SYSTEM_PROMPT)

    def test_prompt_requires_employer_and_title_fidelity(self):
        self.assertIn("EMPLOYER NAME FIDELITY", app.SYSTEM_PROMPT)
        self.assertIn("ROLE TITLE FIDELITY", app.SYSTEM_PROMPT)

    def test_all_languages_kept_despite_interleaved_two_column_layout(self):
        # LANGUAGES sidebar interleaved with work history by PDF extraction: all
        # three declared languages must survive (not just the first), and be
        # canonicalised. A language with no source line is still dropped.
        cv_text = (
            "LANGUAGES Head of Modern Trade 07/2023 - 05/2024\n"
            "Reckitt\n"
            "Chinese (Professional Working):\n"
            "• Reignited MT business with quarter-to-quarter growth\n"
            "Malay (Professional Working):\n"
            "• Led both MT key account management teams\n"
            "English (Professional Working):\n"
        )
        parsed = {"candidate": {"languages": "Chinese, Malay, English, French"}}
        out = app._normalize_candidate_languages(parsed, cv_text)
        result = out["candidate"]["languages"]
        self.assertIn("English", result)
        self.assertIn("Bahasa Malaysia", result)
        self.assertIn("Chinese", result)
        self.assertNotIn("French", result)  # no source line -> dropped


if __name__ == "__main__":
    unittest.main()
