"""Focused legacy-DOC and coordinate-highlight preview regressions."""

import io
import os
from pathlib import Path
import tempfile
import unittest
from unittest import mock

from docx import Document
from PIL import Image
from owner_build_tools.build_protected import write_test_receipt


ROOT = Path(__file__).resolve().parents[1]
_MODULE_TEMPORARY = tempfile.TemporaryDirectory(prefix="cvstudio-preview-highlight-")
_ORIGINAL_DATABASE_OVERRIDE = os.environ.get("CVSTUDIO_DB_PATH")
os.environ["CVSTUDIO_DB_PATH"] = str(
    Path(_MODULE_TEMPORARY.name) / "state" / "cv_studio.sqlite3"
)
write_test_receipt(ROOT)
try:
    import app
finally:
    if _ORIGINAL_DATABASE_OVERRIDE is None:
        os.environ.pop("CVSTUDIO_DB_PATH", None)
    else:
        os.environ["CVSTUDIO_DB_PATH"] = _ORIGINAL_DATABASE_OVERRIDE


def _converted_docx_bytes():
    output = io.BytesIO()
    document = Document()
    document.add_paragraph("Yan Yen Fen")
    document.add_paragraph("Experience")
    document.add_paragraph(
        "Finance and accounting professional with consolidation, reporting, "
        "audit, tax, reconciliation, and full-set account experience."
    )
    document.add_paragraph("Education")
    document.add_paragraph("Bachelor Degree in Accounting and Finance")
    document.save(output)
    return output.getvalue()


class LegacyDocTrustedPreviewTests(unittest.TestCase):
    def setUp(self):
        self.legacy_doc = bytes.fromhex("d0cf11e0a1b11ae1") + (b"\0" * 2048)
        self.docx = _converted_docx_bytes()

    def test_converter_runs_only_after_completed_antiword_rejection(self):
        converter = mock.Mock(return_value=(self.docx, "microsoft-word"))
        with (
            mock.patch.object(app, "_require_verified_antiword", return_value="verified"),
            mock.patch.object(
                app,
                "_run_verified_antiword",
                return_value=mock.Mock(returncode=1, stdout=b""),
            ),
            mock.patch.object(app, "_convert_legacy_doc_to_docx_bytes", converter),
        ):
            result = app._spider_extract_legacy_doc_for_preview(self.legacy_doc)

        self.assertEqual(result["converter"], "microsoft-word")
        self.assertEqual(result["docx_bytes"], self.docx)
        self.assertIn("Finance and accounting professional", result["text"])
        self.assertIn("converted via microsoft-word", result["source"])
        converter.assert_called_once_with(self.legacy_doc)

    def test_execution_time_trust_failure_never_reaches_converter(self):
        converter = mock.Mock()
        with (
            mock.patch.object(app, "_require_verified_antiword", return_value="verified"),
            mock.patch.object(
                app,
                "_run_verified_antiword",
                side_effect=app.AntiwordDependencyError("runtime-lock-failed"),
            ),
            mock.patch.object(app, "_convert_legacy_doc_to_docx_bytes", converter),
        ):
            with self.assertRaises(app.AntiwordDependencyError) as raised:
                app._spider_extract_legacy_doc_for_preview(self.legacy_doc)

        self.assertEqual(raised.exception.reason, "runtime-lock-failed")
        converter.assert_not_called()

    def test_converted_doc_gets_visual_html_and_search_text_without_libreoffice(self):
        preview_result = {
            "text": "Yan Yen Fen Finance and accounting Experience Education",
            "source": "legacy DOC converted via microsoft-word",
            "converter": "microsoft-word",
            "docx_bytes": self.docx,
        }
        with (
            mock.patch.object(
                app,
                "_spider_extract_legacy_doc_for_preview",
                return_value=preview_result,
            ),
            mock.patch.object(
                app,
                "_office_bytes_to_pdf_preview",
                return_value=(None, "LibreOffice unavailable"),
            ) as office_preview,
        ):
            payload = app._spider_visual_preview_payload(
                self.legacy_doc,
                "application/msword",
                "Yan_YenFen_Selangor_7.doc",
            )

        self.assertEqual(payload["visual_mode"], "html")
        self.assertIn("Yan Yen Fen", payload["html"])
        self.assertIn("Finance and accounting", payload["visual_search_text"])
        self.assertIn("microsoft-word", payload["source"])
        office_preview.assert_called_once_with(
            self.docx,
            ".docx",
            cancel_check=None,
            legacy_doc_gate_complete=True,
        )


class OCRCoordinateLayerTests(unittest.TestCase):
    class _Image:
        size = (1000, 2000)

    class _Output:
        DICT = "dict"

    class _Tesseract:
        Output = None

        @staticmethod
        def image_to_data(image, **kwargs):
            return {
                "text": ["Finance", "accounting", "reporting"],
                "conf": ["96", "94", "91"],
                "left": [100, 260, 100],
                "top": [200, 200, 260],
                "width": [140, 190, 160],
                "height": [40, 40, 40],
                "page_num": [1, 1, 1],
                "block_num": [1, 1, 1],
                "par_num": [1, 1, 1],
                "line_num": [1, 1, 2],
            }

    def test_tesseract_words_become_normalized_visual_boxes(self):
        self._Tesseract.Output = self._Output
        layer = app._spider_ocr_text_layer_from_image(
            self._Image(),
            pytesseract_module=self._Tesseract,
        )

        self.assertTrue(layer["ocr"])
        self.assertEqual(layer["text"], "Finance accounting\nreporting")
        self.assertEqual(len(layer["items"]), 3)
        self.assertEqual(layer["items"][0]["s"], 0)
        self.assertEqual(layer["items"][0]["e"], 7)
        self.assertAlmostEqual(layer["items"][0]["x"], 0.1)
        self.assertAlmostEqual(layer["items"][0]["y"], 0.1)
        self.assertLessEqual(layer["items"][2]["x"] + layer["items"][2]["w"], 1.0)

    def test_pdf_renderer_fallback_adds_ocr_boxes_only_when_requested(self):
        ocr_layer = {
            "text": "Finance accounting",
            "items": [
                {
                    "t": "Finance",
                    "s": 0,
                    "e": 7,
                    "x": 0.1,
                    "y": 0.1,
                    "w": 0.2,
                    "h": 0.05,
                }
            ],
            "truncated": False,
            "ocr": True,
        }

        def rendered_image():
            return Image.new("RGB", (240, 320), "white")

        with (
            mock.patch.object(
                app,
                "_spider_pdfium_preview_pages",
                side_effect=RuntimeError("force fallback"),
            ),
            mock.patch.object(app, "_pdf_page_count", return_value=1),
            mock.patch.object(
                app,
                "_render_pdf_page_images",
                side_effect=lambda *args, **kwargs: [rendered_image()],
            ),
            mock.patch.object(
                app,
                "_spider_ocr_text_layer_from_image",
                return_value=ocr_layer,
            ) as ocr_boxes,
        ):
            foreground = app._spider_pdf_pages_preview_payload(
                b"%PDF-1.7 synthetic",
                "synthetic PDF",
                ocr_missing_text_layers=True,
            )
            background = app._spider_pdf_pages_preview_payload(
                b"%PDF-1.7 synthetic",
                "synthetic PDF",
                ocr_missing_text_layers=False,
            )

        self.assertTrue(foreground["visual_text_layers"][0]["ocr"])
        self.assertEqual(foreground["visual_search_text"], "Finance accounting")
        self.assertIn("through OCR", foreground["note"])
        self.assertIsNone(background["visual_text_layers"][0])
        self.assertEqual(background["visual_search_text"], "")
        ocr_boxes.assert_called_once()


if __name__ == "__main__":
    unittest.main()
