#!/usr/bin/env python3
"""
Hyppies CV Formatter — Local Server
Run: python app.py  (source build) or python app.pyc (source-removed build)
Then open: http://localhost:5000
"""
import time as _boot_time
_BOOT_T0 = _boot_time.time()

# ── Machine-bound installer receipt gate ────────────────────────────
# A valid Authy/TOTP installation creates a signed receipt outside the app
# folder.  Launchers and this backend all verify the same receipt so extracting
# the ZIP and running CV Studio.bat or app.py directly cannot skip first-time
# authorization.  This is still a local deterrent, not server licensing/DRM.
import os as _receipt_os
import sys as _receipt_sys
import json as _receipt_json
import hmac as _receipt_hmac
import hashlib as _receipt_hashlib
import platform as _receipt_platform
import subprocess as _receipt_subprocess
import re as _receipt_re

_INSTALL_RECEIPT_SCHEMA = 2
_INSTALL_RECEIPT_PRODUCT = "TheGuoLab-CVStudio"
_INSTALL_RECEIPT_VERSION = "v24.6.281"
_INSTALL_RECEIPT_MASK = bytes([147, 57, 36, 83, 116, 245, 122, 57, 165, 162, 176, 168, 249, 50, 204, 128, 45, 174, 232, 56])
_INSTALL_RECEIPT_MASKED = bytes([49, 16, 244, 145, 19, 123, 118, 27, 71, 171, 180, 177, 120, 122, 255, 68, 100, 150, 118, 10])


def _install_receipt_totp_secret():
    if len(_INSTALL_RECEIPT_MASK) != len(_INSTALL_RECEIPT_MASKED):
        raise RuntimeError("Installer verifier material is invalid")
    return bytes(a ^ b for a, b in zip(_INSTALL_RECEIPT_MASK, _INSTALL_RECEIPT_MASKED))


def _install_receipt_signing_key():
    secret = bytearray(_install_receipt_totp_secret())
    try:
        return _receipt_hashlib.sha256(bytes(secret) + b"|CVStudio|install-receipt-v1").digest()
    finally:
        for i in range(len(secret)):
            secret[i] = 0


def _install_receipt_path():
    if _receipt_os.name == "nt":
        base = _receipt_os.environ.get("LOCALAPPDATA") or _receipt_os.path.expanduser(r"~\AppData\Local")
        return _receipt_os.path.join(base, "TheGuoLab", "CVStudio", "install_receipt.json")
    return _receipt_os.path.expanduser("~/.guo_lab_cv_studio/install_receipt.json")


def _install_receipt_machine_source():
    if _receipt_os.name == "nt":
        value = ""
        try:
            import winreg as _receipt_winreg
            with _receipt_winreg.OpenKey(
                _receipt_winreg.HKEY_LOCAL_MACHINE,
                r"SOFTWARE\Microsoft\Cryptography",
                0,
                _receipt_winreg.KEY_READ | getattr(_receipt_winreg, "KEY_WOW64_64KEY", 0),
            ) as key:
                value = str(_receipt_winreg.QueryValueEx(key, "MachineGuid")[0] or "").strip()
        except Exception:
            value = ""
        if not value:
            value = str(_receipt_os.environ.get("COMPUTERNAME") or _receipt_platform.node() or "unknown").strip()
        return "windows|" + value.lower()

    if _receipt_sys.platform == "darwin":
        value = ""
        try:
            proc = _receipt_subprocess.run(
                ["ioreg", "-rd1", "-c", "IOPlatformExpertDevice"],
                capture_output=True,
                text=True,
                timeout=4,
                check=False,
            )
            match = _receipt_re.search(r'"IOPlatformUUID"\s*=\s*"([^"]+)"', proc.stdout or "", _receipt_re.I)
            if match:
                value = match.group(1).strip()
        except Exception:
            value = ""
        if not value:
            value = str(_receipt_platform.node() or "unknown").strip()
        return "macos|" + value.lower()

    value = ""
    for candidate in ("/etc/machine-id", "/var/lib/dbus/machine-id"):
        try:
            with open(candidate, "r", encoding="utf-8") as handle:
                value = handle.read().strip()
            if value:
                break
        except Exception:
            continue
    if not value:
        value = str(_receipt_platform.node() or "unknown").strip()
    return "linux|" + value.lower()


def _install_receipt_machine_hash():
    return _receipt_hashlib.sha256(_install_receipt_machine_source().encode("utf-8")).hexdigest()


def _install_package_root():
    """Return the outer cv_formatter root for source and Nuitka standalone builds."""
    here = _receipt_os.path.realpath(_receipt_os.path.dirname(_receipt_os.path.abspath(__file__)))
    if _receipt_os.path.basename(here).lower() == "native" and _receipt_os.path.basename(_receipt_os.path.dirname(here)).lower() == "runtime":
        here = _receipt_os.path.dirname(_receipt_os.path.dirname(here))
    return _receipt_os.path.normcase(_receipt_os.path.realpath(here))


def _install_package_root_hash():
    return _receipt_hashlib.sha256(_install_package_root().encode("utf-8", errors="surrogatepass")).hexdigest()


def _install_receipt_message(data):
    return "{}|{}|{}|{}|{}|{}".format(
        int(data.get("schema") or 0),
        str(data.get("product") or ""),
        str(data.get("machine") or ""),
        str(data.get("version") or ""),
        str(data.get("rootHash") or ""),
        str(data.get("issuedAt") or ""),
    )


def _install_receipt_status():
    path = _install_receipt_path()
    try:
        with open(path, "r", encoding="utf-8-sig") as handle:
            data = _receipt_json.load(handle)
    except FileNotFoundError:
        return False, "Installation receipt is missing", path
    except Exception as exc:
        return False, "Installation receipt could not be read: {}".format(str(exc)[:160]), path
    if not isinstance(data, dict):
        return False, "Installation receipt has an invalid format", path
    if int(data.get("schema") or 0) != _INSTALL_RECEIPT_SCHEMA:
        return False, "Installation receipt schema is unsupported", path
    if str(data.get("product") or "") != _INSTALL_RECEIPT_PRODUCT:
        return False, "Installation receipt belongs to another product", path
    if not _receipt_hmac.compare_digest(str(data.get("machine") or ""), _install_receipt_machine_hash()):
        return False, "Installation receipt belongs to another computer", path
    if str(data.get("version") or "") != _INSTALL_RECEIPT_VERSION:
        return False, "Installation receipt belongs to another CV Studio version", path
    if not _receipt_hmac.compare_digest(str(data.get("rootHash") or ""), _install_package_root_hash()):
        return False, "Installation receipt belongs to another extracted folder", path
    signature = str(data.get("signature") or "").lower()
    expected = _receipt_hmac.new(
        _install_receipt_signing_key(),
        _install_receipt_message(data).encode("utf-8"),
        _receipt_hashlib.sha256,
    ).hexdigest()
    if not signature or not _receipt_hmac.compare_digest(signature, expected):
        return False, "Installation receipt signature is invalid", path
    return True, "Installation receipt is valid", path


def _enforce_install_receipt_or_exit():
    ok, reason, path = _install_receipt_status()
    if ok:
        return
    message = (
        "CV Studio has not been authorized on this computer. "
        "Run INSTALL.bat (Windows) or install.sh (macOS), enter the current Authy code, "
        "then launch CV Studio again. [{}] Receipt: {}"
    ).format(reason, path)
    try:
        _receipt_sys.stderr.write(message + "\n")
    except Exception:
        pass
    try:
        root = _receipt_os.path.dirname(_receipt_os.path.abspath(__file__))
        with open(_receipt_os.path.join(root, "startup_authorization_error.txt"), "w", encoding="utf-8") as handle:
            handle.write(message + "\n")
    except Exception:
        pass
    raise SystemExit(13)


if "--check-install-receipt" in _receipt_sys.argv:
    _receipt_ok, _receipt_reason, _receipt_file = _install_receipt_status()
    print(("OK: " if _receipt_ok else "DENIED: ") + _receipt_reason)
    raise SystemExit(0 if _receipt_ok else 13)

_enforce_install_receipt_or_exit()

# A restarted Windows process is spawned before the old process releases port
# 5000. Honour a short owner-controlled delay so the replacement binds only
# after the previous runtime exits. The variable is removed immediately so it
# cannot leak into child tools.
_restart_delay_raw = _receipt_os.environ.pop("CVSTUDIO_RESTART_DELAY_SECONDS", "")
if _restart_delay_raw:
    try:
        _boot_time.sleep(max(0.0, min(10.0, float(_restart_delay_raw))))
    except Exception:
        pass

from flask import request, jsonify, send_from_directory, send_file, g, after_this_request
import urllib.request, urllib.error, urllib.parse, json, os, subprocess, tempfile, traceback, sys, re, socket, ssl, ast, platform, base64, hashlib, threading, atexit, logging, secrets, plistlib, zipfile, html, functools
try:
    import certifi
except Exception:
    certifi = None

def safe_json(raw, fallback=None):
    """Parse JSON bytes/str safely, returning fallback on empty or invalid input."""
    if fallback is None:
        fallback = {}
    if not raw:
        return fallback
    text = raw.decode() if isinstance(raw, bytes) else raw
    if not text.strip():
        return fallback
    try:
        return json.loads(text)
    except Exception:
        return fallback
import io, time
from pathlib import Path
from cvstudio_storage import (
    BROWSER_SETTING_KEYS,
    BrowserSettingsRepository,
    CVStudioStorage,
    LeadContactCacheRepository,
    LeadTitleCacheRepository,
    OneNoteSavedLinkRepository,
    OneNoteTransferRepository,
    PPCMetadataRepository,
    SalaryComponentCacheRepository,
    StorageCorruptionError,
    StorageError,
    StorageMigrationError,
    UsageHistoryRepository,
)
from cvstudio_clients import (
    AIProviderClient,
    ExternalServiceError,
    ExternalServiceHTTPError,
    JobAdderClient,
    MicrosoftGraphClient,
)
from cvstudio_jobs import (
    PersistentJobError,
    PersistentJobStore,
    deterministic_job_id,
)
from cvstudio_ai_costs import (
    MODEL_PRICING_USD_PER_MILLION as _PHASE5B_MODEL_PRICING,
    cost_details as _phase5b_cost_details,
    enforce_request_guardrail as _phase5b_enforce_request_guardrail,
    extract_provider_billing_result as _phase5b_extract_provider_billing_result,
    merge_usage as _phase5b_merge_usage,
    normalize_provider as _phase5b_normalize_provider,
    normalize_usage as _phase5b_normalize_usage,
    paid_failure_reconciliation as _phase5b_paid_failure_reconciliation,
    pricing_for_model as _phase5b_pricing_for_model,
    unavailable_external_billing as _phase5b_unavailable_external_billing,
    usage_int as _phase5b_usage_int,
)
from cvstudio_storage_bridge import (
    StorageBridge,
    phase2a_ppc_metadata as _phase4_phase2a_ppc_metadata,
    phase2a_usage_records as _phase4_phase2a_usage_records,
    phase2a_usage_safe_value as _phase4_phase2a_usage_safe_value,
    phase2a_usage_secret_field as _phase4_phase2a_usage_secret_field,
    phase2b_browser_setting_keys as _phase4_phase2b_browser_setting_keys,
    phase2b_browser_settings as _phase4_phase2b_browser_settings,
    phase2b_record_array as _phase4_phase2b_record_array,
)
from cvstudio_diagnostics import (
    DiagnosticsService,
    dependency_status as _phase4_dependency_status,
    redact_support_text as _phase4_redact_support_text,
    sanitize_browser_diagnostics as _phase4_sanitize_browser_diagnostics,
    system_memory_status as _phase4_system_memory_status,
)
from cvstudio_document_safety import (
    MAX_IMAGE_PIXELS as _PHASE4_MAX_IMAGE_PIXELS,
    MAX_OCR_PAGES as _PHASE4_MAX_OCR_PAGES,
    MAX_PDF_PAGES as _PHASE4_MAX_PDF_PAGES,
    MAX_ZIP_ENTRIES as _PHASE4_MAX_ZIP_ENTRIES,
    MAX_ZIP_EXPANDED_BYTES as _PHASE4_MAX_ZIP_EXPANDED_BYTES,
    MAX_ZIP_SINGLE_ENTRY_BYTES as _PHASE4_MAX_ZIP_SINGLE_ENTRY_BYTES,
    OCR_TOTAL_DEADLINE_SECONDS as _PHASE4_OCR_TOTAL_DEADLINE_SECONDS,
    document_validation_status as _phase4_document_validation_status,
    ocr_image_text as _phase4_ocr_image_text,
    ocr_pdf_pagewise as _phase4_ocr_pdf_pagewise,
    pdf_page_count as _phase4_pdf_page_count,
    render_pdf_page_images as _phase4_render_pdf_page_images,
    safe_image_open as _phase4_safe_image_open,
    validate_zip_payload as _phase4_validate_zip_payload,
)
from cvstudio_antiword import (
    AntiwordDependencyError,
    antiword_health as _verified_antiword_health,
    antiword_subprocess_env as _verified_antiword_subprocess_env,
    find_verified_antiword as _verified_antiword_finder,
    require_verified_antiword as _require_verified_antiword_runtime,
    run_verified_antiword as _run_verified_antiword_runtime,
)
from cvstudio_tesseract import (
    find_tesseract as _find_mandatory_tesseract,
    tesseract_health as _mandatory_tesseract_health,
)
from cvstudio_architecture import (
    create_modular_monolith_app as _create_modular_monolith_app,
    finalize_modular_monolith_app as _finalize_modular_monolith_app,
)
from cvstudio_startup import StartupService
from cvstudio_runtime import RuntimeService
from cvstudio_web_assets import WebAssetsService
from cvstudio_secrets import SecretsService
from cvstudio_jobadder_read import JobAdderReadService
from cvstudio_jobadder_write import JobAdderWriteService

_CVSTUDIO_VERSION = "v24.6.281"
_CVSTUDIO_ROOT = _install_package_root()
_CVSTUDIO_ROOT_HASH = hashlib.sha256(_CVSTUDIO_ROOT.encode("utf-8", errors="surrogatepass")).hexdigest()
_CVSTUDIO_INSTANCE_ID = _CVSTUDIO_ROOT_HASH[:24]

_CVSTUDIO_STORAGE = CVStudioStorage()
_CVSTUDIO_USAGE_REPOSITORY = UsageHistoryRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_LEAD_TITLE_REPOSITORY = LeadTitleCacheRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_LEAD_CONTACT_REPOSITORY = LeadContactCacheRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_SALARY_REPOSITORY = SalaryComponentCacheRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_PPC_REPOSITORY = PPCMetadataRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_ONENOTE_TRANSFER_REPOSITORY = OneNoteTransferRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_ONENOTE_LINK_REPOSITORY = OneNoteSavedLinkRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_BROWSER_SETTINGS_REPOSITORY = BrowserSettingsRepository(_CVSTUDIO_STORAGE)
_CVSTUDIO_STORAGE_STARTUP_ERROR = None
try:
    _CVSTUDIO_STORAGE.initialize()
except StorageError as _storage_startup_error:
    # Keep non-storage routes available. Affected routes and diagnostics return
    # the structured request-ID recovery contract once Flask is initialised.
    _CVSTUDIO_STORAGE_STARTUP_ERROR = _storage_startup_error

_CVSTUDIO_JOBS = PersistentJobStore()
_CVSTUDIO_JOBS_STARTUP_ERROR = None
try:
    _CVSTUDIO_JOBS.initialize()
except PersistentJobError as _jobs_startup_error:
    # Keep unrelated routes available. Persistent-job operations surface the
    # same failure through the structured request-ID recovery contract.
    _CVSTUDIO_JOBS_STARTUP_ERROR = _jobs_startup_error


def _cvstudio_legacy_json_read(path, expected_type):
    """Read a transition JSON store without renaming, deleting or rewriting it."""
    try:
        with open(path, "rb") as handle:
            raw = handle.read()
        data = json.loads(raw.decode("utf-8-sig"))
        if not isinstance(data, expected_type):
            return None, ""
        return data, hashlib.sha256(raw).hexdigest()
    except Exception:
        return None, ""


def _cvstudio_legacy_json_write(path, data, *, indent=None):
    """Keep the v24.6.217 JSON shape backward-readable for one release."""
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    tmp_path = path + ".tmp"
    with open(tmp_path, "w", encoding="utf-8") as handle:
        json.dump(data, handle, ensure_ascii=False, indent=indent)
    os.replace(tmp_path, path)


def _resolve_cvstudio_port():
    """Use port 5000 normally; allow an owner-build-only ephemeral smoke port."""
    smoke_mode = str(os.environ.pop("CVSTUDIO_OWNER_BUILD_SMOKE", "") or "").strip() == "1"
    raw_port = str(os.environ.pop("CVSTUDIO_OWNER_BUILD_SMOKE_PORT", "") or "").strip()
    if not smoke_mode or not raw_port:
        return 5000
    try:
        port = int(raw_port)
    except (TypeError, ValueError):
        return 5000
    return port if 1024 <= port <= 65535 else 5000


_CVSTUDIO_PORT = _resolve_cvstudio_port()
_CVSTUDIO_BASE_URL = "http://localhost:{}".format(_CVSTUDIO_PORT)

# Native Windows no-console builds have no stdout/stderr handles. Route Flask,
# Werkzeug and traceback output to a private, continuously bounded local log.
_RUNTIME_STREAM_LOG = None
_RUNTIME_LOG_MAX_BYTES = 5 * 1024 * 1024
_RUNTIME_LOG_BACKUPS = 1
_RUNTIME_STATE_DIR = os.path.join(
    os.environ.get("LOCALAPPDATA") or os.path.expanduser(r"~\AppData\Local"),
    "TheGuoLab", "CVStudio",
)
_RUNTIME_LOG_PATH = os.environ.get("CVSTUDIO_RUNTIME_LOG_PATH") or os.path.join(_RUNTIME_STATE_DIR, "runtime.log")
_RUNTIME_PID_PATH = os.path.join(_RUNTIME_STATE_DIR, "cvstudio.{}.pid.json".format(_CVSTUDIO_INSTANCE_ID))
_RUNTIME_LEGACY_PID_PATH = os.path.join(_RUNTIME_STATE_DIR, "cvstudio.pid.json")


def _rotate_runtime_log(path):
    """Keep runtime diagnostics bounded to one active file plus one backup."""
    try:
        if os.path.isfile(path) and os.path.getsize(path) >= _RUNTIME_LOG_MAX_BYTES:
            backup = path + ".1"
            try:
                if os.path.exists(backup):
                    os.remove(backup)
            except Exception:
                pass
            os.replace(path, backup)
    except Exception:
        pass


class _CVStudioRotatingTextStream:
    """Line-buffered stdout/stderr stream that stays bounded during long sessions."""
    encoding = "utf-8"
    errors = "replace"
    _TRUNCATION_MARKER = "[CV Studio log entry truncated to keep runtime.log bounded]\n"

    def __init__(self, path):
        self.path = path
        self._lock = threading.RLock()
        self._handle = None
        self._bytes_written = 0
        self._open()

    def _open(self):
        os.makedirs(os.path.dirname(self.path), exist_ok=True)
        _rotate_runtime_log(self.path)
        self._handle = open(self.path, "a", encoding="utf-8", buffering=1)
        try:
            self._bytes_written = os.path.getsize(self.path)
        except Exception:
            self._bytes_written = 0

    def _maybe_rotate(self, incoming=0):
        try:
            incoming = max(0, int(incoming or 0))
            if self._bytes_written + incoming < _RUNTIME_LOG_MAX_BYTES:
                return
            try:
                self._handle.flush()
            except Exception:
                pass
            try:
                self._handle.close()
            except Exception:
                pass
            # Rotate before the next write even when the active file is still
            # just below the threshold; otherwise the incoming write would
            # push it beyond the configured bound.
            try:
                if os.path.isfile(self.path) and os.path.getsize(self.path) > 0:
                    backup = self.path + ".1"
                    if os.path.exists(backup):
                        os.remove(backup)
                    os.replace(self.path, backup)
            except Exception:
                pass
            self._open()
        except Exception:
            pass

    def _redact_sensitive_text(self, text):
        # Runtime diagnostics must not become a second store for candidate PII
        # or integration credentials. Keep enough structure for debugging.
        text = re.sub(r'(?i)(authorization\s*[:=]\s*bearer\s+)[^\s,;]+', r'\1[redacted]', text)
        text = re.sub(
            r'(?i)(["\']?(?:access_token|refresh_token|client_secret|api_key|authorization_code)["\']?\s*[:=]\s*["\']?)[^"\'\s,}]+',
            r'\1[redacted]',
            text,
        )
        text = re.sub(r'(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b', '[email redacted]', text)
        return text

    def _bounded_text(self, text):
        text = self._redact_sensitive_text(text)
        encoded = text.encode("utf-8", errors="replace")
        if len(encoded) < _RUNTIME_LOG_MAX_BYTES:
            return text, len(encoded)
        # A single untrusted traceback/request line must never defeat rotation.
        marker = self._TRUNCATION_MARKER.encode("utf-8")
        available = max(0, _RUNTIME_LOG_MAX_BYTES - len(marker) - 1)
        tail_limit = min(max(0, _RUNTIME_LOG_MAX_BYTES // 2), available)
        tail = encoded[-tail_limit:].decode("utf-8", errors="replace") if tail_limit else ""
        bounded = self._TRUNCATION_MARKER + tail
        bounded_bytes = bounded.encode("utf-8", errors="replace")
        if len(bounded_bytes) >= _RUNTIME_LOG_MAX_BYTES:
            bounded_bytes = bounded_bytes[:max(0, _RUNTIME_LOG_MAX_BYTES - 1)]
            bounded = bounded_bytes.decode("utf-8", errors="ignore")
            bounded_bytes = bounded.encode("utf-8", errors="replace")
        return bounded, len(bounded_bytes)

    def write(self, value):
        text, byte_count = self._bounded_text(str(value or ""))
        with self._lock:
            self._maybe_rotate(byte_count)
            result = self._handle.write(text)
            self._bytes_written += byte_count
            return result

    def flush(self):
        with self._lock:
            try:
                self._handle.flush()
            except Exception:
                pass

    def close(self):
        with self._lock:
            try:
                self._handle.flush()
            except Exception:
                pass
            try:
                self._handle.close()
            except Exception:
                pass

    def isatty(self): return False
    def fileno(self): return self._handle.fileno()


_RUNTIME_LOG_ENV_REQUESTED = bool(str(os.environ.get("CVSTUDIO_RUNTIME_LOG_PATH") or "").strip())
_RUNTIME_NO_CONSOLE = bool(
    os.name == "nt"
    and (getattr(sys, "stdout", None) is None or getattr(sys, "stderr", None) is None)
)
if _RUNTIME_LOG_ENV_REQUESTED or _RUNTIME_NO_CONSOLE:
    try:
        _RUNTIME_STREAM_LOG = _CVStudioRotatingTextStream(_RUNTIME_LOG_PATH)
        # macOS launchers explicitly opt in so long-running native/source runs
        # receive the same live rotation and redaction as Windows no-console
        # builds. On Windows without the opt-in, replace only missing handles.
        if _RUNTIME_LOG_ENV_REQUESTED or getattr(sys, "stdout", None) is None:
            sys.stdout = _RUNTIME_STREAM_LOG
        if _RUNTIME_LOG_ENV_REQUESTED or getattr(sys, "stderr", None) is None:
            sys.stderr = _RUNTIME_STREAM_LOG
    except Exception:
        _RUNTIME_STREAM_LOG = None


class _CVStudioAccessLogFilter(logging.Filter):
    """Suppress noisy probes and redact query strings from Werkzeug logs."""
    def filter(self, record):
        try:
            message = record.getMessage()
            if re.search(r'"(?:GET|POST) /(?:ping|heartbeat)(?:[ ?])', message):
                return False
            message = re.sub(
                r'("[A-Z]+ )([^ ?"]+)\?[^ ]+( HTTP/\d(?:\.\d)?")',
                r'\1\2?[redacted]\3',
                message,
            )
            record.msg = message
            record.args = ()
        except Exception:
            pass
        return True


def _write_runtime_pid_file():
    """Record this exact package instance so STOP never targets another build."""
    if os.name != "nt":
        return
    try:
        os.makedirs(_RUNTIME_STATE_DIR, exist_ok=True)
        payload = {
            "pid": os.getpid(),
            "root": _CVSTUDIO_ROOT,
            "root_hash": _CVSTUDIO_ROOT_HASH,
            "instance_id": _CVSTUDIO_INSTANCE_ID,
            "executable": os.path.abspath(str(sys.executable or "")),
            "entry": os.path.abspath(str(sys.argv[0] or "")),
            "version": _CVSTUDIO_VERSION,
        }
        tmp = _RUNTIME_PID_PATH + ".tmp"
        with open(tmp, "w", encoding="utf-8") as handle:
            json.dump(payload, handle, separators=(",", ":"))
        os.replace(tmp, _RUNTIME_PID_PATH)
        # Delete only a legacy PID file that already identifies this same process.
        try:
            if os.path.exists(_RUNTIME_LEGACY_PID_PATH):
                with open(_RUNTIME_LEGACY_PID_PATH, "r", encoding="utf-8-sig") as legacy:
                    old = json.load(legacy)
                if int(old.get("pid") or 0) == os.getpid():
                    os.remove(_RUNTIME_LEGACY_PID_PATH)
        except Exception:
            pass
    except Exception:
        pass


def _remove_runtime_pid_file():
    if os.name != "nt":
        return
    try:
        with open(_RUNTIME_PID_PATH, "r", encoding="utf-8-sig") as handle:
            payload = json.load(handle)
        if int(payload.get("pid") or 0) == os.getpid() and payload.get("instance_id") == _CVSTUDIO_INSTANCE_ID:
            os.remove(_RUNTIME_PID_PATH)
    except Exception:
        pass


_write_runtime_pid_file()
atexit.register(_remove_runtime_pid_file)

import xml.etree.ElementTree as ET
from datetime import date, datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
from collections import OrderedDict
import calendar

# ── Heartbeat / auto-shutdown / keep-alive ───────────────────────────
_last_heartbeat   = time.time()
_last_activity    = time.time()   # updated on any API call too
_HEARTBEAT_TIMEOUT = 0  # Disabled: browser heartbeats are diagnostic only.
_shutdown_armed   = False         # becomes True once the first heartbeat arrives

def _watchdog():
    """
    Background maintenance thread.

    Browser-heartbeat shutdown is intentionally disabled while
    _HEARTBEAT_TIMEOUT is zero. A lightweight self-ping runs every 120 seconds
    to keep the local process responsive during long idle periods.
    """
    global _shutdown_armed
    _self_ping_interval = 120  # seconds between self-pings
    _last_self_ping = time.time()

    while True:
        time.sleep(2)
        now = time.time()

        # ── Job 1: shutdown if browser gone ──────────────────────────
        if _HEARTBEAT_TIMEOUT > 0 and _shutdown_armed and (now - _last_heartbeat) > _HEARTBEAT_TIMEOUT:
            os._exit(0)

        # ── Job 2: self-ping to stay alive ───────────────────────────
        if now - _last_self_ping >= _self_ping_interval:
            _last_self_ping = now
            try:
                req = urllib.request.Request(
                    _CVSTUDIO_BASE_URL + "/ping",
                    method="GET"
                )
                urllib.request.urlopen(req, timeout=3)
            except Exception:
                pass  # if it fails we're probably shutting down anyway

threading.Thread(target=_watchdog, daemon=True).start()

app = _create_modular_monolith_app(__name__, static_folder=None)
logging.getLogger("werkzeug").addFilter(_CVStudioAccessLogFilter())


def _cvstudio_current_request_id():
    try:
        return str(getattr(g, "cvstudio_request_id", "") or "")
    except Exception:
        return ""


def _cvstudio_safe_error_message(value, fallback="CV Studio request failed"):
    """Return bounded user-facing error text without leaking bearer credentials."""
    message = str(value or fallback).strip() or fallback
    message = re.sub(r"(?i)(authorization\s*[:=]\s*bearer\s+)[^\s,;]+", r"\1[redacted]", message)
    message = re.sub(r"(?i)([?&](?:access_token|refresh_token|client_secret|api_key|code|password)=)[^&\s]+", r"\1[redacted]", message)
    message = re.sub(r"(?i)(access_token|refresh_token|client_secret|api_key|password)(\s*[:=]\s*)[^\s,;}]+", r"\1\2[redacted]", message)
    return message[:2000]


def _cvstudio_classify_error(status, message, path=""):
    """Map legacy route failures into a stable machine-readable contract."""
    status = int(status or 500)
    low = str(message or "").lower()
    path = str(path or "").lower()
    code = "CVSTUDIO_ERROR"
    retryable = False
    action = ""

    if status == 404:
        return "NOT_FOUND", False, ""
    if status == 405:
        return "METHOD_NOT_ALLOWED", False, ""
    if status == 413:
        return "REQUEST_TOO_LARGE", False, "use_smaller_file"
    if "ai cost guardrail configuration" in low:
        return "AI_COST_GUARDRAIL_CONFIG_INVALID", False, "review_ai_cost_limit"
    if "ai cost guardrail blocked" in low:
        return "AI_COST_GUARDRAIL_BLOCKED", False, "review_ai_cost_limit"
    if "storage database is corrupt" in low:
        return "STORAGE_CORRUPT", False, "restore_storage_backup"
    if "storage migration" in low:
        return "STORAGE_MIGRATION_FAILED", False, "restore_storage_backup"
    if "local storage is busy" in low:
        return "STORAGE_BUSY", True, "retry"
    if "local storage" in low and "newer cv studio" in low:
        return "STORAGE_VERSION_UNSUPPORTED", False, "restore_storage_backup"
    if "local storage" in low and "unavailable" in low:
        return "STORAGE_UNAVAILABLE", True, "check_storage_access"
    if status == 429 or "rate limit" in low or "too many requests" in low:
        return "RATE_LIMITED", True, "retry_later"
    if status in (408, 504) or "timed out" in low or "timeout" in low:
        return "REQUEST_TIMEOUT", True, "retry"
    if "superseded" in low or "stale request" in low:
        return "REQUEST_SUPERSEDED", False, ""
    if "locked" in low or "unlock" in low:
        return "FEATURE_LOCKED", False, "unlock_feature"
    if "jobadder" in path or "jobadder" in low:
        if status in (401, 403) or "not authenticated" in low or "reconnect" in low or "invalid_grant" in low:
            return "JOBADDER_RECONNECT_REQUIRED", False, "open_jobadder_settings"
        code = "JOBADDER_REQUEST_FAILED"
        retryable = status >= 500
        action = "retry" if retryable else ""
    elif "onenote" in path or "onenote" in low:
        if status in (401, 403) or "not connected" in low or "reconnect" in low:
            return "ONENOTE_RECONNECT_REQUIRED", False, "open_onenote_settings"
        code = "ONENOTE_REQUEST_FAILED"
        retryable = status >= 500
        action = "retry" if retryable else ""
    elif "outlook" in path or "outlook" in low:
        if status in (401, 403) or "not connected" in low or "reconnect" in low:
            return "OUTLOOK_RECONNECT_REQUIRED", False, "open_outlook_settings"
        code = "OUTLOOK_REQUEST_FAILED"
        retryable = status >= 500
        action = "retry" if retryable else ""
    elif "deepseek" in low or "anthropic" in low or "openai" in low or "api provider" in low:
        if status in (401, 403) or "api key" in low or "unauthorized" in low or "authentication" in low:
            return "AI_PROVIDER_AUTH_REQUIRED", False, "open_ai_settings"
        if status == 402 or "insufficient balance" in low or "insufficient_quota" in low or "exceeded your current quota" in low:
            return "AI_PROVIDER_INSUFFICIENT_BALANCE", False, "top_up_ai_balance"
        if status == 422 or "invalid model" in low or "model not found" in low or "unknown model" in low:
            return "AI_PROVIDER_INVALID_MODEL", False, "reset_ai_model"
        code = "AI_PROVIDER_REQUEST_FAILED"
        retryable = status >= 429
        action = "retry_or_switch_provider" if retryable else "open_ai_settings"
    elif status == 401:
        code, retryable, action = "AUTHENTICATION_REQUIRED", False, "open_settings"
    elif status == 403:
        code, retryable, action = "REQUEST_FORBIDDEN", False, ""
    elif status >= 500:
        code, retryable, action = "INTERNAL_SERVER_ERROR", True, "retry"
    elif status >= 400:
        code, retryable, action = "INVALID_REQUEST", False, ""
    return code, retryable, action


def _cvstudio_normalize_error_data(data, status=500, path=""):
    """Normalise any legacy JSON failure while retaining all original fields."""
    source = dict(data) if isinstance(data, dict) else {}
    message = _cvstudio_safe_error_message(source.get("message") or source.get("error") or source.get("detail"))
    inferred_code, inferred_retryable, inferred_action = _cvstudio_classify_error(status, message, path)
    source["ok"] = False
    source["error"] = message                 # legacy readers
    source["message"] = message               # v24.6.217 contract
    source["code"] = str(source.get("code") or inferred_code)
    source["retryable"] = bool(source.get("retryable", inferred_retryable))
    source["request_id"] = str(source.get("request_id") or _cvstudio_current_request_id())
    source["severity"] = str(source.get("severity") or ("warning" if source["retryable"] or int(status or 500) < 500 else "error"))
    if not source.get("action") and inferred_action:
        source["action"] = inferred_action
    return source


def _cvstudio_error_payload(code, message, status=400, *, retryable=False, action="", details=None, extra=None):
    """Return the stable local API error shape without breaking legacy readers."""
    payload = {
        "ok": False,
        "error": _cvstudio_safe_error_message(message),
        "message": _cvstudio_safe_error_message(message),
        "code": str(code or "CVSTUDIO_ERROR"),
        "retryable": bool(retryable),
        "request_id": _cvstudio_current_request_id(),
        "severity": "warning" if retryable or int(status or 400) < 500 else "error",
    }
    if action:
        payload["action"] = str(action)
    if details not in (None, ""):
        payload["details"] = details
    if isinstance(extra, dict):
        payload.update(extra)
    return jsonify(payload), int(status or 400)


@app.errorhandler(StorageError)
def _cvstudio_storage_error(error):
    return _cvstudio_error_payload(
        getattr(error, "code", "STORAGE_UNAVAILABLE"),
        getattr(error, "public_message", "Local storage is unavailable."),
        503,
        retryable=bool(getattr(error, "retryable", False)),
        action=getattr(error, "recovery_action", "restore_storage_backup"),
        details={
            "recovery": getattr(error, "public_message", "Local storage is unavailable.")
        },
    )


@app.errorhandler(PersistentJobError)
def _cvstudio_persistent_job_error(error):
    return _cvstudio_error_payload(
        getattr(error, "code", "JOB_STATE_UNAVAILABLE"),
        getattr(
            error,
            "public_message",
            "Persistent background task state is unavailable.",
        ),
        int(getattr(error, "http_status", 503) or 503),
        retryable=bool(getattr(error, "retryable", False)),
        action=getattr(error, "recovery_action", "retry"),
        details={
            "recovery": getattr(
                error,
                "public_message",
                "Persistent background task state is unavailable.",
            )
        },
    )


@app.errorhandler(ExternalServiceError)
@app.errorhandler(ExternalServiceHTTPError)
def _cvstudio_external_service_error(error):
    structured = error.structured()
    status = int(structured.get("status") or 503)
    if status < 400 or status > 599:
        status = 503
    message = _cvstudio_safe_error_message(str(error), "External service request failed")
    return _cvstudio_error_payload(
        structured.get("code") or "EXTERNAL_SERVICE_REQUEST_FAILED",
        message,
        status,
        retryable=bool(structured.get("retryable")),
        action=structured.get("action") or "",
        details={"service": structured.get("service"), "upstream": structured.get("detail")},
        extra={"service": structured.get("service"), "upstream_status": structured.get("status") or 0},
    )


@app.before_request
def _assign_cvstudio_request_id():
    supplied = str(request.headers.get("X-CV-Studio-Request-ID") or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9_.:-]{8,80}", supplied):
        supplied = secrets.token_hex(8)
    g.cvstudio_request_id = supplied


@app.errorhandler(413)
def _request_too_large(_error):
    return _cvstudio_error_payload(
        "REQUEST_TOO_LARGE",
        "File or request is too large",
        413,
        action="use_smaller_file",
        extra={"limit_mb": 80},
    )


_MAX_PDF_PAGES = _PHASE4_MAX_PDF_PAGES
_MAX_OCR_PAGES = _PHASE4_MAX_OCR_PAGES
_MAX_IMAGE_PIXELS = _PHASE4_MAX_IMAGE_PIXELS
_MAX_ZIP_ENTRIES = _PHASE4_MAX_ZIP_ENTRIES
_MAX_ZIP_EXPANDED_BYTES = _PHASE4_MAX_ZIP_EXPANDED_BYTES
_MAX_ZIP_SINGLE_ENTRY_BYTES = _PHASE4_MAX_ZIP_SINGLE_ENTRY_BYTES
_OCR_SEMAPHORE = threading.BoundedSemaphore(1)
_OCR_TOTAL_DEADLINE_SECONDS = _PHASE4_OCR_TOTAL_DEADLINE_SECONDS


def _document_validation_status(exc):
    return _phase4_document_validation_status(exc)


def _validate_zip_payload(file_bytes, label="document"):
    return _phase4_validate_zip_payload(
        file_bytes,
        label,
        max_entries=_MAX_ZIP_ENTRIES,
        max_expanded_bytes=_MAX_ZIP_EXPANDED_BYTES,
        max_single_entry_bytes=_MAX_ZIP_SINGLE_ENTRY_BYTES,
    )


def _safe_image_open(file_bytes):
    return _phase4_safe_image_open(
        file_bytes,
        max_image_pixels=_MAX_IMAGE_PIXELS,
    )


def _pdf_page_count(file_bytes):
    return _phase4_pdf_page_count(
        file_bytes,
        max_pdf_pages=_MAX_PDF_PAGES,
    )


def _ocr_image_text(pytesseract, image, lang="eng", timeout=35):
    return _phase4_ocr_image_text(
        pytesseract,
        image,
        lang=lang,
        timeout=timeout,
        ocr_semaphore=_OCR_SEMAPHORE,
    )


def _render_pdf_page_images(
    file_bytes,
    first_page=1,
    last_page=None,
    dpi=220,
    poppler_path=None,
    timeout=45,
):
    return _phase4_render_pdf_page_images(
        file_bytes,
        first_page=first_page,
        last_page=last_page,
        dpi=dpi,
        poppler_path=poppler_path,
        timeout=timeout,
    )


def _ocr_pdf_pagewise(file_bytes, pytesseract, poppler_path=None, dpi=220):
    return _phase4_ocr_pdf_pagewise(
        file_bytes,
        pytesseract,
        poppler_path=poppler_path,
        dpi=dpi,
        pdf_page_count=lambda payload: _pdf_page_count(payload),
        render_pdf_page_images=lambda *args, **kwargs: _render_pdf_page_images(
            *args, **kwargs
        ),
        ocr_semaphore=_OCR_SEMAPHORE,
        max_ocr_pages=_MAX_OCR_PAGES,
        max_image_pixels=_MAX_IMAGE_PIXELS,
        deadline_seconds=_OCR_TOTAL_DEADLINE_SECONDS,
        monotonic=lambda: time.monotonic(),
    )


@app.before_request
def _reject_declared_oversize_request():
    # Reject on Content-Length before route code can catch Werkzeug's 413 and
    # accidentally turn it into a generic 500 response.
    declared = request.content_length
    limit = int(app.config.get("MAX_CONTENT_LENGTH") or 0)
    if declared is not None and limit > 0 and declared > limit:
        return _cvstudio_error_payload("REQUEST_TOO_LARGE", "File or request is too large", 413, action="use_smaller_file", extra={"limit_mb": limit // (1024 * 1024)})


_ALLOWED_LOCAL_HOST_HEADERS = frozenset({
    "localhost", "localhost:{}".format(_CVSTUDIO_PORT),
    "127.0.0.1", "127.0.0.1:{}".format(_CVSTUDIO_PORT),
    "[::1]", "[::1]:{}".format(_CVSTUDIO_PORT),
})


def _cvstudio_host_header_allowed(value):
    """Reject DNS-rebinding hosts before any token-backed route executes."""
    host = str(value or "").strip().lower().rstrip(".")
    return host in _ALLOWED_LOCAL_HOST_HEADERS


@app.before_request
def _reject_non_local_host_header():
    # Binding to loopback alone does not prevent a browser DNS-rebinding page
    # from sending a hostile Host header to 127.0.0.1. Enforce the exact local
    # origins used by CV Studio before all routes, including JobAdder proxies.
    if not _cvstudio_host_header_allowed(request.host):
        return _cvstudio_error_payload("INVALID_LOCAL_HOST", "Invalid local host header", 403, action="reopen_local_app")


_LOCAL_BROWSER_ORIGIN_RE = re.compile(r"^http://(?:localhost|127\.0\.0\.1|\[::1\]):{}$".format(_CVSTUDIO_PORT), re.I)
_EXTENSION_ORIGIN_RE = re.compile(r"^(?:chrome|moz)-extension://[A-Za-z0-9_-]+$", re.I)


def _origin_is_local(value):
    return bool(_LOCAL_BROWSER_ORIGIN_RE.match(str(value or "").strip().rstrip("/")))


# A random, per-process HttpOnly browser-session cookie protects every route
# capable of spending configured AI-provider credits. Host/CSRF validation
# already blocks remote websites; this additional gate prevents stale tabs,
# copied requests, and ordinary localhost callers that did not first load this
# exact CV Studio process from invoking paid routes accidentally. As with any
# fully local application, this is hardening rather than protection from malware
# already running as the same OS user.
_AI_SPEND_SESSION_COOKIE = "cvstudio_ai_spend_session"
_AI_SPEND_SESSION_TOKEN = secrets.token_urlsafe(32)
_AI_SPEND_EXACT_PATHS = frozenset({
    "/test",
    "/parse",
    "/generate-ai",
    "/blind",
    "/jobadder/onenote_log_screening",
    "/lead-finder/search",
    "/lead-finder/find-people",
    "/lead-finder/find-emails",
    "/salary-comparison/api/rules/preview",
})


def _ai_spend_session_allowed():
    supplied = str(request.cookies.get(_AI_SPEND_SESSION_COOKIE) or "")
    return bool(supplied) and secrets.compare_digest(supplied, _AI_SPEND_SESSION_TOKEN)


@app.before_request
def _require_ai_spend_browser_session():
    if request.method == "POST" and request.path in _AI_SPEND_EXACT_PATHS:
        if not _ai_spend_session_allowed():
            return _cvstudio_error_payload(
                "AI_SPEND_SESSION_REQUIRED",
                "Paid AI request was not authorised for this CV Studio browser session",
                403,
                action="reload_browser",
                extra={"next_step": "Reload CV Studio in this browser tab and try again."},
            )
    return None


@app.before_request
def _reject_cross_site_unsafe_request():
    """Block ordinary cross-site forms from changing localhost state."""
    if request.method in {"GET", "HEAD", "OPTIONS"}:
        return None
    if str(request.remote_addr or "") not in {"127.0.0.1", "::1"}:
        return _cvstudio_error_payload("UNSAFE_LOCAL_REQUEST", "Unsafe local request was not authorised", 403)

    origin = str(request.headers.get("Origin") or "").strip().rstrip("/")
    referer = str(request.headers.get("Referer") or "").strip()
    fetch_site = str(request.headers.get("Sec-Fetch-Site") or "").strip().lower()

    # The companion browser extension is permitted only on the dedicated OCR
    # route. It cannot use this exception to reach JobAdder, Outlook or caches.
    if request.path == "/ocr" and origin and (_EXTENSION_ORIGIN_RE.match(origin) or _origin_is_local(origin)):
        return None

    if origin:
        if not _origin_is_local(origin):
            return _cvstudio_error_payload("CROSS_SITE_LOCALHOST_BLOCKED", "Cross-site localhost request blocked", 403)
        if fetch_site and fetch_site not in {"same-origin", "same-site", "none"}:
            return _cvstudio_error_payload("CROSS_SITE_LOCALHOST_BLOCKED", "Cross-site localhost request blocked", 403)
        return None

    if referer:
        try:
            parsed = urllib.parse.urlsplit(referer)
            if _origin_is_local("{}://{}".format(parsed.scheme, parsed.netloc)):
                return None
        except Exception:
            pass
        return _cvstudio_error_payload("CROSS_SITE_LOCALHOST_BLOCKED", "Cross-site localhost request blocked", 403)

    # Non-browser owner tools and the explicit restart action can opt in with a
    # non-simple custom header that normal cross-site HTML forms cannot set.
    if request.headers.get("X-CV-Studio-Request") == "1" or request.headers.get("X-CV-Studio-Restart") == "1":
        return None
    return _cvstudio_error_payload("UNSAFE_LOCAL_REQUEST", "Unsafe local request was not authorised", 403)


@app.after_request
def _cvstudio_security_headers(response):
    """Normalise JSON failures, apply browser hardening and disable sensitive caching."""
    # This completes the additive error contract without rewriting hundreds of
    # mature routes at once. Existing ``error`` readers still work, while every
    # JSON failure gains code/message/retry/action/request-id fields.
    try:
        is_json_response = str(getattr(response, "mimetype", "") or "").lower() == "application/json"
        if is_json_response and not getattr(response, "direct_passthrough", False):
            payload = response.get_json(silent=True)
            is_failure = isinstance(payload, dict) and (
                int(getattr(response, "status_code", 200) or 200) >= 400
                or payload.get("ok") is False
                or (bool(str(payload.get("error") or "").strip()) and payload.get("ok") is not True)
            )
            if is_failure:
                normalised = _cvstudio_normalize_error_data(payload, response.status_code, request.path)
                response.set_data(json.dumps(normalised, ensure_ascii=False, separators=(",", ":")))
                response.headers["Content-Type"] = "application/json; charset=utf-8"
    except Exception:
        # Error normalisation must never replace the route's original response.
        pass
    request_id = _cvstudio_current_request_id()
    if request_id:
        response.headers.setdefault("X-CV-Studio-Request-ID", request_id)
    # The Salary Comparison page is embedded as a same-origin iframe inside the
    # CV Studio shell, so it is allowed to be framed by this origin only; every
    # other route keeps the strict anti-clickjacking default.
    if request.path.startswith("/salary-comparison"):
        response.headers.setdefault("Content-Security-Policy", "frame-ancestors 'self'")
        response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    else:
        response.headers.setdefault("Content-Security-Policy", "frame-ancestors 'none'")
        response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("Referrer-Policy", "no-referrer")
    response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
    # JSON and integration responses can contain candidate data, connection
    # metadata, draft links, or transient OAuth status. Never allow a browser,
    # proxy, or service worker to cache them. Static UI/assets retain their
    # normal cache behaviour.
    is_json = str(getattr(response, "mimetype", "") or "").lower() == "application/json"
    is_attachment = "attachment" in str(response.headers.get("Content-Disposition") or "").lower()
    if is_json or is_attachment or request.path.startswith(("/jobadder/", "/onenote/", "/ppc/outlook/", "/secure-secrets/", "/diagnostics/")):
        response.headers["Cache-Control"] = "no-store, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response


SYSTEM_PROMPT = """You are a professional CV parser. Your ONLY job is to output valid JSON.

CRITICAL: Output ONLY a valid JSON object. No markdown fences. No explanation. No preamble.

Input: raw CV text
Output: JSON object matching this exact schema (no variations):

{
  "candidate": {
    "name": "Full Name in Title Case (capitalize first letter of each word) or empty string",
    "email": "candidate@example.com or empty string",
    "phone": "mobile or phone number as written, or empty string",
    "linkedin": "full LinkedIn profile URL if present, or empty string",
    "industry": "one of: Digital & E-Commerce, Financial Services, FMCG, Industrial/Manufacturing, Information Technology & Services, Life Science/Medical, Property & Construction, Professional Services, Education, Government Sector — pick the single best match based on the candidate current and past employers, or empty string if unclear",
    "industry_sub": "one of the valid sub-categories that best matches the candidate industry and employer background, or empty string if unclear",
    "address": {
      "city": "city name or empty string",
      "state": "state/province or empty string",
      "countryCode": "2-letter ISO country code e.g. AU, MY, SG, or empty string"
    },
    "notice_period": "verbatim notice/availability text from the CV (e.g. '1 month', 'Immediate', 'Available from 1 August 2026 onwards') or empty string",
    "current_position": "Job Title or empty string",
    "current_company": "Company Name or empty string",
    "is_employed": true,
    "languages": "Language names only, comma-separated, no proficiency or fluency descriptions. Standardize Malay/Bahasa Melayu/BM as Bahasa Malaysia; Mandarin/Cantonese/Hokkien/Hakka/Chinese dialects as Chinese. Only include languages explicitly stated in the CV."
  },
  "summary_bullets": [],
  "work_experiences": [
    {
      "date_range": "Mon YYYY to Mon YYYY or empty",
      "company": "Company Name",
      "roles": [
        {
          "title": "Job Title exactly as stated in the source, or empty string when absent",
          "date_range": "Mon YYYY to Mon YYYY or empty",
          "reason_for_leaving": "Reason text or empty string",
          "bullets": [
            "plain bullet text",
            { "heading": "Subheader Label", "bullets": ["bullet under heading", "bullet under heading"] }
          ]
        }
      ]
    }
  ],
  "education": [
    {
      "date_range": "YYYY to YYYY or empty",
      "institution": "Institution Name",
      "degree": "Degree Name",
      "cgpa": "CGPA/GPA exactly as written in the source CV (e.g. \"CGPA 4.0\", \"GPA 3.8/4.0\"), or empty string if not mentioned",
      "honors": "Academic distinctions for this qualification exactly as written -- e.g. First Class Honours, Magna Cum Laude, Dean's List, academic scholarships/awards -- comma-separated if more than one, or empty string if none mentioned",
      "description": "Thesis title and/or project/dissertation description for this qualification, copied VERBATIM from the source CV with no rewording, shortening, or paraphrasing -- or empty string if the source CV does not mention one"
    }
  ],
  "certifications": ["cert 1", "cert 2"],
  "skills": [
    { "category": "Category Name", "items": "Skill1, Skill2, Skill3" }
  ]
}

RULES:
- Work experiences in reverse chronological order (most recent first)
- If older/early-career roles have no reliable dates in the source CV, do NOT create separate company rows with empty dates such as " | Company". Group them under a single company entry named "Earlier Career" and preserve each role as bullets in the style "Job Title – Company".
- If the source CV contains a work-history table with columns like Dates / Organization / Role, treat those rows as AUTHORITATIVE for employer name, date range, role title, and order. Reproduce every row from that table as a work experience role; do not skip, merge, or invent rows.
- Do NOT use Project Experience fields such as "Organization", "Client", "Project Name", or "Location" to overwrite the employer/date/title from the work-history table. Project Organization can be the delivery/vendor/client context, not necessarily the employer row.
- Group multiple roles under one company ONLY when the source work-history rows explicitly show the same employer name for those roles. Never group roles across different employers, clients, projects, or non-adjacent rows.
- If a company has multiple roles, list under same entry with multiple role objects
- Within the same company, roles MUST be ordered newest/current first (reverse chronological). Example: Manager (Oct 2022 to Present), then Senior Associate (Oct 2019 to Sep 2022), then Associate (May 2016 to Sep 2019). Never output oldest role first.
- For project-based CVs, preserve every explicit project/client heading instead of flattening all project bullets into one anonymous list. Put the role title first, then each project heading, then that project's bullets. Never place a project heading before the role title.
- Represent each project group inside the relevant role as {"heading": "Project Name – Client", "bullets": ["project bullet", "project bullet"]}. Assign projects to the employment role covering the project start date; if a project spans several promotions or is ongoing across the full employer period, place it under the newest/current applicable role. Preserve the project wording and bullet content.
- Within each role, order dated project groups chronologically from the role's earliest project to its latest project. Put broad ongoing/ad-hoc support or business-proposal groups last.
- If an education entry has no stated degree or qualification, leave `degree` empty. Never output placeholder labels such as "No Degree", "Not specified", or "N/A". If the source states a non-degree qualification such as SPM, PMR, UPSR, a certificate, or a training programme, return the qualification name directly without a "No Degree:" prefix.
- For a company with multiple roles, the top-level date_range MUST span from the earliest role start date to the latest role end date (e.g. if roles are "Sep 2024 to Nov 2025" and "Nov 2025 to Present", the company date_range is "Sep 2024 to Present"). The company date_range must never end before it starts.
- SEPARATE STINTS AT THE SAME EMPLOYER: only group roles under one work-experience entry when they are ONE continuous, back-to-back tenure (a promotion path with adjoining dates). If the candidate worked at the same employer across DISTINCT, non-adjacent periods (a gap between them, or a later return), emit each period as its OWN separate work-experience entry in reverse-chronological position — do not merge them into a single entry. It is correct for the same company name to appear more than once in work_experiences.
- candidate.email: extract the candidate's email address (look for patterns like name@domain.com anywhere in the CV, including in header lines mixed with phone/location/linkedin separated by | or spaces). Return exactly as written, or empty string if not found
- candidate.phone: extract the candidate's mobile or phone number exactly as written, or empty string if not found
- candidate.linkedin: extract the LinkedIn profile URL if present.
- candidate.industry: classify the candidate into exactly one industry from the allowed list based on their employers and role history
- candidate.industry_sub: pick the most specific sub-category that matches. Valid sub-categories:
  Financial Services: FSI - Asset Management, FSI – Banking, FSI - Hedge/Mutual Funds, FSI - Investment Bank, FSI - Insurance, FSI - Private Equity/Venture Capital
  FMCG: FMCG – Consumer Durables, FMCG – Consumer Electronics, FMCG – Food & Beverage, FMCG – Fashion & Apparel/Accessories, FMCG – Luxury Goods, FMCG – MLM, FMCG – Personal Care/Cosmetics, FMCG – Retail, FMCG – Sporting Goods
  Industrial/Manufacturing: Industrial/Manufacturing - Aerospace, Agriculture – Agriculture Supply, Agriculture – Crop/Plantation, Agriculture – Livestock/Feed Mill, Industrial/Manufacturing - Automation, Industrial/Manufacturing - Automotive, Chemicals – Food/Health/Additives, Chemicals – Paint/Coat/Adhesive, Chemicals – Plastic/Package/Print, Chemicals – Textiles, Distributions & Logistics, Industrial/Manufacturing - Diversified Manufacturing, Industrial/Manufacturing - Electronic Components/EMS/CMS, Energy/Power/Utilities – EPC, Energy/Power/Utilities – Mining, Oil & Gas – Upstream, Oil & Gas – DownStream, Energy/Power/Utilities - Renewable/Power, Industrial/Manufacturing - Heavy Industry/Equipment, Industrial/Manufacturing - Marine & Shipping, Industrial/Manufacturing - Metalwork, Industrial/Manufacturing - Railway & Transportation, Industrial/Manufacturing - Semiconductor
  Life Science/Medical: Life Science/Medical - Biotechnology, Life Science/Medical - Hospitals/Nursing, Life Science/Medical - Labs/CROs, Life Science/Medical - Medical devices/Equipment, Life Science/Medical - OTC/Healthcare, Life Science/Medical - Pharmaceutical, Life Science/Medical - Veterinary
  Property & Construction: P&C – Architecture, P&C – Building Materials, P&C – Construction, P&C – General Practice, P&C – Hospitality & Tourism, P&C – Interior Design, P&C – Property Development, P&C – Real Estate Consultancy, P&C – REITS
  Information Technology & Services: Tech - Call Centre/BPO, Tech - Gaming, Tech – ERP/CRM/HCM/Application, Tech - Hardware, Tech - Networks, Tech - Security, Tech – Software/Software House, Tech - Cloud/SaaS, Tech - Sys Integrator/Svc Provider, Tech – Telecommunications, Tech - Blockchain
  Digital & E-Commerce: Digital – FinTech/Payment Getaway/E-Wallet, Digital – MarTech, Digital – EduTech, Digital - InsurTech, Digital - HealthTech/MedTech, Digital - LogTech, E-Commerce – Marketplace/Digital Platform, E-Commerce – Fulfillment/Enabler, E-Commerce – Food & Grocery Delivery
  Professional Services: Professional Services - Accounting Outsourcing, Professional Services- Advertising & Marketing, Professional Services - Corporate Secretary, Professional Services - Financial Advisory, Professional Services - IT Outsourcing, Professional Services - Recruitment & Staffing, Professional Services - Consulting Always include https:// prefix — if the CV shows "www.linkedin.com/in/..." or "linkedin.com/in/..." add "https://" to make it "https://www.linkedin.com/in/...". Return empty string if not found.
- candidate.address.city: city from address section, or empty string
- candidate.address.state: state/province from address, or empty string
- candidate.address.countryCode: 2-letter ISO country code from address (e.g. AU for Australia, MY for Malaysia), or empty string
- candidate.name: the person's full name ONLY — never the address, city, country, email, phone, or a job title. It is normally the most prominent line at the very top of the CV (contact lines such as address/phone/email usually follow it, sometimes preceded by icon glyphs). Always output in Title Case (capitalize first letter of each word, lowercase the rest), e.g. "MOHAMMAD AJMER BIN ABDUL HASAN" → "Mohammad Ajmer Bin Abdul Hasan"
- Work experience date ranges must use exactly this style: "Mon YYYY to Mon YYYY" or "Mon YYYY to Present". Convert "Till Date", "Current", hyphens/dashes, and ALL-CAPS months into this style (e.g. "OCT 2022 - Till Date" → "Oct 2022 to Present").
- DATE ACCURACY: transcribe every start and end month/year EXACTLY as written in the source. Never shift, round, guess, or invent a month or year. Only "Present"/"Current"/"Till date" (or no end at all) may become "to Present" — a stated end date such as a specific month/year is NEVER rewritten as "Present".
- Work experience company names and role titles must not be returned in ALL CAPS unless they are genuine acronyms (e.g. COGNIZANT → Cognizant, DATA ECONOMY → Data Economy, WOLTERS KLUWER → Wolters Kluwer, but CGI/AWS/SQL/SAP stay uppercase).
- EMPLOYER NAME FIDELITY: return each employer name as written in the source. Do not invent, swap, translate, expand an abbreviation, contract a full name, or borrow a name from a different role/company. You may drop a trailing legal form (Sdn Bhd, Pte Ltd, PT, Tbk, Inc, LLC) and fix casing, but the core name must match the source exactly. Never output a company name that does not appear in the CV.
- ROLE TITLE FIDELITY: return each role title exactly as written in the source. Do not paraphrase, generalise, "clean up", shorten, or invent a title (e.g. do NOT turn "Dispatcher Technical Support" into "Customer Care Consultant", or "Senior Linux System Administrator" into "Senior Systems Administrator"). Keep every qualifier the source states (Linux, Non-Wintel, AVP, etc.).
- candidate.is_employed: set to true ONLY if the candidate has a role that explicitly says "Present", "Current", "Till date", "To date", or similar — AND that role is a full-time, permanent, contract, or consulting role with a company (not freelance or self-employed). If ALL roles have a definite end date (e.g. "Dec 2025", "Feb 2025"), set is_employed to false even if the end date is very recent. If the only "Present" role is freelance or self-employment, set is_employed to false.
- candidate.current_company: populate this from the most recent employer when the source states one.
- candidate.current_position: populate this only when the source explicitly states the corresponding role title. Never infer a title from responsibilities, achievements, industry, seniority, dates, surrounding context, or common career patterns. If the title is absent, return an empty string.
- work_experiences[].roles[].title: copy only an explicitly stated source title. Never invent, infer, imply, annotate, or explain a title. Use an empty string when the source provides no title.
- reason_for_leaving: only if stated in input, else empty string ""
- summary_bullets: always empty array []
- bullets: preserve original wording exactly, no paraphrasing. If a role has named sub-sections or categories within it (e.g. "Key responsibilities", "Key achievements", "Data Engineering", "Database", or "Regional Finance Oversight:"), represent each as a JSON object, never as JSON serialized inside a string: { "heading": "...", "bullets": ["did x", "did y"], "kind": "section" }. Plain bullets with no sub-section go in as plain strings. Mix freely — e.g. ["plain bullet", { "heading": "Data Engineering", "bullets": ["did x", "did y"], "kind": "section" }, "another plain bullet"]
- education[].cgpa, education[].honors, and education[].description: do NOT omit these even though they are easy to skip. If the source CV states a CGPA/GPA for a qualification, it MUST be captured in that entry's "cgpa" field exactly as written. If the source CV mentions academic distinctions -- First Class Honours, cum laude variants, Dean's List, academic scholarships/awards, or similar -- for that qualification, it MUST be captured in "honors" exactly as written. If the source CV includes a thesis title, dissertation topic, or a description of a capstone/major project for that qualification, it MUST be captured in "description" — copied verbatim, not summarized or shortened, even if it is several sentences long.
- All text: fix spelling/typos silently, preserve everything else
- Missing fields: use empty string "" or empty array []
- JSON MUST be valid and properly formatted
- NO markdown fences
- NO backticks
- NO explanation before or after
- Output starts with { and ends with }

SECTION MAPPING RULES — very important:
- Any section labelled Summary, Professional Summary, Profile, Objective, Career Objective, Overview → add as a skills category { "category": "Summary", "items": "<full text, preserve wording exactly>" } so it appears under Additional Information
- Any section labelled Skills, Core Skills, Technical Skills, Key Skills, Competencies, Core Competencies, Strengths, Key Strengths, Areas of Expertise, Expertise, Specialties → map into the "skills" array. Each distinct sub-category or skill group becomes its own { "category": "...", "items": "..." } entry
- If a skills section has sub-headings with bullet points (e.g. "Oracle:" followed by bullets, "PostgreSQL:" followed by bullets), map EACH sub-heading as its own skill category, and join its bullet points into the "items" field separated by newline characters (\n). Preserve the full detail of every bullet point exactly.
- Any section labelled Certifications, Certificates, Licenses, Accreditations, Professional Certifications → map into "certifications" array
- Any section labelled Training, Trainings, Professional Development, Courses, Short Courses, Workshops → map each training/course as an item in "certifications" array (they appear together under Additional Information)
- Preserve every item under explicitly labelled source sections such as "Project Involvement History" and "Participated Training Programme". Never truncate these later sections because they appear near the bottom of a long CV.
- CERTIFICATION/TRAINING DATES — KEEP THEM: when a certification or training states a date or year (e.g. "12/2023", "2019", "Nov 2019"), preserve it in the certification item exactly as given. Do not drop the date. Keep the issuing body/institution too.
- Any section labelled Languages → map to candidate.languages field ONLY when it clearly belongs to the candidate. Extract EVERY language listed — capture all of them, not just the first (a two-column CV may interleave the Languages sidebar with other sections, e.g. "Chinese (Professional Working)", "Malay (Professional Working)", "English (Professional Working)" separated by other lines; include all three). Extract language names ONLY — strip all proficiency levels, fluency descriptions, written/spoken qualifiers. Standardize language names: English; Bahasa Malaysia for Malay/Bahasa Melayu/BM; Chinese for Mandarin/Cantonese/Hokkien/Hakka/other Chinese dialects. Only include languages explicitly stated in the CV.
- REDACTED LANGUAGE GUARD: If a Languages section or candidate.languages value is explicitly redacted/masked/withheld (e.g. "[redacted]", "redacted", "confidential", "masked", "hidden", "***", "xxx", "████"), do not infer or invent languages from it. Leave candidate.languages as an empty string in that case. Do not use broad template/filler assumptions to remove languages; only ignore explicit redaction/masking.
- Any section labelled Awards, Achievements, Honours, Accomplishments → add as a skills category: { "category": "Awards & Achievements", "items": "Award 1\nAward 2" }
- Any section labelled Volunteer, Volunteering, Community, Extracurricular → add as a skills category: { "category": "Volunteer & Community", "items": "Activity 1\nActivity 2" }
- NEVER put summary, skills, competencies, strengths, or training content inside work experience bullets
- The ONLY things that go in work_experiences are actual jobs with dates, company names, job titles, and role-specific bullet points
- If a skills/competency section has plain comma-separated items with no sub-headings, put them all in items as a single comma-separated string
- If a skills section has sub-headings with detailed bullet points, preserve ALL bullet points in full — do not summarize or truncate

OMISSION & EXTRA SECTIONS RULES:
- Do NOT place phone numbers, email addresses, physical addresses, LinkedIn URLs, or location/relocation info into the formatted CV body (skills, summary, work experience bullets, education, additional information). HOWEVER, still extract these into candidate.email, candidate.phone, candidate.linkedin, and candidate.address in the JSON — the app uses them for JobAdder matching and upload.
- SALARY / REMUNERATION — OMIT ENTIRELY: Never include any salary, remuneration, or compensation detail anywhere in the formatted CV body or JSON. This covers current/present salary, expected/asking/desired/target salary, base pay, monthly or annual salary figures, bonus, commission, incentives, allowances (transport/housing/meal/etc.), EPF/KWSP/benefits amounts, and total package/CTC figures — whether written inline or under a section labelled Current Remuneration, Remuneration, Compensation, Salary, Expected Salary, Current Salary, Salary Expectation, Package, or similar. Drop the whole section and every such figure; do NOT map it to a skills category, Additional Information, or any other field. This overrides the catch-all mapping rule below. (Notice period is NOT salary — keep mapping it to candidate.notice_period per the rule below.)
- RECRUITMENT-SYSTEM METADATA — OMIT ENTIRELY: Never include source-routing/application metadata such as "Position: Retrieved Resumes (SiVA folder: ...); Date Applied: ...", JobStreet/SiVA folder labels, retrieval status, or application dates anywhere in the JSON or formatted CV.
- KEEP & map to Additional Information skills categories:
  - GitHub links, personal websites, portfolio URLs, Behance, Dribbble → { "category": "Portfolio & Links", "items": "GitHub: https://... | Website: https://..." }, but ONLY when the exact link is explicitly present in the source CV. Never invent, infer, complete, or emit a placeholder URL such as `https://github.com/unknown`; if no source link exists, omit the category.
  - Patents section → { "category": "Patents", "items": "Patent title (Patent number, Year)\nPatent title (Patent number, Year)" } — preserve each patent as a separate line with full detail
  - Publications, Research Papers → { "category": "Publications", "items": "Title (Journal, Year)\nTitle (Journal, Year)" }
  - Any other section that does not fit work_experiences, education, certifications, or skills (e.g. Interests, References, Projects, Open Source, Speaking Engagements, Conference Talks) → add as its own skills category with a sensible category name, preserve all content. EXCEPTION: never apply this catch-all to salary/remuneration/compensation content — that is omitted entirely per the SALARY / REMUNERATION rule above.
  - Relocate/Open to relocation info → omit
  - Notice period if stated in the CV → map to candidate.notice_period VERBATIM. Copy the exact wording the CV uses (a month count like "1 month", "Immediate", or an availability date like "Available from 1 August 2026 onwards"). Never infer, round, or invent a value, and never convert an availability date into a month count. If no notice/availability is stated, return an empty string."""


_AI_PROVIDER_CLIENT = AIProviderClient()


def call_anthropic(api_key, payload_dict):
    # Internal-only control key. Lets long-running tools fail gracefully instead
    # of leaving the browser waiting until its own timeout fires.
    payload_dict = dict(payload_dict or {})
    timeout_seconds = payload_dict.pop("_timeout_seconds", 180)
    try:
        timeout_seconds = max(15, int(timeout_seconds))
    except Exception:
        timeout_seconds = 180

    _phase5b_enforce_request_guardrail("anthropic", payload_dict)
    return _AI_PROVIDER_CLIENT.request(
        "anthropic",
        api_key,
        payload_dict,
        timeout=timeout_seconds,
    )


# ── Multi-provider LLM support (v22.0) ──────────────────────────────────
# call_anthropic() above remains the default compatibility entry point.
# call_llm() below is a thin dispatcher every route now goes through; when
# `provider` is empty/unset/unrecognized it delegates straight to
# call_anthropic() with the exact same payload, so nothing changes for
# anyone who never touches the new provider selector.

_LLM_PROVIDER_INFO = {
    "anthropic": {"label": "Claude (Anthropic)", "supports_web_search": True},
    "deepseek": {"label": "DeepSeek", "supports_web_search": False},
    "openai": {"label": "GPT (OpenAI)", "supports_web_search": True},
}


def _provider_short_label(provider):
    p = str(provider or "anthropic").strip().lower()
    if p == "deepseek":
        return "DeepSeek"
    if p in ("openai", "gpt"):
        return "GPT/OpenAI"
    return "Claude"

def _provider_down_hint(provider, msg):
    text = str(msg or "")
    low = text.lower()
    if re.search(r"invalid api key|api key required|authentication|unauthorized|forbidden|permission|credit|quota|billing|insufficient|model.*not|unsupported|bad request", low):
        return ""
    if re.search(r"timed?\s*out|timeout|temporar(?:y|ily)|unavailable|overloaded|server error|bad gateway|gateway|service unavailable|connection reset|connection aborted|connection refused|network|upstream|rate limit|too many requests|\b429\b|\b500\b|\b502\b|\b503\b|\b504\b|\b529\b", low):
        return f"{_provider_short_label(provider)} appears down, overloaded, or temporarily unreachable. Try again shortly or switch provider in Settings."
    return ""

def _provider_error_message(provider, msg):
    hint = _provider_down_hint(provider, msg)
    msg = str(msg or "API provider error")
    return (hint + "\n\nDetails: " + msg) if hint else f"API provider error: {msg}"


class _LLMProviderNoWebSearch(Exception):
    """Raised when a request needs live web search but the selected provider
    cannot do it and no search-provider results are available to classify
    instead. Surfaced to the user as a clear, actionable error rather than
    silently sending the request anyway and risking a hallucinated
    'search-based' answer with no actual search behind it."""
    pass


def _call_deepseek(api_key, payload_dict, timeout_seconds=180):
    """DeepSeek's Anthropic-compatible endpoint (api.deepseek.com/anthropic)
    accepts the same Messages API request shape as Anthropic itself, so no
    request/response translation is needed -- just a different host. DeepSeek
    has no equivalent to Anthropic's built-in web_search server tool, so that
    tool is stripped before sending; a request that truly needs live web
    search raises _LLMProviderNoWebSearch instead of silently proceeding
    without it.
    """
    payload_dict = dict(payload_dict or {})
    if payload_dict.get("tools"):
        payload_dict = {k: v for k, v in payload_dict.items() if k != "tools"}
        raise _LLMProviderNoWebSearch(
            "DeepSeek cannot browse the web itself (no built-in web search tool). "
            "Configure a Job Search Provider (Tavily or SerpAPI) so DeepSeek can classify "
            "already-fetched results instead, or switch this task back to Claude/GPT."
        )
    # DeepSeek V4 enables chain-of-thought "thinking" by default, which makes
    # every call spend its output budget on hidden reasoning first -- observed
    # as 30-160s latency (and mixed reasoning_content that can break strict-JSON
    # parsing) on CV Studio's structured extraction/generation tasks: /parse,
    # Blind JD, Company Profile, The Owl and AI Crawler. None of these benefit
    # from thinking, so disable it explicitly via the Anthropic-compatible field
    # (accepted values: "adaptive"/"enabled"/"disabled"). A caller may still opt
    # in by setting "thinking" itself. Only the final text content is returned.
    if "thinking" not in payload_dict:
        payload_dict["thinking"] = {"type": "disabled"}
    _phase5b_enforce_request_guardrail("deepseek", payload_dict)
    return _AI_PROVIDER_CLIENT.request(
        "deepseek",
        api_key,
        payload_dict,
        timeout=timeout_seconds,
    )


def _llm_usage_int(usage, *keys):
    return _phase5b_usage_int(usage, *keys)


def _normalize_llm_usage(usage, api_calls=None):
    """Return one provider-neutral usage shape without discarding native counters."""
    return _phase5b_normalize_usage(usage, api_calls)


def _merge_llm_usage(*usages):
    return _phase5b_merge_usage(*usages)


def _openai_text_from_content(content):
    """Flatten Anthropic-style message content blocks into text for OpenAI Responses."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for block in content:
            if isinstance(block, dict):
                if block.get("type") in ("text", "input_text", "output_text"):
                    parts.append(str(block.get("text") or ""))
            else:
                parts.append(str(block))
        return "".join(parts)
    return str(content or "")


def _openai_payload_from_anthropic_shape(payload_dict):
    """Translate an Anthropic-Messages-shaped request (model/messages/system/
    max_tokens/tools) into an OpenAI Responses API request.

    Preserve multi-turn role order for repair/continuation calls. v22.0
    concatenated user+assistant+user turns into one string, which could make
    OpenAI continuation repair behave differently from the original Claude path.
    """
    model = payload_dict.get("model") or "gpt-5.5"
    messages = payload_dict.get("messages") or []
    input_items = []
    for m in messages:
        if not isinstance(m, dict):
            continue
        role = str(m.get("role") or "user").strip().lower()
        if role not in ("user", "assistant", "developer", "system"):
            role = "user"
        text = _openai_text_from_content(m.get("content"))
        if text:
            input_items.append({"role": role, "content": text})
    # Responses API accepts a plain string for simple one-shot requests and a
    # message array for conversation-style input. Keep one-shot payloads compact,
    # but preserve roles whenever there is more than one turn.
    if len(input_items) == 1 and input_items[0].get("role") == "user":
        openai_input = input_items[0].get("content", "")
    else:
        openai_input = input_items
    max_tokens = payload_dict.get("max_tokens") or 4096
    out = {
        "model": model,
        "input": openai_input,
        "max_output_tokens": max_tokens,
    }
    system_text = payload_dict.get("system")
    if isinstance(system_text, str) and system_text.strip():
        out["instructions"] = system_text
    elif isinstance(system_text, list):
        # Anthropic also allows a list of content blocks for system.
        sys_joined = "".join(_openai_text_from_content(b) for b in system_text if isinstance(b, dict))
        if sys_joined.strip():
            out["instructions"] = sys_joined
    if payload_dict.get("tools"):
        # Anthropic's web_search_20250305 -> OpenAI Responses API's web_search.
        out["tools"] = [{"type": "web_search"}]
    return out


def _openai_response_to_anthropic_shape(data):
    """Translate an OpenAI Responses API response back into the
    {"content": [...], "usage": {...}} shape every existing caller in this
    file already expects, so no downstream parsing code needs to know or
    care which provider actually answered.

    OpenAI examples expose a convenient output_text field on SDK responses,
    while raw REST responses are usually an output[] list of message blocks.
    Read both shapes so a live API response cannot look empty just because it
    arrived in the convenience form.
    """
    text_out = str(data.get("output_text") or "")
    if not text_out:
        for item in (data.get("output") or []):
            if not isinstance(item, dict):
                continue
            if item.get("type") in ("output_text", "text"):
                text_out += str(item.get("text") or "")
                continue
            if item.get("type") == "message":
                for block in (item.get("content") or []):
                    if isinstance(block, dict) and block.get("type") in ("output_text", "text"):
                        text_out += str(block.get("text") or "")
    usage = dict(data.get("usage") or {})
    return {
        "content": [{"type": "text", "text": text_out}],
        "usage": dict(
            usage,
            input_tokens=usage.get("input_tokens", 0),
            output_tokens=usage.get("output_tokens", 0),
        ),
    }


def _call_openai(api_key, payload_dict, timeout_seconds=180):
    """Uses OpenAI's Responses API (not the older Chat Completions API) for
    every call, whether or not web search is requested -- Responses is the
    only OpenAI endpoint with a native web search tool, and using it
    uniformly avoids juggling two different OpenAI request/response shapes
    depending on whether a given call needs search.
    """
    req_payload = _openai_payload_from_anthropic_shape(payload_dict)
    guard_payload = dict(req_payload)
    if "max_tokens" in payload_dict:
        guard_payload["max_output_tokens"] = payload_dict.get("max_tokens")
    _phase5b_enforce_request_guardrail("openai", guard_payload)
    data = _AI_PROVIDER_CLIENT.request(
        "openai",
        api_key,
        req_payload,
        timeout=timeout_seconds,
    )
    translated = _openai_response_to_anthropic_shape(data)
    billing_result = _phase5b_extract_provider_billing_result("openai", data)
    billing = billing_result.get("billing")
    if billing:
        translated["_cvstudio_provider_billing"] = billing
    if billing_result.get("error"):
        translated["_cvstudio_provider_billing_error"] = billing_result["error"]
    return translated


_DEEPSEEK_RETIRED_MODELS = {
    "deepseek-chat": "deepseek-v4-flash",
    "deepseek-reasoner": "deepseek-v4-flash",
    "deepseek-v4-flahs": "deepseek-v4-flash",
}


def _migrate_deepseek_model(model):
    """Map retired/misspelled DeepSeek model names to the current default.

    Idempotent last line of defence at the transport boundary: a stale
    browser setting (e.g. ``deepseek-chat``) or a typo should not silently
    fail every DeepSeek feature call. Only known-retired values are rewritten;
    unrecognised values are left untouched so genuinely new models still pass.
    """
    key = str(model or "").strip()
    return _DEEPSEEK_RETIRED_MODELS.get(key.lower(), key)


def call_llm(provider, api_key, payload_dict):
    """Dispatch a Messages-API-shaped request and normalize returned usage.

    Every successful provider response represents one chargeable API call.
    Canonical counters let every feature calculate DeepSeek cache discounts
    consistently, including retry and repair paths.
    """
    provider = str(provider or "anthropic").strip().lower()
    if provider in ("", "anthropic", "claude"):
        data = call_anthropic(api_key, payload_dict)
    elif provider == "deepseek":
        payload_dict = dict(payload_dict or {})
        if payload_dict.get("model"):
            payload_dict["model"] = _migrate_deepseek_model(payload_dict.get("model"))
        timeout_seconds = payload_dict.pop("_timeout_seconds", 180)
        try:
            timeout_seconds = max(15, int(timeout_seconds))
        except Exception:
            timeout_seconds = 180
        data = _call_deepseek(api_key, payload_dict, timeout_seconds)
    elif provider in ("openai", "gpt"):
        payload_dict = dict(payload_dict or {})
        timeout_seconds = payload_dict.pop("_timeout_seconds", 180)
        try:
            timeout_seconds = max(15, int(timeout_seconds))
        except Exception:
            timeout_seconds = 180
        data = _call_openai(api_key, payload_dict, timeout_seconds)
    else:
        data = call_anthropic(api_key, payload_dict)
    if isinstance(data, dict):
        data = dict(data)
        billing_result = _phase5b_extract_provider_billing_result(provider, data)
        billing = billing_result.get("billing")
        billing_error = (
            billing_result.get("error")
            or data.get("_cvstudio_provider_billing_error")
        )
        # Raw billing envelopes may contain provider account references or
        # malformed fields. Never return or persist them; retain only the
        # normalized allowlist or a sanitized failure status in usage.
        data.pop("provider_billing", None)
        data.pop("_cvstudio_provider_billing", None)
        data.pop("_cvstudio_provider_billing_error", None)
        usage = _normalize_llm_usage(data.get("usage"), api_calls=1)
        if billing:
            usage["_cvstudio_provider_billing"] = billing
        if billing_error:
            usage["_cvstudio_provider_billing_error"] = billing_error
        data["usage"] = usage
    return data


# ── Output normalization: keep single and batch CV exports consistent ──────────
# The AI may preserve source casing differently between runs/providers (e.g.
# "OCT 2022 - Till Date | COGNIZANT" vs "Oct 2022 to Present | Cognizant").
# These small deterministic cleanup helpers are intentionally conservative and
# apply to header/date/company/title fields only — bullets are preserved verbatim.
from cvstudio_cv_normalize import (
    _CV_LEADING_BULLET_MARKER_RE,
    _CV_LANGUAGE_ALIASES,
    _CV_LANGUAGE_ALIAS_TO_STANDARD,
    _CV_MONTH_NUMBER,
    _CV_REDACTED_LANGUAGE_RE,
    _CV_ROLE_DENSE_MARKER_RE,
    _WORK_TABLE_DATE_RE,
    _canonical_cv_section_heading,
    _canonical_language_name,
    _cv_combine_date_ranges,
    _cv_company_span_from_roles,
    _cv_date_parts,
    _cv_date_sort_point,
    _cv_pretranslate_iso_dates,
    _cv_lang_alias_re,
    _cv_match_key,
    _cv_parse_backend_timeout_seconds,
    _cv_plain_bullet_texts,
    _cv_project_group_sort_key,
    _cv_text_similarity,
    _cv_token_overlap_score,
    _cv_token_set,
    _language_evidence_text,
    _language_has_source_evidence,
    _normalize_candidate_language_value,
    _normalize_candidate_languages,
    _correct_mistagged_candidate_name,
    _normalize_cv_bullet_items,
    _normalize_cv_data_for_output,
    _normalize_cv_date_range,
    _normalize_cv_structured_content,
    _normalize_month_token,
    _smart_title_text,
    _smart_word_case,
    _split_candidate_language_names,
    _strip_cv_inferred_title,
)
# Brand-like company tokens that are often damaged by generic title-casing.
# Keep this small and deterministic; it is only for well-known casing patterns.











from cvstudio_cv_reconcile import (
    _role_plain_bullets,
    _collapse_incomplete_earlier_career,
    _source_has_redacted_language_block,
    _clean_candidate_languages_from_redaction,
    _order_same_company_roles_newest_first,
    _extract_explicit_project_blocks,
    _restore_explicit_project_headings,
    _extract_authoritative_work_rows,
    _flatten_parsed_work_roles,
    _score_authoritative_row_match,
    _reconcile_work_experience_with_authoritative_table,
)
from cvstudio_cv_fidelity import (
    evaluate_cv_fidelity,
    summarize_fidelity_warning,
)










for _std_lang, _aliases in _CV_LANGUAGE_ALIASES.items():
    for _alias in _aliases:
        _CV_LANGUAGE_ALIAS_TO_STANDARD[_alias] = _std_lang





































































_CVSTUDIO_STARTUP_SERVICE = StartupService(
    jsonify=lambda payload: jsonify(payload),
    root_path=lambda: _CVSTUDIO_ROOT,
    instance_id=lambda: _CVSTUDIO_INSTANCE_ID,
)


@app.route("/startup/status", methods=["GET"])
def startup_status():
    return _CVSTUDIO_STARTUP_SERVICE.status()

@app.route("/startup/enable", methods=["POST"])
def startup_enable():
    return _CVSTUDIO_STARTUP_SERVICE.enable()

@app.route("/startup/disable", methods=["POST"])
def startup_disable():
    return _CVSTUDIO_STARTUP_SERVICE.disable()

_CVSTUDIO_RUNTIME_SERVICE = RuntimeService(
    jsonify=lambda payload: jsonify(payload),
    version=lambda: _CVSTUDIO_VERSION,
    root_path=lambda: _CVSTUDIO_ROOT,
    root_hash=lambda: _CVSTUDIO_ROOT_HASH,
    instance_id=lambda: _CVSTUDIO_INSTANCE_ID,
    port=lambda: _CVSTUDIO_PORT,
)


@app.route("/status", methods=["GET"])
def status():
    return _CVSTUDIO_RUNTIME_SERVICE.status()

@app.route("/instance", methods=["GET"])
def instance_info():
    return _CVSTUDIO_RUNTIME_SERVICE.instance_info()

@app.route("/instance-id", methods=["GET"])
def instance_id_text():
    return _CVSTUDIO_RUNTIME_SERVICE.instance_id_text()

@app.route("/ping", methods=["GET"])
def ping():
    return _CVSTUDIO_RUNTIME_SERVICE.ping()

# ── JobAdder OAuth and protected credential store ────────────────────
_ja_creds_store = {}
_ja_oauth_sessions = {}
_ja_oauth_lock = threading.RLock()
_ja_store_lock = threading.RLock()
_ja_refresh_lock = threading.Lock()
_ja_critical_write_state_lock = threading.RLock()
_ja_critical_writes_active = 0
_ja_account_transition_in_progress = False
_JA_SESSION_TTL = 10 * 60


def _ja_begin_account_transition():
    """Prevent sign-out/account replacement from racing an unsafe JA write."""
    global _ja_account_transition_in_progress
    with _ja_critical_write_state_lock:
        if _ja_account_transition_in_progress:
            return "transition"
        if _ja_critical_writes_active:
            return "write"
        _ja_account_transition_in_progress = True
        return ""


def _ja_end_account_transition():
    global _ja_account_transition_in_progress
    with _ja_critical_write_state_lock:
        _ja_account_transition_in_progress = False


def _ja_account_transition_conflict(reason):
    if reason == "write":
        return _cvstudio_error_payload(
            "JOBADDER_WRITE_IN_PROGRESS",
            "A critical JobAdder write or upload is still in progress. Try Sign Out again after it finishes.",
            409,
            retryable=True,
            action="retry_later",
        )
    return _cvstudio_error_payload(
        "JOBADDER_ACCOUNT_TRANSITION",
        "Another JobAdder account operation is still in progress.",
        409,
        retryable=True,
        action="retry_later",
    )


def _ja_critical_write_route(route_function):
    """Track unsafe JA writes without changing their established call contract."""
    @functools.wraps(route_function)
    def guarded(*args, **kwargs):
        global _ja_critical_writes_active
        with _ja_critical_write_state_lock:
            if _ja_account_transition_in_progress:
                return _ja_account_transition_conflict("transition")
            _ja_critical_writes_active += 1
        try:
            return route_function(*args, **kwargs)
        finally:
            with _ja_critical_write_state_lock:
                _ja_critical_writes_active = max(
                    0, _ja_critical_writes_active - 1
                )
    return guarded


def _ja_cleanup_sessions():
    now = time.time()
    with _ja_oauth_lock:
        for sid in list(_ja_oauth_sessions):
            if float((_ja_oauth_sessions.get(sid) or {}).get("expires_at") or 0) <= now:
                _ja_oauth_sessions.pop(sid, None)


def _ja_finish_oauth_exchange_if_active(session_id, updates, remove_keys=()):
    """Finish only the still-live exchange claimed before remote transport."""
    with _ja_oauth_lock:
        current = _ja_oauth_sessions.get(session_id)
        if not isinstance(current, dict) or current.get("status") != "exchanging":
            return False
        current = dict(current)
        current.update(dict(updates or {}))
        for key in remove_keys:
            current.pop(key, None)
        _ja_oauth_sessions[session_id] = current
        return True


def _ja_public_info():
    with _ja_store_lock:
        connected = bool(str(_ja_creds_store.get("access_token") or "").strip())
        private_cache_namespace = str(
            _ja_creds_store.get("cache_namespace") or ""
        ).strip()
        if connected and not private_cache_namespace:
            legacy_material = "|".join(
                str(_ja_creds_store.get(key) or "")
                for key in (
                    "api_url",
                    "client_id",
                    "refresh_token",
                    "access_token",
                )
            )
            private_cache_namespace = hashlib.sha256(
                legacy_material.encode("utf-8", errors="ignore")
            ).hexdigest()
        browser_cache_namespace = (
            hashlib.sha256(
                ("cvstudio-browser-cache-v1|" + private_cache_namespace).encode(
                    "utf-8", errors="ignore"
                )
            ).hexdigest()
            if connected and private_cache_namespace
            else ""
        )
        return {
            "api_url": _ja_normalize_api_base(_ja_creds_store.get("api_url")),
            "connected": connected,
            "expires_at": int(_ja_creds_store.get("expires_at") or 0),
            "client_id": str(_ja_creds_store.get("client_id") or ""),
            "needs_reconnect": bool(_ja_creds_store.get("_needs_reconnect")),
            # The browser needs to know whether reconnect can reuse the secret,
            # but the secret itself must never be returned to JavaScript.
            "client_secret_configured": bool(str(_ja_creds_store.get("client_secret") or "").strip()),
            # This one-way, connection-scoped value lets browser caches reject
            # cross-account reuse without exposing the protected namespace.
            "account_cache_namespace": browser_cache_namespace,
            "storage": str(_ja_creds_store.get("_storage") or "backend_secure_store"),
        }


def _ja_save_store():
    with _ja_store_lock:
        record = {k: _ja_creds_store.get(k) for k in (
            "client_id", "client_secret", "access_token", "refresh_token", "api_url", "expires_at", "cache_namespace"
        ) if _ja_creds_store.get(k) not in (None, "")}
        if record:
            kind = _cv_secure_save("jobadder", record)
            _ja_creds_store["_storage"] = kind
        else:
            _cv_secure_delete("jobadder")


def _ja_exchange_token(params):
    return _JOBADDER_CLIENT.exchange_token(params, timeout=20)


def _ja_apply_token(token, client_id=None, client_secret=None, clear_spider_cache=False):
    expires_in = int(token.get("expires_in") or 3300)
    with _ja_store_lock:
        previous = dict(_ja_creds_store)
        try:
            next_api_url = _ja_normalize_api_base(token.get("api") or token.get("api_url") or _ja_creds_store.get("api_url"))
            next_client_id = str(client_id or _ja_creds_store.get("client_id") or "")
            # cache_namespace is intentionally stable across ordinary access-token
            # refreshes. A newly completed OAuth connection gets a fresh random
            # namespace, preventing tenant crossover without tying cache reuse to
            # the rotating bearer token.
            cache_namespace = str(_ja_creds_store.get("cache_namespace") or "").strip()
            account_changed = bool(clear_spider_cache)
            if previous.get("api_url") and next_api_url and _ja_normalize_api_base(previous.get("api_url")) != next_api_url:
                account_changed = True
            if previous.get("client_id") and next_client_id and str(previous.get("client_id")) != next_client_id:
                account_changed = True
            if account_changed or not cache_namespace:
                cache_namespace = secrets.token_hex(24)
            _ja_creds_store.update({
                "access_token": str(token.get("access_token") or ""),
                "refresh_token": str(token.get("refresh_token") or _ja_creds_store.get("refresh_token") or ""),
                "api_url": next_api_url,
                "expires_at": int(time.time() + max(60, expires_in)),
                "cache_namespace": cache_namespace,
                "_needs_reconnect": False,
            })
            if client_id:
                _ja_creds_store["client_id"] = str(client_id)
            if client_secret:
                _ja_creds_store["client_secret"] = str(client_secret)
            _ja_save_store()
            if clear_spider_cache:
                try:
                    _spider_resume_text_cache_clear()
                except Exception:
                    pass
                try:
                    _ppc_detail_cache_clear()
                except Exception:
                    pass
        except Exception:
            _ja_creds_store.clear(); _ja_creds_store.update(previous)
            raise


def _ja_mark_reconnect_required():
    """Invalidate a rejected access token without discarding reusable OAuth registration."""
    with _ja_store_lock:
        _ja_creds_store["access_token"] = ""
        _ja_creds_store["expires_at"] = 0
        _ja_creds_store["_needs_reconnect"] = True
        try:
            _ja_save_store()
        except Exception:
            pass
    # The crawler cache is tenant-bound. Drop it immediately when authentication
    # becomes uncertain, but remain safe if this helper runs before crawler setup.
    try:
        clearer = globals().get("_spider_resume_text_cache_clear")
        if callable(clearer):
            clearer()
    except Exception:
        pass


def _ja_refresh_access_token(force=False):
    with _ja_refresh_lock:
        token = str(_ja_creds_store.get("access_token") or "").strip()
        expires_at = int(_ja_creds_store.get("expires_at") or 0)
        if token and not force and (not expires_at or expires_at > time.time() + 120):
            return token
        refresh = str(_ja_creds_store.get("refresh_token") or "").strip()
        client_id = str(_ja_creds_store.get("client_id") or "").strip()
        client_secret = str(_ja_creds_store.get("client_secret") or "").strip()
        if not refresh or not client_id or not client_secret:
            return token
        try:
            response = _ja_exchange_token({
                "grant_type": "refresh_token",
                "refresh_token": refresh,
                "client_id": client_id,
                "client_secret": client_secret,
            })
            _ja_apply_token(response, client_id, client_secret)
        except (urllib.error.HTTPError, urllib.error.URLError, RuntimeError, TimeoutError, OSError):
            # invalid_grant, revoked/expired refresh tokens, network-level token
            # failures, malformed token responses and token-application failures
            # must never escape from route entry points as generic HTTP 500s.
            _ja_mark_reconnect_required()
            return ""
        return str(_ja_creds_store.get("access_token") or "")


@app.route("/jobadder/store_creds", methods=["POST"])
def jobadder_store_creds():
    """Create a session-bound OAuth attempt; credentials never enter browser storage.

    A reconnect may omit the secret only when the secure backend vault already
    contains one for the exact same Client ID. Changing Client ID always
    requires the corresponding secret, which prevents silently pairing a new
    application ID with an old application's secret.
    """
    data = request.get_json(silent=True) or {}
    client_id = str(data.get("client_id") or "").strip()
    client_secret = str(data.get("client_secret") or "").strip()
    if not client_id:
        return jsonify({"error": "Enter the JobAdder Client ID"}), 400
    if not client_secret:
        with _ja_store_lock:
            stored_client_id = str(_ja_creds_store.get("client_id") or "").strip()
            stored_secret = str(_ja_creds_store.get("client_secret") or "").strip()
        if stored_secret and secrets.compare_digest(stored_client_id, client_id):
            client_secret = stored_secret
    if not client_secret:
        return jsonify({"error": "Enter the JobAdder Client Secret. A saved secret can only be reused with the same Client ID."}), 400
    if bool(data.get("save_only")):
        with _ja_store_lock:
            stored_client_id = str(_ja_creds_store.get("client_id") or "").strip()
        registration_changed = not stored_client_id or not secrets.compare_digest(
            stored_client_id, client_id
        )
        transition_started = False
        if registration_changed:
            conflict = _ja_begin_account_transition()
            if conflict:
                return _ja_account_transition_conflict(conflict)
            transition_started = True
        try:
            if registration_changed:
                # A different OAuth application registration cannot inherit any
                # account/tenant state from the previous registration.
                with _ja_refresh_lock:
                    _spider_preview_cancel_persistent_work()
                    with _ja_oauth_lock:
                        with _ja_store_lock:
                            previous = dict(_ja_creds_store)
                            try:
                                _ja_creds_store.clear()
                                _ja_creds_store.update({
                                    "client_id": client_id,
                                    "client_secret": client_secret,
                                })
                                _ja_save_store()
                            except Exception:
                                _ja_creds_store.clear()
                                _ja_creds_store.update(previous)
                                raise
                        _ja_oauth_sessions.clear()
                    _spider_resume_text_cache_clear()
                    _ppc_detail_cache_clear()
            else:
                with _ja_store_lock:
                    previous = dict(_ja_creds_store)
                    try:
                        _ja_creds_store["client_id"] = client_id
                        _ja_creds_store["client_secret"] = client_secret
                        _ja_save_store()
                    except Exception:
                        _ja_creds_store.clear()
                        _ja_creds_store.update(previous)
                        raise
            result = _ja_public_info()
            result["ok"] = True
            return jsonify(result)
        except Exception:
            return _cvstudio_error_payload(
                "JOBADDER_SETTINGS_SAVE_FAILED",
                "Could not save JobAdder settings in the protected credential store.",
                500,
                action="retry",
            )
        finally:
            if transition_started:
                _ja_end_account_transition()
    _ja_cleanup_sessions()
    session_id = secrets.token_urlsafe(24)
    state = secrets.token_urlsafe(32)
    # Serialize the brief session insertion with sign-out's transition flag.
    # A session created first is then cleared by sign-out; a later attempt gets
    # a visible conflict instead of surviving the completed sign-out.
    with _ja_critical_write_state_lock:
        if _ja_account_transition_in_progress:
            return _ja_account_transition_conflict("transition")
        with _ja_oauth_lock:
            _ja_oauth_sessions[session_id] = {
                "state": state,
                "client_id": client_id,
                "client_secret": client_secret,
                "created_at": time.time(),
                "expires_at": time.time() + _JA_SESSION_TTL,
                "status": "pending",
            }
    return jsonify({"login_session_id": session_id, "state": state, "expires_in": _JA_SESSION_TTL})


@app.route("/jobadder/callback", methods=["GET"])
def jobadder_callback():
    """Validate OAuth state, exchange the code, and bind result to one login session."""
    code = str(request.args.get("code") or "").strip()
    state = str(request.args.get("state") or "").strip()
    if not code or not state:
        return "<h3>JobAdder login failed: missing code or state.</h3>", 400
    _ja_cleanup_sessions()
    session_id = ""
    session = None
    with _ja_oauth_lock:
        for sid, candidate in _ja_oauth_sessions.items():
            # OAuth state is single-use. Mark it consumed while still holding
            # the lock so a callback reload or two simultaneous callbacks
            # cannot exchange the same authorization code twice or overwrite a
            # completed result with a later error.
            if candidate.get("status") == "pending" and secrets.compare_digest(str(candidate.get("state") or ""), state):
                session_id = sid
                session = dict(candidate)
                candidate["status"] = "exchanging"
                candidate["state"] = ""
                _ja_oauth_sessions[sid] = candidate
                break
    if not session_id or not session:
        return "<h3>JobAdder login expired, already completed, or state validation failed. Start Connect again.</h3>", 400
    try:
        token = _ja_exchange_token({
            "grant_type": "authorization_code",
            "code": code,
            "redirect_uri": "http://localhost:5000/jobadder/callback",
            "client_id": session["client_id"],
            "client_secret": session["client_secret"],
        })
        # Keep the token bound to this one login session until the initiating
        # browser tab claims it. A second callback cannot overwrite the active
        # JobAdder account before its own tab finishes polling.
        completed = _ja_finish_oauth_exchange_if_active(
            session_id,
            {
                "status": "complete",
                "token": token,
                "activated": False,
                "expires_at": time.time() + 300,
            },
        )
        if not completed:
            return "<h3>JobAdder login was cancelled or expired. Start Connect again.</h3>", 409
        return """<html><body style="font-family:sans-serif;text-align:center;padding-top:80px;">
            <h2>&#x2705; JobAdder Connected!</h2><p>You can close this window and return to CV Studio.</p>
            <script>window.close();</script></body></html>"""
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode(errors="replace")[:800]
        active = _ja_finish_oauth_exchange_if_active(
            session_id,
            {
                "status": "error",
                "error": "JobAdder token exchange failed",
                "detail": detail,
                "expires_at": time.time() + 300,
            },
            remove_keys=("client_secret",),
        )
        if not active:
            return "<h3>JobAdder login was cancelled or expired. Start Connect again.</h3>", 409
        return "<h3>JobAdder token exchange failed.</h3><pre>{}</pre>".format(html.escape(detail)), exc.code
    except Exception as exc:
        active = _ja_finish_oauth_exchange_if_active(
            session_id,
            {
                "status": "error",
                "error": str(exc),
                "expires_at": time.time() + 300,
            },
            remove_keys=("client_secret",),
        )
        if not active:
            return "<h3>JobAdder login was cancelled or expired. Start Connect again.</h3>", 409
        return "<h3>JobAdder token exchange failed.</h3>", 500


@app.route("/jobadder/poll_token", methods=["POST"])
def jobadder_poll_token():
    """Return only public connection status for the initiating login session."""
    data = request.get_json(silent=True) or {}
    session_id = str(data.get("login_session_id") or "").strip()
    if not session_id:
        return jsonify({"error": "Missing login_session_id"}), 400
    _ja_cleanup_sessions()
    with _ja_oauth_lock:
        session = dict(_ja_oauth_sessions.get(session_id) or {})
        if not session:
            return jsonify({"error": "JobAdder login session expired"}), 404
        status = session.get("status") or "pending"
        if status == "pending":
            return jsonify({"status": "pending"}), 202
        if status == "error":
            return jsonify({"status": "error", "error": session.get("error") or "JobAdder login failed", "detail": session.get("detail") or ""}), 400
        if not session.get("activated"):
            token = dict(session.get("token") or {})
            if not token.get("access_token"):
                return jsonify({"status": "error", "error": "JobAdder login result is incomplete"}), 500
            _ja_apply_token(token, session.get("client_id"), session.get("client_secret"), clear_spider_cache=True)
            session["activated"] = True
            session["activated_at"] = time.time()
            session.pop("token", None)
            session.pop("client_secret", None)
            session["result"] = _ja_public_info()
            _ja_oauth_sessions[session_id] = session
        result = dict(session.get("result") or _ja_public_info())
    result.update({"status": "complete", "connected": True})
    return jsonify(result)


def _ja_normalize_api_base(value):
    """Return a safe JobAdder API base without guessing a tenant region.

    OAuth token responses normally provide the tenant-specific API URL. Keep
    that exact JobAdder-owned HTTPS host when present; otherwise use JobAdder's
    generic documented endpoint. Never silently force AU3, which breaks EU/US
    tenants, and never allow a migrated browser value to redirect bearer tokens
    to a non-JobAdder host.
    """
    raw = str(value or "").strip().rstrip("/") or "https://api.jobadder.com/v2"
    try:
        parsed = urllib.parse.urlsplit(raw)
        host = str(parsed.hostname or "").lower().rstrip(".")
        if parsed.scheme.lower() != "https" or not (host == "jobadder.com" or host.endswith(".jobadder.com")):
            return "https://api.jobadder.com/v2"
        path = (parsed.path or "").rstrip("/")
        if not path:
            path = "/v2"
        return "https://{}{}".format(host, path)
    except Exception:
        return "https://api.jobadder.com/v2"


def _ja_api(path):
    """Build JobAdder API URL using the token-provided regional base."""
    base = _ja_normalize_api_base(_ja_creds_store.get("api_url"))
    return base + "/" + path.lstrip("/")


_JOBADDER_CLIENT = JobAdderClient(
    api_base_provider=lambda: _ja_api(""),
    token_provider=lambda force=False: _ja_refresh_access_token(force=force),
    reconnect_handler=lambda: _ja_mark_reconnect_required(),
)

_CVSTUDIO_JOBADDER_READ_SERVICE = JobAdderReadService(
    jsonify=lambda payload: jsonify(payload),
    query_arg=lambda name, default="": request.args.get(name, default),
    refresh_token=lambda force=False: _ja_refresh_access_token(force=force),
    client=lambda: _JOBADDER_CLIENT,
    public_info=lambda: _ja_public_info(),
    creds_store=lambda: _ja_creds_store,
    ja_api=lambda path: _ja_api(path),
)

_CVSTUDIO_JOBADDER_WRITE_SERVICE = JobAdderWriteService(
    jsonify=lambda payload: jsonify(payload),
    request_json=lambda: request.get_json(silent=True) or {},
    refresh_token=lambda force=False: _ja_refresh_access_token(force=force),
    client=lambda: _JOBADDER_CLIENT,
)


@app.route("/jobadder/restore_token", methods=["POST"])
@_ja_critical_write_route
def jobadder_restore_token():
    """One-time migration of legacy browser tokens into protected backend storage."""
    data = request.get_json(silent=True) or {}
    token = str(data.get("access_token") or "").strip()
    refresh = str(data.get("refresh_token") or "").strip()
    if not token and not refresh:
        return jsonify({"ok": True, "connected": bool(_ja_creds_store.get("access_token"))})
    with _ja_store_lock:
        previous = dict(_ja_creds_store)
        try:
            if token:
                _ja_creds_store["access_token"] = token
            if refresh:
                _ja_creds_store["refresh_token"] = refresh
            _ja_creds_store["_needs_reconnect"] = False
            for key in ("api_url", "client_id", "client_secret"):
                if data.get(key):
                    _ja_creds_store[key] = _ja_normalize_api_base(data.get(key)) if key == "api_url" else str(data.get(key))
            try:
                _ja_creds_store["expires_at"] = int(data.get("expires_at") or time.time() + 1800)
            except Exception:
                _ja_creds_store["expires_at"] = int(time.time() + 1800)
            _ja_save_store()
            try:
                _spider_resume_text_cache_clear()
            except Exception:
                pass
        except Exception:
            _ja_creds_store.clear(); _ja_creds_store.update(previous)
            raise
    return jsonify({"ok": True, "connected": True})


@app.route("/jobadder/refresh_token", methods=["POST"])
def jobadder_refresh_token():
    try:
        token = _ja_refresh_access_token(force=True)
        info = _ja_public_info()
        if not token and info.get("needs_reconnect"):
            info["error"] = "JobAdder connection expired. Reconnect JobAdder in Settings."
            return jsonify(info), 401
        return jsonify(info)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/jobadder/disconnect", methods=["POST"])
def jobadder_disconnect():
    with _ja_store_lock:
        _ja_creds_store.clear()
        _cv_secure_delete("jobadder")
    with _ja_oauth_lock:
        _ja_oauth_sessions.clear()
    try:
        _spider_resume_text_cache_clear()
    except Exception:
        pass
    return jsonify({"ok": True, "connected": False})


@app.route("/jobadder/sign_out", methods=["POST"])
def jobadder_sign_out():
    """Sign out locally while retaining the protected application registration."""
    conflict = _ja_begin_account_transition()
    if conflict:
        return _ja_account_transition_conflict(conflict)
    try:
        # Wait for an already-running refresh instead of cancelling its network
        # request. New unsafe writes are blocked by the transition flag.
        with _ja_refresh_lock:
            # Persistent prefetch cancellation must be durable before the
            # credential record changes. A journal write failure is visible.
            _spider_preview_cancel_persistent_work()
            with _ja_oauth_lock:
                with _ja_store_lock:
                    previous = dict(_ja_creds_store)
                    reduced = {
                        key: str(_ja_creds_store.get(key) or "").strip()
                        for key in ("client_id", "client_secret")
                        if str(_ja_creds_store.get(key) or "").strip()
                    }
                    try:
                        _ja_creds_store.clear()
                        _ja_creds_store.update(reduced)
                        _ja_save_store()
                    except Exception:
                        _ja_creds_store.clear()
                        _ja_creds_store.update(previous)
                        raise
                _ja_oauth_sessions.clear()
            _spider_resume_text_cache_clear()
            _ppc_detail_cache_clear()
        return jsonify({
            "ok": True,
            "connected": False,
            "client_id": str(_ja_creds_store.get("client_id") or ""),
            "client_secret_configured": bool(
                str(_ja_creds_store.get("client_secret") or "").strip()
            ),
        })
    except Exception:
        return _cvstudio_error_payload(
            "JOBADDER_SIGN_OUT_FAILED",
            "Could not sign out from JobAdder in CV Studio. The protected account state was not silently discarded.",
            500,
            action="retry",
        )
    finally:
        _ja_end_account_transition()


@app.route("/jobadder/api_info", methods=["GET"])
def jobadder_api_info():
    return _CVSTUDIO_JOBADDER_READ_SERVICE.api_info()

@app.route("/jobadder/search_candidate", methods=["GET"])
def jobadder_search_candidate():
    return _CVSTUDIO_JOBADDER_READ_SERVICE.search_candidate()



# ── Microsoft OneNote (Graph) import connector ────────────────────────
def _ms_ssl_context():
    return ssl.create_default_context(cafile=certifi.where()) if certifi else None

# The OneNote Graph connection layer (token store, lifecycle, shared Graph
# client, JSON/bytes helpers, device sessions and the six connection handlers)
# lives in the OneNoteGraphService in cvstudio_msgraph (Phase 7B). It is
# instantiated further below once the secure-store and token-endpoint helpers
# exist; app-level aliases (_ms_graph_json/_ms_graph_store/etc.) keep the
# OneNote content handlers and the external MS-status surface unchanged. The
# OutlookService is also imported here. This import never pulls app into the
# module.
from cvstudio_msgraph import (
    OneNoteGraphService,
    OutlookService,
    _ms_outlook_account_normalize,
    _ms_outlook_error_payload,
    _ms_outlook_validate_draft_input,
    _ms_safe_tenant,
    _onenote_parse_date_bound,
)

from cvstudio_onenote_text import (
    _OneNoteHtmlTextParser,
    _onenote_html_to_text,
    _onenote_graph_id,
    _onenote_normalize_date_mode,
    _onenote_date_filtered_items,
    _onenote_text_search_items,
    _onenote_extract_urls,
    _onenote_url_parts,
    _onenote_manual_search_terms,
    _onenote_match_notebook_from_manual,
    _onenote_link_hrefs,
    _onenote_manual_candidates,
    _onenote_match_section_from_manual,
    _onenote_iso,
    _onenote_desktop_parse_sections_xml,
    _onenote_xml_nodes,
    _onenote_page_in_date_range,
)

def _ms_token_response(req_data, tenant="common", timeout=20):
    return _ONENOTE_GRAPH_CLIENT.token_request(
        req_data,
        tenant=_ms_safe_tenant(tenant),
        endpoint="token",
        timeout=timeout,
    )


# The PPC rich-Outlook-draft token store, lifecycle and handlers live in the
# OutlookService in cvstudio_msgraph (Phase 7B). The generic secret vault below
# still shares two low-level primitives with it — the Outlook token-store
# cryptography and the macOS Keychain availability/account naming — so those
# stay here. The Outlook crypto keystream/DPAPI transform are imported for the
# generic _cv_secure_* store.
from cvstudio_outlook_crypto import (
    keystream as _ms_outlook_keystream,
    dpapi_transform as _ms_outlook_dpapi_transform,
)


def _ms_outlook_keychain_account():
    return "cvstudio-" + _install_receipt_machine_hash()[:24]


def _ms_outlook_keychain_available():
    if sys.platform != "darwin":
        return False
    try:
        proc = subprocess.run(["security", "help"], capture_output=True, timeout=5, check=False)
        return proc.returncode in (0, 1)
    except Exception:
        return False


# Generic backend secret vault for JobAdder, OneNote and AI/search keys. It uses
# Windows DPAPI, macOS Keychain, or an authenticated machine-bound fallback.
def _cv_secure_service_name(service):
    clean = re.sub(r"[^a-z0-9_-]", "", str(service or "").lower())
    if clean not in {"jobadder", "onenote", "ai"}:
        raise ValueError("Unsupported secure-store service")
    return clean


def _cv_secure_paths(service):
    name = _cv_secure_service_name(service)
    base = os.path.dirname(_install_receipt_path())
    return (
        os.path.join(base, "{}_secrets_v1.dpapi".format(name)),
        os.path.join(base, "{}_secrets_v1.json".format(name)),
    )


def _cv_secure_fallback_key(service):
    material = ("CVStudio|SecureStore|{}|{}".format(_cv_secure_service_name(service), _install_receipt_machine_hash())).encode("utf-8")
    return _receipt_hmac.new(_install_receipt_signing_key(), material, _receipt_hashlib.sha256).digest()


def _cv_secure_protect(service, record):
    key = _cv_secure_fallback_key(service)
    nonce = os.urandom(16)
    plain = json.dumps(record, ensure_ascii=False, separators=(",", ":"), sort_keys=True).encode("utf-8")
    stream = _ms_outlook_keystream(key, nonce, len(plain))
    cipher = bytes(a ^ b for a, b in zip(plain, stream))
    mac = _receipt_hmac.new(key, nonce + cipher, _receipt_hashlib.sha256).hexdigest()
    return {"schema": 1, "nonce": base64.urlsafe_b64encode(nonce).decode("ascii"), "ciphertext": base64.urlsafe_b64encode(cipher).decode("ascii"), "mac": mac}


def _cv_secure_unprotect(service, payload):
    key = _cv_secure_fallback_key(service)
    nonce = base64.urlsafe_b64decode(str((payload or {}).get("nonce") or "").encode("ascii"))
    cipher = base64.urlsafe_b64decode(str((payload or {}).get("ciphertext") or "").encode("ascii"))
    expected = _receipt_hmac.new(key, nonce + cipher, _receipt_hashlib.sha256).hexdigest()
    received = str((payload or {}).get("mac") or "")
    if not received or not _receipt_hmac.compare_digest(received, expected):
        raise ValueError("Secure-store integrity check failed")
    plain = bytes(a ^ b for a, b in zip(cipher, _ms_outlook_keystream(key, nonce, len(cipher))))
    data = json.loads(plain.decode("utf-8"))
    return data if isinstance(data, dict) else {}


def _cv_secure_load(service):
    name = _cv_secure_service_name(service)
    dpapi_path, fallback_path = _cv_secure_paths(name)
    try:
        if os.name == "nt" and os.path.exists(dpapi_path):
            with open(dpapi_path, "rb") as handle:
                data = json.loads(_ms_outlook_dpapi_transform(handle.read(), protect=False).decode("utf-8"))
            if isinstance(data, dict):
                data["_storage"] = "windows_dpapi"
                return data
        if sys.platform == "darwin" and _ms_outlook_keychain_available():
            proc = subprocess.run(["security", "find-generic-password", "-a", _ms_outlook_keychain_account(), "-s", "TheGuoLab.CVStudio." + name, "-w"], capture_output=True, text=True, timeout=10, check=False)
            if proc.returncode == 0:
                data = json.loads((proc.stdout or "{}").strip())
                if isinstance(data, dict):
                    data["_storage"] = "macos_keychain"
                    return data
        if os.path.exists(fallback_path):
            with open(fallback_path, "r", encoding="utf-8") as handle:
                data = _cv_secure_unprotect(name, json.load(handle))
            data["_storage"] = "machine_bound_file"
            return data
    except Exception:
        pass
    return {}


def _cv_secure_save(service, record):
    name = _cv_secure_service_name(service)
    clean = {k: v for k, v in dict(record or {}).items() if not str(k).startswith("_")}
    dpapi_path, fallback_path = _cv_secure_paths(name)
    text = json.dumps(clean, ensure_ascii=False, separators=(",", ":"), sort_keys=True)
    if os.name == "nt":
        os.makedirs(os.path.dirname(dpapi_path), exist_ok=True)
        temp = dpapi_path + ".tmp"
        with open(temp, "wb") as handle:
            handle.write(_ms_outlook_dpapi_transform(text.encode("utf-8"), protect=True))
        os.replace(temp, dpapi_path)
        try: os.remove(fallback_path)
        except Exception: pass
        return "windows_dpapi"
    if sys.platform == "darwin" and _ms_outlook_keychain_available():
        proc = subprocess.run(["security", "add-generic-password", "-U", "-a", _ms_outlook_keychain_account(), "-s", "TheGuoLab.CVStudio." + name, "-w", text], capture_output=True, text=True, timeout=15, check=False)
        if proc.returncode == 0:
            try: os.remove(fallback_path)
            except Exception: pass
            return "macos_keychain"
    os.makedirs(os.path.dirname(fallback_path), exist_ok=True)
    temp = fallback_path + ".tmp"
    with open(temp, "w", encoding="utf-8") as handle:
        json.dump(_cv_secure_protect(name, clean), handle, separators=(",", ":"))
        handle.flush()
        try:
            os.fsync(handle.fileno())
        except Exception:
            pass
    try:
        os.chmod(temp, 0o600)
    except Exception:
        pass
    os.replace(temp, fallback_path)
    try:
        os.chmod(fallback_path, 0o600)
    except Exception:
        pass
    return "machine_bound_file"


def _cv_secure_delete(service):
    name = _cv_secure_service_name(service)
    dpapi_path, fallback_path = _cv_secure_paths(name)
    for path in (dpapi_path, fallback_path):
        try: os.remove(path)
        except Exception: pass
    if sys.platform == "darwin":
        try:
            subprocess.run(["security", "delete-generic-password", "-a", _ms_outlook_keychain_account(), "-s", "TheGuoLab.CVStudio." + name], capture_output=True, timeout=10, check=False)
        except Exception:
            pass


# Load server-side integration secrets only after all vault primitives exist.
_loaded_ja_store = _cv_secure_load("jobadder")
if isinstance(_loaded_ja_store, dict) and _loaded_ja_store.get("api_url"):
    _loaded_ja_store["api_url"] = _ja_normalize_api_base(_loaded_ja_store.get("api_url"))
_ja_creds_store.update(_loaded_ja_store if isinstance(_loaded_ja_store, dict) else {})
# ── OneNote Graph connection service (Phase 7B) ────────────────────────────
# Instantiated here (after the secure store and token endpoint exist). The
# service owns the OneNote token store and loads it from the backend secure
# store in __init__. App-level aliases keep the OneNote content handlers, the
# shared _ms_token_response helper and the external MS-status surface working
# against the service's shared Graph client and store.
_ONENOTE_SERVICE = OneNoteGraphService(
    jsonify=lambda payload: jsonify(payload),
    request_json=lambda: request.get_json(silent=True) or {},
    request_args=lambda: request.args,
    token_response=_ms_token_response,
    secure_save=lambda service, record: _cv_secure_save(service, record),
    secure_load=lambda service: _cv_secure_load(service),
    secure_delete=lambda service: _cv_secure_delete(service),
    graph_client_factory=lambda token_provider, reconnect_handler: MicrosoftGraphClient(
        token_provider=token_provider,
        not_connected_message="Microsoft OneNote is not connected",
        reconnect_handler=reconnect_handler,
        context_provider=_ms_ssl_context,
    ),
)

# Shared aliases: the OneNote content handlers and desktop cluster still call
# these graph helpers, _ms_token_response uses the shared client, and the
# external MS-status surface reads the store (by reference, so it stays live).
_ONENOTE_GRAPH_CLIENT = _ONENOTE_SERVICE._graph_client
_ms_graph_store = _ONENOTE_SERVICE._store
_ms_graph_store_lock = _ONENOTE_SERVICE._store_lock
_ms_graph_device_store = _ONENOTE_SERVICE._device_store
_ms_graph_device_lock = _ONENOTE_SERVICE._device_lock
_ms_graph_json = _ONENOTE_SERVICE.graph_json
_ms_graph_post_json = _ONENOTE_SERVICE.graph_post_json
_ms_graph_bytes = _ONENOTE_SERVICE.graph_bytes
# The OneNote desktop route handlers (desktop_pages / manual_pages) that remain
# in the shell still call these page-listing helpers, now owned by the service.
_onenote_list_notebook_pages = _ONENOTE_SERVICE._list_notebook_pages
_onenote_list_section_pages = _ONENOTE_SERVICE._list_section_pages
_onenote_get_all_notebooks_and_sections = _ONENOTE_SERVICE._get_all_notebooks_and_sections
_ai_secret_store = _cv_secure_load("ai")

_AI_SECRET_SLOTS = frozenset({
    "main_anthropic", "main_deepseek", "main_openai",
    "lead_anthropic", "lead_deepseek", "lead_openai",
    "search_tavily", "search_serpapi", "enrichment_apollo",
})


def _ai_secret_slot(provider, scope="main"):
    p = str(provider or "anthropic").strip().lower()
    if p in {"claude", "anthropic"}: p = "anthropic"
    elif p in {"gpt", "openai"}: p = "openai"
    elif p != "deepseek": p = "anthropic"
    slot = "{}_{}".format("lead" if scope == "lead" else "main", p)
    return slot


def _resolve_request_api_key(body, scope="main"):
    body = body if isinstance(body, dict) else {}
    supplied = str(body.get("api_key") or "").strip()
    if supplied and supplied != "__BACKEND_SECURE__":
        return supplied
    slot = str(body.get("api_key_slot") or "").strip().lower()
    if slot not in _AI_SECRET_SLOTS:
        slot = _ai_secret_slot(body.get("provider") or body.get("llm_provider"), scope)
    return str(_ai_secret_store.get(slot) or "").strip()


_CVSTUDIO_SECRETS_SERVICE = SecretsService(
    jsonify=lambda payload: jsonify(payload),
    request_json=lambda: request.get_json(silent=True) or {},
    secret_store=lambda: _ai_secret_store,
    secret_slots=lambda: _AI_SECRET_SLOTS,
    secure_save=lambda service, record: _cv_secure_save(service, record),
    secure_delete=lambda service: _cv_secure_delete(service),
)


@app.route("/secure-secrets/info", methods=["GET"])
def secure_secrets_info():
    return _CVSTUDIO_SECRETS_SERVICE.info()


@app.route("/secure-secrets/save", methods=["POST"])
def secure_secrets_save():
    return _CVSTUDIO_SECRETS_SERVICE.save()


@app.route("/secure-secrets/clear", methods=["POST"])
def secure_secrets_clear():
    return _CVSTUDIO_SECRETS_SERVICE.clear()

# ── Outlook rich-draft service (Phase 7B) ──────────────────────────────────
# The Outlook token store, lifecycle, MS-Graph calls, device-login sessions,
# draft idempotency cache and the eight /ppc/outlook/* handler bodies live in
# the OutlookService. The web shell owns the token endpoint, the graph client
# factory and the install-receipt identity, and keeps the Flask routes.
_OUTLOOK_SERVICE = OutlookService(
    jsonify=lambda payload: jsonify(payload),
    request_json=lambda: request.get_json(silent=True) or {},
    token_response=_ms_token_response,
    graph_client_factory=lambda token_provider, reconnect_handler: MicrosoftGraphClient(
        token_provider=token_provider,
        not_connected_message="Microsoft Outlook drafts are not connected",
        reconnect_handler=reconnect_handler,
        context_provider=_ms_ssl_context,
    ),
    receipt_path=_install_receipt_path,
    machine_hash=_install_receipt_machine_hash,
    signing_key=_install_receipt_signing_key,
)


# Two Outlook helpers are called elsewhere in the web shell (the combined MS
# status surface and the diagnostics probe); keep thin app-level aliases so
# those call sites stay unchanged.
def _ms_outlook_public_info():
    return _OUTLOOK_SERVICE.public_info()


def _ms_outlook_graph_json(path, body=None, method="GET", timeout=25):
    return _OUTLOOK_SERVICE.graph_json(path, body=body, method=method, timeout=timeout)


@app.route("/ppc/outlook/api_info", methods=["GET"])
def ppc_outlook_api_info():
    return _OUTLOOK_SERVICE.api_info()


@app.route("/ppc/outlook/test_connection", methods=["POST"])
def ppc_outlook_test_connection():
    return _OUTLOOK_SERVICE.test_connection()


@app.route("/ppc/outlook/refresh_token", methods=["POST"])
def ppc_outlook_refresh_token():
    return _OUTLOOK_SERVICE.refresh_token()


@app.route("/ppc/outlook/disconnect", methods=["POST"])
def ppc_outlook_disconnect():
    return _OUTLOOK_SERVICE.disconnect()


@app.route("/ppc/outlook/device_start", methods=["POST"])
def ppc_outlook_device_start():
    return _OUTLOOK_SERVICE.device_start()


@app.route("/ppc/outlook/device_poll", methods=["POST"])
def ppc_outlook_device_poll():
    return _OUTLOOK_SERVICE.device_poll()


@app.route("/ppc/outlook/create_draft", methods=["POST"])
def ppc_outlook_create_draft():
    return _OUTLOOK_SERVICE.create_draft()


@app.route("/ppc/outlook/create_test_draft", methods=["POST"])
def ppc_outlook_create_test_draft():
    return _OUTLOOK_SERVICE.create_test_draft()


@app.route("/onenote/api_info", methods=["GET"])
def onenote_api_info():
    return _ONENOTE_SERVICE.api_info()


@app.route("/onenote/store_token", methods=["POST"])
def onenote_store_token():
    return _ONENOTE_SERVICE.store_token()


@app.route("/onenote/refresh_token", methods=["POST"])
def onenote_refresh_token():
    return _ONENOTE_SERVICE.refresh_token()


@app.route("/onenote/disconnect", methods=["POST"])
def onenote_disconnect():
    return _ONENOTE_SERVICE.disconnect()


@app.route("/onenote/device_start", methods=["POST"])
def onenote_device_start():
    return _ONENOTE_SERVICE.device_start()


@app.route("/onenote/device_poll", methods=["POST"])
def onenote_device_poll():
    return _ONENOTE_SERVICE.device_poll()












@app.route("/onenote/notebooks", methods=["GET"])
def onenote_notebooks():
    return _ONENOTE_SERVICE.notebooks()

@app.route("/onenote/sections", methods=["GET"])
def onenote_sections():
    return _ONENOTE_SERVICE.sections()






@app.route("/onenote/section_pages", methods=["GET"])
def onenote_section_pages():
    return _ONENOTE_SERVICE.section_pages()



from cvstudio_onenote_desktop import (
    _onenote_com_error_label,
    _onenote_desktop_app,
    _onenote_desktop_available,
    _onenote_desktop_bstr_out,
    _onenote_desktop_get_hierarchy,
    _onenote_desktop_get_hierarchy_powershell,
    _onenote_desktop_get_page_content,
    _onenote_desktop_get_page_content_powershell,
    _onenote_desktop_local_notebook_sections,
    _onenote_desktop_open_hierarchy,
    _onenote_desktop_open_hierarchy_powershell,
    _onenote_desktop_page_text,
    _onenote_desktop_pages,
    _onenote_desktop_powershell,
    _onenote_desktop_resolve_manual_section,
    _onenote_desktop_sections,
    _onenote_match_desktop_section_from_manual,
)

@app.route("/onenote/desktop_sections", methods=["GET"])
def onenote_desktop_sections_route():
    """List desktop sections, including every section in a pasted local notebook folder."""
    manual = str(request.args.get("input") or request.args.get("link") or request.args.get("q") or "").strip()
    notebooks, sections, manual_warnings = [], [], []
    root_exc = None

    # Always resolve a pasted local notebook-folder link independently. This
    # supplements root COM enumeration, which can return only the first section
    # on some Microsoft 365 desktop OneNote installations.
    manual_notebooks, manual_sections = [], []
    if manual:
        try:
            manual_notebooks, manual_sections, manual_warnings = _onenote_desktop_local_notebook_sections(manual)
        except Exception as exc:
            manual_warnings = [str(exc)]

    try:
        notebooks, sections = _onenote_desktop_sections()
    except RuntimeError as exc:
        root_exc = exc
        notebooks, sections = [], []

    def merge_unique(base, extra, is_section=False):
        out, seen = [], set()
        for item in list(base or []) + list(extra or []):
            item = dict(item or {})
            key = str(item.get("id") or item.get("path") or item.get("displayName") or "").strip().lower()
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(item)
        return out

    notebooks = merge_unique(manual_notebooks, notebooks)
    sections = merge_unique(manual_sections, sections, True)
    for sec in sections:
        nbname = str((sec or {}).get("_parentNotebookName") or "").strip()
        label = str((sec or {}).get("displayName") or "Untitled section")
        group = str((sec or {}).get("_sectionGroupPath") or "").strip()
        if not sec.get("_label"):
            sec["_label"] = label + ((" · " + group) if group else "") + ((" · " + nbname) if nbname else " · Desktop OneNote")
        sec["_source"] = sec.get("_source") or "desktop"

    if sections:
        return jsonify({
            "items": sections,
            "notebooks": notebooks,
            "raw_count": len(sections),
            "source": "desktop",
            "manual_notebook_detected": bool(manual_sections),
            "manual_sections_count": len(manual_sections),
            "manual_warnings": manual_warnings[:20],
            "root_warning": str(root_exc) if root_exc else "",
        })

    # A full local Section.one link can still be resolved directly even if no
    # notebook-folder sections were found and root enumeration failed.
    if manual:
        try:
            sec = _onenote_desktop_resolve_manual_section(manual)
            if sec and sec.get("id"):
                sec["_label"] = (sec.get("displayName") or "Desktop OneNote section") + " · Direct section link"
                return jsonify({
                    "items": [sec], "notebooks": [], "raw_count": 1,
                    "source": "desktop", "resolved_directly": True,
                    "root_warning": str(root_exc) if root_exc else "",
                    "manual_warnings": manual_warnings[:20],
                })
        except Exception as direct_exc:
            if root_exc:
                msg = "{} | Direct section-link fallback also failed: {}".format(root_exc, direct_exc)
            else:
                msg = str(direct_exc)
            return jsonify({"error": msg, "desktop_available": False, "hint": "Keep Microsoft 365 desktop OneNote open. Paste the notebook-folder or full Section.one link, then click Load Notebooks / Sections."}), 500

    if root_exc:
        msg = str(root_exc)
        status = 503 if ("pywin32" in msg.lower() or "only available on windows" in msg.lower()) else 500
        return jsonify({"error": msg, "desktop_available": False, "hint": "Keep Microsoft 365 desktop OneNote open. Paste the local notebook-folder or full Section.one link before clicking Load Notebooks / Sections."}), status
    return jsonify({"items": [], "notebooks": [], "raw_count": 0, "source": "desktop", "manual_warnings": manual_warnings[:20]})

@app.route("/onenote/desktop_pages", methods=["GET"])
def onenote_desktop_pages_route():
    manual = str(request.args.get("input") or request.args.get("link") or request.args.get("q") or "").strip()
    direct_section_id = str(request.args.get("section_id") or "").strip()
    if direct_section_id.lower().startswith("desktop:"):
        direct_section_id = direct_section_id.split(":", 1)[1].strip()
    try:
        top = max(1, min(int(request.args.get("top") or 50), 100))
    except Exception:
        top = 50
    search = str(request.args.get("search") or "").strip()
    date_from_iso = _onenote_parse_date_bound(request.args.get("date_from") or request.args.get("from"), end_of_day=False)
    date_to_iso = _onenote_parse_date_bound(request.args.get("date_to") or request.args.get("to"), end_of_day=True)
    date_mode = _onenote_normalize_date_mode(request.args.get("date_mode") or request.args.get("date_type") or request.args.get("mode"))
    try:
        if direct_section_id:
            items, raw_count = _onenote_desktop_pages(direct_section_id, top, search, date_from_iso, date_to_iso, date_mode)
            return jsonify({"items":items, "raw_count":raw_count, "section_id":direct_section_id, "resolved_as":"selected desktop section", "source":"desktop"})

        # A notebook-folder link is a container, not a section. Do not
        # silently scan the first child section; return the choices so callers
        # must select the intended section id.
        if manual and (_onenote_manual_candidates(manual).get("local_notebook_paths") or []):
            folder_notebooks, folder_sections, folder_warnings = _onenote_desktop_local_notebook_sections(manual)
            if len(folder_sections) > 1:
                return jsonify({
                    "error": "Local OneNote notebook folder contains multiple sections. Choose a section from the Section dropdown before scanning.",
                    "requires_section_selection": True,
                    "sections": folder_sections,
                    "notebooks": folder_notebooks,
                    "raw_count": len(folder_sections),
                    "warnings": folder_warnings[:20],
                    "source": "desktop",
                }), 409
            if len(folder_sections) == 1 and folder_sections[0].get("id"):
                only = folder_sections[0]
                items, raw_count = _onenote_desktop_pages(only.get("id"), top, search, date_from_iso, date_to_iso, date_mode)
                return jsonify({
                    "items": items, "raw_count": raw_count,
                    "section_id": only.get("id"),
                    "resolved_as": "desktop notebook folder / only section: {}".format(only.get("displayName") or "section"),
                    "source": "desktop", "section_path": only.get("path") or "",
                    "warnings": folder_warnings[:20],
                })

        # Resolve local onenote:///C:\...\Section.one links directly first.
        # This path does not depend on successfully enumerating the root
        # hierarchy and therefore still works when GetHierarchy('', hsSections)
        # is unavailable on a particular Office/COM registration.
        direct_error = None
        if manual:
            try:
                direct = _onenote_desktop_resolve_manual_section(manual)
                if direct and direct.get("id"):
                    items, raw_count = _onenote_desktop_pages(direct.get("id"), top, search, date_from_iso, date_to_iso, date_mode)
                    return jsonify({
                        "items": items,
                        "raw_count": raw_count,
                        "section_id": direct.get("id"),
                        "resolved_as": "desktop section link: {}".format(direct.get("displayName") or "section"),
                        "source": "desktop",
                        "resolved_directly": True,
                        "section_path": direct.get("path") or "",
                    })
            except Exception as exc:
                direct_error = str(exc)

        notebooks, sections = _onenote_desktop_sections()
        if manual:
            sec = _onenote_match_desktop_section_from_manual(manual, sections)
            if sec and sec.get("id"):
                items, raw_count = _onenote_desktop_pages(sec.get("id"), top, search, date_from_iso, date_to_iso, date_mode)
                return jsonify({"items":items, "raw_count":raw_count, "section_id":sec.get("id"), "resolved_as":"desktop section: {}".format(sec.get("displayName") or "section"), "source":"desktop"})
            nb = _onenote_match_notebook_from_manual(manual, notebooks)
            if nb:
                all_pages, raw_count, seen = [], 0, set()
                nbname = str(nb.get("displayName") or "").lower()
                for sec2 in sections:
                    if nbname and nbname not in str(sec2.get("_parentNotebookName") or "").lower():
                        continue
                    try:
                        items, cnt = _onenote_desktop_pages(sec2.get("id"), top, search, date_from_iso, date_to_iso, date_mode)
                        raw_count += cnt
                        for it in items:
                            if it.get("id") not in seen:
                                it["_parentSectionName"] = sec2.get("displayName") or ""
                                seen.add(it.get("id")); all_pages.append(it)
                    except Exception:
                        continue
                all_pages.sort(key=lambda it: str((it or {}).get("lastModifiedDateTime") or (it or {}).get("createdDateTime") or ""), reverse=True)
                return jsonify({"items":all_pages[:top], "raw_count":raw_count, "notebook_id":nb.get("id"), "resolved_as":"desktop notebook: {}".format(nb.get("displayName") or "notebook"), "source":"desktop"})
        return jsonify({"error":"Could not resolve this desktop OneNote input. Try pasting the desktop onenote: section link, or type the exact section tab name.", "direct_link_error": direct_error or "", "desktop_notebooks_seen": len(notebooks), "desktop_sections_seen": len(sections), "desktop_sections_preview": [str((x or {}).get("displayName") or "") for x in sections[:12]], "terms_seen": _onenote_manual_search_terms(manual)[:8]}), 404
    except RuntimeError as e:
        msg = str(e)
        status = 503 if ("pywin32" in msg.lower() or "only available on windows" in msg.lower()) else 500
        return jsonify({"error": msg, "desktop_available": False, "hint": "Paste the full desktop onenote:///...Section.one link and click Scan Source. Re-run INSTALL.bat only if the detailed error says pywin32 or COM is unavailable."}), status
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/onenote/manual_pages", methods=["GET"])
def onenote_manual_pages():
    """Resolve a pasted OneNote section/page link, Graph id, or section name, then list pages for review."""
    manual = str(request.args.get("input") or request.args.get("link") or request.args.get("q") or "").strip()
    if not manual:
        return jsonify({"error":"Paste a OneNote section/page link, Graph ID, or section name first"}), 400
    try:
        top = max(1, min(int(request.args.get("top") or 50), 100))
    except Exception:
        top = 50
    search = str(request.args.get("search") or "").strip()
    date_from_iso = _onenote_parse_date_bound(request.args.get("date_from") or request.args.get("from"), end_of_day=False)
    date_to_iso = _onenote_parse_date_bound(request.args.get("date_to") or request.args.get("to"), end_of_day=True)
    date_mode = _onenote_normalize_date_mode(request.args.get("date_mode") or request.args.get("date_type") or request.args.get("mode"))
    try:
        cand = _onenote_manual_candidates(manual)
        for pid in cand.get("page_ids") or []:
            try:
                d = _ms_graph_json("me/onenote/pages/{}".format(_onenote_graph_id(pid)), params={"$select":"id,title,createdDateTime,lastModifiedDateTime,links"}, timeout=25)
                item = d if isinstance(d, dict) else {"id": pid, "title": "Manual OneNote page"}
                if _onenote_page_in_date_range(item, date_from_iso, date_to_iso, date_mode) and _onenote_text_search_items([item], search):
                    return jsonify({"items":[item], "raw_count":1, "resolved_as":"manual page link/id", "filters":{"search":search,"date_from":date_from_iso,"date_to":date_to_iso,"date_mode":date_mode,"requested_top":top}})
            except Exception:
                continue
        for sid in cand.get("section_ids") or []:
            try:
                items, raw_count = _onenote_list_section_pages(sid, top, search, date_from_iso, date_to_iso, date_mode)
                return jsonify({"items":items, "raw_count":raw_count, "section_id":sid, "resolved_as":"manual section link/id", "filters":{"search":search,"date_from":date_from_iso,"date_to":date_to_iso,"date_mode":date_mode,"requested_top":top}})
            except Exception:
                continue
        notebooks, sections = _onenote_get_all_notebooks_and_sections(200)
        for web_url in (cand.get("web_urls") or []):
            try:
                nb = _ms_graph_post_json("me/onenote/notebooks/GetNotebookFromWebUrl", {"webUrl": web_url}, timeout=25)
                if isinstance(nb, dict) and nb.get("id"):
                    items, raw_count = _onenote_list_notebook_pages(str(nb.get("id")), top, search, date_from_iso, date_to_iso, date_mode)
                    return jsonify({"items":items, "raw_count":raw_count, "notebook_id":str(nb.get("id")), "resolved_as":"notebook web URL: {}".format(nb.get("displayName") or "notebook"), "filters":{"search":search,"date_from":date_from_iso,"date_to":date_to_iso,"date_mode":date_mode,"requested_top":top}})
            except Exception:
                pass
        match = _onenote_match_section_from_manual(manual, sections)
        if match and match.get("id"):
            sid = str(match.get("id"))
            items, raw_count = _onenote_list_section_pages(sid, top, search, date_from_iso, date_to_iso, date_mode)
            label = str(match.get("displayName") or "section")
            nb = str(match.get("_parentNotebookName") or "")
            return jsonify({"items":items, "raw_count":raw_count, "section_id":sid, "resolved_as":"section match: {}{}".format(label, (" · " + nb) if nb else ""), "filters":{"search":search,"date_from":date_from_iso,"date_to":date_to_iso,"date_mode":date_mode,"requested_top":top}})
        nb_match = _onenote_match_notebook_from_manual(manual, notebooks)
        if nb_match and nb_match.get("id"):
            nbid = str(nb_match.get("id"))
            items, raw_count = _onenote_list_notebook_pages(nbid, top, search, date_from_iso, date_to_iso, date_mode)
            label = str(nb_match.get("displayName") or "notebook")
            return jsonify({"items":items, "raw_count":raw_count, "notebook_id":nbid, "resolved_as":"notebook match: {}".format(label), "filters":{"search":search,"date_from":date_from_iso,"date_to":date_to_iso,"date_mode":date_mode,"requested_top":top}})
        terms = _onenote_manual_search_terms(manual)
        return jsonify({
            "error":"Could not resolve this OneNote input. If this is a OneDrive notebook-root link, open the notebook/section in OneNote web and copy a section/page link, or paste the desktop onenote: link that includes the notebook/section name. Local-only desktop notebooks cannot be read automatically.",
            "notebooks_seen": len(notebooks),
            "sections_seen": len(sections),
            "terms_seen": terms[:8],
        }), 404
    except PermissionError as e:
        return jsonify({"error": str(e)}), 401
    except urllib.error.HTTPError as e:
        body = e.read().decode(errors="replace")
        return jsonify({"error":"Microsoft OneNote manual link failed", "status": e.code, "detail": body[:1000]}), e.code
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/onenote/import_selected", methods=["POST"])
def onenote_import_selected():
    return _ONENOTE_SERVICE.import_selected()

@app.route("/onenote/pages", methods=["GET"])
def onenote_pages():
    return _ONENOTE_SERVICE.pages()

@app.route("/onenote/page_content", methods=["GET"])
def onenote_page_content():
    return _ONENOTE_SERVICE.page_content()




@app.route("/onenote/import_recent", methods=["POST"])
def onenote_import_recent():
    return _ONENOTE_SERVICE.import_recent()


# ── OneNote → JobAdder Screening Call Notes ─────────────────────────
def _ja_post_json(path, payload, timeout=15):
    """POST JSON to a JobAdder API path using the current server-side token."""
    return _JOBADDER_CLIENT.request_json(
        path,
        method="POST",
        body=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
        },
        timeout=timeout,
        raw_fallback=True,
        safe_to_retry=False,
        retries=0,
    )

def _ja_get_json(path, timeout=15):
    """GET JSON from a JobAdder API path using the current server-side token."""
    return _JOBADDER_CLIENT.request_json(
        path,
        method="GET",
        headers={"Accept": "application/json"},
        timeout=timeout,
        raw_fallback=True,
    )


def _ja_put_json(path, payload, timeout=20):
    """PUT JSON to a JobAdder API path using the current server-side token."""
    return _JOBADDER_CLIENT.request_json(
        path,
        method="PUT",
        body=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
        },
        timeout=timeout,
        raw_fallback=True,
        safe_to_retry=False,
        retries=0,
    )


import cvstudio_ja_salary_ai as _cvstudio_salary_ai_module
from cvstudio_ja_salary_ai import (
    _ONENOTE_JA_PERMANENT_WORKTYPE_ID,
    _ja_compact_error_body,
    _ja_validate_currency_option_with_api,
    _ja_salary_cost_details,
    _ja_salary_ai_extract,
    _ja_build_salary_notice_canonical,
    _ja_submit_employment_from_candidate,
    _ja_candidate_salary_notice_payload,
    _ja_update_candidate_salary_notice,
)
_cvstudio_salary_ai_module.bind_salary_ai_dependencies(
    call_llm=lambda *a, **k: call_llm(*a, **k),
    _onenote_clean_field_value=lambda *a, **k: _onenote_clean_field_value(*a, **k),
    _llm_cost_details=lambda *a, **k: _llm_cost_details(*a, **k),
    _llm_paid_failure_fields=lambda *a, **k: _llm_paid_failure_fields(*a, **k),
    safe_json=lambda *a, **k: safe_json(*a, **k),
    _ja_get_json=lambda *a, **k: _ja_get_json(*a, **k),
    _ja_put_json=lambda *a, **k: _ja_put_json(*a, **k),
    _ja_salary_ai_cache_get=lambda *a, **k: _ja_salary_ai_cache_get(*a, **k),
    _ja_salary_ai_cache_put=lambda *a, **k: _ja_salary_ai_cache_put(*a, **k),
)


from cvstudio_ja_answers import (
    _ONENOTE_JA_SCREENING_SETTING_ID,
    _ja_maybe_int,
    _ja_answer_model_qid_list_payloads,
    _ja_rating_map_answer_payloads,
    _ja_candidate_screening_call_base_payload,
    _ja_answer_model_qid_list_payloads_v105,
    _ja_controlled_official_activity_payload,
)


# v24.6.97-v24.6.116: user captured the real JobAdder web payload for the tenant's
# Activities > Screening > Candidate Screening Call form.  Prefer the exact
# activity setting/question-ID shape before generic API guesses.  Do not include
# the captured activityID (it looks like JobAdder creates that per activity).
#
# The live API diagnostic then confirmed the valid endpoint is
# candidates/{id}/activities, but JobAdder rejected answers[] because that API
# expects answers to be an object (ActivityAnswerListModel), unlike the web UI
# devtools payload.  Therefore transfer now tries object-wrapped answer shapes
# first and keeps the browser-like array shape only as late diagnostics.
from cvstudio_ja_answers import (
    _ONENOTE_JA_PRESENTABILITY_QUESTION_IDS,
    _ja_presentability_answer,
    _ja_answers_object_variants,
    _ja_answer_is_presentability,
    _ja_answer_model_bundle,
    _ja_split_text_rating_answers,
)
_ONENOTE_JA_SCREENING_QUESTIONS_BASE = [
    (41172, "Brief Overview of Experience", "brief_overview"),
    (41173, "Reason For Leaving", "reason_leaving"),
    (41174, "Looking For", "looking_for"),
    (41175, "Current Salary Breakdown", "current_salary_breakdown"),
    (41176, "Expected Salary", "expected_salary"),
    (41177, "Notice Period", "notice_period"),
    (35900, "Leads - Follow format: Active (Client - Role, Stage, Y/N)  Y denotes agency representation", "leads"),
    (41178, "Remarks and things to take note", "remarks"),
]


def _onenote_screening_remarks_value(fields, note_text=""):
    remarks = _onenote_clean_field_value(
        (fields or {}).get("remarks")
        or (fields or {}).get("red_flags")
        or (fields or {}).get("next_steps")
        or (fields or {}).get("raw_presentability")
        or ""
    )
    if remarks:
        return remarks
    # Keep remarks blank unless the note explicitly has a remarks-ish field.
    # Full raw note is not copied here to avoid dumping excessive text into a
    # tenant-specific field that is often meant for brief recruiter reminders.
    return ""


def _onenote_presentability_rating_int(fields):
    """Return a strict Presentability rating from 1 to 4, or ``None``.

    Accept the normal one-digit value, an explicit ``3/4`` / ``3 out of 4``
    score, or a labelled ``rating 3`` form.  Do not pull a stray 1-4 digit from
    values such as ``14``, ``40``, ``10/4`` or ``4 years``.
    """
    text = str((fields or {}).get("presentability_rating") or "").strip()
    if not text:
        return None
    m = re.fullmatch(
        r"(?:presentability|rating|score)?\s*[:=\-–—]?\s*([1-4])"
        r"(?:\s*(?:/\s*4|out\s+of\s+4))?"
        r"(?:\s*[-–—:]\s*[^\d].*)?",
        text,
        re.I,
    )
    if not m:
        m = re.fullmatch(r"\s*[^\d]{0,60}([1-4])\s*(?:/\s*4|out\s+of\s+4)\s*", text, re.I)
    return int(m.group(1)) if m else None




def _ja_screening_call_answers(fields, note_text="", presentability_question_id=None, presentability_mode="numeric"):
    fields = fields or {}
    rating = _onenote_presentability_rating_int(fields)
    values = {
        "brief_overview": _onenote_clean_field_value(fields.get("brief_overview", ""), max_len=8000),
        "reason_leaving": _onenote_clean_field_value(fields.get("reason_leaving", "")),
        "looking_for": _onenote_clean_field_value(fields.get("looking_for", "")),
        "current_salary_breakdown": _onenote_clean_field_value(fields.get("current_salary_breakdown", "")),
        "expected_salary": _onenote_clean_field_value(fields.get("expected_salary", "")),
        "notice_period": _onenote_clean_field_value(fields.get("notice_period", "")),
        "leads": _onenote_clean_field_value(fields.get("leads", "")),
        "remarks": _onenote_screening_remarks_value(fields, note_text),
    }
    answers = []
    for qid, qtext, key in _ONENOTE_JA_SCREENING_QUESTIONS_BASE:
        answers.append({
            "questionID": qid,
            "questionText": qtext,
            "textValue": values.get(key, ""),
            "startDateValue": None,
            "endDateValue": None,
            "numberValue": None,
            "numericValue": None,
            "decimalValue": None,
            "booleanValue": None,
            "singleSelectValue": None,
            "multiSelectValue": [],
        })
    answers.append(_ja_presentability_answer(
        presentability_question_id or _ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[0],
        "Presentability (Confidence, Comms, Business Awareness)",
        rating or "",
        mode=presentability_mode,
    ))
    return answers





















def _ja_browser_activity_answer_list_payload_variants(candidate_id, fields, note_text, reference):
    """Build browser-observed answer-list payloads for JobAdder API.

    v24.6.116: the user's successful JobAdder SPA save shows the tenant stores
    Screening Call fields as a flat answers array. The public candidate activity
    API rejects answers as a raw array, but its .NET error reveals that inside
    ActivityAnswerListModel, textAnswers must be an IList[ActivityTextAnswerModel]
    rather than a questionID map. Try object-wrapping the exact browser answer
    models as answers.textAnswers = [ ... ], including Presentability 62988 with
    numericValue + textValue.
    """
    raw_answers = _ja_screening_call_answers(
        fields,
        note_text,
        presentability_question_id=_ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[0],
        presentability_mode="number",
    )
    rating_int = _onenote_presentability_rating_int(fields)
    answer_list = []
    text_list = []
    rating_list = []
    for a in raw_answers:
        qid = a.get("questionID") or a.get("questionId")
        if not qid:
            continue
        qid_int = _ja_maybe_int(qid)
        is_rating = _ja_answer_is_presentability(a)
        model = {
            "questionID": qid_int,
            "questionText": a.get("questionText") or ("Presentability (Confidence, Comms, Business Awareness)" if is_rating else ""),
            "textValue": str(rating_int) if is_rating else (a.get("textValue") or ""),
            "numericValue": rating_int if is_rating else None,
            "startDateValue": None,
            "endDateValue": None,
        }
        answer_list.append(dict(model))
        if is_rating:
            rating_list.append(dict(model))
        else:
            text_list.append(dict(model))

    answers_payloads = [
        {"textAnswers": answer_list},
        {"textAnswers": text_list, "ratingAnswers": rating_list},
        {"activityTextAnswers": answer_list},
        {"activityAnswers": answer_list},
        {"questionAnswers": answer_list},
        {"items": answer_list},
        {"answers": answer_list},
        {"values": answer_list},
        {"TextAnswers": answer_list},
        {"TextAnswers": text_list, "RatingAnswers": rating_list},
        {"ActivityAnswers": answer_list},
        {"Items": answer_list},
    ]
    variants = []
    for answers_payload in answers_payloads:
        payload = _ja_candidate_screening_call_base_payload(candidate_id, reference=reference)
        payload["answers"] = answers_payload
        variants.append(payload)
    return variants


def _ja_precise_qid_activity_payload_variants(candidate_id, fields, note_text, reference):
    """Build precise v105 JobAdder answer-map attempts before older guesses."""
    answers = _ja_screening_call_answers(
        fields,
        note_text,
        presentability_question_id=_ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[0],
        presentability_mode="number",
    )
    rating = _onenote_presentability_rating_int(fields)
    variants = []
    for answers_payload in _ja_answer_model_qid_list_payloads_v105(answers, rating):
        payload = _ja_candidate_screening_call_base_payload(candidate_id, reference=reference)
        payload["answers"] = answers_payload
        variants.append(payload)
    return variants

def _ja_candidate_screening_call_payload(candidate_id, fields, note_text, reference=None, activity_id=None, answers_shape="items", presentability_question_id=None, presentability_mode="numeric"):
    cid_value = _ja_maybe_int(candidate_id)
    answers = _ja_screening_call_answers(fields, note_text, presentability_question_id=presentability_question_id, presentability_mode=presentability_mode)
    rating_required_validation = None
    if answers_shape.startswith("top_rating_"):
        rating = _onenote_presentability_rating_int(fields)
        text_answers, rating_answers, rating_by_qid = _ja_split_text_rating_answers(answers, rating)
        if answers_shape == "top_rating_items":
            answers_payload = {"items": text_answers}
        elif answers_shape == "top_rating_text_answers":
            answers_payload = {"textAnswers": text_answers, "items": text_answers}
        elif answers_shape == "top_rating_map":
            answers_payload = {str(a.get("questionID")): {"questionID": a.get("questionID"), "textValue": a.get("textValue") or ""} for a in text_answers}
        elif answers_shape == "top_rating_questions":
            answers_payload = {"questions": text_answers, "values": text_answers}
        else:
            answers_payload = {"items": text_answers}
        rating_required_validation = {"rating_answers": rating_answers, "rating_by_qid": rating_by_qid}
    elif answers_shape == "array":
        answers_payload = answers
    elif answers_shape == "map_full":
        answers_payload = {str(a.get("questionID")): dict(a) for a in answers}
    elif answers_shape == "map_text":
        answers_payload = {str(a.get("questionID")): (a.get("textValue") or "") for a in answers}
    elif answers_shape == "map_value":
        answers_payload = {}
        for a in answers:
            qid = str(a.get("questionID"))
            val = a.get("textValue")
            if not val and (a.get("ratingValue") is not None or a.get("numericValue") is not None or a.get("numberValue") is not None):
                val = a.get("ratingValue") if a.get("ratingValue") is not None else (a.get("numericValue") if a.get("numericValue") is not None else a.get("numberValue"))
            answers_payload[qid] = val if val is not None else ""
    elif answers_shape == "map_compact":
        answers_payload = {}
        for a in answers:
            qid = str(a.get("questionID"))
            if a.get("ratingValue") is not None or a.get("numericValue") is not None or a.get("numberValue") is not None:
                answers_payload[qid] = {
                    "questionID": a.get("questionID"),
                    "questionId": a.get("questionID"),
                    "ratingValue": a.get("ratingValue") if a.get("ratingValue") is not None else a.get("numericValue"),
                    "rating": a.get("rating") if a.get("rating") is not None else a.get("numericValue"),
                    "value": a.get("value") if a.get("value") is not None else a.get("numericValue"),
                    "numericValue": a.get("numericValue"),
                    "numberValue": a.get("numberValue"),
                    "textValue": a.get("textValue") or str(a.get("numericValue") or a.get("numberValue") or ""),
                }
            else:
                answers_payload[qid] = {
                    "questionID": a.get("questionID"),
                    "questionId": a.get("questionID"),
                    "textValue": a.get("textValue") or "",
                }
    elif answers_shape.startswith("rating_map_"):
        rating = _onenote_presentability_rating_int(fields)
        payloads = _ja_rating_map_answer_payloads(answers, rating)
        try:
            idx = int(answers_shape.rsplit("_", 1)[-1])
        except Exception:
            idx = 0
        answers_payload = payloads[idx] if 0 <= idx < len(payloads) else payloads[0]
    elif answers_shape in ("answer_bundle", "answer_bundle_object", "answer_bundle_pascal", "answer_bundle_maps", "answer_bundle_values"):
        rating = _onenote_presentability_rating_int(fields)
        bundle_mode = {
            "answer_bundle": "camel",
            "answer_bundle_object": "camel_object",
            "answer_bundle_pascal": "pascal",
            "answer_bundle_maps": "maps",
            "answer_bundle_values": "dollar_values",
        }.get(answers_shape, "camel")
        answers_payload = _ja_answer_model_bundle(answers, rating, bundle_mode=bundle_mode)
    else:
        wrappers = {
            "items": {"items": answers},
            "answers": {"answers": answers},
            "activityAnswers": {"activityAnswers": answers},
            "values": {"values": answers},
            "questions": {"questions": answers},
        }
        answers_payload = wrappers.get(answers_shape) or {"items": answers}
    payload = {
        "activitySettingID": _ONENOTE_JA_SCREENING_SETTING_ID,
        "actionName": "Candidate Screening Call",
        "activityType": "Screening",
        "answers": answers_payload,
        "attachments": [],
        "entity": {"entityID": cid_value},
        "entityID": cid_value,
        "mentionedUserIDs": [],
        "status": None,
        "task": None,
    }
    if rating_required_validation:
        rating_answers = rating_required_validation.get("rating_answers") or []
        rating_by_qid = rating_required_validation.get("rating_by_qid") or {}
        payload.update({
            "ratingAnswers": rating_answers,
            "ratings": rating_answers,
            "ratingQuestionAnswers": rating_answers,
            "activityRatingAnswers": rating_answers,
            "ratingValues": rating_answers,
            "ratingsByQuestionID": rating_by_qid,
            "ratingAnswersByQuestionID": rating_by_qid,
        })
    if reference:
        payload["reference"] = reference
    # Optional only: if a future diagnostic captures a required draft activityID,
    # the UI/backend can pass it without hard-coding the user's sample ID.
    if activity_id:
        payload["activityID"] = _ja_maybe_int(activity_id)
    return payload


def _ja_activity_payload_variants(candidate_id, activity_type, subject, note_text, fields, reference, now_utc):
    """Build JobAdder candidate activity payload variants.

    v24.6.97: Live diagnostics show candidates/{id}/activities is the
    reachable endpoint, but it rejects answers[] and expects an answers object
    (ActivityAnswerListModel). Therefore the first variants use exact tenant
    IDs with answers wrapped as objects. Browser-like answers[] remains last
    only for diagnostics.
    """
    cid = str(candidate_id or "")
    exact_variants = []
    # v24.6.116: Presentability is a mandatory 1-4 button/rating in JobAdder,
    # not a free-text field.  Try rating/numeric/select payloads before the
    # older text-style fallback.  Keep the attempt list short enough to avoid
    # slow transfers/timeouts.
    # Keep attempts limited: the user's captured browser payload confirms
    # Presentability questionID 62988, while 52988 is only a historical fallback.
    primary_presentability_qid = _ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[0]
    # v24.6.116: the successful JobAdder SPA payload is a flat answers array.
    # The API rejects answers as a raw array, but its model binder says
    # answers.textAnswers must be an IList. Try answers.textAnswers = [models]
    # first, including Presentability 62988 as numericValue/textValue.
    exact_variants.extend(_ja_browser_activity_answer_list_payload_variants(cid, fields, note_text, reference))
    # v24.6.106 fallback: qid-map/list variants retained for diagnostics.
    exact_variants.extend(_ja_precise_qid_activity_payload_variants(cid, fields, note_text, reference))
    # v24.6.116: try top-level rating collections before answer-object
    # variants, because live JobAdder says the mandatory rating question is
    # still missing even though the endpoint and activity setting are valid.
    # First try answer-object shapes that explicitly expose rating values by
    # question ID. This targets JobAdder's ActivityAnswerListModel binding where
    # normal answers are object-wrapped but rating controls require a separate
    # rating map/typed answer.
    for p_mode in ("rating_full", "rating_value", "number", "single_int"):
        for shape in ("rating_map_0", "rating_map_1", "rating_map_2", "rating_map_3", "rating_map_4", "rating_map_5", "rating_map_6"):
            exact_variants.append(_ja_candidate_screening_call_payload(
                cid, fields, note_text, reference=reference,
                answers_shape=shape, presentability_question_id=primary_presentability_qid,
                presentability_mode=p_mode,
            ))
    for p_mode in ("rating_full", "rating_value", "number", "single_int"):
        for shape in ("top_rating_items", "top_rating_text_answers", "top_rating_map", "top_rating_questions"):
            exact_variants.append(_ja_candidate_screening_call_payload(
                cid, fields, note_text, reference=reference,
                answers_shape=shape, presentability_question_id=primary_presentability_qid,
                presentability_mode=p_mode,
            ))
    # v24.6.116: diagnostics still say the mandatory rating question is not
    # bound. Try richer ActivityAnswerListModel-style objects before older maps.
    for p_mode in ("rating_full", "rating_value", "number", "single_int"):
        for shape in ("answer_bundle", "answer_bundle_object", "answer_bundle_pascal", "answer_bundle_maps", "answer_bundle_values", "map_full", "map_compact", "map_value", "items"):
            exact_variants.append(_ja_candidate_screening_call_payload(
                cid, fields, note_text, reference=reference,
                answers_shape=shape, presentability_question_id=primary_presentability_qid,
                presentability_mode=p_mode,
            ))
    # Fallbacks only if the tenant still expects a different legacy ID/shape.
    for presentability_qid in _ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[1:]:
        for shape in ("map_full", "map_compact", "items"):
            exact_variants.append(_ja_candidate_screening_call_payload(
                cid, fields, note_text, reference=reference,
                answers_shape=shape, presentability_question_id=presentability_qid,
                presentability_mode="rating_full",
            ))
    # Late diagnostics only: retained for tenants that unexpectedly accept the
    # original browser-like array or text maps.
    for shape in ("answers", "map_text", "array"):
        exact_variants.append(_ja_candidate_screening_call_payload(
            cid, fields, note_text, reference=reference,
            answers_shape=shape, presentability_question_id=primary_presentability_qid,
            presentability_mode="text",
        ))

    base = {
        "type": activity_type,
        "activityType": activity_type,
        "activityTypeName": activity_type,
        "subject": subject,
        "title": subject,
        "description": note_text,
        "text": note_text,
        "comments": note_text,
        "comment": note_text,
        "note": note_text,
        "completed": True,
        "completedAt": now_utc,
        "date": now_utc,
        "reference": reference,
        "fields": fields,
        "answers": {"items": _ja_screening_call_answers(fields, note_text)},
    }
    with_candidate = dict(base)
    with_candidate.update({
        "entityType": "candidate",
        "candidateId": cid,
        "candidate_id": cid,
        "entityId": cid,
        "entityID": _ja_maybe_int(cid),
        "entity": {"entityID": _ja_maybe_int(cid)},
        "candidate": {"candidateId": cid, "id": cid},
    })
    setting_named = dict(base)
    setting_named.update({
        "candidateId": cid,
        "activitySettingID": _ONENOTE_JA_SCREENING_SETTING_ID,
        "activitySetting": {"name": "Candidate Screening Call", "id": _ONENOTE_JA_SCREENING_SETTING_ID},
        "activitySettingName": "Candidate Screening Call",
        "actionName": "Candidate Screening Call",
        "activityType": "Screening",
        "isCompleted": True,
    })
    return exact_variants + [setting_named, with_candidate, base]



def _onenote_clean_note_text(value, max_len=25000):
    text = re.sub(r"\r\n?", "\n", str(value or ""))
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if len(text) > max_len:
        text = text[:max_len].rstrip() + "\n\n[CV Studio truncated this note because it was too long.]"
    return text


# v24.6.80: user clarified the OneNote Screening Call transfer may leave
# normal text fields blank when the candidate did not provide the information.
# Only Presentability must contain a 1-4 rating before transfer.
_ONENOTE_SCREENING_REQUIRED_FIELDS = [
    ("presentability_rating", "Presentability rating 1-4"),
]

def _onenote_clean_field_value(value, max_len=4000):
    text = re.sub(r"\r\n?", "\n", str(value or ""))
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) > max_len:
        text = text[:max_len].rstrip() + "…"
    return text

def _onenote_normalize_screening_fields(value):
    if not isinstance(value, dict):
        return {}
    out = {}
    for key in [
        "brief_overview", "reason_leaving", "looking_for", "current_salary_breakdown",
        "expected_salary", "notice_period", "leads", "remarks", "presentability_rating", "name", "email",
        "phone", "role", "location", "red_flags", "next_steps", "raw_presentability"
    ]:
        out[key] = _onenote_clean_field_value(value.get(key, ""), max_len=8000 if key == "brief_overview" else 4000)
    # Accept a few legacy/earlier field keys so older UI/state doesn't silently fail.
    if not out.get("current_salary_breakdown"):
        out["current_salary_breakdown"] = _onenote_clean_field_value(value.get("current_salary") or value.get("current_package") or "")
    if not out.get("looking_for"):
        out["looking_for"] = _onenote_clean_field_value(value.get("motivation") or value.get("interest") or value.get("target_role") or "")
    rating_int = _onenote_presentability_rating_int(out)
    out["presentability_rating"] = str(rating_int) if rating_int in (1, 2, 3, 4) else ""
    return out

def _onenote_missing_screening_fields(fields):
    missing = []
    for key, label in _ONENOTE_SCREENING_REQUIRED_FIELDS:
        if key == "presentability_rating":
            # Presentability is not a free-text field in JobAdder. It is a
            # mandatory 1/2/3/4 rating button. Treat anything else as missing
            # locally so we do not send a blank rating and receive JobAdder's
            # confusing "All rating mandatory questions are required" error.
            if _onenote_presentability_rating_int(fields) not in (1, 2, 3, 4):
                missing.append(label)
            continue
        if not _onenote_clean_field_value(fields.get(key, "")):
            missing.append(label)
    return missing



def _jobadder_web_base():
    api_url = (_ja_creds_store.get("api_url") or "").strip().rstrip("/")
    m = re.search(r"https?://(?:api\.([a-z0-9]+\.jobadder\.com)|([a-z0-9]+)api\.jobadder\.com)", api_url, re.I)
    if m:
        region = m.group(1) or ((m.group(2) or "") + ".jobadder.com")
        return "https://" + region
    return "https://app.jobadder.com"


def _ja_spa_screening_call_answers(fields, note_text="", include_extended=False):
    """Return the flat answer array observed from JobAdder's SPA UI.

    v24.6.116: the user captured a successful *new* Candidate Screening Call
    create request. The SPA create route posts to
    /spa/api/candidates/{id}/activities/standardactivity and includes only the
    standard six text questions plus Presentability. Older edit requests can
    include additional Leads/Remarks questions, but sending them during create
    is less safe, so the browser-create helper defaults to the captured create
    shape.
    """
    fields = fields or {}
    rating = _onenote_presentability_rating_int(fields)
    rating = int(rating) if str(rating or "").strip().isdigit() else None
    values = {
        "brief_overview": _onenote_clean_field_value(fields.get("brief_overview", ""), max_len=8000),
        "reason_leaving": _onenote_clean_field_value(fields.get("reason_leaving", "")),
        "looking_for": _onenote_clean_field_value(fields.get("looking_for", "")),
        "current_salary_breakdown": _onenote_clean_field_value(fields.get("current_salary_breakdown", "")),
        "expected_salary": _onenote_clean_field_value(fields.get("expected_salary", "")),
        "notice_period": _onenote_clean_field_value(fields.get("notice_period", "")),
        "leads": _onenote_clean_field_value(fields.get("leads", "")),
        "remarks": _onenote_screening_remarks_value(fields, note_text),
    }
    question_rows = [
        (41172, "Brief Overview of Experience", "brief_overview"),
        (41173, "Reason For Leaving", "reason_leaving"),
        (41174, "Looking For?", "looking_for"),
        (41175, "Current Salary Breakdown", "current_salary_breakdown"),
        (41176, "Expected Salary", "expected_salary"),
        (41177, "Notice Period", "notice_period"),
    ]
    if include_extended:
        question_rows.extend([
            (35900, "Leads - Follow format: Active (Client - Role, Stage, Y/N)  Y denotes agency representation", "leads"),
            (41178, "Remarks and things to take note", "remarks"),
        ])
    answers = []
    # v24.6.116: user clarified that if Brief Overview / Summary is empty,
    # leave Brief Overview blank instead of stuffing the raw OneNote note there.
    # Other standard create fields still use N/A when absent to avoid vague
    # JobAdder 500s on empty required text fields.
    for qid, qtext, key in question_rows:
        text_value = values.get(key, "")
        if not text_value and key != "brief_overview":
            text_value = "N/A"
        answers.append({
            "questionID": qid,
            "questionText": qtext,
            "textValue": text_value,
            "startDateValue": None,
            "endDateValue": None,
            "numericValue": None,
        })
    answers.append({
        "questionID": _ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[0],
        "questionText": "Presentability (Confidence, Comms, Business Awareness)",
        "textValue": str(rating or ""),
        "startDateValue": None,
        "endDateValue": None,
        "numericValue": rating,
    })
    return answers


def _ja_spa_screening_call_payload(candidate_id, fields, note_text="", activity_id=None, include_extended=False):
    """Build the JobAdder browser-SPA payload captured from manual Network requests.

    This payload is not sent by the Flask backend. It is used to generate a
    console helper script that the user runs inside their already-authenticated
    JobAdder browser tab, so cookies/tokens never need to be copied into CV Studio.
    """
    cid_value = _ja_maybe_int(candidate_id)
    payload = {
        "activityID": _ja_maybe_int(activity_id) if activity_id else 0,
        "activitySettingID": _ONENOTE_JA_SCREENING_SETTING_ID,
        "actionName": "Candidate Screening Call",
        "activityType": "Screening",
        "answers": _ja_spa_screening_call_answers(fields, note_text, include_extended=include_extended),
        "attachments": None,
        "entity": {"entityID": cid_value},
        "entityID": cid_value,
        "mentionedUserIDs": None,
        "status": None,
        "task": "",
    }
    return payload




from cvstudio_ja_salary_notice import (
    _JA_CURRENCY_CUSTOM_FIELD_ID,
    _JA_SALARY_ALLOWANCE_COMPONENT_RE,
    _JA_SALARY_FIXED_RECURRING_RE,
    _JA_SALARY_BASE_HINT_RE,
    _ja_spa_salary_update_helper_js,
    _ja_candidate_existing_currency,
    _ja_currency_option_for_code,
    _ja_currency_selection,
    _ja_salary_currency,
    _ja_extract_guaranteed_salary_months,
    _ja_calc_fixed_salary_detailed,
    _ja_expected_means_same_as_current,
    _ja_expected_is_open,
    _ja_build_expected_canonical,
    _ja_expected_salary_range,
    _ja_existing_current_salary,
    _ja_valid_iso_date,
    _ja_notice_availability,
    _ja_add_months_iso,
    _ja_notice_spa_payload,
    _ja_ai_number_matches_token,
    _ja_validate_ai_current_components,
    _ja_validate_ai_expected_components,
    _ja_current_from_ai_components,
    _ja_expected_from_ai_components,
)

def _ja_spa_browser_bridge(candidate_id, fields, note_text="", email="", salary_canonical=None, spelling_correction=True):
    """Return a safe copy/paste browser helper for the JobAdder SPA endpoint.

    v24.6.116: create-only. The user clarified OneNote transfer must always
    create a new Screening Call and must not update an existing Screening Call.
    The helper therefore tries likely JobAdder SPA create routes only. If all
    create routes fail, it prints a concise diagnostic and asks the user to
    capture the Network request from manually creating a brand-new Screening
    Call (not editing an existing one).
    """
    cid = str(candidate_id or "").strip()
    cid_q = urllib.parse.quote(cid, safe="")
    endpoint = "/spa/api/candidates/{}/activities/standardactivity".format(cid_q)
    profile_path = "/candidates/{}?tab=2&activityView=Activity".format(cid_q)
    # v24.6.116: captured successful create request uses standardactivity,
    # activityID: 0, attachments/mentionedUserIDs as null, task as empty string.
    # Add an in-browser sanitizer too, because older row.browser_bridge objects can
    # linger in the UI and because JobAdder returns only a vague HTTP 500 when a
    # required standard text field is blank.
    payload = _ja_spa_screening_call_payload(cid, fields, note_text, include_extended=False)
    if salary_canonical is None:
        salary_canonical = _ja_build_salary_notice_canonical(fields, {}, spelling_correction=spelling_correction)
    payload_json = json.dumps(payload, ensure_ascii=False, indent=2)
    compact_payload_json = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
    script = """(async () => {
  const helperVersion = 'v24.6.281';
  const candidateId = %s;
  const payload = %s;
  const profilePath = %s;
  const headers = {
    'Accept': 'application/json, text/plain, */*',
    'Content-Type': 'application/json',
    'x-jobadder-spa': '1'
  };
  function cleanText(value, fallback) {
    const s = String(value == null ? '' : value).replace(/\\s+/g, ' ').trim();
    return s || fallback || 'N/A';
  }
  function answerByQuestion(qid) {
    return (Array.isArray(payload.answers) ? payload.answers : []).find(a => Number(a && a.questionID) === Number(qid)) || {};
  }
  function textAnswer(qid, qtext) {
    const a = answerByQuestion(qid);
    const fallback = Number(qid) === 41172 ? '' : 'N/A';
    return {
      questionID: qid,
      questionText: qtext,
      textValue: cleanText(a.textValue, fallback),
      startDateValue: null,
      endDateValue: null,
      numericValue: null
    };
  }
  function ratingAnswer() {
    const a = answerByQuestion(62988);
    let rating = Number(a.numericValue || a.textValue || 0);
    if (![1, 2, 3, 4].includes(rating)) {
      throw new Error('Presentability must be 1, 2, 3 or 4 before creating the Screening Call.');
    }
    return {
      questionID: 62988,
      questionText: 'Presentability (Confidence, Comms, Business Awareness)',
      textValue: String(rating),
      startDateValue: null,
      endDateValue: null,
      numericValue: rating
    };
  }
  // Match the successful manual-create shape exactly: six text answers + rating.
  const createPayload = {
    activityID: 0,
    activitySettingID: 8225,
    actionName: 'Candidate Screening Call',
    activityType: 'Screening',
    answers: [
      textAnswer(41172, 'Brief Overview of Experience'),
      textAnswer(41173, 'Reason For Leaving'),
      textAnswer(41174, 'Looking For?'),
      textAnswer(41175, 'Current Salary Breakdown'),
      textAnswer(41176, 'Expected Salary'),
      textAnswer(41177, 'Notice Period'),
      ratingAnswer()
    ],
    attachments: null,
    entity: {entityID: Number(candidateId) || candidateId},
    entityID: Number(candidateId) || candidateId,
    mentionedUserIDs: null,
    status: null,
    task: ''
  };
  async function readText(res) {
    try { return await res.text(); } catch (e) { return String(e || ''); }
  }
%s
  console.log('CV Studio OneNote → CREATE-ONLY helper', helperVersion);
  console.log('CV Studio OneNote → sanitized JobAdder CREATE payload', createPayload);
  const url = '/spa/api/candidates/' + encodeURIComponent(candidateId) + '/activities/standardactivity';
  const res = await fetch(url, {
    method: 'POST',
    credentials: 'same-origin',
    headers,
    body: JSON.stringify(createPayload)
  });
  const text = await readText(res);
  console.log('CV Studio OneNote → create-only response', res.status, text);
  if (res.ok) {
    let salaryResult = null;
    try {
      salaryResult = await cvStudioOneNoteSalaryUpdate();
      console.log('CV Studio OneNote → salary/notice update result', salaryResult);
    } catch (salaryErr) {
      salaryResult = { ok: false, reason: String(salaryErr && salaryErr.message || salaryErr) };
      console.warn('CV Studio OneNote → salary/notice update failed after Screening Call create', salaryResult);
    }
    const salaryMsg = salaryResult && salaryResult.ok
      ? 'Salary/notice update also completed.'
      : 'Screening Call created, but salary/notice update needs review in console.';
    alert('CV Studio Screening Call created as a NEW activity. ' + salaryMsg + ' Refresh the Activity tab if you do not see it immediately.');
    if (location.pathname !== profilePath.split('?')[0]) location.href = profilePath;
    return;
  }
  console.error('CV Studio OneNote → CREATE-ONLY failed. No existing Screening Call was updated.', {status: res.status, response: text, payload: createPayload});
  alert([
    'CV Studio could not create a NEW Candidate Screening Call via JobAdder standardactivity.',
    'No existing Screening Call was updated.',
    '',
    'Please copy the failed Network request Payload and Preview/Response only.'
  ].join(String.fromCharCode(10)));
  throw new Error('Create-only JobAdder SPA attempt failed: HTTP ' + res.status);
})();""" % (json.dumps(cid), compact_payload_json, json.dumps(profile_path), _ja_spa_salary_update_helper_js(salary_canonical))
    return {
        "method": "CREATE-ONLY browser SPA helper using /standardactivity; never updates existing Screening Calls",
        "endpoint": endpoint,
        "fallback": "No update fallback. Uses the captured JobAdder SPA create route /activities/standardactivity. If it fails, capture the successful Network request from manually creating a brand-new Candidate Screening Call.",
        "profile_url": _jobadder_candidate_activity_url(cid),
        "payload": payload,
        "payload_json": payload_json,
        "script": script,
        "warning": "Run this only in the JobAdder candidate page DevTools Console while logged into JobAdder. It uses your browser session automatically and does not expose cookies to CV Studio. This script never updates an existing Screening Call; it posts only to /activities/standardactivity to create a new one. Blank standard text fields are filled with N/A for create compatibility.",
        "email": email,
        "candidate_id": cid,
        "salary_canonical": salary_canonical,
    }

def _jobadder_candidate_activity_url(candidate_id):
    cid = urllib.parse.quote(str(candidate_id or ""), safe="")
    return _jobadder_web_base() + "/candidates/{}?tab=2&activityView=Activity".format(cid)



# v24.6.137: production OneNote → JobAdder path based on JobAdder's official
# OpenAPI v2 contract. Source of truth:
# https://api.jobadder.com/v2/docs#operation/AddCandidateActivity
#
# POST/write model:
#   textAnswers [{questionId, text}]
#   ratingValueAnswers [{questionId, rating}]
# GET/read model is intentionally different and must not be reused for writes.
def _ja_official_screening_activity_payload(fields, note_text=""):
    fields = fields or {}
    rating = _onenote_presentability_rating_int(fields)
    if rating not in (1, 2, 3, 4):
        raise ValueError("Presentability rating must be 1-4")
    values = {
        "brief_overview": _onenote_clean_field_value(fields.get("brief_overview", ""), max_len=8000),
        "reason_leaving": _onenote_clean_field_value(fields.get("reason_leaving", "")),
        "looking_for": _onenote_clean_field_value(fields.get("looking_for", "")),
        "current_salary_breakdown": _onenote_clean_field_value(fields.get("current_salary_breakdown", "")),
        "expected_salary": _onenote_clean_field_value(fields.get("expected_salary", "")),
        "notice_period": _onenote_clean_field_value(fields.get("notice_period", "")),
        "leads": _onenote_clean_field_value(fields.get("leads", "")),
        "remarks": _onenote_screening_remarks_value(fields, note_text),
    }
    text_answers = []
    for question_id, _question_text, key in _ONENOTE_JA_SCREENING_QUESTIONS_BASE:
        value = values.get(key, "")
        # Normal Screening Call text fields are optional in this tenant. Omit
        # unanswered questions rather than inserting artificial N/A text.
        if value:
            text_answers.append({"questionId": int(question_id), "text": value})
    return {
        "activitySettingId": _ONENOTE_JA_SCREENING_SETTING_ID,
        "answers": {
            "textAnswers": text_answers,
            "listValueAnswers": [],
            "dateRangeValueAnswers": [],
            "ratingValueAnswers": [
                {"questionId": _ONENOTE_JA_PRESENTABILITY_QUESTION_IDS[0], "rating": int(rating)}
            ],
        },
    }



# JobAdder candidate custom field 4 is the single-value Currency field, confirmed
# by Kano's API-backed Panel implementation. The public OAuth API represents the
# same field through UpdateCandidate.custom [{fieldId, value}].
from cvstudio_salary_parse import (
    _JA_CURRENCY_OPTIONS,
    _JA_CURRENCY_ALIAS_PATTERNS,
    _JA_SALARY_ANNUAL_COMPONENT_RE,
    _JA_SALARY_EXCLUDED_COMPONENT_RE,
    _ja_round_currency,
    _ja_salary_currency_candidates_from_field,
    _ja_salary_currency_from_field,
    _ja_salary_plain_number,
    _ja_salary_has_actual_amount,
    _ja_salary_component_parts,
    _ja_salary_amount_tokens,
    _ja_salary_component_is_excluded,
    _ja_salary_currency_label,
    _ja_salary_display,
    _ja_salary_parse_range,
    _ja_salary_parse_percent_increment,
    _ja_salary_ai_num,
    _ja_salary_ai_sanitize,
    _ja_salary_usage_int,
    _ja_salary_cost_provenance,
    _ja_salary_llm_text,
)














# Recruiter salary notes are often typed quickly and contain predictable
# spelling variants.  Normalize only a conservative allow-list inside the two
# salary fields; never run broad fuzzy autocorrection that could silently turn
# unrelated prose into salary components.  The original raw note remains the
# stored/displayed breakdown.









































def _ja_bool_setting(value, default=True):
    if value is None:
        return bool(default)
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() not in ("0", "false", "off", "no", "disabled")









_SALARY_AI_CACHE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "runtime", "salary_ai_component_cache.json")
_SALARY_AI_CACHE_LOCK = threading.Lock()


from cvstudio_ja_salary_ai import SalaryAiCacheService as _SalaryAiCacheService

_SALARY_AI_CACHE_SERVICE = _SalaryAiCacheService(
    repository=lambda: _CVSTUDIO_SALARY_REPOSITORY,
    cache_path=lambda: _SALARY_AI_CACHE_PATH,
    lock=_SALARY_AI_CACHE_LOCK,
    legacy_json_read=_cvstudio_legacy_json_read,
    legacy_json_write=_cvstudio_legacy_json_write,
    storage_error=StorageError,
)


def _ja_salary_ai_cache_load():
    return _SALARY_AI_CACHE_SERVICE.load()


def _ja_salary_ai_cache_import_legacy_locked():
    return _SALARY_AI_CACHE_SERVICE.import_legacy_locked()


def _ja_salary_ai_cache_get(cache_key):
    return _SALARY_AI_CACHE_SERVICE.get(cache_key)


def _ja_salary_ai_cache_put(cache_key, components, provider, model):
    return _SALARY_AI_CACHE_SERVICE.put(cache_key, components, provider, model)

































@app.route("/jobadder/onenote_browser_bridge", methods=["POST"])
def jobadder_onenote_browser_bridge():
    """Generate a JobAdder-browser console helper for Screening Call creation.

    This does not contact JobAdder and does not require or read browser cookies.
    The generated script must be run by the user inside JobAdder's own page, so
    the request is same-origin and uses their normal JobAdder session safely.
    """
    data = request.get_json(silent=True) or {}
    candidate_id = str(data.get("candidate_id") or data.get("candidateId") or "").strip()
    email = str(data.get("email") or "").strip()
    if not candidate_id:
        return jsonify({"error": "Missing candidate_id"}), 400
    spelling_correction = _ja_bool_setting(data.get("spelling_correction"), default=True)
    fields = _onenote_normalize_screening_fields(data.get("fields") or {})
    if spelling_correction:
        fields = _ja_correct_screening_field_typos(fields, enabled=True)
        fields["notice_period"] = _ja_correct_notice_typos(fields.get("notice_period", ""), enabled=True)
    missing = _onenote_missing_screening_fields(fields)
    if missing:
        return jsonify({"error": "Missing mandatory Screening Call fields", "missing": missing}), 400
    note_text = _onenote_clean_note_text(data.get("note_text") or data.get("text") or "")
    # Never trust a client-supplied canonical salary object.  Rebuild from the
    # labelled fields so emergency profile values remain server-calculated.
    salary_canonical = _ja_build_salary_notice_canonical(fields, {}, spelling_correction=spelling_correction)
    bridge = _ja_spa_browser_bridge(candidate_id, fields, note_text, email=email, salary_canonical=salary_canonical, spelling_correction=spelling_correction)
    return jsonify({"ok": True, "browser_bridge": bridge, "salary_canonical": salary_canonical})

@app.route("/jobadder/onenote_log_screening", methods=["POST"])
@_ja_critical_write_route
def jobadder_onenote_log_screening():
    """Create a native Candidate Screening Call from a OneNote screening note.

    v24.6.137 uses JobAdder's documented AddCandidateActivity OAuth schema as
    the only normal backend activity write. It creates a NEW activity, never
    edits an existing Screening Call, and never creates a Candidate Note.
    After creation it separately updates parsed current/expected salary and
    notice-period profile fields through the documented UpdateCandidate model.
    """
    token = _ja_refresh_access_token(force=False)
    if not token:
        return jsonify({"error": "Not authenticated"}), 401
    data = request.get_json(silent=True) or {}
    candidate_id = str(data.get("candidate_id") or data.get("candidateId") or "").strip()
    email = str(data.get("email") or "").strip()
    if not candidate_id:
        return jsonify({"error": "Missing candidate_id"}), 400
    spelling_correction = _ja_bool_setting(data.get("spelling_correction"), default=True)
    fields = _onenote_normalize_screening_fields(data.get("fields") or {})
    if spelling_correction:
        fields = _ja_correct_screening_field_typos(fields, enabled=True)
        fields["notice_period"] = _ja_correct_notice_typos(fields.get("notice_period", ""), enabled=True)
    missing = _onenote_missing_screening_fields(fields)
    if missing:
        return jsonify({"error": "Missing mandatory Screening Call fields", "missing": missing}), 400
    note_text = _onenote_clean_note_text(data.get("note_text") or data.get("text") or "")
    if not note_text:
        return jsonify({"error": "Missing note_text"}), 400

    candidate_id_q = urllib.parse.quote(candidate_id, safe="")
    create_path = "candidates/{}/activities".format(candidate_id_q)
    activity_url = _jobadder_candidate_activity_url(candidate_id)
    payload = _ja_official_screening_activity_payload(fields, note_text)
    salary_ai_config = dict(data.get("salary_ai") or {}) if isinstance(data.get("salary_ai"), dict) else {}
    if salary_ai_config:
        # Browser requests carry only a backend-secure sentinel after the localStorage
        # migration. Resolve the real provider key inside the trusted backend before
        # invoking the narrowly scoped salary component extractor.
        salary_ai_config["api_key"] = _resolve_request_api_key(salary_ai_config, "main")
    ai_components, ai_processing = _ja_salary_ai_extract(fields, salary_ai_config)
    preview_salary_canonical = _ja_build_salary_notice_canonical(fields, {}, ai_components, ai_processing, spelling_correction=spelling_correction)
    try:
        status, result = _ja_post_json(create_path, payload, timeout=30)
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        browser_bridge = _ja_spa_browser_bridge(candidate_id, fields, note_text, email=email, salary_canonical=preview_salary_canonical)
        http_status = int(e.code)
        response_status = http_status if http_status in (401, 403) else 502
        return jsonify({
            "error": "Could not create JobAdder Candidate Screening Call through the official OAuth API. No Candidate Note was created.",
            "why": _ja_compact_error_body(body) or "JobAdder rejected AddCandidateActivity.",
            "next_step": "Use Copy Create Script only as the emergency fallback. It creates a new Screening Call through the logged-in JobAdder browser and never edits an existing Screening Call.",
            "candidate_id": candidate_id,
            "email": email,
            "activity_url": activity_url,
            "official_reference": "https://api.jobadder.com/v2/docs#operation/AddCandidateActivity",
            "browser_bridge": browser_bridge,
            "salary_canonical": preview_salary_canonical,
            "attempts": [{
                "path": create_path,
                "status": http_status,
                "detail": _ja_compact_error_body(body),
                "raw": body[:3000],
            }],
        }), response_status
    except PermissionError as e:
        return jsonify({"error": str(e)}), 401
    except Exception as e:
        browser_bridge = _ja_spa_browser_bridge(candidate_id, fields, note_text, email=email, salary_canonical=preview_salary_canonical)
        return jsonify({
            "error": "Could not create JobAdder Candidate Screening Call through the official OAuth API. No Candidate Note was created.",
            "why": str(e)[:1000],
            "candidate_id": candidate_id,
            "email": email,
            "activity_url": activity_url,
            "official_reference": "https://api.jobadder.com/v2/docs#operation/AddCandidateActivity",
            "browser_bridge": browser_bridge,
            "salary_canonical": preview_salary_canonical,
            "attempts": [{"path": create_path, "error": str(e)[:1000]}],
        }), 502

    activity_id = None
    if isinstance(result, dict):
        for key in ("activityId", "activityID"):
            try:
                activity_id = int(result.get(key))
                break
            except Exception:
                pass
    profile_update = _ja_update_candidate_salary_notice(candidate_id, fields, ai_components, ai_processing, spelling_correction=spelling_correction)
    warning_parts = []
    if not profile_update.get("ok"):
        warning_parts.append("Screening Call was created, but salary/notice profile update failed: {}".format(
            profile_update.get("detail") or profile_update.get("error") or "review required"
        ))
    elif profile_update.get("warning"):
        warning_parts.append("Screening Call was created; some profile values were not parsed: {}".format(profile_update.get("warning")))
    candidate_read_warning = profile_update.get("candidate_read_warning")
    if candidate_read_warning and not any(candidate_read_warning in str(part) for part in warning_parts):
        warning_parts.append(candidate_read_warning)
    warning = " ".join(p for p in warning_parts if p).strip()
    return jsonify({
        "ok": True,
        "mode": "candidate_activity",
        "path": create_path,
        "status": status,
        "candidate_id": candidate_id,
        "email": email,
        "activity_id": activity_id,
        "activity_url": activity_url,
        "screening_fields": fields,
        "result": result,
        "profile_update": profile_update,
        "salary_canonical": profile_update.get("salary_canonical"),
        "warning": warning or None,
        "official_reference": "https://api.jobadder.com/v2/docs#operation/AddCandidateActivity",
    })



def _ja_activity_diagnostic_response_headers(headers):
    """Keep useful response metadata while excluding cookie/authentication headers."""
    blocked = {"set-cookie", "set-cookie2", "authorization", "proxy-authorization"}
    try:
        return {str(k): str(v) for k, v in headers.items() if str(k).lower() not in blocked}
    except Exception:
        return {}


def _ja_activity_diagnostic_network_error(error):
    """Preserve the legacy diagnostic field with shared-client redaction."""
    if isinstance(error, ExternalServiceError):
        return str(getattr(error, "safe_detail", "") or error)
    return str(getattr(error, "reason", error))


def _ja_activity_diagnostic_get(path, timeout=25):
    """Perform one exact read-only JobAdder OAuth GET and preserve the raw response.

    The returned report deliberately excludes the OAuth token and request
    Authorization header. Response bodies are not shortened so the user can
    provide the complete model-binding evidence for review.
    """
    token = str(_ja_creds_store.get("access_token") or "").strip()
    if not token:
        raise PermissionError("Not authenticated")
    url = _ja_api(path)
    started = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    try:
        response = _JOBADDER_CLIENT.request_raw(
            path,
            method="GET",
            token=token,
            timeout=timeout,
            headers={"Accept": "application/json, text/plain, */*"},
        )
        raw_text = response.body.decode("utf-8", errors="replace")
        headers = _ja_activity_diagnostic_response_headers(response.headers)
        parsed = None
        if raw_text.strip():
            try:
                parsed = json.loads(raw_text)
            except Exception:
                parsed = None
        return {
            "method": "GET",
            "path": path,
            "url": url,
            "started_utc": started,
            "ok": True,
            "status": int(response.status or 200),
            "response_headers": headers,
            "response_body": raw_text,
            "response_json": parsed,
        }
    except urllib.error.HTTPError as e:
        raw_bytes = e.read()
        raw_text = raw_bytes.decode("utf-8", errors="replace")
        headers = _ja_activity_diagnostic_response_headers(e.headers)
        parsed = None
        if raw_text.strip():
            try:
                parsed = json.loads(raw_text)
            except Exception:
                parsed = None
        return {
            "method": "GET",
            "path": path,
            "url": url,
            "started_utc": started,
            "ok": False,
            "status": int(e.code),
            "reason": str(getattr(e, "reason", "") or ""),
            "response_headers": headers,
            "response_body": raw_text,
            "response_json": parsed,
        }
    except (urllib.error.URLError, ExternalServiceError) as e:
        return {
            "method": "GET",
            "path": path,
            "url": url,
            "started_utc": started,
            "ok": False,
            "status": None,
            "network_error": _ja_activity_diagnostic_network_error(e),
            "response_headers": {},
            "response_body": "",
            "response_json": None,
        }


def _ja_activity_diagnostic_summary(request_results):
    """Return a conservative summary without guessing undocumented API behavior."""
    statuses = [r.get("status") for r in (request_results or [])]
    numeric_statuses = [s for s in statuses if isinstance(s, int)]
    if any(200 <= s < 300 for s in numeric_statuses):
        return "At least one OAuth activity GET succeeded. Review the complete response body for the saved Screening Call and Presentability answer model."
    if len(numeric_statuses) == 2 and all(s in (404, 405) for s in numeric_statuses):
        return "Both OAuth activity GET routes returned 404/405. This is evidence that these read routes are unavailable, but the exact POST response must still be reviewed separately before making a final endpoint conclusion."
    if any(s in (401, 403) for s in numeric_statuses):
        return "JobAdder rejected at least one read request as unauthorised. Reconnect JobAdder OAuth before interpreting endpoint availability."
    return "The read requests did not expose a successful activity model. Review each full status and response body before deciding the next controlled test."


@app.route("/jobadder/onenote_activity_diagnostic", methods=["POST"])
def jobadder_onenote_activity_diagnostic():
    """Run the two exact read-only OAuth activity checks requested by the user.

    Safety boundary: this route only performs GET requests. It never creates or
    edits a Screening Call, Candidate Note, salary, notice period, or profile.
    """
    if not _ja_creds_store.get("access_token"):
        return jsonify({"error": "JobAdder is not connected. Reconnect OAuth in CV Studio first."}), 401
    data = request.get_json(silent=True) or {}
    candidate_id = str(data.get("candidate_id") or "").strip()
    activity_id = str(data.get("activity_id") or "").strip()
    if not re.fullmatch(r"\d+", candidate_id):
        return jsonify({"error": "Candidate ID must contain digits only."}), 400
    if not re.fullmatch(r"\d+", activity_id):
        return jsonify({"error": "Activity ID must contain digits only."}), 400

    cid_q = urllib.parse.quote(candidate_id, safe="")
    aid_q = urllib.parse.quote(activity_id, safe="")
    paths = [
        "candidates/{}/activities".format(cid_q),
        "candidates/{}/activities/{}".format(cid_q, aid_q),
    ]
    results = [_ja_activity_diagnostic_get(path, timeout=25) for path in paths]
    generated = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    report = {
        "diagnostic": "CV Studio JobAdder OAuth Candidate Activity Read Test",
        "cv_studio_version": _CVSTUDIO_VERSION,
        "generated_utc": generated,
        "safety": {
            "read_only": True,
            "http_methods_used": ["GET"],
            "created_or_updated_anything": False,
            "oauth_token_included_in_report": False,
        },
        "candidate_id": candidate_id,
        "activity_id": activity_id,
        "api_base": _ja_normalize_api_base(_ja_creds_store.get("api_url")),
        "requests": results,
        "summary": _ja_activity_diagnostic_summary(results),
        "next_boundary": "No OAuth POST was sent. Review this read report before authorising one controlled create test.",
    }
    filename = "jobadder_activity_diagnostic_{}_{}_{}.json".format(
        candidate_id,
        activity_id,
        re.sub(r"[^0-9]", "", generated)[:14],
    )
    return jsonify({"ok": True, "filename": filename, "report": report})


_JA_ACTIVITY_CREATE_DIAG_USED = set()


def _ja_activity_diagnostic_post(path, payload, timeout=25):
    """Send one OAuth POST and preserve the complete response for diagnostics.

    Authentication/cookie headers are never copied into the returned report.
    """
    token = _ja_refresh_access_token(force=False)
    if not token:
        raise PermissionError("JobAdder is not connected")
    url = _ja_api(path)
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    started = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    base = {
        "method": "POST",
        "path": path,
        "url": url,
        "started_utc": started,
        "request_payload": payload,
    }
    try:
        response = _JOBADDER_CLIENT.request_raw(
            path,
            method="POST",
            body=body,
            token=token,
            timeout=timeout,
            headers={
                "Accept": "application/json, text/plain, */*",
                "Content-Type": "application/json; charset=utf-8",
            },
            safe_to_retry=False,
            retries=0,
        )
        raw_text = response.body.decode("utf-8", errors="replace")
        parsed = None
        if raw_text.strip():
            try:
                parsed = json.loads(raw_text)
            except Exception:
                parsed = None
        base.update({
            "ok": True,
            "status": int(response.status or 200),
            "response_headers": _ja_activity_diagnostic_response_headers(response.headers),
            "response_body": raw_text,
            "response_json": parsed,
        })
        return base
    except urllib.error.HTTPError as e:
        raw_bytes = e.read()
        raw_text = raw_bytes.decode("utf-8", errors="replace")
        parsed = None
        if raw_text.strip():
            try:
                parsed = json.loads(raw_text)
            except Exception:
                parsed = None
        base.update({
            "ok": False,
            "status": int(e.code),
            "reason": str(getattr(e, "reason", "") or ""),
            "response_headers": _ja_activity_diagnostic_response_headers(e.headers),
            "response_body": raw_text,
            "response_json": parsed,
        })
        return base
    except (urllib.error.URLError, ExternalServiceError) as e:
        base.update({
            "ok": False,
            "status": None,
            "network_error": _ja_activity_diagnostic_network_error(e),
            "response_headers": {},
            "response_body": "",
            "response_json": None,
        })
        return base


def _ja_activity_ids_from_list_result(result):
    parsed = (result or {}).get("response_json")
    if not isinstance(parsed, dict):
        return []
    items = parsed.get("items")
    if not isinstance(items, list):
        return []
    out = []
    for item in items:
        if not isinstance(item, dict):
            continue
        value = item.get("activityId")
        if value is None:
            value = item.get("activityID")
        try:
            out.append(int(value))
        except Exception:
            pass
    return out


def _ja_activity_id_from_post_result(result):
    parsed = (result or {}).get("response_json")
    if isinstance(parsed, dict):
        for key in ("activityId", "activityID"):
            value = parsed.get(key)
            try:
                return int(value)
            except Exception:
                pass
        for container_key in ("activity", "item", "result", "data"):
            child = parsed.get(container_key)
            if isinstance(child, dict):
                for key in ("activityId", "activityID"):
                    value = child.get(key)
                    try:
                        return int(value)
                    except Exception:
                        pass
    headers = (result or {}).get("response_headers") or {}
    for key, value in headers.items():
        if str(key).lower() == "location":
            m = re.search(r"/activities/(\d+)(?:$|[/?#])", str(value))
            if m:
                return int(m.group(1))
    return None




@app.route("/jobadder/onenote_activity_create_diagnostic", methods=["POST"])
@_ja_critical_write_route
def jobadder_onenote_activity_create_diagnostic():
    """Create exactly one controlled Screening Call on the Max Low dummy fixture.

    This is intentionally separate from the normal OneNote transfer route. It
    never creates Candidate Notes, never edits an existing activity, never
    updates salary/notice/profile fields, and never changes the F12 helper.
    """
    if not _ja_creds_store.get("access_token"):
        return jsonify({"error": "JobAdder is not connected. Reconnect OAuth in CV Studio first."}), 401
    data = request.get_json(silent=True) or {}
    candidate_id = str(data.get("candidate_id") or "").strip()
    confirmation = str(data.get("confirmation") or "").strip()
    if candidate_id != "41262878":
        return jsonify({"error": "This controlled diagnostic is locked to Max Low candidate ID 41262878."}), 400
    if confirmation != "CREATE ONE MAX LOW TEST":
        return jsonify({"error": "Type CREATE ONE MAX LOW TEST exactly before running the controlled POST."}), 400

    guard_key = (_CVSTUDIO_VERSION, candidate_id)
    if guard_key in _JA_ACTIVITY_CREATE_DIAG_USED:
        return jsonify({"error": "The one-shot controlled POST has already been run in this CV Studio session. Restarting is intentionally required before any repeat test."}), 409
    # Mark before the network call so a timeout/double-click cannot emit a second POST.
    _JA_ACTIVITY_CREATE_DIAG_USED.add(guard_key)

    cid_q = urllib.parse.quote(candidate_id, safe="")
    list_path = "candidates/{}/activities".format(cid_q)
    create_path = list_path
    before = _ja_activity_diagnostic_get(list_path, timeout=25)
    before_ids = set(_ja_activity_ids_from_list_result(before))
    payload = _ja_controlled_official_activity_payload(candidate_id)
    post_result = _ja_activity_diagnostic_post(create_path, payload, timeout=30)
    after = _ja_activity_diagnostic_get(list_path, timeout=25)
    after_ids = set(_ja_activity_ids_from_list_result(after))
    new_ids = sorted(after_ids - before_ids)
    created_activity_id = _ja_activity_id_from_post_result(post_result)
    if created_activity_id is None and len(new_ids) == 1:
        created_activity_id = new_ids[0]
    detail = None
    if created_activity_id is not None:
        detail = _ja_activity_diagnostic_get(
            "candidates/{}/activities/{}".format(cid_q, urllib.parse.quote(str(created_activity_id), safe="")),
            timeout=25,
        )

    generated = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    report = {
        "diagnostic": "CV Studio JobAdder OAuth Official AddCandidateActivity Create Test",
        "cv_studio_version": _CVSTUDIO_VERSION,
        "generated_utc": generated,
        "candidate_fixture": {
            "name": "Max Low",
            "candidate_id": candidate_id,
            "previous_successful_activity_id": 503608,
        },
        "test_hypothesis": "Use JobAdder's official AddCandidateActivity v2 write contract exactly: activitySettingId 8225; textAnswers use questionId + text; ratingValueAnswers uses questionId 62988 + rating 3; listValueAnswers and dateRangeValueAnswers are empty; no SPA-only or response-only fields.",
        "official_reference": "https://api.jobadder.com/v2/docs#operation/AddCandidateActivity",
        "safety": {
            "oauth_post_count": 1,
            "candidate_notes_created": False,
            "existing_activity_edited": False,
            "salary_or_notice_profile_updated": False,
            "f12_helper_changed": False,
            "oauth_token_included_in_report": False,
            "one_shot_guard_active_for_session": True,
        },
        "preflight_activity_list": before,
        "controlled_post": post_result,
        "postflight_activity_list": after,
        "activity_ids_before": sorted(before_ids),
        "activity_ids_after": sorted(after_ids),
        "new_activity_ids_detected": new_ids,
        "created_activity_id": created_activity_id,
        "created_activity_detail": detail,
        "result_summary": (
            "Controlled POST returned success and a new activity was detected. Review created_activity_detail to confirm text and Presentability 3/4."
            if post_result.get("ok") and created_activity_id is not None
            else "Controlled POST did not produce a confirmed new activity. Review the complete response body; no second POST was attempted."
        ),
        "next_boundary": "Do not repeat or alter the payload until this exact report is reviewed.",
    }
    filename = "jobadder_activity_create_official_schema_diag_{}_{}.json".format(
        candidate_id,
        re.sub(r"[^0-9]", "", generated)[:14],
    )
    return jsonify({"ok": True, "filename": filename, "report": report})


@app.route("/jobadder/onenote_activity_probe", methods=["GET"])
def jobadder_onenote_activity_probe():
    """Read-only diagnostics for JobAdder candidate activity availability.

    This does not create notes or activities. It helps identify the tenant's
    activity settings/endpoints after a transfer fails.
    """
    token = _ja_refresh_access_token(force=False)
    if not token:
        return jsonify({"error": "Not authenticated"}), 401
    candidate_id = str(request.args.get("candidate_id") or "").strip()
    cid = urllib.parse.quote(candidate_id, safe="") if candidate_id else ""
    paths = [
        "candidateActivitySettings",
        "candidate-activity-settings",
        "activities/settings",
        "activities/settings?entityType=candidate",
        "activities/candidates/settings",
        "activitySettings?entityType=candidate",
    ]
    if cid:
        paths.extend([
            "candidateActivities?candidateId={}".format(cid),
            "activities?candidateId={}".format(cid),
            "candidates/{}/activities".format(cid),
            "candidates/{}/activity".format(cid),
            "candidate/{}/activities".format(cid),
            "candidate/{}/activity".format(cid),
        ])
    results = []
    for path in paths:
        try:
            status, data = _ja_get_json(path, timeout=12)
            snippet = data
            try:
                snippet = json.dumps(data, ensure_ascii=False)[:1500]
            except Exception:
                snippet = str(data)[:1500]
            results.append({"path": path, "status": status, "ok": True, "data": snippet})
        except urllib.error.HTTPError as e:
            body = e.read().decode(errors="replace")
            results.append({"path": path, "status": e.code, "ok": False, "detail": _ja_compact_error_body(body), "raw": body[:1500]})
        except Exception as e:
            results.append({"path": path, "ok": False, "error": str(e)[:1000]})
    return jsonify({"ok": True, "candidate_id": candidate_id, "results": results})


from cvstudio_spider_boolean import (
    _SPIDER_COUNTRY_CODE_LOOKUP,
    _SPIDER_COUNTRY_DEFINITIONS,
    _SPIDER_COUNTRY_NAME_LOOKUP,
    _spider_boolean_expression_match,
    _spider_boolean_fallback_atoms,
    _spider_boolean_fallback_terms,
    _spider_boolean_not_terms,
    _spider_boolean_positive_terms,
    _spider_boolean_terms,
    _spider_boolean_tokens,
    _spider_candidate_blob,
    _spider_candidate_id,
    _spider_clean_doc_text_for_preview,
    _spider_country_aliases,
    _spider_country_match,
    _spider_country_profile,
    _spider_discovery_keyword_match,
    _spider_flatten,
    _spider_has_any,
    _spider_hit_terms,
    _spider_location_identity,
    _spider_max_years_value,
    _spider_min_years_value,
    _spider_normalized_record_label,
    _spider_parse_employment_month,
    _spider_pick_field_text,
    _spider_residential_classes,
    _spider_residential_status_text,
    _spider_status_aliases,
    _spider_status_target,
    _spider_structured_employment_years,
    _spider_term_coverage,
    _spider_term_matches,
    _spider_terms,
    _spider_terms_for_fit,
    _spider_visible_years,
    _spider_work_setup_aliases,
    _spider_years_bounds,
)

from cvstudio_spider_summary import (
    _spider_card_fields,
    _spider_custom_field_value,
    _spider_custom_record_key,
    _spider_merge_candidate_summary_and_detail,
    _spider_merge_missing_json,
    _spider_merge_record_lists,
    _spider_notice_snapshot,
    _spider_number,
    _spider_salary_snapshot,
    _spider_work_record_key,
)

from cvstudio_spider_score import (
    _SPIDER_JD_HEADING_PREFIX_RE,
    _spider_context_only_fit_term,
    _spider_extract_option_values,
    _spider_ignored_fit_term,
    _spider_item_score,
    _spider_jd_heading_section,
    _spider_jd_relevance_terms,
    _spider_jd_scoring_lines,
    _spider_match_fit_percent,
    _spider_option_fallbacks,
    _spider_strip_context_fit_term,
    _spider_strip_jd_heading_prefix,
    _spider_weighted_coverage,
)

from cvstudio_ja_typos import (
    _JA_RECRUITMENT_FUZZY_TERMS,
    _JA_RECRUITMENT_PROTECTED_WORDS,
    _JA_RECRUITMENT_TYPO_ALIASES,
    _JA_SALARY_TYPO_RULES,
    _ja_case_like,
    _ja_correct_notice_typos,
    _ja_correct_recruitment_typos,
    _ja_correct_screening_field_typos,
    _ja_edit_distance_limited,
    _ja_recruitment_typo_target,
    _ja_salary_normalize_recruiter_typos,
)







































def _spider_fetch_candidate_detail(token, candidate_id):
    if not token or not candidate_id:
        return None
    try:
        raw, _ctype, _dispo = _spider_get_ja_raw(
            token,
            "candidates/{}".format(urllib.parse.quote(str(candidate_id), safe="")),
            timeout=8,
            accept_json=True,
        )
        parsed = safe_json(raw, {})
        return parsed if isinstance(parsed, dict) else None
    except _SpiderJobAdderReconnectRequired:
        raise
    except Exception:
        return None


_SPIDER_RESUME_TEXT_CACHE = {}
_SPIDER_RESUME_TEXT_CACHE_LOCK = threading.RLock()
_SPIDER_RESUME_TEXT_CACHE_TTL = 15 * 60

# Final rendered AI Crawler preview payloads are expensive: PDFium/WebP, OCR and
# legacy Word conversion can all be involved. Keep a bounded account-scoped LRU
# of successful JSON payloads. Every request still performs the cheap JobAdder
# latest-attachment listing and includes that listing in the cache fingerprint,
# so replacing a resume invalidates the cached visual automatically.
_SPIDER_PREVIEW_PAYLOAD_CACHE = OrderedDict()
_SPIDER_PREVIEW_PAYLOAD_CACHE_LOCK = threading.RLock()
_SPIDER_PREVIEW_PAYLOAD_CACHE_TTL = 30 * 60
_SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_ENTRIES = 120
_SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_BYTES = 256 * 1024 * 1024
_SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES = 0

# Background preview work is cooperatively cancellable. Foreground opens and
# Clear/new-search increment this generation; PDFium page loops and LibreOffice
# conversion stop at the next safe checkpoint. Expensive rendering is also
# serialized so an abandoned prefetch cannot compete with the active preview.
_SPIDER_PREVIEW_CANCEL_LOCK = threading.RLock()
_SPIDER_PREVIEW_BACKGROUND_GENERATION = 0
_SPIDER_PREVIEW_RENDER_LOCK = threading.Lock()
_SPIDER_PREVIEW_PREFETCH_JOB_KIND = "ai_crawler_preview_prefetch"


class _SpiderPreviewCancelled(RuntimeError):
    pass


def _spider_preview_background_generation():
    with _SPIDER_PREVIEW_CANCEL_LOCK:
        return int(_SPIDER_PREVIEW_BACKGROUND_GENERATION)


def _spider_preview_cancel_background_work():
    global _SPIDER_PREVIEW_BACKGROUND_GENERATION
    with _SPIDER_PREVIEW_CANCEL_LOCK:
        _SPIDER_PREVIEW_BACKGROUND_GENERATION += 1
        return int(_SPIDER_PREVIEW_BACKGROUND_GENERATION)


def _spider_preview_cancel_persistent_work():
    generation = _spider_preview_cancel_background_work()
    _CVSTUDIO_JOBS.request_cancel_kind(_SPIDER_PREVIEW_PREFETCH_JOB_KIND)
    return generation


def _spider_preview_cancel_check(generation):
    if int(generation) != _spider_preview_background_generation():
        raise _SpiderPreviewCancelled("Preview preload superseded by foreground work")


@app.errorhandler(_SpiderPreviewCancelled)
def _spider_preview_cancelled_response(_exc):
    return _cvstudio_error_payload("PREVIEW_PREFETCH_SUPERSEDED", "Preview preload superseded", 409, retryable=True, extra={"prefetch_cancelled": True})


def _ja_cache_namespace():
    """Return a stable opaque namespace for the currently connected JobAdder account.

    Access tokens rotate routinely, so they must not define cache identity. The
    namespace is generated on OAuth connection, persisted in the protected backend
    store, and replaced when a new connection/account is completed.
    """
    with _ja_store_lock:
        namespace = str(_ja_creds_store.get("cache_namespace") or "").strip()
        if namespace:
            return namespace
        material = "{}|{}|{}".format(
            _ja_normalize_api_base(_ja_creds_store.get("api_url")),
            str(_ja_creds_store.get("client_id") or ""),
            str(_ja_creds_store.get("refresh_token") or ""),
        )
        if not material.strip("|"):
            return ""
        namespace = hashlib.sha256(material.encode("utf-8", errors="ignore")).hexdigest()
        _ja_creds_store["cache_namespace"] = namespace
        try:
            _ja_save_store()
        except Exception:
            pass
        return namespace


def _spider_resume_cache_key(token, candidate_id):
    """Return an opaque stable account-scoped cache key."""
    cid = str(candidate_id or "").strip()
    if not cid or not str(token or "").strip():
        return ""
    namespace = _ja_cache_namespace()
    if not namespace:
        return ""
    api_base = _ja_normalize_api_base(_ja_creds_store.get("api_url"))
    material = "{}|{}|{}".format(api_base, namespace, cid)
    return hashlib.sha256(material.encode("utf-8", errors="ignore")).hexdigest()


def _spider_resume_text_cache_clear():
    global _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES
    with _SPIDER_RESUME_TEXT_CACHE_LOCK:
        _SPIDER_RESUME_TEXT_CACHE.clear()
    with _SPIDER_PREVIEW_PAYLOAD_CACHE_LOCK:
        _SPIDER_PREVIEW_PAYLOAD_CACHE.clear()
        _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES = 0


def _spider_download_content_identity(raw):
    if not isinstance(raw, (bytes, bytearray)) or not raw:
        return ""
    return hashlib.sha256(bytes(raw)).hexdigest()


def _spider_resume_text_cache_get(token, candidate_id, content_sha256=""):
    key = _spider_resume_cache_key(token, candidate_id)
    expected_content = str(content_sha256 or "").strip().lower()
    if not key or not expected_content:
        return None
    now = time.time()
    with _SPIDER_RESUME_TEXT_CACHE_LOCK:
        item = _SPIDER_RESUME_TEXT_CACHE.get(key)
        if not item:
            return None
        if now - float(item.get("saved_at") or 0) > _SPIDER_RESUME_TEXT_CACHE_TTL:
            _SPIDER_RESUME_TEXT_CACHE.pop(key, None)
            return None
        # Empty/error results are never valid cache entries.  This also discards
        # any stale in-memory negative entry left by an older process version.
        if not str(item.get("text") or "").strip():
            _SPIDER_RESUME_TEXT_CACHE.pop(key, None)
            return None
        if not _receipt_hmac.compare_digest(
            str(item.get("content_sha256") or "").strip().lower(),
            expected_content,
        ):
            return None
        return dict(item)


def _spider_resume_text_cache_put(
    token,
    candidate_id,
    text,
    source="",
    *,
    content_sha256="",
    content_kind="",
    antiword_verified=False,
):
    clean_text = str(text or "").strip()
    content_identity = str(content_sha256 or "").strip().lower()
    key = _spider_resume_cache_key(token, candidate_id)
    if not key or not clean_text or not content_identity:
        return
    with _SPIDER_RESUME_TEXT_CACHE_LOCK:
        if len(_SPIDER_RESUME_TEXT_CACHE) >= 500:
            oldest = sorted(_SPIDER_RESUME_TEXT_CACHE.items(), key=lambda kv: float((kv[1] or {}).get("saved_at") or 0))[:80]
            for old_key, _ in oldest:
                _SPIDER_RESUME_TEXT_CACHE.pop(old_key, None)
        _SPIDER_RESUME_TEXT_CACHE[key] = {
            "text": clean_text[:30000],
            "source": str(source or "")[:160],
            "content_sha256": content_identity,
            "content_kind": str(content_kind or "")[:32],
            "antiword_verified": bool(antiword_verified),
            "saved_at": time.time(),
        }


def _spider_attachment_fingerprint(records):
    """Stable validator for the current JobAdder resume attachment set."""
    canonical = []
    for record in records or []:
        if not isinstance(record, dict):
            continue
        canonical.append({
            "id": str(record.get("attachmentId") or record.get("id") or ""),
            "type": str(record.get("type") or ""),
            "file": str(record.get("fileName") or record.get("filename") or ""),
            "created": str(record.get("createdAt") or ""),
            "updated": str(record.get("updatedAt") or record.get("modifiedAt") or ""),
            "size": str(record.get("size") or record.get("fileSize") or record.get("sizeBytes") or ""),
        })
    canonical.sort(key=lambda item: (item["type"], item["id"], item["updated"], item["created"], item["file"]))
    raw = json.dumps(canonical, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(raw.encode("utf-8", errors="ignore")).hexdigest()


def _spider_preview_payload_cache_key(token, candidate_id, content_sha256, variant="full"):
    account_candidate = _spider_resume_cache_key(token, candidate_id)
    content_identity = str(content_sha256 or "").strip().lower()
    variant = "prefetch" if str(variant or "").lower() == "prefetch" else "full"
    if not account_candidate or not content_identity:
        return ""
    return hashlib.sha256((account_candidate + "|" + content_identity + "|" + variant).encode("ascii", errors="ignore")).hexdigest()


def _spider_preview_payload_cache_cleanup(now=None):
    global _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES
    now = float(now or time.time())
    expired = []
    for key, item in list(_SPIDER_PREVIEW_PAYLOAD_CACHE.items()):
        if now - float((item or {}).get("saved_at") or 0) > _SPIDER_PREVIEW_PAYLOAD_CACHE_TTL:
            expired.append(key)
    for key in expired:
        item = _SPIDER_PREVIEW_PAYLOAD_CACHE.pop(key, None) or {}
        _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES = max(0, _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES - int(item.get("bytes") or 0))


def _spider_preview_payload_cache_get(token, candidate_id, content_sha256, variant="full"):
    key = _spider_preview_payload_cache_key(token, candidate_id, content_sha256, variant)
    if not key:
        return None
    with _SPIDER_PREVIEW_PAYLOAD_CACHE_LOCK:
        _spider_preview_payload_cache_cleanup()
        item = _SPIDER_PREVIEW_PAYLOAD_CACHE.pop(key, None)
        if not item:
            return None
        _SPIDER_PREVIEW_PAYLOAD_CACHE[key] = item
        try:
            payload = json.loads(item.get("json") or "{}")
        except Exception:
            return None
        if not isinstance(payload, dict) or not payload.get("ok"):
            return None
        payload["_cache_content_kind"] = str(
            item.get("content_kind") or ""
        )
        payload["_cache_antiword_verified"] = bool(
            item.get("antiword_verified")
        )
        payload["preview_cache_hit"] = True
        return payload


def _spider_preview_payload_cache_put(
    token,
    candidate_id,
    content_sha256,
    payload,
    variant="full",
    *,
    content_kind="",
    antiword_verified=False,
):
    global _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES
    if not isinstance(payload, dict) or not payload.get("ok"):
        return
    key = _spider_preview_payload_cache_key(token, candidate_id, content_sha256, variant)
    if not key:
        return
    clean = dict(payload)
    clean.pop("preview_cache_hit", None)
    try:
        encoded = json.dumps(clean, ensure_ascii=False, separators=(",", ":"))
    except Exception:
        return
    size = len(encoded.encode("utf-8", errors="ignore"))
    if size <= 0 or size > _SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_BYTES // 2:
        return
    with _SPIDER_PREVIEW_PAYLOAD_CACHE_LOCK:
        _spider_preview_payload_cache_cleanup()
        previous = _SPIDER_PREVIEW_PAYLOAD_CACHE.pop(key, None) or {}
        _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES = max(0, _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES - int(previous.get("bytes") or 0))
        _SPIDER_PREVIEW_PAYLOAD_CACHE[key] = {
            "json": encoded,
            "bytes": size,
            "content_kind": str(content_kind or "")[:32],
            "antiword_verified": bool(antiword_verified),
            "saved_at": time.time(),
        }
        _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES += size
        while (_SPIDER_PREVIEW_PAYLOAD_CACHE and (
            len(_SPIDER_PREVIEW_PAYLOAD_CACHE) > _SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_ENTRIES
            or _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES > _SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_BYTES
        )):
            _old_key, old = _SPIDER_PREVIEW_PAYLOAD_CACHE.popitem(last=False)
            _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES = max(0, _SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES - int((old or {}).get("bytes") or 0))



def _cvstudio_system_memory_status():
    return _phase4_system_memory_status()


def _spider_preview_cache_stats():
    with _SPIDER_PREVIEW_PAYLOAD_CACHE_LOCK:
        _spider_preview_payload_cache_cleanup()
        preview_entries = len(_SPIDER_PREVIEW_PAYLOAD_CACHE)
        preview_bytes = int(_SPIDER_PREVIEW_PAYLOAD_CACHE_BYTES)
    with _SPIDER_RESUME_TEXT_CACHE_LOCK:
        resume_entries = len(_SPIDER_RESUME_TEXT_CACHE)
        resume_chars = sum(len(str((item or {}).get("text") or "")) for item in _SPIDER_RESUME_TEXT_CACHE.values())
    return {
        "backend_preview_entries": preview_entries,
        "backend_preview_bytes": preview_bytes,
        "backend_preview_max_entries": int(_SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_ENTRIES),
        "backend_preview_max_bytes": int(_SPIDER_PREVIEW_PAYLOAD_CACHE_MAX_BYTES),
        "resume_text_entries": resume_entries,
        "resume_text_characters": resume_chars,
        "preview_ttl_seconds": int(_SPIDER_PREVIEW_PAYLOAD_CACHE_TTL),
        "resume_ttl_seconds": int(_SPIDER_RESUME_TEXT_CACHE_TTL),
    }


def _cvstudio_dependency_status():
    status = _phase4_dependency_status(
        _find_soffice_binary,
        _find_antiword_binary,
        lambda: _verified_antiword_health(
            _CVSTUDIO_ROOT,
            os.path.dirname(os.path.abspath(__file__)),
        ),
    )
    tesseract_health = _mandatory_tesseract_health(_CVSTUDIO_ROOT)
    status["tesseract"] = bool(
        tesseract_health.get("available")
        and tesseract_health.get("functional")
        and tesseract_health.get("english_language_data")
    )
    status["tesseract_health"] = tesseract_health
    return status


def _cvstudio_connection_status_redacted():
    try:
        jobadder = _ja_public_info()
    except Exception:
        jobadder = {}
    try:
        outlook = _ms_outlook_public_info()
    except Exception:
        outlook = {}
    try:
        with _ms_graph_store_lock:
            onenote = {
                "connected": bool(str(_ms_graph_store.get("access_token") or _ms_graph_store.get("refresh_token") or "").strip()),
                "tenant_configured": bool(str(_ms_graph_store.get("tenant") or "").strip()),
                "client_id_configured": bool(str(_ms_graph_store.get("client_id") or "").strip()),
                "storage": str(_ms_graph_store.get("_storage") or "backend_secure_store"),
            }
    except Exception:
        onenote = {}
    configured_ai = {}
    try:
        configured_ai = {slot: bool(str(_ai_secret_store.get(slot) or "").strip()) for slot in sorted(_AI_SECRET_SLOTS)}
    except Exception:
        pass
    return {
        "jobadder": {
            "connected": bool(jobadder.get("connected")),
            "needs_reconnect": bool(jobadder.get("needs_reconnect")),
            "client_id_configured": bool(jobadder.get("client_id")),
            "client_secret_configured": bool(jobadder.get("client_secret_configured")),
            "storage": str(jobadder.get("storage") or ""),
        },
        "onenote": onenote,
        "outlook": {
            "connected": bool(outlook.get("connected")),
            "client_id_configured": bool(outlook.get("client_id")),
            "tenant_configured": bool(outlook.get("tenant")),
            "storage": str(outlook.get("storage") or ""),
        },
        "ai_keys_configured": configured_ai,
    }


def _cvstudio_runtime_diagnostics_payload():
    executable_name = os.path.basename(str(sys.executable or sys.argv[0] or "")).strip()
    argv0_name = os.path.basename(str(sys.argv[0] or "")).strip()
    native_compiled = bool(getattr(sys, "frozen", False) or executable_name.lower().startswith("cvstudio") or argv0_name.lower().startswith("cvstudio"))
    return {
        "ok": True,
        "product": "CV Studio",
        "version": _CVSTUDIO_VERSION,
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "request_id": _cvstudio_current_request_id(),
        "instance_id": _CVSTUDIO_INSTANCE_ID,
        "root_hash": _CVSTUDIO_ROOT_HASH,
        "pid": os.getpid(),
        "port": _CVSTUDIO_PORT,
        "runtime_mode": "native" if native_compiled else "source",
        "runtime_process": executable_name or argv0_name or "unknown",
        "platform": {"system": platform.system(), "release": platform.release(), "machine": platform.machine(), "python": platform.python_version()},
        "memory": _cvstudio_system_memory_status(),
        "cache": _spider_preview_cache_stats(),
        "durable_storage": _CVSTUDIO_STORAGE.health(),
        "dependencies": _cvstudio_dependency_status(),
        "connections": _cvstudio_connection_status_redacted(),
        "install_receipt": (lambda result: {"valid": bool(result[0]), "status": str(result[1] or "")[:240]})(_install_receipt_status()),
    }


def _cvstudio_redact_support_text(text):
    return _phase4_redact_support_text(
        text,
        (
            (_CVSTUDIO_ROOT, "[cvstudio root]"),
            (os.path.expanduser("~"), "[home]"),
        ),
    )


def _cvstudio_sanitize_browser_diagnostics(payload):
    return _phase4_sanitize_browser_diagnostics(
        payload, _cvstudio_redact_support_text
    )


_OWNER_INTEGRATION_REPORT_LOCK = threading.RLock()
_OWNER_INTEGRATION_LAST_REPORT = {}


def _owner_integration_enabled():
    """Owner source builds only; protected colleague packages cannot enable it."""
    marker = Path(_CVSTUDIO_ROOT) / "KEEP_PRIVATE_PATCH_BASE.txt"
    source_app = Path(_CVSTUDIO_ROOT) / "app.py"
    return bool(marker.is_file() and source_app.is_file() and not getattr(sys, "frozen", False))


def _owner_integration_authorized():
    return _owner_integration_enabled() and request.headers.get("X-CV-Studio-Owner-Test") == "1"


def _owner_integration_test_result(name, started, *, ok, status="passed", detail="", metadata=None):
    result = {
        "name": str(name),
        "ok": bool(ok),
        "status": str(status),
        "duration_ms": max(0, int((time.monotonic() - started) * 1000)),
        "detail": _cvstudio_redact_support_text(str(detail or ""))[:1200],
    }
    if isinstance(metadata, dict):
        result["metadata"] = metadata
    return result


def _owner_integration_local_docx_test():
    started = time.monotonic()
    data_path = out_path = None
    try:
        fixture = {
            "candidate": {"name": "Owner Integration Fixture", "email": "", "phone": "", "linkedin": "", "summary": "Local non-personal integration test."},
            "experience": [], "education": [], "skills": ["CV Studio"], "languages": []
        }
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as handle:
            json.dump(fixture, handle)
            data_path = handle.name
        out_path = data_path[:-5] + ".docx"
        proc = subprocess.run(["node", os.path.join(_CVSTUDIO_ROOT, "generate.js"), data_path, out_path], capture_output=True, text=True, timeout=35)
        if proc.returncode != 0 or not os.path.isfile(out_path) or os.path.getsize(out_path) < 500:
            raise RuntimeError(_summarize_node_failure(proc.stdout, proc.stderr) or "DOCX fixture was not generated")
        with zipfile.ZipFile(out_path) as archive:
            if "word/document.xml" not in archive.namelist():
                raise RuntimeError("Generated DOCX is missing word/document.xml")
        return _owner_integration_test_result("docx_generation", started, ok=True, detail="Local DOCX fixture generated and validated", metadata={"bytes": os.path.getsize(out_path)})
    except Exception as exc:
        return _owner_integration_test_result("docx_generation", started, ok=False, status="failed", detail=str(exc))
    finally:
        for path in (data_path, out_path):
            try:
                if path: os.remove(path)
            except Exception:
                pass


def _owner_integration_jobadder_read(candidate_id):
    started = time.monotonic()
    try:
        candidate_id = str(candidate_id or "").strip()
        if not re.fullmatch(r"\d{1,20}", candidate_id):
            return _owner_integration_test_result("jobadder_candidate_read", started, ok=False, status="skipped", detail="Provide a numeric dedicated JobAdder test candidate ID")
        status, payload = _ja_get_json("candidates/{}".format(candidate_id), timeout=25)
        ok = 200 <= int(status or 0) < 300 and isinstance(payload, dict)
        return _owner_integration_test_result("jobadder_candidate_read", started, ok=ok, status="passed" if ok else "failed", detail="Read-only candidate lookup completed", metadata={"http_status": int(status or 0), "candidate_id_redacted": "…" + candidate_id[-4:]})
    except PermissionError as exc:
        return _owner_integration_test_result("jobadder_candidate_read", started, ok=False, status="skipped", detail=str(exc))
    except Exception as exc:
        return _owner_integration_test_result("jobadder_candidate_read", started, ok=False, status="failed", detail=str(exc))


def _owner_integration_microsoft_read(kind):
    started = time.monotonic()
    name = "onenote_read" if kind == "onenote" else "outlook_read"
    try:
        if kind == "onenote":
            payload = _ms_graph_json("me/onenote/notebooks", params={"$top": 1, "$select": "id,displayName"}, timeout=25)
        else:
            payload = _ms_outlook_graph_json("me?$select=id,displayName,mail,userPrincipalName", timeout=25)
        ok = isinstance(payload, dict)
        return _owner_integration_test_result(name, started, ok=ok, status="passed" if ok else "failed", detail="Read-only Microsoft Graph request completed")
    except PermissionError as exc:
        return _owner_integration_test_result(name, started, ok=False, status="skipped", detail=str(exc))
    except Exception as exc:
        return _owner_integration_test_result(name, started, ok=False, status="failed", detail=str(exc))


def _owner_integration_deepseek_probe(confirmation):
    started = time.monotonic()
    if confirmation != "RUN ONE PAID DEEPSEEK PROBE":
        return _owner_integration_test_result("deepseek_paid_probe", started, ok=False, status="skipped", detail="Exact paid-test confirmation was not supplied")
    api_key = str(_ai_secret_store.get("main_deepseek") or "").strip()
    if not api_key:
        return _owner_integration_test_result("deepseek_paid_probe", started, ok=False, status="skipped", detail="No securely saved main DeepSeek key")
    try:
        model = "deepseek-v4-flash"
        payload = {"model": model, "max_tokens": 12, "messages": [{"role": "user", "content": "Reply only with OK"}]}
        data = _call_deepseek(api_key, payload, timeout_seconds=45)
        usage = _normalize_llm_usage(
            data.get("usage") if isinstance(data, dict) else {},
            api_calls=1,
        )
        cost = _llm_response_cost_fields(model, usage or {}, "deepseek")
        return _owner_integration_test_result(
            "deepseek_paid_probe",
            started,
            ok=True,
            detail="Small paid DeepSeek request completed",
            metadata={
                "model": model,
                "usage": usage or {},
                "cost_usd": cost.get("cost", 0),
                "cost_details": cost.get("cost_details") or {},
            },
        )
    except Exception as exc:
        return _owner_integration_test_result(
            "deepseek_paid_probe",
            started,
            ok=False,
            status="failed",
            detail=str(exc),
            metadata=_llm_paid_failure_fields(
                exc,
                "deepseek-v4-flash",
                "deepseek",
                attempted=True,
            ),
        )


@app.route("/owner/integration/status", methods=["GET"])
def owner_integration_status():
    if not _owner_integration_enabled():
        return jsonify({"ok": True, "enabled": False, "version": APP_VERSION})
    with _OWNER_INTEGRATION_REPORT_LOCK:
        last = dict(_OWNER_INTEGRATION_LAST_REPORT)
    return jsonify({
        "ok": True,
        "enabled": True,
        "owner_source_only": True,
        "request_id": _cvstudio_current_request_id(),
        "connections": _cvstudio_connection_status_redacted(),
        "available_tests": ["local_health", "docx_generation", "jobadder_candidate_read", "onenote_read", "outlook_read", "deepseek_paid_probe"],
        "existing_controlled_activity_fixture": {"route": "/jobadder/onenote_activity_create_diagnostic", "candidate_id": "41262878", "confirmation": "CREATE ONE MAX LOW TEST", "session_limited": True},
        "last_report": last,
    })


@app.route("/owner/integration/run", methods=["POST"])
def owner_integration_run():
    if not _owner_integration_authorized():
        return _cvstudio_error_payload("OWNER_TEST_MODE_FORBIDDEN", "Owner integration test mode requires the private source build and explicit owner-test header", 403)
    body = request.get_json(silent=True) or {}
    requested = body.get("tests") if isinstance(body.get("tests"), list) else ["local_health", "docx_generation"]
    requested = [str(item or "").strip() for item in requested][:12]
    fixtures = body.get("fixtures") if isinstance(body.get("fixtures"), dict) else {}
    confirmations = body.get("confirmations") if isinstance(body.get("confirmations"), dict) else {}
    results = []
    if "local_health" in requested:
        started = time.monotonic()
        payload = _cvstudio_runtime_diagnostics_payload()
        healthy = payload.get("version") == _CVSTUDIO_VERSION and bool(payload.get("install_receipt", {}).get("valid"))
        results.append(_owner_integration_test_result("local_health", started, ok=healthy, status="passed" if healthy else "failed", detail="Runtime, receipt, dependencies and redacted connections inspected", metadata={"version": payload.get("version"), "runtime_mode": payload.get("runtime_mode"), "dependencies": payload.get("dependencies")}))
    if "docx_generation" in requested:
        results.append(_owner_integration_local_docx_test())
    if "jobadder_candidate_read" in requested:
        results.append(_owner_integration_jobadder_read(fixtures.get("jobadder_candidate_id")))
    if "onenote_read" in requested:
        results.append(_owner_integration_microsoft_read("onenote"))
    if "outlook_read" in requested:
        results.append(_owner_integration_microsoft_read("outlook"))
    if "deepseek_paid_probe" in requested:
        results.append(_owner_integration_deepseek_probe(str(confirmations.get("deepseek") or "")))
    passed = sum(1 for item in results if item.get("status") == "passed")
    failed = sum(1 for item in results if item.get("status") == "failed")
    skipped = sum(1 for item in results if item.get("status") == "skipped")
    report = {
        "ok": failed == 0,
        "product": "CV Studio owner integration test",
        "version": _CVSTUDIO_VERSION,
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "request_id": _cvstudio_current_request_id(),
        "summary": {"passed": passed, "failed": failed, "skipped": skipped, "total": len(results)},
        "results": results,
    }
    with _OWNER_INTEGRATION_REPORT_LOCK:
        _OWNER_INTEGRATION_LAST_REPORT.clear()
        _OWNER_INTEGRATION_LAST_REPORT.update(report)
    status = 200 if failed == 0 else 424
    return jsonify(report), status


_PHASE2A_USAGE_MAX_RECORDS = 250000
_PHASE2A_USAGE_SECRET_FIELDS = {
    "access_token", "refresh_token", "client_secret", "api_key", "authorization",
    "authorization_code", "password", "credential", "credentials", "token",
}
_PHASE2A_USAGE_SECRET_SUFFIXES = {
    "accesstoken", "refreshtoken", "clientsecret", "apikey", "authorizationcode",
    "password", "credential", "credentials", "bearertoken", "oauthtoken",
    "oauth2token", "authtoken", "idtoken", "sessiontoken", "csrftoken", "secret",
}
_PHASE2A_USAGE_DROP = object()


def _phase2a_usage_secret_field(key):
    return _phase4_phase2a_usage_secret_field(
        key,
        secret_fields=_PHASE2A_USAGE_SECRET_FIELDS,
        secret_suffixes=_PHASE2A_USAGE_SECRET_SUFFIXES,
    )


def _phase2a_usage_safe_value(value, depth=0):
    return _phase4_phase2a_usage_safe_value(
        value,
        depth,
        drop=_PHASE2A_USAGE_DROP,
        secret_field=lambda key: _phase2a_usage_secret_field(key),
    )


def _phase2a_usage_records(payload):
    return _phase4_phase2a_usage_records(
        payload,
        maximum=_PHASE2A_USAGE_MAX_RECORDS,
        safe_value=lambda value: _phase2a_usage_safe_value(value),
    )


def _phase2a_ppc_metadata(payload):
    return _phase4_phase2a_ppc_metadata(payload)


_CVSTUDIO_STORAGE_BRIDGE = StorageBridge(
    request_json=lambda: request.get_json(silent=True) or {},
    jsonify=lambda payload: jsonify(payload),
    error_payload=lambda *args, **kwargs: _cvstudio_error_payload(
        *args, **kwargs
    ),
    request_id=lambda: _cvstudio_current_request_id(),
    usage_repository=lambda: _CVSTUDIO_USAGE_REPOSITORY,
    ppc_repository=lambda: _CVSTUDIO_PPC_REPOSITORY,
    transfer_repository=lambda: _CVSTUDIO_ONENOTE_TRANSFER_REPOSITORY,
    link_repository=lambda: _CVSTUDIO_ONENOTE_LINK_REPOSITORY,
    browser_settings_repository=lambda: _CVSTUDIO_BROWSER_SETTINGS_REPOSITORY,
    usage_records=lambda payload: _phase2a_usage_records(payload),
    ppc_metadata=lambda payload: _phase2a_ppc_metadata(payload),
    record_array=lambda payload, maximum, normalizer: _phase2b_record_array(
        payload, maximum, normalizer
    ),
    browser_settings=lambda payload: _phase2b_browser_settings(payload),
    browser_setting_keys=lambda payload: _phase2b_browser_setting_keys(payload),
)


@app.route("/storage/usage-history", methods=["GET"])
def phase2a_usage_history_read():
    return _CVSTUDIO_STORAGE_BRIDGE.usage_history_read()


@app.route("/storage/usage-history/import", methods=["POST"])
def phase2a_usage_history_import():
    return _CVSTUDIO_STORAGE_BRIDGE.usage_history_import()


@app.route("/storage/usage-history/upsert", methods=["POST"])
def phase2a_usage_history_upsert():
    return _CVSTUDIO_STORAGE_BRIDGE.usage_history_upsert()


@app.route("/storage/usage-history/clear", methods=["POST"])
def phase2a_usage_history_clear():
    return _CVSTUDIO_STORAGE_BRIDGE.usage_history_clear()


@app.route("/storage/ppc-metadata", methods=["GET"])
def phase2a_ppc_metadata_read():
    return _CVSTUDIO_STORAGE_BRIDGE.ppc_metadata_read()


@app.route("/storage/ppc-metadata/import", methods=["POST"])
def phase2a_ppc_metadata_import():
    return _CVSTUDIO_STORAGE_BRIDGE.ppc_metadata_import()


@app.route("/storage/ppc-metadata/upsert", methods=["POST"])
def phase2a_ppc_metadata_upsert():
    return _CVSTUDIO_STORAGE_BRIDGE.ppc_metadata_upsert()


@app.route("/storage/ppc-metadata/clear", methods=["POST"])
def phase2a_ppc_metadata_clear():
    return _CVSTUDIO_STORAGE_BRIDGE.ppc_metadata_clear()


def _phase2b_record_array(payload, maximum, normalizer):
    return _phase4_phase2b_record_array(payload, maximum, normalizer)


def _phase2b_browser_settings(payload):
    return _phase4_phase2b_browser_settings(
        payload,
        BROWSER_SETTING_KEYS,
        lambda key, value: BrowserSettingsRepository.normalize_value(key, value),
    )


def _phase2b_browser_setting_keys(payload):
    return _phase4_phase2b_browser_setting_keys(
        payload,
        BROWSER_SETTING_KEYS,
    )


@app.route("/storage/onenote-transfer-records", methods=["GET"])
def phase2b_onenote_transfer_read():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_transfer_read()


@app.route("/storage/onenote-transfer-records/import", methods=["POST"])
def phase2b_onenote_transfer_import():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_transfer_import()


@app.route("/storage/onenote-transfer-records/replace", methods=["POST"])
def phase2b_onenote_transfer_replace():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_transfer_replace()


@app.route("/storage/onenote-transfer-records/clear", methods=["POST"])
def phase2b_onenote_transfer_clear():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_transfer_clear()


@app.route("/storage/onenote-saved-links", methods=["GET"])
def phase2b_onenote_links_read():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_links_read()


@app.route("/storage/onenote-saved-links/import", methods=["POST"])
def phase2b_onenote_links_import():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_links_import()


@app.route("/storage/onenote-saved-links/replace", methods=["POST"])
def phase2b_onenote_links_replace():
    return _CVSTUDIO_STORAGE_BRIDGE.onenote_links_replace()


@app.route("/storage/browser-settings", methods=["GET"])
def phase2b_browser_settings_read():
    return _CVSTUDIO_STORAGE_BRIDGE.browser_settings_read()


@app.route("/storage/browser-settings/import", methods=["POST"])
def phase2b_browser_settings_import():
    return _CVSTUDIO_STORAGE_BRIDGE.browser_settings_import()


@app.route("/storage/browser-settings/upsert", methods=["POST"])
def phase2b_browser_settings_upsert():
    return _CVSTUDIO_STORAGE_BRIDGE.browser_settings_upsert()


@app.route("/storage/browser-settings/delete", methods=["POST"])
def phase2b_browser_settings_delete():
    return _CVSTUDIO_STORAGE_BRIDGE.browser_settings_delete()


_CVSTUDIO_DIAGNOSTICS_SERVICE = DiagnosticsService(
    request_json=lambda: request.get_json(silent=True) or {},
    jsonify=lambda payload: jsonify(payload),
    send_file=lambda *args, **kwargs: send_file(*args, **kwargs),
    request_id=lambda: _cvstudio_current_request_id(),
    runtime_payload=lambda: _cvstudio_runtime_diagnostics_payload(),
    cancel_preview_work=lambda: _spider_preview_cancel_persistent_work(),
    clear_preview_cache=lambda: _spider_resume_text_cache_clear(),
    preview_cache_stats=lambda: _spider_preview_cache_stats(),
    redact_text=lambda value: _cvstudio_redact_support_text(value),
    sanitize_browser=lambda payload: _cvstudio_sanitize_browser_diagnostics(
        payload
    ),
    archive_timestamp=lambda: time.strftime(
        "%Y%m%d_%H%M%S", time.localtime()
    ),
    version=lambda: _CVSTUDIO_VERSION,
    root_path=lambda: _CVSTUDIO_ROOT,
    runtime_log_path=lambda: _RUNTIME_LOG_PATH,
)


@app.route("/diagnostics/runtime", methods=["GET"])
def cvstudio_runtime_diagnostics():
    return _CVSTUDIO_DIAGNOSTICS_SERVICE.runtime()


@app.route("/diagnostics/clear_preview_cache", methods=["POST"])
def cvstudio_clear_preview_cache():
    return _CVSTUDIO_DIAGNOSTICS_SERVICE.clear_preview()


@app.route("/diagnostics/support_bundle", methods=["POST"])
def cvstudio_support_bundle():
    return _CVSTUDIO_DIAGNOSTICS_SERVICE.support_bundle()


def _spider_candidate_attachment_records(
    token,
    candidate_id,
    *,
    include_status=False,
):
    """Return official resume records, optionally with discovery success."""
    if not token or not candidate_id:
        return ([], False) if include_status else []
    cid = urllib.parse.quote(str(candidate_id), safe="")
    params = urllib.parse.urlencode([
        ("Type", "Resume"),
        ("Type", "FormattedResume"),
        ("Latest", "true"),
        ("Limit", "4"),
    ])
    try:
        raw, _ctype, _dispo = _spider_get_ja_raw(
            token,
            "candidates/{}/attachments?{}".format(cid, params),
            timeout=10,
            accept_json=True,
        )
        payload = safe_json(raw, None)
        items = payload.get("items") if isinstance(payload, dict) else None
        if not isinstance(items, list):
            return ([], False) if include_status else []
        records = [x for x in items if isinstance(x, dict)]
        # Stable two-stage sort: newest within each attachment type, with the
        # original Resume type ahead of FormattedResume.
        records.sort(key=lambda item: str(item.get("createdAt") or ""), reverse=True)
        records.sort(key=lambda item: 0 if str(item.get("type") or "").lower() == "resume" else 1 if str(item.get("type") or "").lower() == "formattedresume" else 2)
        records = records[:4]
        return (records, True) if include_status else records
    except _SpiderJobAdderReconnectRequired:
        raise
    except Exception:
        return ([], False) if include_status else []


def _spider_fetch_candidate_resume_text(token, candidate_id):
    """Fetch and extract the latest JobAdder resume for match-fit scoring.

    The Keywords endpoint searches the latest resume, while candidate detail JSON
    often omits that resume text.  This helper closes that evidence gap without
    exposing the full extracted CV in the candidate-card response.
    """
    cid = urllib.parse.quote(str(candidate_id), safe="")
    # One latest attachment keeps discovery request volume bounded. JobAdder's
    # Latest=true response is already ordered with Resume ahead of FormattedResume.
    records, discovery_ok = _spider_candidate_attachment_records(
        token,
        candidate_id,
        include_status=True,
    )
    records = records[:1]
    if not records:
        # A confirmed deletion and a transient listing failure must both exclude
        # candidate-only cache reuse: neither supplies current document bytes.
        return "", (
            "resume attachment unavailable"
            if discovery_ok
            else "resume attachment discovery failed"
        )
    for record in records:
        attach_id = record.get("attachmentId") or record.get("id")
        if not attach_id:
            continue
        try:
            raw, ctype, dispo = _spider_get_ja_raw(
                token,
                "candidates/{}/attachments/{}".format(cid, urllib.parse.quote(str(attach_id), safe="")),
                timeout=12,
            )
            filename = record.get("fileName") or dispo or ""
            content_sha256 = _spider_download_content_identity(raw)
            content_kind = _document_content_kind(raw)
            if _document_is_legacy_doc(raw, ctype, filename):
                content_kind = "legacy_doc"
            cached = _spider_resume_text_cache_get(
                token,
                candidate_id,
                content_sha256,
            )
            if cached is not None and str(
                cached.get("content_kind") or ""
            ) == content_kind:
                if content_kind == "legacy_doc":
                    if not cached.get("antiword_verified"):
                        cached = None
                    else:
                        _require_verified_antiword()
                if cached is not None:
                    return (
                        cached.get("text") or "",
                        cached.get("source") or "cached resume",
                    )
            text, source = _spider_extract_text_from_download(raw, ctype, filename, ocr_page_limit=2, ocr_deadline_seconds=18)
            if text:
                _spider_resume_text_cache_put(
                    token,
                    candidate_id,
                    text,
                    source,
                    content_sha256=content_sha256,
                    content_kind=content_kind,
                    antiword_verified=content_kind == "legacy_doc",
                )
                return text, source
        except _SpiderJobAdderReconnectRequired:
            raise
        except AntiwordDependencyError:
            raise
        except Exception:
            continue
    return "", "resume text unavailable"


def _spider_fetch_candidate_bundle(token, candidate_id, include_resume=True):
    detail = _spider_fetch_candidate_detail(token, candidate_id)
    resume_text, resume_source = ("", "")
    if include_resume:
        resume_text, resume_source = _spider_fetch_candidate_resume_text(token, candidate_id)
    return {"detail": detail, "resume_text": resume_text, "resume_source": resume_source}


def _ai_crawler_lock_allowed(body=None):
    """Compatibility hook retained after removing the AI Crawler password."""
    return True


def _ai_crawler_locked_response():
    return _cvstudio_error_payload("AI_CRAWLER_ACCESS_UNAVAILABLE", "AI Crawler access is unavailable.", 503, action="retry")

@app.route("/jobadder/spider_options", methods=["GET"])
def jobadder_spider_options():
    """Best-effort JobAdder dropdown/options lookup for AI Crawler filters, with safe fallbacks."""
    if not _ai_crawler_lock_allowed():
        return _ai_crawler_locked_response()
    name = str(request.args.get("name") or "").strip().lower()
    fallback = _spider_option_fallbacks(name)
    token = _ja_refresh_access_token(force=False)
    if not token:
        return jsonify({"items": fallback, "fallback": True})
    candidate_lists = {
        "industry": ["industry", "industries", "candidateIndustry", "candidateIndustries"],
        "industries": ["industry", "industries", "candidateIndustry", "candidateIndustries"],
        "it_skills": ["skills", "skill", "candidateSkills", "candidateSkill", "itSkills", "it_skills", "technologySkills"],
        "skills": ["skills", "skill", "candidateSkills", "candidateSkill", "itSkills", "technologySkills"],
        "qualifications": ["qualifications", "qualification", "candidateQualifications", "candidateQualification", "certifications"],
        "qualification": ["qualifications", "qualification", "candidateQualifications", "candidateQualification", "certifications"],
    }.get(name, [name])
    errors = []
    for list_name in candidate_lists:
        try:
            _status, payload = _JOBADDER_CLIENT.request_json(
                "lists/" + urllib.parse.quote(list_name, safe=""),
                token=token,
                timeout=8,
                fallback={},
            )
            vals = _spider_extract_option_values(payload)
            if vals:
                return jsonify({"items": vals, "source": list_name, "fallback": False})
        except urllib.error.HTTPError as e:
            errors.append("{}:{}".format(list_name, e.code))
        except Exception as e:
            errors.append("{}:{}".format(list_name, str(e)[:80]))
    return jsonify({"items": fallback, "fallback": True, "tried": candidate_lists, "errors": errors[:6]})



def _spider_preview_name(candidate_id, detail=None):
    if isinstance(detail, dict):
        for key in ["name", "fullName", "displayName"]:
            if detail.get(key):
                return str(detail.get(key))[:160]
        first = str(detail.get("firstName") or "").strip()
        last = str(detail.get("lastName") or "").strip()
        if first or last:
            return (first + " " + last).strip()[:160]
    return "Candidate " + str(candidate_id or "")[:40]


def _spider_preview_text_from_detail(detail):
    """Build a readable side preview from a JobAdder candidate record if no raw resume is available."""
    if not isinstance(detail, dict):
        return ""
    sections = []
    name = _spider_preview_name("", detail)
    title = detail.get("currentPosition") or detail.get("position") or detail.get("jobTitle") or ""
    company = detail.get("currentCompany") or detail.get("employer") or ""
    location = detail.get("location") or detail.get("city") or detail.get("country") or ""
    header = []
    if name and name != "Candidate ":
        header.append("Name: " + str(name))
    if title or company:
        header.append("Current: " + " @ ".join([str(x) for x in [title, company] if x]))
    if location:
        header.append("Location: " + str(location))
    if header:
        sections.append("\n".join(header))

    preferred = [
        ("Summary", ["summary", "profile", "candidateSummary", "objective"]),
        ("Skills", ["skills", "skill", "itSkills", "technologySkills"]),
        ("Qualifications", ["qualifications", "qualification", "certifications", "certification"]),
        ("Employment", ["employment", "workExperience", "experience", "positions", "history"]),
        ("Education", ["education", "educations"]),
        ("Resume / CV", ["resume", "cv", "curriculumVitae"]),
        ("Notes", ["notes", "comments"]),
    ]

    def compact(v):
        if v is None or v == "" or v == [] or v == {}:
            return ""
        if isinstance(v, str):
            out = v
        else:
            try:
                out = json.dumps(v, ensure_ascii=False, indent=2)
            except Exception:
                out = str(v)
        out = re.sub(r"\n{3,}", "\n\n", out)
        return out.strip()[:5000]

    used = set()
    lower_map = {str(k).lower(): k for k in detail.keys()}
    for title_label, keys in preferred:
        vals = []
        for key in keys:
            for lk, real in lower_map.items():
                if key.lower() in lk and real not in used:
                    txt = compact(detail.get(real))
                    if txt:
                        vals.append(txt)
                        used.add(real)
        if vals:
            sections.append(title_label + "\n" + "\n\n".join(vals)[:6500])
    if not sections:
        blob = _spider_candidate_blob(detail)
        return blob[:9000]
    return "\n\n".join(sections)[:12000]



def _spider_is_legacy_word_doc_bytes(raw):
    """Detect legacy binary Word .doc downloads even when JobAdder labels them octet-stream.

    Browser preview cannot embed .doc directly, so Spider must first recognise old
    OLE Word files, then convert or fall back to extracted text.
    """
    raw = raw or b""
    if not raw.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        return False
    try:
        import olefile as _olefile
        ole = _olefile.OleFileIO(io.BytesIO(raw))
        try:
            return ole.exists('WordDocument')
        finally:
            ole.close()
    except Exception:
        # Still let the verified Antiword path inspect it; JobAdder often serves
        # legacy .doc as generic application/octet-stream without a usable name.
        return True


def _document_content_kind(raw):
    """Return a strong byte-signature classification before metadata fallback."""
    raw = raw or b""
    if _spider_is_legacy_word_doc_bytes(raw):
        return "legacy_doc"
    if raw.startswith(b"%PDF"):
        return "pdf"
    if raw[:2] == b"PK":
        return "zip"
    if (
        raw.startswith(b"\x89PNG\r\n\x1a\n")
        or raw.startswith(b"\xff\xd8\xff")
        or raw.startswith((b"II*\x00", b"MM\x00*", b"BM", b"RIFF"))
    ):
        return "image"
    return ""


def _document_is_legacy_doc(raw, content_type="", filename=""):
    """Apply strong byte identity before legacy-DOC metadata fallback."""

    content_kind = _document_content_kind(raw)
    if content_kind:
        return content_kind == "legacy_doc"
    ctype = str(content_type or "").lower()
    name = _spider_filename_from_disposition(
        filename,
        filename,
    ).lower()
    return "msword" in ctype or name.endswith(".doc")


def _document_image_extension(raw):
    raw = raw or b""
    if raw.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if raw.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    if raw.startswith((b"II*\x00", b"MM\x00*")):
        return ".tif"
    if raw.startswith(b"BM"):
        return ".bmp"
    if raw.startswith(b"RIFF"):
        return ".webp"
    return ".img"




def _spider_doc_text_quality_ok(text):
    text = _spider_clean_doc_text_for_preview(text)
    if len(text) < 80:
        return False
    sample = text[:6000]
    visible = sum(1 for c in sample if not c.isspace()) or 1
    bad = sum(1 for c in sample if (ord(c) > 0xFFFF) or (ord(c) >= 0x3000 and not ('\u4e00' <= c <= '\u9fff')))
    signals = len(re.findall(r'(?i)\b(name|email|e-mail|phone|mobile|experience|education|skills|work|project|position|summary|profile)\b|@', sample))
    return (bad / visible) < 0.25 and signals >= 1


def _spider_extract_legacy_doc_text_for_preview(raw):
    """Verified Antiword text extraction for legacy .doc Spider previews.

    LibreOffice and native parsing remain defense-in-depth probes after a
    document-specific Antiword failure, but their text is never returned as a
    verified legacy-.doc result.
    """
    raw = raw or b''
    if not raw:
        raise AntiwordDependencyError(
            "document-extraction-failed",
            "This legacy .doc is empty or unreadable. Convert it to DOCX or "
            "PDF and retry.",
        )

    # 1) The only accepted success path is the pinned, function-verified runtime.
    _require_verified_antiword()
    try:
        import tempfile as _tmp, os as _os
        with _tmp.TemporaryDirectory() as td:
            doc_path = _os.path.join(td, 'input.doc')
            with open(doc_path, 'wb') as fh:
                fh.write(raw)
            res = _run_verified_antiword((doc_path,), timeout=20)
            if res.returncode == 0:
                for enc in ('utf-8', 'cp1252', 'latin-1'):
                    txt = _spider_clean_doc_text_for_preview(res.stdout.decode(enc, errors='ignore'))
                    if _spider_doc_text_quality_ok(txt):
                        return txt[:12000]
    except Exception:
        pass

    # 2) LibreOffice remains a bounded defense-in-depth probe only.
    try:
        import tempfile as _tmp, subprocess as _sp, os as _os
        soffice = _find_soffice_binary()
        if soffice:
            with _tmp.TemporaryDirectory() as td:
                doc_path = _os.path.join(td, 'input.doc')
                with open(doc_path, 'wb') as fh:
                    fh.write(raw)
                _sp.run([soffice, '--headless', '--convert-to', 'txt:Text', '--outdir', td, doc_path], capture_output=True, timeout=45)
                txt_path = _os.path.join(td, 'input.txt')
                if _os.path.exists(txt_path):
                    data = open(txt_path, 'rb').read()
                    for enc in ('utf-8', 'utf-16', 'cp1252', 'latin-1'):
                        txt = _spider_clean_doc_text_for_preview(data.decode(enc, errors='ignore'))
                        if _spider_doc_text_quality_ok(txt):
                            raise AntiwordDependencyError(
                                "document-extraction-failed",
                                "Verified Antiword could not decode this legacy "
                                ".doc. Convert it to DOCX or PDF and retry.",
                            )
    except AntiwordDependencyError:
        raise
    except Exception:
        pass

    # 3) Native OLE parsing remains defense-in-depth and cannot satisfy success.
    try:
        import struct as _struct
        import olefile as _olefile
        ole = _olefile.OleFileIO(io.BytesIO(raw))
        try:
            if ole.exists('WordDocument'):
                word = ole.openstream('WordDocument').read()
                flags = _struct.unpack_from('<H', word, 0x0A)[0]
                table_name = '1Table' if (flags & 0x0200) else '0Table'
                if ole.exists(table_name) and len(word) >= 0x01AA:
                    table = ole.openstream(table_name).read()
                    fc_clx = _struct.unpack_from('<I', word, 0x01A2)[0]
                    lcb_clx = _struct.unpack_from('<I', word, 0x01A6)[0]
                    clx = table[fc_clx:fc_clx + lcb_clx] if lcb_clx and fc_clx < len(table) else b''
                    pos = 0
                    pieces = []
                    while pos < len(clx):
                        b = clx[pos]
                        if b == 0x01 and pos + 3 <= len(clx):
                            cb = _struct.unpack_from('<H', clx, pos + 1)[0]
                            pos += 3 + cb
                            continue
                        if b != 0x02 or pos + 5 > len(clx):
                            pos += 1
                            continue
                        lcb = _struct.unpack_from('<I', clx, pos + 1)[0]
                        pcdt = clx[pos + 5:pos + 5 + lcb]
                        if len(pcdt) < 16:
                            break
                        n = (len(pcdt) - 4) // 12
                        if n <= 0:
                            break
                        cps = [_struct.unpack_from('<I', pcdt, i * 4)[0] for i in range(n + 1)]
                        base = 4 * (n + 1)
                        for i in range(n):
                            cp0, cp1 = cps[i], cps[i + 1]
                            if cp1 <= cp0:
                                continue
                            off = base + (i * 8)
                            if off + 6 > len(pcdt):
                                continue
                            fc_comp = _struct.unpack_from('<I', pcdt, off + 2)[0]
                            compressed = bool(fc_comp & 0x40000000)
                            fc = fc_comp & 0x3FFFFFFF
                            chars = cp1 - cp0
                            if compressed:
                                piece = word[fc // 2:(fc // 2) + chars].decode('cp1252', errors='ignore')
                            else:
                                piece = word[fc:fc + (chars * 2)].decode('utf-16-le', errors='ignore')
                            if piece:
                                pieces.append(piece)
                        break
                    txt = _spider_clean_doc_text_for_preview(''.join(pieces))
                    if _spider_doc_text_quality_ok(txt):
                        raise AntiwordDependencyError(
                            "document-extraction-failed",
                            "Verified Antiword could not decode this legacy "
                            ".doc. Convert it to DOCX or PDF and retry.",
                        )
        finally:
            ole.close()
    except AntiwordDependencyError:
        raise
    except Exception:
        pass

    # 4) Raw scans are also diagnostic-only and never become accepted output.
    for enc in ('utf-16-le', 'cp1252', 'latin-1'):
        try:
            txt = _spider_clean_doc_text_for_preview(raw[:300000].decode(enc, errors='ignore'))
            if _spider_doc_text_quality_ok(txt):
                raise AntiwordDependencyError(
                    "document-extraction-failed",
                    "Verified Antiword could not decode this legacy .doc. "
                    "Convert it to DOCX or PDF and retry.",
                )
        except AntiwordDependencyError:
            raise
        except Exception:
            pass
    raise AntiwordDependencyError(
        "document-extraction-failed",
        "Verified Antiword could not decode this legacy .doc. Convert it to "
        "DOCX or PDF and retry.",
    )


def _spider_scan_doc_text_runs(buf, min_run=6):
    """'strings'-style readable-text recovery from a raw byte buffer.

    A last-resort fallback for legacy .doc files whose Word piece table (CLX)
    cannot be parsed structurally (fast-saved documents, unusual or truncated
    CLX layouts). Scans for runs of printable characters stored two ways:

      * 8-bit cp1252/Latin-1 text (contiguous printable bytes), and
      * ASCII/Latin text stored as UTF-16LE (printable low byte + 0x00 high
        byte), which is exactly how a Unicode-body Word 97 file stores prose.

    Returns cleaned plain text or ''. This output is never Antiword-verified and
    is only reachable from the explicit opt-in recovery path.
    """
    buf = buf or b""
    if not buf:
        return ""
    low = rb"\x09\x0a\x0d\x20-\x7e\xa0-\xff"
    try:
        ascii_parts = re.findall(rb"[" + low + rb"]{%d,}" % int(min_run), buf)
        ascii_text = _spider_clean_doc_text_for_preview(
            "\n".join(p.decode("cp1252", errors="ignore") for p in ascii_parts)
        )
        utf16_parts = re.findall(
            rb"(?:[" + low + rb"]\x00){%d,}" % int(min_run), buf
        )
        utf16_text = _spider_clean_doc_text_for_preview(
            "\n".join(p.decode("utf-16-le", errors="ignore") for p in utf16_parts)
        )
    except Exception:
        return ""

    def _alpha(text):
        return sum(1 for c in text if c.isalpha())

    # Prefer whichever layout yielded more real alphabetic content.
    return utf16_text if _alpha(utf16_text) >= _alpha(ascii_text) else ascii_text


def _spider_recover_legacy_doc_text_unverified(raw):
    """Best-effort native text recovery from a legacy OLE .doc.

    Used ONLY behind an explicit opt-in when verified Antiword cannot decode a
    document it should be able to (e.g. a Word 97 file whose body is stored as
    UTF-16 Unicode). Parses the Word piece table (CLX) and decodes each piece as
    compressed cp1252 or 16-bit UTF-16LE. Returns cleaned plain text or ''.

    This output is NOT Antiword-verified: callers must label it accordingly and
    must never cache or persist it as a verified legacy-.doc result. It cannot be
    reached by the default extraction path and therefore does not weaken the
    verified-Antiword guarantee.
    """
    raw = raw or b''
    if not raw:
        return ""
    try:
        import struct as _struct
        import olefile as _olefile
        ole = _olefile.OleFileIO(io.BytesIO(raw))
    except Exception:
        return ""
    try:
        if not ole.exists('WordDocument'):
            return ""
        word = ole.openstream('WordDocument').read()
        if len(word) < 0x0200:
            return ""
        flags = _struct.unpack_from('<H', word, 0x0A)[0]
        table_name = '1Table' if (flags & 0x0200) else '0Table'
        if not ole.exists(table_name):
            return ""
        table = ole.openstream(table_name).read()
        fc_clx = _struct.unpack_from('<I', word, 0x01A2)[0]
        lcb_clx = _struct.unpack_from('<I', word, 0x01A6)[0]
        if not lcb_clx or fc_clx >= len(table):
            return ""
        clx = table[fc_clx:fc_clx + lcb_clx]
        pos = 0
        pieces = []
        while pos < len(clx):
            b = clx[pos]
            if b == 0x01 and pos + 3 <= len(clx):
                cb = _struct.unpack_from('<H', clx, pos + 1)[0]
                pos += 3 + cb
                continue
            if b != 0x02 or pos + 5 > len(clx):
                pos += 1
                continue
            lcb = _struct.unpack_from('<I', clx, pos + 1)[0]
            pcdt = clx[pos + 5:pos + 5 + lcb]
            if len(pcdt) < 16:
                break
            n = (len(pcdt) - 4) // 12
            if n <= 0:
                break
            cps = [_struct.unpack_from('<I', pcdt, i * 4)[0] for i in range(n + 1)]
            base = 4 * (n + 1)
            for i in range(n):
                cp0, cp1 = cps[i], cps[i + 1]
                if cp1 <= cp0:
                    continue
                off = base + (i * 8)
                if off + 6 > len(pcdt):
                    continue
                fc_comp = _struct.unpack_from('<I', pcdt, off + 2)[0]
                compressed = bool(fc_comp & 0x40000000)
                fc = fc_comp & 0x3FFFFFFF
                chars = cp1 - cp0
                if compressed:
                    piece = word[fc // 2:(fc // 2) + chars].decode('cp1252', errors='ignore')
                else:
                    piece = word[fc:fc + (chars * 2)].decode('utf-16-le', errors='ignore')
                if piece:
                    pieces.append(piece)
            break
        cleaned = _spider_clean_doc_text_for_preview(''.join(pieces))
        if _spider_doc_text_quality_ok(cleaned):
            return cleaned
        # The structured piece-table parse produced nothing usable (fast-saved
        # or unusual CLX). Fall back to a printable-run scan of the main text
        # stream, then the whole file, before giving up. Still never verified.
        for source in (word, raw):
            scanned = _spider_scan_doc_text_runs(source)
            if _spider_doc_text_quality_ok(scanned):
                return scanned
        return cleaned
    except Exception:
        return ""
    finally:
        try:
            ole.close()
        except Exception:
            pass


def _spider_unverified_doc_recovery_text(raw, ctype, filename, allow_unverified):
    """Return opt-in unverified recovery text for a legacy .doc, or ''.

    Shared by every AntiwordDependencyError handler in the AI Crawler preview
    flow so the recruiter's explicit ``allow_unverified`` opt-in is honoured no
    matter where the verified-Antiword gate raised. In particular the *visual*
    render path (`_office_bytes_to_pdf_preview`) runs the verified-Antiword gate
    before extraction, so for a legacy .doc it raises before `extract_text` — the
    place the opt-in recovery used to live — is ever called. Centralising the
    decision here closes that gap.

    The result is never Antiword-verified: callers must label it and must never
    cache it. Recovery requires the explicit opt-in and a real legacy .doc, and
    is still bounded by the same `_spider_doc_text_quality_ok` gate, so this
    cannot weaken the verified-Antiword guarantee.
    """
    if not allow_unverified:
        return ""
    if not _document_is_legacy_doc(raw, ctype, filename):
        return ""
    recovered = _spider_recover_legacy_doc_text_unverified(raw)
    if _spider_doc_text_quality_ok(recovered):
        return recovered[:30000]
    return ""


def _spider_tesseract_path():
    return _find_mandatory_tesseract(_CVSTUDIO_ROOT) or ""


def _spider_poppler_path():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(script_dir, "poppler", "Library", "bin"),
        os.path.join(script_dir, "poppler", "bin"),
        os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "poppler", "Library", "bin"),
        os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "poppler", "bin"),
        os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "poppler", "Library", "bin"),
        os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "poppler", "bin"),
    ]
    for candidate in candidates:
        if candidate and os.path.isdir(candidate) and (
            os.path.isfile(os.path.join(candidate, "pdfinfo.exe"))
            or os.path.isfile(os.path.join(candidate, "pdftoppm.exe"))
            or os.path.isfile(os.path.join(candidate, "pdfinfo"))
            or os.path.isfile(os.path.join(candidate, "pdftoppm"))
        ):
            return candidate
    return None


def _spider_is_image_download(raw, content_type="", filename=""):
    ctype = str(content_type or "").lower()
    name = str(filename or "").lower()
    return bool(
        ctype.startswith("image/")
        or name.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"))
        or raw.startswith(b"\x89PNG\r\n\x1a\n")
        or raw.startswith(b"\xff\xd8\xff")
        or raw.startswith((b"II*\x00", b"MM\x00*", b"BM", b"RIFF"))
    )


def _spider_ocr_image_download(raw, timeout=25):
    try:
        import pytesseract
        tess = _spider_tesseract_path()
        if not tess:
            return "", "image resume found but Tesseract OCR is unavailable"
        pytesseract.pytesseract.tesseract_cmd = tess
        image = _safe_image_open(raw)
        try:
            text = _ocr_image_text(pytesseract, image, lang="eng", timeout=max(5, min(35, int(timeout or 25))))
        finally:
            try:
                image.close()
            except Exception:
                pass
        text = _spider_clean_doc_text_for_preview(text)
        if text:
            return text[:30000], "OCR text from image resume"
        return "", "image resume OCR returned no readable text"
    except Exception as exc:
        return "", "image resume OCR unavailable: " + str(exc)[:120]


def _spider_ocr_pdf_download(raw, page_limit=5, deadline_seconds=35):
    try:
        import pytesseract
        tess = _spider_tesseract_path()
        if not tess:
            return "", "scanned PDF found but Tesseract OCR is unavailable"
        pytesseract.pytesseract.tesseract_cmd = tess
        count = _pdf_page_count(raw)
        pages = min(count, max(1, min(_MAX_OCR_PAGES, int(page_limit or 5))))
        if not _OCR_SEMAPHORE.acquire(timeout=2):
            return "", "scanned PDF OCR is busy; preview remains visual only"
        started = time.monotonic()
        parts = []
        try:
            for page_number in range(1, pages + 1):
                remaining = float(deadline_seconds or 35) - (time.monotonic() - started)
                if remaining <= 2:
                    break
                images = _render_pdf_page_images(
                    raw,
                    dpi=150,
                    poppler_path=_spider_poppler_path(),
                    first_page=page_number,
                    last_page=page_number,
                    timeout=max(5, min(20, int(remaining))),
                )
                if not images:
                    continue
                image = images[0]
                try:
                    width, height = image.size
                    if width * height > _MAX_IMAGE_PIXELS:
                        continue
                    parts.append(pytesseract.image_to_string(image, lang="eng", timeout=max(5, min(18, int(remaining)))))
                finally:
                    try:
                        image.close()
                    except Exception:
                        pass
                if sum(len(x) for x in parts) >= 30000:
                    break
        finally:
            _OCR_SEMAPHORE.release()
        text = _spider_clean_doc_text_for_preview("\n".join(parts))
        if text:
            suffix = "" if pages >= count else " (first {} pages)".format(pages)
            return text[:30000], "OCR text from scanned PDF" + suffix
        return "", "scanned PDF OCR returned no readable text"
    except Exception as exc:
        return "", "scanned PDF OCR unavailable: " + str(exc)[:120]


def _spider_extract_text_from_download(raw, content_type="", filename="", ocr_page_limit=5, ocr_deadline_seconds=35, cancel_check=None):
    """Extract searchable resume text without decoding binary files as prose."""
    raw = raw or b""
    if callable(cancel_check):
        cancel_check()
    content_type = str(content_type or "").lower()
    filename = str(filename or "").lower()
    is_legacy_doc = _document_is_legacy_doc(
        raw,
        content_type,
        filename,
    )
    if not raw:
        if is_legacy_doc:
            raise AntiwordDependencyError(
                "document-extraction-failed",
                "This legacy .doc is empty or unreadable. Convert it to DOCX "
                "or PDF and retry.",
            )
        return "", "empty response"

    # JobAdder metadata is not authoritative for the downloaded bytes. Strong
    # signatures take precedence; metadata is used only when the payload has no
    # recognized identity.
    content_kind = _document_content_kind(raw)
    if is_legacy_doc:
        if callable(cancel_check):
            cancel_check()
        text = _spider_extract_legacy_doc_text_for_preview(raw)
        if callable(cancel_check):
            cancel_check()
        if text:
            return text[:30000], "downloaded legacy DOC resume"
        return "", "legacy DOC resume found but preview extraction failed"

    parsed = None
    stripped = raw.lstrip()[:1]
    if not content_kind and (
        "json" in content_type or stripped in (b"{", b"[")
    ):
        try:
            parsed = json.loads(raw.decode("utf-8", errors="replace"))
        except Exception:
            parsed = None
    if isinstance(parsed, (dict, list)):
        if isinstance(parsed, dict):
            for key in ["text", "plainText", "resumeText", "cvText", "content", "body", "description"]:
                if parsed.get(key) and isinstance(parsed.get(key), str):
                    return re.sub(r"\s+", " ", parsed.get(key)).strip()[:30000], "JobAdder JSON text"
            return _spider_preview_text_from_detail(parsed)[:30000], "JobAdder JSON profile"
        return re.sub(r"\s+", " ", " ".join(_spider_flatten(parsed))).strip()[:30000], "JobAdder JSON list"

    if not content_kind and (
        "text" in content_type
        or filename.endswith((".txt", ".html", ".htm", ".csv", ".rtf"))
    ):
        try:
            text = raw.decode("utf-8", errors="replace")
        except Exception:
            text = raw.decode(errors="replace")
        text = re.sub(r"<[^>]+>", " ", text)
        text = _spider_clean_doc_text_for_preview(text)
        if text:
            return text[:30000], "downloaded text/html"
        return "", "downloaded text contained no readable content"

    is_pdf = content_kind == "pdf" or (
        not content_kind
        and ("pdf" in content_type or filename.endswith(".pdf"))
    )
    if is_pdf:
        try:
            import pdfplumber
            parts = []
            started = time.monotonic()
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                if len(pdf.pages) > _MAX_PDF_PAGES:
                    return "", "PDF exceeds the safe page limit"
                for page in pdf.pages:
                    if callable(cancel_check):
                        cancel_check()
                    if time.monotonic() - started > 25:
                        break
                    page_text = page.extract_text() or ""
                    if page_text:
                        parts.append(page_text)
                    if sum(len(x) for x in parts) >= 30000:
                        break
            text = _spider_clean_doc_text_for_preview("\n".join(parts))
            if text:
                return text[:30000], "downloaded PDF resume"
        except Exception:
            pass
        if callable(cancel_check):
            cancel_check()
        result = _spider_ocr_pdf_download(raw, page_limit=ocr_page_limit, deadline_seconds=ocr_deadline_seconds)
        if callable(cancel_check):
            cancel_check()
        return result

    is_image = content_kind == "image" or (
        not content_kind
        and _spider_is_image_download(raw, content_type, filename)
    )
    if is_image:
        if callable(cancel_check):
            cancel_check()
        result = _spider_ocr_image_download(raw, timeout=min(35, int(ocr_deadline_seconds or 25)))
        if callable(cancel_check):
            cancel_check()
        return result

    is_docx = content_kind == "zip" or (
        not content_kind
        and ("word" in content_type or filename.endswith(".docx"))
    )
    if is_docx:
        try:
            if callable(cancel_check):
                cancel_check()
            text = _extract_docx_text_preserve_tables(raw)
            if callable(cancel_check):
                cancel_check()
            text = _spider_clean_doc_text_for_preview(text)
            if text:
                return text[:30000], "downloaded DOCX resume"
        except Exception:
            pass
        # A ZIP/Office payload is binary.  Never fall through to raw decoding.
        if raw[:2] == b"PK" or filename.endswith(".docx"):
            return "", "DOCX resume found but text extraction failed"

    # Unknown octet-streams may genuinely be plain text, but accept them only if
    # they pass the resume-text quality check.  Known binary formats have already
    # returned above and can never be labelled as resume-scored gibberish.
    try:
        text = raw[:60000].decode("utf-8", errors="ignore")
        text = _spider_clean_doc_text_for_preview(text)
        if _spider_doc_text_quality_ok(text):
            return text[:30000], "downloaded plain-text fallback"
    except Exception:
        pass
    return "", "downloaded file is visual/binary and not text-searchable"



def _spider_filename_from_disposition(content_disposition, fallback=""):
    """Extract a filename-ish value from Content-Disposition or an attachment name."""
    txt = str(content_disposition or "")
    # RFC 5987 filename*=UTF-8''file.pdf or plain filename="file.pdf".
    m = re.search(r"filename\*\s*=\s*(?:UTF-8''|)([^;]+)", txt, re.I)
    if not m:
        m = re.search(r"filename\s*=\s*\"?([^\";]+)\"?", txt, re.I)
    if m:
        try:
            return urllib.parse.unquote(m.group(1).strip().strip('"'))[:180]
        except Exception:
            return m.group(1).strip().strip('"')[:180]
    return str(fallback or "")[:180]


def _spider_pdf_text_layer_from_pdfium(textpage, page_width, page_height, max_items=2200):
    """Build a compact selectable text layer using normalized PDF coordinates.

    The browser lays these word boxes over the rendered page image.  Coordinates
    are normalized so the layer remains aligned when the responsive preview scales.
    This is intentionally bounded to keep JobAdder preview responses predictable.
    """
    try:
        text = str(textpage.get_text_range() or "")[:50000]
    except Exception:
        return None
    if not text.strip() or page_width <= 0 or page_height <= 0:
        return None

    items = []
    word_start = None
    left = bottom = right = top = None

    def flush(end_index):
        nonlocal word_start, left, bottom, right, top
        if word_start is None:
            return
        token = text[word_start:end_index]
        if token and token.strip() and None not in (left, bottom, right, top) and len(items) < max_items:
            x = max(0.0, min(1.0, float(left) / float(page_width)))
            y = max(0.0, min(1.0, (float(page_height) - float(top)) / float(page_height)))
            w = max(0.0001, min(1.0 - x, (float(right) - float(left)) / float(page_width)))
            h = max(0.0025, min(1.0 - y, (float(top) - float(bottom)) / float(page_height)))
            items.append({
                "t": token[:160],
                "s": int(word_start),
                "e": int(end_index),
                "x": round(x, 6),
                "y": round(y, 6),
                "w": round(w, 6),
                "h": round(h, 6),
            })
        word_start = None
        left = bottom = right = top = None

    for index, char in enumerate(text):
        if char.isspace():
            flush(index)
            continue
        try:
            box = textpage.get_charbox(index)
            c_left, c_bottom, c_right, c_top = [float(v) for v in box]
        except Exception:
            flush(index)
            continue
        if not all(map(lambda v: abs(v) < 1000000, (c_left, c_bottom, c_right, c_top))):
            flush(index)
            continue
        if word_start is None:
            word_start = index
            left, bottom, right, top = c_left, c_bottom, c_right, c_top
        else:
            left = min(left, c_left)
            bottom = min(bottom, c_bottom)
            right = max(right, c_right)
            top = max(top, c_top)
    flush(len(text))
    if not items:
        return None
    return {"text": text, "items": items, "truncated": len(items) >= max_items}


def _spider_pdfium_preview_pages(raw, first_page=1, last_page=None, dpi=118, cancel_check=None):
    """Render PDF pages and their selectable coordinate text layers in one pass."""
    import pypdfium2 as pdfium
    document = pdfium.PdfDocument(raw)
    pages = []
    document_text_parts = []
    try:
        page_count = len(document)
        if page_count > _MAX_PDF_PAGES:
            raise ValueError("PDF has {} pages; the safe limit is {}".format(page_count, _MAX_PDF_PAGES))
        first_page = max(1, int(first_page or 1))
        if first_page > page_count:
            return page_count, [], ""
        if last_page is None:
            last_page = page_count
        last_page = min(page_count, max(first_page, int(last_page)))
        scale = max(1.0, min(2.2, float(dpi or 118) / 72.0))
        remaining_text_items = 16000
        for page_index in range(first_page - 1, last_page):
            if callable(cancel_check):
                cancel_check()
            page = document[page_index]
            bitmap = None
            textpage = None
            try:
                page_width = float(page.get_width() or 0)
                page_height = float(page.get_height() or 0)
                rotation = int(page.get_rotation() or 0)
                bitmap = page.render(scale=scale)
                if callable(cancel_check):
                    cancel_check()
                image = bitmap.to_pil().copy()
                text_layer = None
                # Extract page text once for all-page Ctrl+F coverage. Rotated
                # pages keep searchable text but intentionally omit coordinate
                # overlays because producer-specific rotations can misalign boxes.
                try:
                    textpage = page.get_textpage()
                    page_text = str(textpage.get_text_range() or "")[:50000]
                    if page_text:
                        document_text_parts.append(page_text)
                    if rotation % 360 == 0 and remaining_text_items > 0:
                        text_layer = _spider_pdf_text_layer_from_pdfium(
                            textpage,
                            page_width,
                            page_height,
                            max_items=min(2200, remaining_text_items),
                        )
                        if text_layer:
                            remaining_text_items -= len(text_layer.get("items") or [])
                except Exception:
                    text_layer = None
                pages.append({"image": image, "text_layer": text_layer, "rotation": rotation})
            finally:
                try:
                    if textpage is not None:
                        textpage.close()
                except Exception:
                    pass
                try:
                    if bitmap is not None:
                        bitmap.close()
                except Exception:
                    pass
                try:
                    page.close()
                except Exception:
                    pass
        # Preserve the old all-page Ctrl+F/Text coverage without running a
        # second pdfplumber pass. Remaining pages are text-only and therefore
        # cheap compared with rendering/encoding them.
        if sum(len(x) for x in document_text_parts) < 30000 and last_page < page_count:
            for page_index in range(last_page, page_count):
                if callable(cancel_check):
                    cancel_check()
                page = document[page_index]
                textpage = None
                try:
                    textpage = page.get_textpage()
                    text = str(textpage.get_text_range() or "")
                    if text:
                        document_text_parts.append(text)
                    if sum(len(x) for x in document_text_parts) >= 30000:
                        break
                except Exception:
                    pass
                finally:
                    try:
                        if textpage is not None:
                            textpage.close()
                    except Exception:
                        pass
                    try:
                        page.close()
                    except Exception:
                        pass
        return page_count, pages, "\n".join(document_text_parts)[:30000]
    finally:
        try:
            document.close()
        except Exception:
            pass


def _spider_pdf_pages_preview_payload(raw, source, note="", max_pages=12, dpi=118, cancel_check=None):
    """Render PDF pages into fast preview images plus a selectable visual text layer.

    The original PDF bytes are never sent to the browser, so native PDF download
    controls remain absent.  PDFium rendering and text extraction happen in one
    pass.  Fast WebP encoding materially reduces the first-preview delay compared
    with the older high-effort encoder.
    """
    raw = raw or b""
    if not raw:
        return None
    page_count = 0
    rendered = []
    visual_search_text = ""
    try:
        page_count, rendered, visual_search_text = _spider_pdfium_preview_pages(
            raw,
            first_page=1,
            last_page=max(1, int(max_pages or 12)),
            dpi=max(96, min(150, int(dpi or 118))),
            cancel_check=cancel_check,
        )
    except _SpiderPreviewCancelled:
        raise
    except Exception:
        # Owner/source fallback for environments without the bundled PDFium wheel.
        try:
            page_count = _pdf_page_count(raw)
            shown_target = min(max(1, int(max_pages or 12)), page_count)
            images = _render_pdf_page_images(
                raw,
                first_page=1,
                last_page=shown_target,
                dpi=max(96, min(150, int(dpi or 118))),
                timeout=45,
            )
            rendered = [{"image": image, "text_layer": None, "rotation": 0} for image in images]
            visual_search_text = ""
        except Exception:
            return None
    if not rendered:
        return None

    page_images = []
    visual_text_layers = []
    try:
        from PIL import Image as _PILImage
        for page_data in rendered:
            if callable(cancel_check):
                cancel_check()
            image = page_data.get("image")
            work = image
            try:
                # A 118-DPI A4 page is already about 975 px wide.  This remains
                # readable while avoiding the much heavier 135-DPI/method-4 path.
                if max(work.size or (0, 0)) > 1450:
                    work.thumbnail((1450, 2200), _PILImage.Resampling.LANCZOS)
                if work.mode not in ("RGB", "L"):
                    if "A" in work.getbands():
                        flattened = _PILImage.new("RGB", work.size, "white")
                        flattened.paste(work, mask=work.getchannel("A"))
                        work = flattened
                    else:
                        work = work.convert("RGB")
                elif work.mode == "L":
                    work = work.convert("RGB")
                out = io.BytesIO()
                mime = "image/webp"
                try:
                    # method=0 is deliberately chosen for interactive preview
                    # latency; it was about 3x faster than method=4 in QA.
                    work.save(out, format="WEBP", quality=78, method=0)
                except Exception:
                    out = io.BytesIO()
                    mime = "image/jpeg"
                    work.save(out, format="JPEG", quality=80)
                page_images.append("data:{};base64,{}".format(mime, base64.b64encode(out.getvalue()).decode("ascii")))
                visual_text_layers.append(page_data.get("text_layer"))
            finally:
                try:
                    if work is not image:
                        work.close()
                except Exception:
                    pass
                try:
                    image.close()
                except Exception:
                    pass
    finally:
        for page_data in rendered[len(page_images):]:
            try:
                page_data.get("image").close()
            except Exception:
                pass

    if not page_images:
        return None
    notes = [str(note or "").strip()]
    if page_count > len(page_images):
        notes.append("Showing the first {} of {} pages in protected preview mode.".format(len(page_images), page_count))
    if any(layer and layer.get("items") for layer in visual_text_layers):
        notes.append("Visual text selection and keyword overlays are available.")
    else:
        notes.append("This file has no coordinate text layer; use Select text for extracted/OCR text.")
    notes.append("Preview download controls are disabled.")
    return {
        "visual_mode": "pdf_pages",
        "page_images": page_images,
        "visual_text_layers": visual_text_layers,
        "visual_search_text": str(visual_search_text or "")[:30000],
        "page_count": page_count,
        "shown_pages": len(page_images),
        "source": str(source or "actual JobAdder PDF resume"),
        "note": " ".join(x for x in notes if x),
        "download_disabled": True,
    }


def _spider_visual_preview_payload(raw, content_type="", filename="", max_pages=12, dpi=118, cancel_check=None):
    """Return actual visual CV preview payload when JobAdder exposes a resume file.

    PDF resumes are rendered into bounded preview-only page images so the browser's
    native PDF download toolbar is never exposed. Image resumes remain image previews.
    Office files are converted to PDF via LibreOffice and then rendered the same way,
    with DOCX HTML fallback. Text/JSON payloads intentionally return None so the caller
    can fall back to extracted/profile text.
    """
    raw = raw or b""
    if not raw:
        return None
    import base64 as _base64, os as _os
    ctype = str(content_type or "").lower()
    fname = _spider_filename_from_disposition(filename, filename).lower()
    ext = _os.path.splitext(fname)[1].lower()
    content_kind = _document_content_kind(raw)
    is_legacy_doc = _document_is_legacy_doc(raw, ctype, fname)
    if len(raw) > 12 * 1024 * 1024:
        if is_legacy_doc:
            _spider_extract_legacy_doc_text_for_preview(raw)
        return None
    stripped = raw.lstrip()[:1]
    if not content_kind and (
        "json" in ctype or stripped in (b"{", b"[")
    ):
        return None
    if content_kind == "pdf" or (
        not content_kind and ("pdf" in ctype or ext == ".pdf")
    ):
        return _spider_pdf_pages_preview_payload(raw, "actual JobAdder PDF resume", max_pages=max_pages, dpi=dpi, cancel_check=cancel_check)
    if content_kind == "image" or (
        not content_kind
        and ("image/" in ctype or ext in (".png", ".jpg", ".jpeg"))
    ):
        if callable(cancel_check):
            cancel_check()
        mime = "image/png" if (raw[:8].startswith(b"\x89PNG") or ext == ".png" or "png" in ctype) else "image/jpeg"
        data = _base64.b64encode(raw).decode("ascii")
        return {"visual_mode": "image", "data_url": "data:" + mime + ";base64," + data, "source": "actual JobAdder image resume", "download_disabled": True, "note": "Preview download controls are disabled."}
    looks_office = content_kind == "zip" or is_legacy_doc or (
        not content_kind
        and (
            "word" in ctype
            or "msword" in ctype
            or "officedocument" in ctype
            or ext in (".docx", ".doc", ".rtf", ".odt")
        )
    )
    if looks_office:
        if is_legacy_doc:
            conv_ext = ".doc"
        elif content_kind == "zip":
            conv_ext = ".docx"
        else:
            conv_ext = ext or (".doc" if "msword" in ctype else ".docx")
        try:
            pdf_bytes, note = _office_bytes_to_pdf_preview(raw, conv_ext, cancel_check=cancel_check)
            if pdf_bytes:
                return _spider_pdf_pages_preview_payload(
                    pdf_bytes,
                    "actual JobAdder resume converted to PDF",
                    note=note,
                    max_pages=max_pages,
                    dpi=dpi,
                    cancel_check=cancel_check,
                )
        except _SpiderPreviewCancelled:
            raise
        except AntiwordDependencyError:
            raise
        except Exception:
            pass
        if conv_ext == ".docx" or raw[:2] == b"PK":
            try:
                return {"visual_mode": "html", "html": _docx_bytes_to_preview_html(raw), "source": "actual JobAdder DOCX resume layout preview", "note": "LibreOffice visual conversion unavailable; showing DOCX layout fallback. Preview download controls are disabled.", "download_disabled": True}
            except Exception:
                return None
    return None

def _spider_search_text_from_visual_payload(visual):
    """Reuse PDFium text and avoid a second PDF parser pass."""
    if isinstance(visual, dict):
        full_text = _spider_clean_doc_text_for_preview(str(visual.get("visual_search_text") or ""))[:30000]
        if full_text:
            return full_text, "PDFium visual/search text layer"
    layers = (visual or {}).get("visual_text_layers") if isinstance(visual, dict) else None
    if not isinstance(layers, list):
        return "", ""
    parts = []
    for layer in layers:
        if not isinstance(layer, dict):
            continue
        text = str(layer.get("text") or "").strip()
        if text:
            parts.append(text)
        if sum(len(part) for part in parts) >= 30000:
            break
    text = _spider_clean_doc_text_for_preview("\n".join(parts))[:30000]
    if text:
        return text, "PDFium visual text layer"
    return "", ""


def _spider_prefetch_should_defer_ocr(raw, content_type="", filename=""):
    """Defer only strongly identified PDF/image bytes or metadata fallbacks."""
    content_kind = _document_content_kind(raw)
    if content_kind:
        return content_kind in {"pdf", "image"}
    ctype = str(content_type or "").lower()
    name = str(filename or "").lower()
    return bool(
        "pdf" in ctype
        or "image/" in ctype
        or name.endswith((".pdf", ".png", ".jpg", ".jpeg"))
    )


def _spider_apply_visual_search_state(visual, search_text, search_source):
    payload = dict(visual or {})
    payload.pop("visual_search_text", None)
    payload["search_text"] = str(search_text or "")[:30000]
    if not payload["search_text"]:
        payload["search_note"] = str(search_source or "Visual CV loaded, but no searchable text could be extracted.")[:200]
    return payload


class _SpiderJobAdderReconnectRequired(RuntimeError):
    pass


def _spider_mark_jobadder_reconnect_required():
    # Backward-compatible crawler alias for the shared JobAdder connection state.
    _ja_mark_reconnect_required()


def _spider_get_ja_raw(token, path, timeout=10, *, accept_json=False):
    """GET one JobAdder resource and retry once after a forced OAuth refresh."""
    try:
        headers = {"Accept": "application/json"} if accept_json else None
        response = _JOBADDER_CLIENT.request_raw(
            path,
            method="GET",
            token=token,
            timeout=timeout,
            headers=headers,
        )
        return (
            response.body,
            response.headers.get("Content-Type", ""),
            response.headers.get("Content-Disposition", ""),
        )
    except urllib.error.HTTPError as exc:
        if exc.code == 401:
            _spider_mark_jobadder_reconnect_required()
            try:
                exc.close()
            except Exception:
                pass
            raise _SpiderJobAdderReconnectRequired("JobAdder rejected the refreshed OAuth token")
        raise


def _spider_attachment_candidates(payload):
    """Return likely resume attachment ids from a JobAdder attachments/documents payload."""
    out = []
    def walk(x, depth=0):
        if depth > 5 or x is None:
            return
        if isinstance(x, list):
            for item in x[:80]:
                walk(item, depth + 1)
        elif isinstance(x, dict):
            flat = " ".join(str(v) for v in _spider_flatten(x)[:80]).lower()
            looks_resume = any(t in flat for t in ["resume", "cv", "curriculum", "formattedresume", "formatted resume"])
            if looks_resume:
                ids = []
                for key in ["id", "attachmentId", "attachmentID", "documentId", "documentID", "fileId", "fileID", "entityId"]:
                    if x.get(key):
                        ids.append(str(x.get(key)))
                name = str(x.get("fileName") or x.get("filename") or x.get("name") or x.get("title") or "")
                out.append({"ids": ids, "name": name})
            for v in x.values():
                walk(v, depth + 1)
    walk(payload)
    return out[:8]


@app.route("/jobadder/spider_preview_cancel_prefetch", methods=["POST"])
def jobadder_spider_preview_cancel_prefetch():
    # Cancellation exposes no candidate data and remains protected by the local
    # Host/Origin guards. Allow it even after the crawler has just been locked so
    # an already-running background PDFium/LibreOffice job can still be stopped.
    _spider_preview_cancel_persistent_work()
    return "", 204


@app.route("/jobadder/spider_candidate_preview", methods=["GET"])
def jobadder_spider_candidate_preview():
    """Side-panel preview with bounded full/partial caches and cancellation."""
    if not _ai_crawler_lock_allowed():
        return _ai_crawler_locked_response()
    token = _ja_refresh_access_token(force=False)
    candidate_id = str(request.args.get("candidate_id") or "").strip()
    if not token:
        return jsonify({"error": "Connect to JobAdder first", "needs_reconnect": True}), 401
    if not candidate_id:
        return jsonify({"error": "Missing candidate_id"}), 400

    prefetch_mode = str(request.args.get("prefetch") or "").strip().lower() in ("1", "true", "yes", "on")
    validate_only = str(request.args.get("validate") or "").strip().lower() in ("1", "true", "yes", "on")
    # Opt-in only: recover text from a legacy .doc that verified Antiword cannot
    # decode, using the in-process native parser. The recruiter must explicitly
    # request it (the UI sends this after the verified path returns a decode
    # failure); the recovered text is labelled unverified and never cached.
    allow_unverified_doc = str(request.args.get("allow_unverified") or "").strip().lower() in ("1", "true", "yes", "on")
    background_generation = _spider_preview_background_generation()
    prefetch_job_id = ""
    prefetch_job_store = None
    if prefetch_mode:
        opaque_cache_key = _spider_resume_cache_key(token, candidate_id)
        if not opaque_cache_key:
            opaque_cache_key = hashlib.sha256(
                candidate_id.encode("utf-8", errors="surrogatepass")
            ).hexdigest()
        prefetch_job_id = deterministic_job_id(
            _SPIDER_PREVIEW_PREFETCH_JOB_KIND,
            opaque_cache_key,
        )
        prefetch_job_store = _CVSTUDIO_JOBS
        prefetch_job_store.begin(
            prefetch_job_id,
            kind=_SPIDER_PREVIEW_PREFETCH_JOB_KIND,
            safety="safe_idempotent_read",
            request_id=_cvstudio_current_request_id(),
            stage="attachment_metadata",
        )

        @after_this_request
        def finalize_persistent_prefetch(response):
            try:
                payload = response.get_json(silent=True)
                if (
                    response.status_code == 409
                    and isinstance(payload, dict)
                    and payload.get("prefetch_cancelled")
                ):
                    prefetch_job_store.mark_cancelled(prefetch_job_id)
                elif 200 <= int(response.status_code or 0) < 300:
                    prefetch_job_store.complete(prefetch_job_id)
                else:
                    status = int(response.status_code or 500)
                    retryable = status >= 500 or status == 429
                    prefetch_job_store.fail(
                        prefetch_job_id,
                        error_code="PREFETCH_HTTP_{}".format(status),
                        error_summary=(
                            "Preview prefetch returned HTTP status {}."
                        ).format(status),
                        retryable=retryable,
                        recovery_action=(
                            "retry_same_request" if retryable else ""
                        ),
                    )
            except PersistentJobError as error:
                return app.make_response(
                    _cvstudio_persistent_job_error(error)
                )
            return response

    if not prefetch_mode and not validate_only:
        # Foreground opens cancel any old background render before they start
        # their own expensive PDFium/OCR/LibreOffice work.
        _spider_preview_cancel_persistent_work()

    def job_progress(stage, current):
        if prefetch_job_store is not None:
            prefetch_job_store.progress(
                prefetch_job_id,
                stage=stage,
                current=current,
            )

    def cancel_check():
        if prefetch_mode:
            _spider_preview_cancel_check(background_generation)

    candidate_id_safe = urllib.parse.quote(candidate_id, safe="")
    tried = []
    cancel_check()
    attachment_records = _spider_candidate_attachment_records(token, candidate_id)
    cancel_check()
    job_progress("attachment_metadata", 15)
    attachment_fingerprint = _spider_attachment_fingerprint(attachment_records) if attachment_records else ""

    if validate_only:
        return jsonify({
            "ok": True,
            "candidate_id": candidate_id,
            "attachment_fingerprint": attachment_fingerprint[:24],
            "has_attachment_metadata": bool(attachment_records),
        })

    # Attachment metadata is not a content identity. The bounded preview cache
    # is therefore consulted only after current bytes have been downloaded and
    # hashed; profile fallbacks never acquire a reusable content key.
    job_progress("cache_lookup", 25)

    preview_max_pages = 2 if prefetch_mode else 12
    preview_dpi = 104 if prefetch_mode else 118

    def get_raw(path, timeout, *, accept_json=False):
        cancel_check()
        result = _spider_get_ja_raw(
            token,
            path,
            timeout=timeout,
            accept_json=accept_json,
        )
        cancel_check()
        return result

    def render_visual(raw, ctype, filename):
        cancel_check()
        with _SPIDER_PREVIEW_RENDER_LOCK:
            cancel_check()
            return _spider_visual_preview_payload(
                raw,
                ctype,
                filename,
                max_pages=preview_max_pages,
                dpi=preview_dpi,
                cancel_check=cancel_check if prefetch_mode else None,
            )

    def extract_text(raw, ctype, filename):
        cancel_check()
        with _SPIDER_PREVIEW_RENDER_LOCK:
            cancel_check()
            try:
                result = _spider_extract_text_from_download(raw, ctype, filename, cancel_check=cancel_check if prefetch_mode else None)
            except AntiwordDependencyError:
                # Default path is unchanged: re-raise so verified-Antiword failures
                # still return 424. Only when the recruiter explicitly opted in do
                # we recover the text natively for a legacy .doc, clearly labelled
                # and never cached. The native parser does not use Antiword at all,
                # so recovery is offered for *any* Antiword failure reason on a
                # legacy .doc — not only "document-extraction-failed" (verified
                # Antiword decoded nothing) but also cases where the verified
                # runtime could not run (e.g. runtime-missing/functional-*): the
                # recruiter still gets an explicit, labelled recovery option.
                recovered = _spider_unverified_doc_recovery_text(
                    raw, ctype, filename, allow_unverified_doc
                )
                if recovered:
                    cache_provenance["unverified_recovery"] = True
                    cancel_check()
                    return (
                        recovered,
                        "legacy .doc recovered without Antiword (unverified — review before use)",
                    )
                raise
            cancel_check()
            return result

    def recover_unverified_doc_response(raw, ctype, filename):
        """Serve opt-in unverified recovery when a legacy-.doc AntiwordDependency
        error surfaced from any preview stage — including the visual render,
        whose verified-Antiword gate raises before ``extract_text`` runs. Returns
        a ready 200 response or None to fall through to the 424. The recovered
        text is labelled unverified and never cached (``cacheable=False`` plus the
        provenance flag)."""
        recovered = _spider_unverified_doc_recovery_text(
            raw, ctype, filename, allow_unverified_doc
        )
        if not recovered:
            return None
        cache_provenance["unverified_recovery"] = True
        content_sha256 = _spider_download_content_identity(raw)
        return respond(
            {
                "ok": True,
                "candidate_id": candidate_id,
                "name": _spider_preview_name(candidate_id, detail),
                "mode": "resume_text",
                "source": "legacy .doc recovered without Antiword (unverified — review before use)",
                "note": filename or "legacy .doc recovered without Antiword",
                "preview_text": recovered,
                "search_text": recovered,
                "tried": tried[:12],
            },
            content_sha256=content_sha256,
            cacheable=False,
        )

    def visual_search_text(raw, ctype, filename, visual):
        """Reuse visual text and keep background prefetch genuinely lightweight.

        PDF/image OCR can occupy Tesseract for many seconds and cannot be stopped
        safely in the middle of one page. A background preload therefore renders
        only the first two visual pages and reuses any PDFium text layer already
        available. Scanned/image-only OCR is deferred until the recruiter actually
        opens the candidate, when the full foreground request has priority.
        Cheap text/DOCX extraction remains allowed during prefetch.
        """
        search_text, search_source = _spider_search_text_from_visual_payload(visual)
        if search_text:
            return search_text, search_source
        if prefetch_mode and _spider_prefetch_should_defer_ocr(
            raw,
            ctype,
            filename,
        ):
            return "", "Searchable OCR text will be prepared when this candidate is opened."
        return extract_text(raw, ctype, filename)

    cache_provenance = {"legacy_doc": False}

    def cached_preview(raw, ctype="", filename=""):
        content_sha256 = _spider_download_content_identity(raw)
        is_legacy_doc = _document_is_legacy_doc(
            raw,
            ctype,
            filename,
        )
        cache_provenance["legacy_doc"] = is_legacy_doc
        if not content_sha256:
            return "", None
        if prefetch_mode:
            payload = _spider_preview_payload_cache_get(
                token,
                candidate_id,
                content_sha256,
                "full",
            )
            if payload is None:
                payload = _spider_preview_payload_cache_get(
                    token,
                    candidate_id,
                    content_sha256,
                    "prefetch",
                )
        else:
            payload = _spider_preview_payload_cache_get(
                token,
                candidate_id,
                content_sha256,
                "full",
            )
        if payload is not None:
            cached_kind = str(
                payload.pop("_cache_content_kind", "") or ""
            )
            cached_antiword = bool(
                payload.pop("_cache_antiword_verified", False)
            )
            if is_legacy_doc:
                if (
                    cached_kind != "legacy_doc"
                    or not cached_antiword
                ):
                    return content_sha256, None
                _require_verified_antiword()
            payload = dict(payload)
            payload["attachment_fingerprint"] = attachment_fingerprint[:24]
        return content_sha256, payload

    def cache_resume_text(raw, ctype, filename, text, source):
        # Never persist unverified native-recovery text: it must not be served to
        # a later request as if it were a verified legacy-.doc result.
        if cache_provenance.get("unverified_recovery"):
            return
        content_sha256 = _spider_download_content_identity(raw)
        content_kind = _document_content_kind(raw)
        if _document_is_legacy_doc(raw, ctype, filename):
            content_kind = "legacy_doc"
        _spider_resume_text_cache_put(
            token,
            candidate_id,
            text,
            source,
            content_sha256=content_sha256,
            content_kind=content_kind,
            antiword_verified=content_kind == "legacy_doc",
        )

    cancel_check()
    detail = _spider_fetch_candidate_detail(token, candidate_id)
    cancel_check()
    job_progress("candidate_profile", 40)

    def respond(
        payload,
        status=200,
        cacheable=True,
        *,
        content_sha256="",
    ):
        if isinstance(payload, dict):
            payload = dict(payload)
            page_count = int(payload.get("page_count") or 0)
            shown_pages = int(payload.get("shown_pages") or 0)
            is_partial = bool(prefetch_mode and page_count and shown_pages and shown_pages < page_count)
            store_variant = "prefetch" if is_partial else "full"
            payload["attachment_fingerprint"] = attachment_fingerprint[:24]
            payload["preview_cache_hit"] = False
            payload["preview_partial"] = is_partial
            payload["preview_variant"] = store_variant
            if (
                status == 200
                and cacheable
                and content_sha256
                and payload.get("ok")
            ):
                _spider_preview_payload_cache_put(
                    token,
                    candidate_id,
                    content_sha256,
                    payload,
                    store_variant,
                    content_kind=(
                        "legacy_doc"
                        if cache_provenance["legacy_doc"]
                        else ""
                    ),
                    antiword_verified=cache_provenance["legacy_doc"],
                )
        return jsonify(payload), status

    # Official JobAdder candidate-attachment API first.
    job_progress("official_attachment", 50)
    for record in attachment_records:
        cancel_check()
        attach_id = record.get("attachmentId") or record.get("id")
        if not attach_id:
            continue
        path = "candidates/{}/attachments/{}".format(candidate_id_safe, urllib.parse.quote(str(attach_id), safe=""))
        try:
            raw, ctype, dispo = get_raw(path, timeout=12)
            tried.append(path)
            filename = record.get("fileName") or dispo or ""
            content_sha256, cached_payload = cached_preview(
                raw,
                ctype,
                filename,
            )
            if cached_payload is not None:
                return jsonify(cached_payload)
            visual = render_visual(raw, ctype, filename)
            search_text, search_source = visual_search_text(raw, ctype, filename, visual)
            if search_text:
                cache_resume_text(
                    raw,
                    ctype,
                    filename,
                    search_text,
                    search_source,
                )
            if visual:
                visual = _spider_apply_visual_search_state(visual, search_text, search_source)
                visual.update({
                    "ok": True,
                    "candidate_id": candidate_id,
                    "name": _spider_preview_name(candidate_id, detail),
                    "mode": "resume_attachment_file",
                    "note": filename or "latest resume attachment",
                    "tried": tried,
                })
                return respond(
                    visual,
                    content_sha256=content_sha256,
                )
            if search_text:
                return respond({
                    "ok": True,
                    "candidate_id": candidate_id,
                    "name": _spider_preview_name(candidate_id, detail),
                    "mode": "resume_attachment_text",
                    "source": search_source,
                    "note": filename or "latest resume attachment",
                    "preview_text": search_text,
                    "search_text": search_text,
                    "tried": tried,
                }, content_sha256=content_sha256)
        except _SpiderPreviewCancelled:
            raise
        except AntiwordDependencyError as exc:
            recovered_response = recover_unverified_doc_response(raw, ctype, filename)
            if recovered_response is not None:
                return recovered_response
            return respond(
                _antiword_dependency_payload(exc),
                424,
                cacheable=False,
            )
        except urllib.error.HTTPError as exc:
            tried.append(path + ":" + str(exc.code))
        except Exception as exc:
            tried.append(path + ":" + str(exc)[:80])

    # Direct resume/document endpoints used by legacy tenant integrations.
    job_progress("direct_resume", 65)
    direct_paths = [
        "candidates/{}/resume".format(candidate_id_safe),
        "candidates/{}/attachments/Resume".format(candidate_id_safe),
        "candidates/{}/attachments/FormattedResume".format(candidate_id_safe),
    ]
    for path in direct_paths:
        cancel_check()
        try:
            raw, ctype, dispo = get_raw(path, timeout=10)
            tried.append(path)
            content_sha256, cached_payload = cached_preview(
                raw,
                ctype,
                dispo,
            )
            if cached_payload is not None:
                return jsonify(cached_payload)
            visual = render_visual(raw, ctype, dispo)
            if visual:
                search_text, search_source = visual_search_text(raw, ctype, dispo, visual)
                if search_text:
                    cache_resume_text(
                        raw,
                        ctype,
                        dispo,
                        search_text,
                        search_source,
                    )
                visual = _spider_apply_visual_search_state(visual, search_text, search_source)
                visual.update({
                    "ok": True,
                    "candidate_id": candidate_id,
                    "name": _spider_preview_name(candidate_id, detail),
                    "mode": "resume_file",
                    "tried": tried,
                })
                return respond(
                    visual,
                    content_sha256=content_sha256,
                )
            txt, src = extract_text(raw, ctype, dispo)
            if txt:
                cache_resume_text(raw, ctype, dispo, txt, src)
                return respond({
                    "ok": True,
                    "candidate_id": candidate_id,
                    "name": _spider_preview_name(candidate_id, detail),
                    "mode": "resume_text",
                    "source": src,
                    "preview_text": txt,
                    "search_text": txt,
                    "tried": tried,
                }, content_sha256=content_sha256)
        except _SpiderPreviewCancelled:
            raise
        except AntiwordDependencyError as exc:
            recovered_response = recover_unverified_doc_response(raw, ctype, dispo)
            if recovered_response is not None:
                return recovered_response
            return respond(
                _antiword_dependency_payload(exc),
                424,
                cacheable=False,
            )
        except urllib.error.HTTPError as exc:
            tried.append(path + ":" + str(exc.code))
        except Exception as exc:
            tried.append(path + ":" + str(exc)[:80])

    # Attachments/documents listing, then likely resume IDs.
    job_progress("attachment_discovery", 80)
    for list_path, kind in [
        ("candidates/{}/attachments".format(candidate_id_safe), "attachments"),
        ("candidates/{}/documents".format(candidate_id_safe), "documents"),
    ]:
        cancel_check()
        try:
            raw, _ctype, _dispo = get_raw(
                list_path,
                timeout=10,
                accept_json=True,
            )
            tried.append(list_path)
            listing = safe_json(raw, {})
            for att in _spider_attachment_candidates(listing):
                cancel_check()
                for att_id in att.get("ids") or []:
                    att_safe = urllib.parse.quote(str(att_id), safe="")
                    for suffix in ["", "/download", "/file"]:
                        cancel_check()
                        path = "candidates/{}/{}/{}{}".format(candidate_id_safe, kind, att_safe, suffix)
                        try:
                            raw2, ctype2, dispo2 = get_raw(path, timeout=12)
                            tried.append(path)
                            filename = att.get("name") or dispo2
                            content_sha256, cached_payload = cached_preview(
                                raw2,
                                ctype2,
                                filename,
                            )
                            if cached_payload is not None:
                                return jsonify(cached_payload)
                            visual = render_visual(raw2, ctype2, filename)
                            if visual:
                                search_text, search_source = visual_search_text(raw2, ctype2, filename, visual)
                                if search_text:
                                    cache_resume_text(
                                        raw2,
                                        ctype2,
                                        filename,
                                        search_text,
                                        search_source,
                                    )
                                visual = _spider_apply_visual_search_state(visual, search_text, search_source)
                                visual.update({
                                    "ok": True,
                                    "candidate_id": candidate_id,
                                    "name": _spider_preview_name(candidate_id, detail),
                                    "mode": "resume_attachment_file",
                                    "note": att.get("name") or visual.get("note") or "resume attachment",
                                    "tried": tried[:12],
                                })
                                return respond(
                                    visual,
                                    content_sha256=content_sha256,
                                )
                            txt, src = extract_text(raw2, ctype2, filename)
                            if txt:
                                cache_resume_text(
                                    raw2,
                                    ctype2,
                                    filename,
                                    txt,
                                    src,
                                )
                                return respond({
                                    "ok": True,
                                    "candidate_id": candidate_id,
                                    "name": _spider_preview_name(candidate_id, detail),
                                    "mode": "resume_attachment_text",
                                    "source": src,
                                    "note": att.get("name") or "resume attachment",
                                    "preview_text": txt,
                                    "search_text": txt,
                                    "tried": tried[:12],
                                }, content_sha256=content_sha256)
                        except _SpiderPreviewCancelled:
                            raise
                        except AntiwordDependencyError as exc:
                            recovered_response = recover_unverified_doc_response(raw2, ctype2, filename)
                            if recovered_response is not None:
                                return recovered_response
                            return respond(
                                _antiword_dependency_payload(exc),
                                424,
                                cacheable=False,
                            )
                        except urllib.error.HTTPError as exc:
                            tried.append(path + ":" + str(exc.code))
                        except Exception as exc:
                            tried.append(path + ":" + str(exc)[:80])
        except _SpiderPreviewCancelled:
            raise
        except AntiwordDependencyError as exc:
            return respond(
                _antiword_dependency_payload(exc),
                424,
                cacheable=False,
            )
        except urllib.error.HTTPError as exc:
            tried.append(list_path + ":" + str(exc.code))
        except Exception as exc:
            tried.append(list_path + ":" + str(exc)[:80])

    cancel_check()
    job_progress("profile_fallback", 90)
    if isinstance(detail, dict) and detail:
        preview_text = _spider_preview_text_from_detail(detail)
        return respond({
            "ok": True,
            "candidate_id": candidate_id,
            "name": _spider_preview_name(candidate_id, detail),
            "mode": "profile",
            "source": "JobAdder profile fields",
            "note": "Resume file not exposed by JobAdder API; showing profile fields instead.",
            "preview_text": preview_text,
            "search_text": preview_text,
            "tried": tried[:12],
        })
    return respond({
        "ok": False,
        "candidate_id": candidate_id,
        "error": "Candidate preview unavailable",
        "note": "JobAdder did not return a resume/profile preview for this candidate.",
        "tried": tried[:12],
    }, 404, cacheable=False)


def _spider_candidate_search_key(candidate):
    """Stable key for set algebra and de-duplication across JobAdder keyword probes."""
    cid = _spider_candidate_id(candidate)
    if cid:
        return "id:" + str(cid)
    if isinstance(candidate, dict):
        email = str(candidate.get("email") or candidate.get("emailAddress") or "").strip().lower()
        if email:
            return "email:" + email
        name = str(candidate.get("name") or candidate.get("fullName") or "").strip().lower()
        if name:
            return "name:" + name
    return "raw:" + hashlib.sha1(json.dumps(candidate, sort_keys=True, default=str).encode("utf-8", errors="ignore")).hexdigest()


def _spider_parse_candidate_items(payload):
    if isinstance(payload, dict):
        data_obj = payload.get("data")
        items = payload.get("items") or payload.get("candidates") or (data_obj.get("items") if isinstance(data_obj, dict) else None) or []
        total = None
        # Preserve an explicit totalCount=0 instead of losing it through truthiness.
        for key in ("totalCount", "total", "count"):
            if key in payload and payload.get(key) is not None:
                total = payload.get(key)
                break
        if total is None and isinstance(data_obj, dict):
            for key in ("totalCount", "total", "count"):
                if key in data_obj and data_obj.get(key) is not None:
                    total = data_obj.get(key)
                    break
    elif isinstance(payload, list):
        items, total = payload, len(payload)
    else:
        items, total = [], 0
    return (items if isinstance(items, list) else []), total






def _spider_jobadder_keyword_items(token, keyword, max_items=3000, page_size=1000, location="", embed=True):
    """Search JobAdder latest resumes with complete, defensive pagination.

    ``Location`` is city/state in JobAdder. Country is deliberately never sent
    here and remains a local eligibility filter. Short pages do not end a search
    while JobAdder's ``totalCount`` says more rows exist.
    """
    keyword = re.sub(r"[\x00-\x1f\x7f]", " ", str(keyword or ""))
    keyword = re.sub(r"\s+", " ", keyword).strip()[:1500]
    city_state = re.sub(r"\s+", " ", str(location or "")).strip()[:120]
    if not keyword:
        return [], {"keyword": keyword, "scanned": 0, "reported_total": 0, "pages": 0, "embed_applied": False}
    max_items = max(1, min(5000, int(max_items or 3000)))
    page_size = max(1, min(1000, int(page_size or 1000)))
    state = {"embed_active": bool(embed), "embed_rejected": False}
    base_params = [("Keywords", keyword)]
    if city_state and city_state.lower() != "any":
        base_params.append(("Location", city_state))
    if state["embed_active"]:
        base_params.extend([("Embed", "skills"), ("Embed", "notes")])

    def fetch_page(page_params):
        query = urllib.parse.urlencode(page_params, doseq=True)
        raw, _ctype, _dispo = _spider_get_ja_raw(
            token,
            "candidates?" + query,
            timeout=25,
            accept_json=True,
        )
        return safe_json(raw, {"items": []})

    def reject_unsupported_embed(error, active_params, _offset, _pages):
        if state["embed_active"] and int(getattr(error, "code", 0) or 0) in (400, 403, 422):
            state["embed_active"] = False
            state["embed_rejected"] = True
            return [(key, value) for key, value in active_params if str(key).lower() != "embed"]
        return None

    def page_signature(_items, keys):
        unique = sorted(set(keys))
        return hashlib.sha1("\n".join(unique).encode("utf-8", errors="ignore")).hexdigest() if unique else "empty"

    out, page_state = _JOBADDER_CLIENT.paginate_offset(
        "candidates",
        params=base_params,
        token=token,
        max_items=max_items,
        page_size=page_size,
        timeout=25,
        parse_page=_spider_parse_candidate_items,
        item_key=_spider_candidate_search_key,
        page_signature=page_signature,
        total_mode="replace",
        deduplicate=True,
        on_http_error=reject_unsupported_embed,
        fetch_page=fetch_page,
    )
    warning_map = {
        "repeated_page": "JobAdder repeated a candidate page before totalCount was reached",
        "no_new_unique": "JobAdder pagination returned no new unique candidates",
        "empty_before_total": "JobAdder returned an empty page before totalCount was reached",
        "no_progress": "JobAdder pagination made no forward progress",
        "missing_at_total": "JobAdder pagination reached totalCount with duplicate or missing candidate rows",
        "cap_reached": "Candidate search reached the {}-candidate safety cap".format(max_items),
    }
    warnings = [warning_map[code] for code in page_state.get("codes") or [] if code in warning_map]
    return out, {
        "keyword": keyword,
        "location": city_state,
        "scanned": len(out),
        "reported_total": page_state.get("reported_total"),
        "pages": page_state.get("pages"),
        "truncated": bool(page_state.get("truncated")),
        "incomplete": bool(page_state.get("incomplete")),
        "warnings": warnings[:5],
        "embed_applied": bool(embed and not state["embed_rejected"]),
        "embed_rejected": bool(state["embed_rejected"]),
    }


def _spider_mark_plain_keyword_candidates(items, query_text, fallback_rule="", fallback_negative_terms=None):
    query_terms = _spider_terms_for_fit(query_text, 24)
    negative_terms = list(fallback_negative_terms or [])
    marked = []
    for item in items or []:
        out = dict(item) if isinstance(item, dict) else {"value": item}
        out["_spiderSearchTerms"] = list(query_terms)
        if fallback_rule:
            out["_spiderBooleanFallback"] = True
            out["_spiderOriginalBooleanRule"] = str(fallback_rule)[:1500]
            out["_spiderBooleanNegativeTerms"] = negative_terms[:24]
        marked.append(out)
    return marked, query_terms


def _spider_plain_keyword_jobadder_candidates(token, query_text, result_pool_limit=1200, location="", fallback_rule="", fallback_negative_terms=None):
    query_text = re.sub(r"[\x00-\x1f\x7f]", " ", str(query_text or ""))
    query_text = re.sub(r"\s+", " ", query_text).strip()[:1500]
    max_items = max(100, min(3000, int(result_pool_limit or 1200)))
    items, search_meta = _spider_jobadder_keyword_items(
        token,
        query_text,
        max_items=max_items,
        page_size=min(1000, max_items),
        location=location,
    )
    marked, query_terms = _spider_mark_plain_keyword_candidates(
        items,
        query_text,
        fallback_rule=fallback_rule,
        fallback_negative_terms=fallback_negative_terms,
    )
    mode = "native_boolean→fallback" if fallback_rule else "plain"
    meta = {
        "mode": mode,
        "query": query_text,
        "returned": len(marked),
        "search": search_meta,
    }
    if fallback_rule:
        meta.update({
            "fallback": "plain_keywords",
            "fallback_terms": query_terms,
            "negative_terms": list(fallback_negative_terms or [])[:24],
            "original_boolean": str(fallback_rule)[:1500],
            "notice": "JobAdder returned zero rows for the complete Boolean expression; positive plain-keyword fallback results are shown. Visible NOT matches are excluded locally.",
        })
    return marked, meta


def _spider_native_boolean_jobadder_candidates(token, rule_text, result_pool_limit=1200, location=""):
    """Attempt the recruiter's complete Boolean expression once in Keywords.

    JobAdder's public specification describes latest-resume keyword search but
    does not guarantee Boolean parsing for every tenant. Returned membership is
    still discovery proof and is never revalidated against partial detail JSON.
    """
    rule = re.sub(r"[\x00-\x1f\x7f]", " ", str(rule_text or ""))
    rule = re.sub(r"\s+", " ", rule).strip()[:1500]
    if not rule:
        return [], {"mode": "native_boolean", "query": "", "returned": 0, "search": {}}
    max_items = max(100, min(3000, int(result_pool_limit or 1200)))
    items, meta = _spider_jobadder_keyword_items(
        token,
        rule,
        max_items=max_items,
        page_size=min(1000, max_items),
        location=location,
    )
    marked = []
    for item in items:
        out = dict(item) if isinstance(item, dict) else {"value": item}
        out["_spiderNativeBooleanMatched"] = True
        out["_spiderBooleanRule"] = rule
        out["_spiderSearchTerms"] = [rule]
        marked.append(out)
    return marked, {
        "mode": "native_boolean",
        "query": rule,
        "returned": len(marked),
        "search": meta,
    }


@app.route("/jobadder/spider_search", methods=["POST"])
def jobadder_spider_search():
    """JobAdder candidate discovery followed by local eligibility and fit scoring."""
    data = request.get_json(silent=True) or {}
    if not _ai_crawler_lock_allowed(data):
        return _ai_crawler_locked_response()
    token = _ja_refresh_access_token(force=False)
    if not token:
        info = _ja_public_info()
        if info.get("needs_reconnect"):
            return jsonify({
                "error": "JobAdder connection expired. Reconnect JobAdder in Settings, then run AI Crawler again.",
                "needs_reconnect": True,
            }), 401
        return jsonify({"error": "Connect to JobAdder first"}), 401
    query = str(data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Missing query"}), 400
    limit = data.get("limit", 200)
    try:
        limit = max(1, min(250, int(limit)))
    except Exception:
        limit = 200
    filters = data.get("filters") if isinstance(data.get("filters"), dict) else {}
    safe_filters = {
        "role": str(filters.get("role") or "")[:120],
        "country": str(filters.get("country") or "")[:80],
        "city_state": str(filters.get("city_state") or filters.get("location") or "")[:120],
        "residential": str(filters.get("residential") or "Any")[:80],
        "industry": str(filters.get("industry") or "")[:300],
        "it_skills": str(filters.get("it_skills") or filters.get("skills") or "")[:1500],
        "qualifications": str(filters.get("qualifications") or "")[:1200],
        "salary": str(filters.get("salary") or "")[:120],
        "years": _spider_min_years_value(filters.get("years_min") if filters.get("years_min") is not None else (filters.get("years") or filters.get("experience_years"))),
        "years_min": _spider_min_years_value(filters.get("years_min") if filters.get("years_min") is not None else (filters.get("years") or filters.get("experience_years"))),
        "years_max": _spider_max_years_value(filters.get("years_max") if filters.get("years_max") is not None else filters.get("experience_years_max")),
        "must": str(filters.get("must") or filters.get("boolean") or "")[:1500],
        "nice": str(filters.get("nice") or "")[:1200],
        "exclude": str(filters.get("exclude") or "")[:1200],
        "targets": str(filters.get("targets") or "")[:1200],
        "jd": str(filters.get("jd") or filters.get("job_description") or "")[:12000],
        "strict": bool(filters.get("strict")),
        "adjacent": bool(filters.get("adjacent")),
    }
    query = re.sub(r"[\x00-\x1f\x7f]", " ", query)
    query = re.sub(r"\s+", " ", query).strip()[:260]
    # Prefer the recruiter-authored rule. Generated plain-keyword OR/AND queries are
    # used when the original field did not itself contain Boolean syntax.
    rule_for_search = safe_filters.get("must") or query
    if not re.search(r"\b(AND|OR|NOT)\b|[()\"']", rule_for_search, re.I) and re.search(r"\b(AND|OR|NOT)\b|[()\"']", query, re.I):
        rule_for_search = query
    use_native_boolean = re.search(r"\b(AND|OR|NOT)\b|[()\"']", rule_for_search, re.I) is not None

    try:
        boolean_meta = None
        parsed = {"items": []}
        search_mode = "native_boolean" if use_native_boolean else "plain"
        city_state = safe_filters.get("city_state") or ""
        if use_native_boolean:
            raw_items, boolean_meta = _spider_native_boolean_jobadder_candidates(
                token,
                rule_for_search,
                result_pool_limit=max(1200, limit * 5),
                location=city_state,
            )
            if not raw_items:
                fallback_terms, fallback_negative_terms = _spider_boolean_fallback_atoms(rule_for_search, 8, 24)
                fallback_query = " ".join(fallback_terms).strip()
                if fallback_query:
                    raw_items, fallback_meta = _spider_plain_keyword_jobadder_candidates(
                        token,
                        fallback_query,
                        result_pool_limit=max(1200, limit * 5),
                        location=city_state,
                        fallback_rule=rule_for_search,
                        fallback_negative_terms=fallback_negative_terms,
                    )
                    fallback_meta["native_attempt"] = boolean_meta
                    boolean_meta = fallback_meta
                    search_mode = "native_boolean→fallback"
                else:
                    boolean_meta = {
                        "mode": "native_boolean_no_fallback",
                        "query": rule_for_search,
                        "returned": 0,
                        "fallback": "skipped_negative_only",
                        "negative_terms": fallback_negative_terms,
                        "notice": "The Boolean expression contains no positive discovery terms. CV Studio did not invert a negative-only query into a plain-keyword search.",
                        "native_attempt": boolean_meta,
                        "search": (boolean_meta or {}).get("search", {}) if isinstance(boolean_meta, dict) else {},
                    }
                    search_mode = "native_boolean_no_fallback"
            parsed = {"items": raw_items}
        else:
            raw_items, boolean_meta = _spider_plain_keyword_jobadder_candidates(
                token,
                query,
                result_pool_limit=max(1200, limit * 5),
                location=city_state,
            )
            search_mode = "plain"
            parsed = {"items": raw_items}

        if not isinstance(raw_items, list):
            raw_items = []

        excluded_count = 0
        detail_excluded_count = 0
        unknown_count = 0
        enriched_count = 0
        resume_scored_count = 0

        def build_scored_candidate(base_item, discovery_index, detail=None, resume_text="", resume_source=""):
            scoring_item = base_item
            if isinstance(detail, dict) and detail:
                scoring_item = _spider_merge_candidate_summary_and_detail(base_item, detail)
                if isinstance(base_item, dict) and base_item.get("_spiderSearchTerms"):
                    scoring_item["_spiderSearchTerms"] = list(base_item.get("_spiderSearchTerms") or [])
            scoring_for_fit = scoring_item
            if resume_text:
                scoring_for_fit = dict(scoring_item) if isinstance(scoring_item, dict) else {"value": scoring_item}
                scoring_for_fit["resumeText"] = str(resume_text)[:30000]
            scoring_result = _spider_item_score(
                scoring_for_fit, safe_filters, enriched=bool(detail or resume_text)
            )
            fit_breakdown = {}
            if len(scoring_result) >= 8:
                keep, score, matched, unknown, excluded, hard_passed, discovery_evidence, fit_breakdown = scoring_result[:8]
            elif len(scoring_result) >= 7:
                keep, score, matched, unknown, excluded, hard_passed, discovery_evidence = scoring_result[:7]
            elif len(scoring_result) == 6:
                keep, score, matched, unknown, excluded, hard_passed = scoring_result
                discovery_evidence = []
            else:
                keep, score, matched, unknown, excluded = scoring_result
                hard_passed = []
                discovery_evidence = []
            if not keep:
                return None
            if isinstance(scoring_item, dict):
                out_item = dict(scoring_item)
                out_item.pop("_spiderDetail", None)
            else:
                out_item = {"value": scoring_item}
            out_item["_spiderScore"] = score
            out_item["_spiderFitPercent"] = score
            out_item["_spiderFitBreakdown"] = fit_breakdown if isinstance(fit_breakdown, dict) else {}
            out_item["_spiderMatchedFilters"] = matched
            out_item["_spiderDiscoveryEvidence"] = discovery_evidence
            out_item["_spiderHardFilters"] = hard_passed
            out_item["_spiderUnknownFilters"] = unknown
            out_item["_spiderEnriched"] = bool(detail)
            out_item["_spiderResumeScored"] = bool(resume_text)
            out_item["_spiderFitSource"] = (resume_source or "latest resume") if resume_text else ("JobAdder profile fields" if detail else "JobAdder list result")
            out_item["_spiderCardFields"] = _spider_card_fields(out_item)
            out_item["_spiderDiscoveryIndex"] = int(discovery_index)
            return out_item

        # Stage 1: score every discovered list row before any per-candidate API
        # calls.  This removes the old hard-coded first-120/first-200 position
        # bias and applies explicit local exclusions across the entire pool.
        preliminary = []
        base_by_id = {}
        for discovery_index, item in enumerate(raw_items):
            scored = build_scored_candidate(item, discovery_index)
            if scored is None:
                excluded_count += 1
                continue
            cid = _spider_candidate_id(item)
            preliminary.append({
                "base": item,
                "id": cid,
                "index": discovery_index,
                "score": int(scored.get("_spiderScore") or 0),
                "scored": scored,
            })
            if cid and cid not in base_by_id:
                base_by_id[cid] = item

        # Stage 2: latest-resume scoring uses a strict request budget selected
        # from the whole discovery pool, not merely the first JobAdder page. Half
        # is merit-ranked; half is evenly spread from first to last so later rows
        # can compete when list data is sparse or tied.
        resume_budget = min(len(preliminary), max(20, min(40, max(20, limit // 5))))
        ranked_preliminary = sorted(
            preliminary,
            key=lambda row: (-int(row.get("score") or 0), int(row.get("index") or 0)),
        )
        resume_ids = []
        resume_seen = set()

        def add_resume_row(row):
            cid = str((row or {}).get("id") or "").strip()
            if cid and cid not in resume_seen and len(resume_ids) < resume_budget:
                resume_seen.add(cid)
                resume_ids.append(cid)

        merit_count = min(len(ranked_preliminary), max(1, resume_budget // 2))
        for row in ranked_preliminary[:merit_count]:
            add_resume_row(row)
        spread_count = max(0, resume_budget - len(resume_ids))
        if spread_count and preliminary:
            if spread_count == 1:
                positions = [len(preliminary) - 1]
            else:
                positions = [round(i * (len(preliminary) - 1) / float(spread_count - 1)) for i in range(spread_count)]
            for position in positions:
                add_resume_row(preliminary[max(0, min(len(preliminary) - 1, int(position)))])
        if len(resume_ids) < resume_budget:
            for row in ranked_preliminary:
                add_resume_row(row)

        resume_map = {}
        if resume_ids:
            try:
                from concurrent.futures import ThreadPoolExecutor, as_completed
                with ThreadPoolExecutor(max_workers=min(5, len(resume_ids))) as pool:
                    future_map = {
                        pool.submit(_spider_fetch_candidate_resume_text, token, cid): cid
                        for cid in resume_ids
                    }
                    for future in as_completed(future_map):
                        cid = future_map[future]
                        try:
                            text, source = future.result()
                            resume_map[cid] = {"text": str(text or ""), "source": str(source or "")}
                        except _SpiderJobAdderReconnectRequired:
                            raise
                        except AntiwordDependencyError:
                            # One candidate's legacy .doc that verified Antiword
                            # cannot decode must not abort the whole sourcing run.
                            # Skip that resume and keep scoring them on profile data.
                            resume_map[cid] = {"text": "", "source": "legacy .doc could not be decoded (convert to DOCX/PDF)"}
                        except Exception:
                            resume_map[cid] = {"text": "", "source": "resume text unavailable"}
            except _SpiderJobAdderReconnectRequired:
                raise
            except Exception:
                for cid in resume_ids:
                    try:
                        text, source = _spider_fetch_candidate_resume_text(token, cid)
                        resume_map[cid] = {"text": str(text or ""), "source": str(source or "")}
                    except _SpiderJobAdderReconnectRequired:
                        raise
                    except AntiwordDependencyError:
                        resume_map[cid] = {"text": "", "source": "legacy .doc could not be decoded (convert to DOCX/PDF)"}
                    except Exception:
                        resume_map[cid] = {"text": "", "source": "resume text unavailable"}

        stage_two = []
        for row in preliminary:
            resume = resume_map.get(row.get("id")) or {}
            rescored = build_scored_candidate(
                row.get("base"), row.get("index"),
                resume_text=resume.get("text") or "",
                resume_source=resume.get("source") or "",
            )
            if rescored is not None:
                stage_two.append({
                    "base": row.get("base"),
                    "id": row.get("id"),
                    "index": row.get("index"),
                    "resume_text": resume.get("text") or "",
                    "resume_source": resume.get("source") or "",
                    "score": int(rescored.get("_spiderScore") or 0),
                    "scored": rescored,
                })

        stage_two.sort(key=lambda row: (-int(row.get("score") or 0), int(row.get("index") or 0)))
        total_kept_before_limit = len(stage_two)

        # Stage 3: fetch full candidate details for the competitive return
        # pool.  With no upper experience cap, a modest ranking buffer is enough.
        # With a maximum selected, scan deeper in bounded batches until the
        # requested shortlist plus a ranking buffer is eligible, or the pool is
        # exhausted.  This prevents senior candidates near the top from hiding
        # eligible junior candidates below a fixed 12-row margin.
        _min_years, _max_years = _spider_years_bounds(safe_filters)
        experience_upper_active = _max_years is not None
        ranking_buffer = min(24, max(0, len(stage_two) - limit))
        desired_eligible = min(len(stage_two), limit + ranking_buffer)
        fixed_detail_ceiling = min(len(stage_two), desired_eligible)
        detail_scan_ceiling = len(stage_two) if experience_upper_active else fixed_detail_ceiling
        detail_batch_size = 40 if experience_upper_active else max(1, fixed_detail_ceiling)
        detail_map = {}
        filtered = []
        detail_attempted = 0
        detail_batches = 0
        cursor = 0

        while cursor < detail_scan_ceiling:
            batch_end = min(detail_scan_ceiling, cursor + detail_batch_size)
            detail_targets = stage_two[cursor:batch_end]
            cursor = batch_end
            detail_batches += 1
            detail_attempted += len(detail_targets)

            detail_ids = []
            for row in detail_targets:
                cid = str(row.get("id") or "").strip()
                if cid and cid not in detail_ids and cid not in detail_map:
                    detail_ids.append(cid)
            if detail_ids:
                try:
                    from concurrent.futures import ThreadPoolExecutor, as_completed
                    with ThreadPoolExecutor(max_workers=min(6, len(detail_ids))) as pool:
                        future_map = {pool.submit(_spider_fetch_candidate_detail, token, cid): cid for cid in detail_ids}
                        for future in as_completed(future_map):
                            cid = future_map[future]
                            try:
                                detail = future.result()
                                detail_map[cid] = detail if isinstance(detail, dict) else None
                            except _SpiderJobAdderReconnectRequired:
                                raise
                            except Exception:
                                detail_map[cid] = None
                except _SpiderJobAdderReconnectRequired:
                    raise
                except Exception:
                    for cid in detail_ids:
                        try:
                            detail_map[cid] = _spider_fetch_candidate_detail(token, cid)
                        except _SpiderJobAdderReconnectRequired:
                            raise
                        except Exception:
                            detail_map[cid] = None

            for row in detail_targets:
                detail = detail_map.get(row.get("id"))
                final_item = build_scored_candidate(
                    row.get("base"), row.get("index"),
                    detail=detail,
                    resume_text=row.get("resume_text") or "",
                    resume_source=row.get("resume_source") or "",
                )
                if final_item is None:
                    detail_excluded_count += 1
                    continue
                if detail:
                    enriched_count += 1
                if row.get("resume_text"):
                    resume_scored_count += 1
                if final_item.get("_spiderUnknownFilters"):
                    unknown_count += 1
                filtered.append(final_item)

            if len(filtered) >= desired_eligible:
                break
            if not experience_upper_active:
                break

        detail_budget = detail_attempted
        experience_backfill_exhausted = bool(
            experience_upper_active and len(filtered) < min(limit, len(stage_two)) and cursor >= len(stage_two)
        )

        filtered.sort(key=lambda item: (-int(item.get("_spiderScore") or 0), int(item.get("_spiderDiscoveryIndex") or 0)))
        filtered = filtered[:limit]
        for item in filtered:
            item.pop("_spiderDiscoveryIndex", None)
        excluded_count += detail_excluded_count
        search_meta = boolean_meta.get("search", {}) if isinstance(boolean_meta, dict) else {}
        try:
            reported_total = int(search_meta.get("reported_total"))
        except Exception:
            reported_total = len(raw_items)
        summary = {
            "query": query,
            "reported_total": reported_total,
            "scanned": len(raw_items),
            "enriched_count": enriched_count,
            "excluded_count": excluded_count,
            "kept": len(filtered),
            "eligible_before_limit": total_kept_before_limit,
            "mode": search_mode,
            # Backward-compatible fields retained for existing renderers/exports.
            "returned_before_filter": len(raw_items),
            "returned_after_filter": len(filtered),
            "total_kept_before_limit": total_kept_before_limit,
            "result_limit": limit,
            "excluded": excluded_count,
            "kept_with_unknown_fields": unknown_count,
            "enriched_candidate_profiles": enriched_count,
            "resume_scored_candidates": resume_scored_count,
            "resume_scoring_budget": resume_budget,
            "detail_enrichment_budget": detail_budget,
            "detail_enrichment_batches": detail_batches,
            "experience_upper_bound_active": experience_upper_active,
            "experience_range_backfill_exhausted": experience_backfill_exhausted,
            "detail_only_excluded": detail_excluded_count,
            "ranking_pool": len(preliminary),
            "estimated_candidate_api_request_ceiling": (resume_budget * 2) + detail_budget,
            "filters_applied": safe_filters,
            "minimum_fit_percent": 10,
            "fit_scoring": "JD + recruiter Boolean/must-haves + role + nice-to-haves; latest resume text where available",
            "workflow": "JobAdder latest-resume keyword discovery, then local eligibility and 0-100 job-scope fit",
            "note": "Country is filtered locally and is never sent as JobAdder Location. The complete recruiter Boolean is attempted once; when it returns zero rows, CV Studio labels and runs a positive-term plain-keyword fallback. NOT operands are never used for discovery and visible NOT matches are excluded conservatively. Fallback membership is not treated as proof that the complete Boolean matched. Resume scoring is request-bounded (one latest attachment per selected candidate) and selected across the full discovery pool using merit plus first-to-last spread, rather than the old first-page cutoff. Full detail is then loaded for the competitive return pool so salary and notice fields can populate the cards; when a maximum experience bound is active, CV Studio scans deeper in batches until the requested eligible shortlist is filled or the discovery pool is exhausted. Missing hard-filter fields on non-enriched candidates remain visible as unknown unless Strict is enabled; explicit mismatches still exclude.",
            "pagination_incomplete": bool(search_meta.get("incomplete")),
            "pagination_warnings": list(search_meta.get("warnings") or [])[:5],
            "pages": search_meta.get("pages"),
        }

        if boolean_meta:
            summary["boolean_search"] = boolean_meta
        response_obj = parsed if isinstance(parsed, dict) else {}
        response_obj["query"] = query
        response_obj["items"] = filtered
        response_obj["filter_summary"] = summary
        return jsonify(response_obj)
    except _SpiderJobAdderReconnectRequired:
        return jsonify({
            "error": "JobAdder connection expired. Reconnect JobAdder in Settings, then run AI Crawler again.",
            "needs_reconnect": True,
            "query": query,
        }), 401
    except AntiwordDependencyError as e:
        return _antiword_dependency_response(e)
    except urllib.error.HTTPError as e:
        body = e.read().decode(errors="replace")
        if e.code == 401:
            _spider_mark_jobadder_reconnect_required()
            return jsonify({
                "error": "JobAdder connection expired. Reconnect JobAdder in Settings, then run AI Crawler again.",
                "needs_reconnect": True,
                "query": query,
            }), 401
        return jsonify({"error": "JobAdder candidate search failed: {}".format(e.code), "detail": body[:600], "query": query}), e.code
    except Exception as e:
        return jsonify({"error": str(e), "query": query}), 500

@app.route("/jobadder/create_candidate", methods=["POST"])
@_ja_critical_write_route
def jobadder_create_candidate():
    return _CVSTUDIO_JOBADDER_WRITE_SERVICE.create_candidate()

@app.route("/jobadder/update_candidate", methods=["POST"])
@_ja_critical_write_route
def jobadder_update_candidate():
    return _CVSTUDIO_JOBADDER_WRITE_SERVICE.update_candidate()

def _safe_jobadder_attachment_filename(value, fallback):
    """Keep uploaded names inside one multipart header line."""
    name = os.path.basename(str(value or "").replace("\r", " ").replace("\n", " ")).strip()
    name = re.sub(r'["\\/:*?<>|]+', "_", name)
    name = re.sub(r"\s+", " ", name).strip(" .")
    if not name:
        name = fallback
    return name[:180]


@app.route("/jobadder/upload_original_cv", methods=["POST"])
@_ja_critical_write_route
def jobadder_upload_original_cv():
    """Upload original (unformatted) CV as Resume attachment.
    Endpoint: POST /candidates/{id}/attachments/Resume
    Field: fileData (binary)
    """
    token = _ja_refresh_access_token(force=False)
    if not token:
        return jsonify({"error": "Not authenticated"}), 401
    candidate_id = request.form.get("candidate_id", "")
    file = request.files.get("file")
    if not candidate_id or not file:
        return jsonify({"error": "Missing candidate_id or file"}), 400
    try:
        file_data = file.read()
        filename  = _safe_jobadder_attachment_filename(file.filename, "Original_CV.pdf")
        boundary  = "----CVStudioBoundary7MA4YWxkTrZu0gW"
        bnd       = boundary.encode()
        crlf      = b"\x0d\x0a"
        body  = b"--" + bnd + crlf
        body += b'Content-Disposition: form-data; name="fileData"; filename="' + filename.encode("utf-8", errors="replace") + b'"' + crlf
        body += b"Content-Type: application/octet-stream" + crlf + crlf
        body += file_data + crlf
        body += b"--" + bnd + b"--" + crlf
        response = _JOBADDER_CLIENT.request_raw(
            "candidates/{}/attachments/Resume".format(candidate_id),
            method="POST",
            body=body,
            token=token,
            headers={
                "Accept": "application/json",
                "Content-Type": "multipart/form-data; boundary=" + boundary
            },
            timeout=30,
            safe_to_retry=False,
            retries=0,
        )
        return jsonify({"ok": True, "response": response.body.decode() if response.body else ""})
    except urllib.error.HTTPError as e:
        err_body = e.read().decode()
        return jsonify({"error": "JobAdder error: {}".format(e.code), "detail": err_body}), e.code
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/jobadder/lists", methods=["GET"])
def jobadder_lists():
    return _CVSTUDIO_JOBADDER_READ_SERVICE.lists()

@app.route("/jobadder/get_candidate", methods=["GET"])
def jobadder_get_candidate():
    return _CVSTUDIO_JOBADDER_READ_SERVICE.get_candidate()

@app.route("/jobadder/debug_endpoints", methods=["GET"])
def jobadder_debug_endpoints():
    return _CVSTUDIO_JOBADDER_READ_SERVICE.debug_endpoints()

@app.route("/jobadder/upload_cv", methods=["POST"])
@_ja_critical_write_route
def jobadder_upload_cv():
    """Upload DOCX as FormattedResume attachment.
    Endpoint: POST /candidates/{id}/attachments/FormattedResume
    Field: fileData (binary)
    """
    token = _ja_refresh_access_token(force=False)
    if not token:
        return jsonify({"error": "Not authenticated"}), 401
    candidate_id = request.form.get("candidate_id", "")
    file = request.files.get("file")
    if not candidate_id or not file:
        return jsonify({"error": "Missing candidate_id or file"}), 400
    try:
        file_data = file.read()
        filename  = _safe_jobadder_attachment_filename(file.filename, "CV.docx")
        boundary  = "----CVStudioBoundary7MA4YWxkTrZu0gW"
        bnd       = boundary.encode()
        crlf      = b"\x0d\x0a"
        body  = b"--" + bnd + crlf
        body += b'Content-Disposition: form-data; name="fileData"; filename="' + filename.encode("utf-8", errors="replace") + b'"' + crlf
        body += b"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document" + crlf + crlf
        body += file_data + crlf
        body += b"--" + bnd + b"--" + crlf
        response = _JOBADDER_CLIENT.request_raw(
            "candidates/{}/attachments/FormattedResume".format(candidate_id),
            method="POST",
            body=body,
            token=token,
            headers={
                "Accept": "application/json",
                "Content-Type": "multipart/form-data; boundary=" + boundary
            },
            timeout=30,
            safe_to_retry=False,
            retries=0,
        )
        return jsonify({"ok": True, "response": response.body.decode() if response.body else ""})
    except urllib.error.HTTPError as e:
        err_body = e.read().decode()
        return jsonify({"error": "JobAdder error: {}".format(e.code), "detail": err_body}), e.code
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _is_authorized_local_restart_request():
    """Accept only the CV Studio page's explicit loopback restart POST.

    A custom header blocks ordinary cross-site forms/images. Origin/Referer and
    the remote address add a second local-app boundary without enabling CORS.
    """
    if request.headers.get("X-CV-Studio-Restart") != "1":
        return False
    if str(request.remote_addr or "") not in {"127.0.0.1", "::1"}:
        return False
    local_origin_re = re.compile(r"^https?://(?:localhost|127\.0\.0\.1|\[::1\])(?::\d+)?$", re.I)
    origin = str(request.headers.get("Origin") or "").strip().rstrip("/")
    if origin and not local_origin_re.match(origin):
        return False
    referer = str(request.headers.get("Referer") or "").strip()
    if referer:
        parsed = urllib.parse.urlsplit(referer)
        ref_origin = "{}://{}".format(parsed.scheme, parsed.netloc).rstrip("/")
        if not local_origin_re.match(ref_origin):
            return False
    fetch_site = str(request.headers.get("Sec-Fetch-Site") or "").lower()
    if fetch_site and fetch_site not in {"same-origin", "same-site", "none"}:
        return False
    return True


@app.route("/restart", methods=["POST"])
def restart():
    """Gracefully restart only after an explicit same-origin CV Studio action."""
    if not _is_authorized_local_restart_request():
        return jsonify({"error": "Restart request was not authorised"}), 403

    def _do_restart():
        time.sleep(0.4)
        # Source mode needs: python app.py [args]. A compiled/native build has
        # sys.argv[0] == sys.executable and must not pass the executable path
        # twice. Keep both launch modes deterministic.
        try:
            same_entry = os.path.abspath(str(sys.argv[0] or "")) == os.path.abspath(str(sys.executable or ""))
        except Exception:
            same_entry = False
        restart_argv = [sys.executable] + (sys.argv[1:] if same_entry else sys.argv)

        if os.name == "nt":
            # os.execv does not replace the current process atomically on
            # Windows. Spawn a detached replacement with a short startup delay,
            # then terminate this process so port 5000 is released first.
            env = os.environ.copy()
            env["CVSTUDIO_RESTART_DELAY_SECONDS"] = "1.2"
            flags = int(getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0))
            flags |= int(getattr(subprocess, "CREATE_NO_WINDOW", 0))
            subprocess.Popen(
                restart_argv,
                cwd=os.path.dirname(os.path.abspath(sys.argv[0] or __file__)),
                env=env,
                creationflags=flags,
                close_fds=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            os._exit(0)

        os.execv(sys.executable, restart_argv)

    threading.Thread(target=_do_restart, daemon=True).start()
    return "", 204

@app.route("/heartbeat", methods=["POST"])
def heartbeat():
    global _last_heartbeat, _last_activity, _shutdown_armed
    _last_heartbeat = time.time()
    _last_activity  = time.time()
    _shutdown_armed = True
    return "", 204


@app.route("/")
def index():
    from flask import make_response
    resp = make_response(send_from_directory(os.path.dirname(os.path.abspath(__file__)), "index.html"))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    resp.set_cookie(
        _AI_SPEND_SESSION_COOKIE,
        _AI_SPEND_SESSION_TOKEN,
        httponly=True,
        samesite="Strict",
        secure=False,  # CV Studio is intentionally loopback HTTP only.
        path="/",
    )
    return resp


_CVSTUDIO_WEB_ASSETS_SERVICE = WebAssetsService(
    jsonify=lambda payload: jsonify(payload),
    send_from_directory=lambda *args, **kwargs: send_from_directory(*args, **kwargs),
    source_dir=lambda: os.path.dirname(os.path.abspath(__file__)),
)


@app.route("/vendor/<path:filename>")
def vendor_asset(filename):
    return _CVSTUDIO_WEB_ASSETS_SERVICE.vendor_asset(filename)


def _local_tool_cors_origin():
    """Return a safe CORS origin for local OCR helper routes.

    v24.6.116: avoid wildcard CORS on the local OCR endpoint. Normal websites
    should not be able to silently POST files to localhost OCR and read the
    extracted text. Allow browser-extension origins and local app origins only.
    """
    origin = (request.headers.get("Origin") or "").strip()
    if not origin:
        return ""
    if re.match(r"^chrome-extension://[A-Za-z0-9_-]+$", origin):
        return origin
    if re.match(r"^moz-extension://[A-Za-z0-9_-]+$", origin):
        return origin
    if re.match(r"^https?://(localhost|127\.0\.0\.1|\[::1\])(?::\d+)?$", origin, re.I):
        return origin
    return ""


def _apply_local_tool_cors(resp, methods="POST, OPTIONS"):
    allowed_origin = _local_tool_cors_origin()
    if allowed_origin:
        resp.headers["Access-Control-Allow-Origin"] = allowed_origin
        resp.headers["Access-Control-Allow-Methods"] = methods
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        resp.headers["Vary"] = "Origin"
    return resp


@app.route("/ocr/health", methods=["GET", "OPTIONS"])
def ocr_health():
    """Health check for the local OCR endpoint used by the JobAdder extension."""
    if request.method == "OPTIONS":
        resp = app.make_response("")
        _apply_local_tool_cors(resp, "GET, OPTIONS")
        return resp, 204

    import shutil
    script_dir = os.path.dirname(os.path.abspath(__file__))

    def first_existing_file(paths):
        for item in paths:
            if item and os.path.exists(item):
                return item
        return None

    def first_poppler_bin():
        candidates = [
            os.path.join(script_dir, "poppler", "Library", "bin"),
            os.path.join(script_dir, "poppler", "bin"),
            os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "poppler", "Library", "bin"),
            os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "poppler", "bin"),
            os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "poppler", "Library", "bin"),
            os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "poppler", "bin"),
            os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "poppler", "Library", "bin"),
            os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "poppler", "bin"),
            "/usr/local/bin",
            "/opt/homebrew/bin",
            "/usr/bin",
        ]
        for folder in candidates:
            if not folder or not os.path.isdir(folder):
                continue
            if os.path.exists(os.path.join(folder, "pdfinfo.exe")) or os.path.exists(os.path.join(folder, "pdfinfo")):
                return folder
        path_hit = shutil.which("pdfinfo")
        return os.path.dirname(path_hit) if path_hit else None

    tesseract_dependency = _mandatory_tesseract_health(_CVSTUDIO_ROOT)
    tess = tesseract_dependency.get("executable") or None

    poppler = first_poppler_bin()
    try:
        import pypdfium2  # noqa: F401
        pdfium_ok = True
    except Exception:
        pdfium_ok = False

    renderer_ok = bool(pdfium_ok or poppler)
    resp = jsonify({
        "ok": bool(tess and renderer_ok),
        "service": "CV Studio OCR Endpoint",
        "endpoint": "/ocr",
        "tesseract": tess or "not found",
        "tesseract_health": tesseract_dependency,
        "pdf_renderer": "built-in PDFium" if pdfium_ok else ("Poppler fallback" if poppler else "not found"),
        "pdfium_bundled": pdfium_ok,
        "poppler_bin": poppler or "optional / not found",
        "hint": None if (tess and renderer_ok) else (
            "Run INSTALL.bat again, then restart CV Studio. Tesseract and the built-in PDFium renderer are required; Poppler is only an optional fallback."
        )
    })
    _apply_local_tool_cors(resp, "GET, OPTIONS")
    return resp


@app.route("/ocr", methods=["POST", "OPTIONS"])
def ocr_endpoint():
    """OCR endpoint for Chrome extension — extracts PDF/image/DOCX text using local Tesseract and the bundled PDFium renderer."""
    if request.method == "OPTIONS":
        resp = app.make_response("")
        _apply_local_tool_cors(resp, "POST, OPTIONS")
        return resp, 204

    try:
        import base64
        import io
        import shutil

        script_dir = os.path.dirname(os.path.abspath(__file__))

        def cors_json(payload, status=200):
            resp = jsonify(payload)
            _apply_local_tool_cors(resp, "POST, OPTIONS")
            return resp, status

        def first_existing_file(paths):
            for item in paths:
                if item and os.path.exists(item):
                    return item
            return None

        def first_poppler_bin():
            candidates = [
                os.path.join(script_dir, "poppler", "Library", "bin"),
                os.path.join(script_dir, "poppler", "bin"),
                os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "poppler", "Library", "bin"),
                os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), "poppler", "bin"),
                os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "poppler", "Library", "bin"),
                os.path.join(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"), "poppler", "bin"),
                os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "poppler", "Library", "bin"),
                os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "poppler", "bin"),
                "/usr/local/bin",
                "/opt/homebrew/bin",
                "/usr/bin",
            ]
            for folder in candidates:
                if not folder or not os.path.isdir(folder):
                    continue
                if os.path.exists(os.path.join(folder, "pdfinfo.exe")) or os.path.exists(os.path.join(folder, "pdfinfo")):
                    return folder
            path_hit = shutil.which("pdfinfo")
            return os.path.dirname(path_hit) if path_hit else None

        # Accept both multipart/form-data and the older JSON/base64 style.
        uploaded = request.files.get("file")
        file_bytes = b""
        fname = "upload"
        if uploaded:
            file_bytes = uploaded.read()
            fname = uploaded.filename or "upload"
        else:
            payload = request.get_json(silent=True) or {}
            b64 = payload.get("base64") or payload.get("fileBase64") or payload.get("data")
            if b64:
                if not isinstance(b64, str):
                    return cors_json({"ok": False, "error": "Invalid base64 upload"}, 400)
                if len(b64) > 56 * 1024 * 1024:
                    return cors_json({"ok": False, "error": "OCR upload is too large", "limit_mb": 40}, 413)
                if "," in b64 and b64.strip().lower().startswith("data:"):
                    b64 = b64.split(",", 1)[1]
                try:
                    file_bytes = base64.b64decode(b64, validate=True)
                except Exception:
                    return cors_json({"ok": False, "error": "Invalid base64 upload"}, 400)
                fname = payload.get("filename") or payload.get("name") or "upload"
            else:
                return cors_json({"ok": False, "error": "No file provided. Send multipart form-data field named 'file'."}, 400)

        if not file_bytes:
            return cors_json({"ok": False, "error": "Uploaded file is empty"}, 400)

        fname_lower = (fname or "").lower()
        text = ""

        # Strong byte signatures take precedence over stale or misleading
        # upload names so genuine OLE .doc content always reaches Antiword.
        content_kind = _document_content_kind(file_bytes)
        is_pdf = content_kind == "pdf" or (
            not content_kind and fname_lower.endswith(".pdf")
        )
        is_docx = content_kind == "zip" or (
            not content_kind and fname_lower.endswith(".docx")
        )
        is_doc = content_kind == "legacy_doc" or (
            not content_kind and fname_lower.endswith(".doc")
        )
        is_image = content_kind == "image" or (
            not content_kind
            and fname_lower.endswith(
                (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp")
            )
        )

        pytesseract = None
        tess_ok = False
        if is_pdf or is_image:
            try:
                import pytesseract
                tess_path = _find_mandatory_tesseract(_CVSTUDIO_ROOT)
                if tess_path:
                    pytesseract.pytesseract.tesseract_cmd = tess_path
                tess_ok = bool(tess_path)
            except ImportError:
                return cors_json({"ok": False, "error": "Python package pytesseract is not installed. Run INSTALL.bat again."}, 500)

        if is_pdf:
            if not tess_ok:
                return cors_json({"ok": False, "error": "Tesseract not found. Run INSTALL.bat again or install Tesseract OCR."}, 500)
            poppler_path = first_poppler_bin()
            try:
                # PDFium is bundled in protected builds. Poppler remains an
                # optional fallback for source environments and older machines.
                if poppler_path:
                    os.environ["PATH"] = poppler_path + os.pathsep + os.environ.get("PATH", "")
                text = _ocr_pdf_pagewise(file_bytes, pytesseract, poppler_path=poppler_path, dpi=220)
            except Exception as e:
                return cors_json({
                    "ok": False,
                    "error": "PDF OCR failed: " + str(e),
                    "renderer": "built-in PDFium with optional Poppler fallback",
                    "hint": "Run INSTALL.bat again if the built-in renderer or Tesseract files are missing."
                }, 500)

        elif is_image:
            if not tess_ok:
                return cors_json({"ok": False, "error": "Tesseract not found. Run INSTALL.bat again or install Tesseract OCR."}, 500)
            img = _safe_image_open(file_bytes)
            try:
                text = _ocr_image_text(pytesseract, img, lang="eng", timeout=35)
            finally:
                img.close()

        elif is_doc:
            try:
                text = _spider_extract_legacy_doc_text_for_preview(file_bytes)
            except AntiwordDependencyError as error:
                return cors_json(_antiword_dependency_payload(error), 424)

        elif is_docx:
            try:
                _validate_zip_payload(file_bytes, "DOCX")
                text = _extract_docx_text_preserve_tables(file_bytes)
            except Exception as e:
                return cors_json({"ok": False, "error": "DOCX text extraction failed: " + str(e)}, 500)

        else:
            # Plain text fallback.
            try:
                text = file_bytes.decode("utf-8", errors="ignore")
            except Exception:
                text = ""
            if not text.strip():
                return cors_json({"ok": False, "error": "Unsupported file type: " + fname_lower}, 400)

        return cors_json({
            "ok": True,
            "text": text.strip(),
            "source": "Local OCR",
            "chars": len(text.strip()),
            "filename": fname,
        })

    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/cv_studio_logo.png")
def app_logo_png():
    return _CVSTUDIO_WEB_ASSETS_SERVICE.app_logo_png()


@app.route("/cv_studio.ico")
def app_icon_ico():
    return _CVSTUDIO_WEB_ASSETS_SERVICE.app_icon_ico()


@app.route("/favicon.ico")
def favicon():
    return _CVSTUDIO_WEB_ASSETS_SERVICE.favicon()


@app.route("/test", methods=["POST"])
def test_key():
    try:
        body = request.get_json(force=True, silent=True) or {}
        api_key = _resolve_request_api_key(body, "main")
        provider = (body.get("provider") or "anthropic").strip().lower()
        if not api_key:
            return jsonify({"ok": False, "error": "No API key provided"})
        if provider in ("", "anthropic", "claude") and not api_key.startswith("sk-ant-"):
            return jsonify({"ok": False, "error": f"Key should start with sk-ant-, got: {api_key[:12]}..."})
        default_model = {"deepseek": "deepseek-v4-flash", "openai": "gpt-5.5", "gpt": "gpt-5.5"}.get(provider, "claude-haiku-4-5-20251001")
        # Test the model the user has actually saved, not a hardcoded one, so a
        # stale/invalid model (e.g. a retired DeepSeek name) is caught here
        # rather than silently failing every real feature call.
        test_model = str(body.get("model") or "").strip() or default_model
        if provider == "deepseek":
            test_model = _migrate_deepseek_model(test_model)
        # A single probe that also exercises strict-JSON output — representative
        # of the structured tasks (/parse, Blind JD, etc.) without a second call.
        data = call_llm(provider, api_key, {
            "model": test_model,
            "max_tokens": 40,
            "messages": [{"role": "user", "content": 'Reply with only this JSON and nothing else: {"ok": true}'}]
        })
        text = ""
        for block in ((data or {}).get("content") or []):
            if isinstance(block, dict) and block.get("type") == "text":
                text += block.get("text") or ""
        json_ok = False
        try:
            match = re.search(r"\{.*\}", text, re.DOTALL)
            if match:
                json_ok = "ok" in json.loads(match.group(0))
        except Exception:
            json_ok = False
        usage = (data or {}).get("usage") or {}
        out = {"ok": True, "json_ok": json_ok, "usage": usage, "model": test_model, "provider": provider}
        out.update(_llm_response_cost_fields(test_model, usage, provider))
        return jsonify(out)
    except urllib.error.HTTPError as e:
        raw = e.read()
        try:
            err = json.loads(raw)
            msg = err.get("error", {}).get("message", "Unknown")
        except Exception:
            msg = raw.decode("utf-8", errors="replace")[:500] or "Unknown"
        out = {"ok": False, "error": f"HTTP {e.code}: {msg}"}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("test_model", ""),
            locals().get("provider", "anthropic"),
            attempted=True,
        ))
        return jsonify(out)
    except Exception as e:
        out = {"ok": False, "error": str(e)}
        if "test_model" in locals():
            out.update(_llm_paid_failure_fields(
                e,
                test_model,
                provider,
                attempted=True,
            ))
        return jsonify(out)


@app.route("/parse", methods=["POST"])
def parse_cv():
    try:
        body = request.get_json(force=True, silent=True)
        if not body:
            return jsonify({"error": "Invalid JSON body"}), 400
        api_key = _resolve_request_api_key(body, "main")
        cv_text = (body.get("cv_text") or "").strip()
        model   = body.get("model") or "claude-sonnet-4-6"
        llm_provider = (body.get("provider") or "anthropic").strip().lower()
        usage_total = _merge_llm_usage()

        if not api_key:
            return jsonify({"error": "API key required"}), 400
        if not cv_text:
            return jsonify({"error": "CV text required"}), 400

        # Compress CV text to reduce token usage — collapse excessive whitespace
        import re as _re
        cv_text = _re.sub(r'\n{3,}', '\n\n', cv_text)   # max 2 blank lines
        cv_text = _re.sub(r'[ \t]{2,}', ' ', cv_text)   # collapse spaces
        cv_text = cv_text.strip()
        # Rewrite ISO YYYY-MM dates to "Mon YYYY" before the model sees them.
        # Providers mis-tag the year-first format (wrong month/year, or a real
        # end date read as "Present"); "Jun 2020" is read reliably. Doing it on
        # cv_text keeps the model prompt and the downstream table reconciliation
        # consistent.
        cv_text = _cv_pretranslate_iso_dates(cv_text)
        parse_timeout_seconds = _cv_parse_backend_timeout_seconds(cv_text)

        data = call_llm(llm_provider, api_key, {
            "model": model,
            "max_tokens": 64000,
            "_timeout_seconds": parse_timeout_seconds,
            "system": SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": f"Parse this raw CV into the JSON schema:\n\n{cv_text}"}]
        })
        usage_total = _merge_llm_usage(usage_total, data.get("usage"))

        raw_text = "".join(b.get("text", "") for b in data.get("content", []))
        raw_text = raw_text.strip()

        # Strip any markdown fences
        if raw_text.startswith("```json"):
            raw_text = raw_text[7:]  # Remove ```json
        if raw_text.startswith("```"):
            raw_text = raw_text[3:]  # Remove ```
        if raw_text.endswith("```"):
            raw_text = raw_text[:-3]  # Remove closing ```
        raw_text = raw_text.strip()

        # Completeness marker. A first-try clean parse, a clean Strategy-2
        # re-send, and a clean Strategy-3 stitched continuation are all lossless.
        # Only the bracket-closing salvage (try_close_json) truncates the JSON at
        # the last valid point and can silently drop trailing roles/bullets, so it
        # is flagged as degraded -- the CV is still fully repaired and returned,
        # but the caller is told completeness is not guaranteed.
        parse_degraded = False
        parse_degraded_reason = ""

        try:
            parsed = json.loads(raw_text)
        except json.JSONDecodeError:

            # ── Strategy 1: Smart bracket-aware scan (up to 2000 chars back) ────
            def try_close_json(s):
                """
                Scan back up to 2000 chars finding the last valid cut-point.
                Uses string-aware bracket counting — ignores { [ inside quoted strings.
                """
                try:
                    def count_brackets_outside_strings(text):
                        """Return (open_braces, open_brackets) ignoring content inside quotes."""
                        in_str = False
                        escape_next = False
                        brace_depth = 0
                        bracket_depth = 0
                        for ch in text:
                            if escape_next:
                                escape_next = False
                                continue
                            if ch == '\\' and in_str:
                                escape_next = True
                                continue
                            if ch == '"':
                                in_str = not in_str
                                continue
                            if in_str:
                                continue
                            if ch == '{':  brace_depth += 1
                            elif ch == '}': brace_depth -= 1
                            elif ch == '[': bracket_depth += 1
                            elif ch == ']': bracket_depth -= 1
                        return brace_depth, bracket_depth

                    scan_limit = max(len(s) - 2000, 0)
                    for cut in range(len(s), scan_limit, -1):
                        candidate = s[:cut].rstrip()
                        if not candidate:
                            continue
                        last_ch = candidate[-1]
                        # Only attempt cuts at sensible boundaries
                        if last_ch not in ('"', '}', ']', '0123456789', 'e', 'l', 'n'):
                            # strip trailing comma or partial value before trying
                            candidate = candidate.rstrip(',').rstrip()
                            if not candidate:
                                continue
                        opens, arr_opens = count_brackets_outside_strings(candidate)
                        if opens < 0 or arr_opens < 0:
                            continue
                        # Close any dangling open string
                        b, a = count_brackets_outside_strings(candidate)
                        # Detect if we are mid-string (odd number of unescaped quotes)
                        def in_open_string(text):
                            in_s = False
                            esc = False
                            for ch in text:
                                if esc: esc = False; continue
                                if ch == '\\' and in_s: esc = True; continue
                                if ch == '"': in_s = not in_s
                            return in_s
                        if in_open_string(candidate):
                            candidate = candidate.rstrip(',').rstrip() + '"'
                        b2, a2 = count_brackets_outside_strings(candidate)
                        closer = ']' * max(a2, 0) + '}' * max(b2, 0)
                        attempt = candidate + closer
                        try:
                            return json.loads(attempt)
                        except Exception:
                            continue
                except Exception:
                    pass
                return None

            repaired = try_close_json(raw_text)
            if repaired:
                parsed = repaired
                parse_degraded = True
                parse_degraded_reason = "truncated_response_bracket_salvage"
            else:
                # ── Strategy 2: Re-send the full CV asking for compact, COMPLETE JSON ──
                # A CV formatter must never silently shorten the candidate's content, so
                # this retry preserves every role/bullet/skill in full and only asks for
                # whitespace-free JSON to save output tokens. If it still truncates,
                # Strategy 3 continues losslessly from the first attempt.
                strategy2_parsed = None
                try:
                    brevity_prompt = (
                        "Your previous response was not valid JSON. Parse this CV into the "
                        "required JSON schema again and return ONLY complete, valid JSON — "
                        "no markdown fences, no explanation. Preserve EVERY role and EVERY "
                        "bullet point in full; do not shorten, summarise, drop, or cap "
                        "bullets, roles, or skills. Emit compact JSON without unnecessary "
                        "whitespace to stay within output limits.\n\n"
                        + cv_text
                    )
                    s2_data = call_llm(llm_provider, api_key, {
                        "model": model,
                        "max_tokens": 64000,
                        "_timeout_seconds": parse_timeout_seconds,
                        "system": SYSTEM_PROMPT,
                        "messages": [{"role": "user", "content": brevity_prompt}]
                    })
                    usage_total = _merge_llm_usage(usage_total, s2_data.get("usage"))
                    s2_raw = "".join(
                        b.get("text", "") for b in s2_data.get("content", [])
                    ).strip().lstrip("```json").lstrip("```").rstrip("```").strip()
                    strategy2_parsed = json.loads(s2_raw)
                except Exception:
                    # Try Strategy 1 on the strategy-2 output too
                    try:
                        strategy2_parsed = try_close_json(s2_raw)
                        if strategy2_parsed:
                            parse_degraded = True
                            parse_degraded_reason = "truncated_retry_bracket_salvage"
                    except Exception:
                        strategy2_parsed = None

                if strategy2_parsed:
                    parsed = strategy2_parsed
                else:
                    # ── Strategy 3: Ask Claude to continue from truncation point ──
                    try:
                        s3_data = call_llm(llm_provider, api_key, {
                            "model": model,
                            "max_tokens": 64000,
                            "_timeout_seconds": parse_timeout_seconds,
                            "messages": [
                                {"role": "user",  "content": f"Parse this CV into the JSON schema:\n\n{cv_text}"},
                                {"role": "assistant", "content": raw_text},
                                {"role": "user",  "content": (
                                    "The JSON was cut off mid-output. "
                                    "Continue EXACTLY from where it stopped — output only the "
                                    "remaining JSON fragment needed to complete the structure. "
                                    "No explanation, no markdown, just the continuation."
                                )}
                            ]
                        })
                        usage_total = _merge_llm_usage(usage_total, s3_data.get("usage"))
                        continuation = "".join(
                            b.get("text", "") for b in s3_data.get("content", [])
                        ).strip().rstrip("```").strip()
                        combined = raw_text + continuation
                        # Run Strategy 1 on the stitched result
                        try:
                            parsed = json.loads(combined)
                        except Exception:
                            parsed = try_close_json(combined)
                            if not parsed:
                                raise ValueError("All three repair strategies failed.")
                            parse_degraded = True
                            parse_degraded_reason = "truncated_continuation_bracket_salvage"
                    except Exception as repair_err:
                        out = {"error": (
                            f"Could not parse CV after 3 repair attempts. "
                            f"Error: {str(repair_err)}\n\n"
                            f"Tip: Try removing the cover page or non-essential sections and re-upload."
                        ), "paid_ai_failure": bool(_llm_usage_int(usage_total, "api_calls")), "usage": usage_total,
                           "model": model, "provider": llm_provider}
                        out.update(_llm_response_cost_fields(model, usage_total, llm_provider))
                        out.update(_llm_paid_failure_fields(
                            repair_err,
                            model,
                            llm_provider,
                            usage_total,
                            attempted=True,
                        ))
                        return jsonify(out), 500
        
        usage = usage_total
        # Email regex fallback — if Claude missed it, scan raw text
        import re as _re
        cand = parsed.get("candidate", {})
        if not cand.get("email") and cv_text:
            m = _re.search(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', cv_text)
            if m:
                cand["email"] = m.group(0)
                parsed["candidate"] = cand
        # Phone regex fallback
        if not cand.get("phone") and cv_text:
            mp = _re.search(r'[+\(]?[\d][\d\s\-\(\)]{7,}[\d]', cv_text)
            if mp:
                cand["phone"] = mp.group(0).strip()
                parsed["candidate"] = cand
        parsed = _correct_mistagged_candidate_name(parsed, cv_text)
        parsed = _reconcile_work_experience_with_authoritative_table(parsed, cv_text)
        parsed = _order_same_company_roles_newest_first(parsed)
        parsed = _restore_explicit_project_headings(parsed, cv_text)
        parsed = _collapse_incomplete_earlier_career(parsed)
        parsed = _clean_candidate_languages_from_redaction(parsed, cv_text)
        parsed = _normalize_candidate_languages(parsed, cv_text)
        # Outline-label nesting must be read from the raw bullet text, so infer it
        # BEFORE _normalize_cv_structured_content strips the labels, and return it
        # to the client to carry to /generate-docx.
        bullet_levels = _infer_label_bullet_levels(parsed)
        parsed = _normalize_cv_structured_content(parsed)
        parsed = _normalize_cv_data_for_output(parsed, cv_text)
        out = {"ok": True, "data": parsed, "usage": usage, "model": model, "provider": llm_provider, "bullet_levels": bullet_levels}
        out.update(_llm_response_cost_fields(model, usage, llm_provider))
        # ── Fidelity audit — observational, never mutates `parsed`. Attaches a
        # report always; only escalates to the degraded/warning fields when the
        # parse was not already flagged as truncated. ──
        fidelity = evaluate_cv_fidelity(parsed, cv_text)
        out["fidelity"] = fidelity
        if not fidelity.get("ok") and not parse_degraded:
            out["degraded"] = True
            out["degraded_reason"] = "fidelity_check"
            out["warning"] = summarize_fidelity_warning(fidelity)
        # The parse still succeeded and the full repaired CV is returned; this
        # only tells the caller the AI response was truncated and salvaged, so
        # some later roles or bullet points may be missing and should be checked.
        if parse_degraded:
            out["degraded"] = True
            out["degraded_reason"] = parse_degraded_reason
            out["warning"] = (
                "This CV was long and the AI response was cut off. It was "
                "automatically repaired and the profile can still be created, but "
                "some later roles or bullet points may be missing — please "
                "check the parsed result against the original CV before continuing."
            )
        return jsonify(out)

    except urllib.error.HTTPError as e:
        err_body = e.read()
        try:
            err = json.loads(err_body)
            msg = err.get("error", {}).get("message", "API error")
        except:
            msg = err_body.decode("utf-8", errors="replace")[:300]
        out = {"error": _provider_error_message(llm_provider, msg)}
        out.update(_llm_paid_failure_fields(
            e,
            model,
            llm_provider,
            locals().get("usage_total"),
            attempted=True,
        ))
        return jsonify(out), 400

    except json.JSONDecodeError as e:
        usage = locals().get("usage_total") or _merge_llm_usage()
        out = {"error": f"Could not parse AI response as JSON: {str(e)}", "paid_ai_failure": bool(_llm_usage_int(usage, "api_calls")),
               "usage": usage, "model": locals().get("model", ""), "provider": locals().get("llm_provider", "anthropic")}
        out.update(_llm_response_cost_fields(out["model"], usage, out["provider"]))
        out.update(_llm_paid_failure_fields(
            e,
            out["model"],
            out["provider"],
            usage,
            attempted=True,
        ))
        return jsonify(out), 500

    except Exception as e:
        traceback.print_exc()
        usage = locals().get("usage_total") or _merge_llm_usage()
        out = {"error": str(e), "paid_ai_failure": bool(_llm_usage_int(usage, "api_calls")),
               "usage": usage, "model": locals().get("model", ""), "provider": locals().get("llm_provider", "anthropic")}
        out.update(_llm_response_cost_fields(out["model"], usage, out["provider"]))
        out.update(_llm_paid_failure_fields(
            e,
            out["model"],
            out["provider"],
            usage,
            attempted=bool(locals().get("api_key") and locals().get("cv_text")),
        ))
        return jsonify(out), 500


@app.route("/generate-ai", methods=["POST"])
def generate_ai():
    """Generic Anthropic proxy for non-CV tools such as Blind JD and Company Profile.
    Keeps API calls server-side instead of exposing the key through browser direct access.
    """
    try:
        body = request.get_json(force=True, silent=True)
        if not body:
            return jsonify({"error": "Invalid JSON body"}), 400

        feature = str(body.get("feature") or "").strip().lower()
        if feature == "the_spider" and not _ai_crawler_lock_allowed(body):
            return _ai_crawler_locked_response()

        api_key = _resolve_request_api_key(body, "main")
        prompt = (body.get("prompt") or "").strip()
        model = body.get("model") or "claude-sonnet-4-6"
        llm_provider = (body.get("provider") or "anthropic").strip().lower()
        max_tokens = int(body.get("max_tokens") or 4000)
        use_tools = bool(body.get("use_tools"))
        max_uses = int(body.get("max_uses") or 3)

        if not api_key:
            return jsonify({"error": "API key required"}), 400
        if not prompt:
            return jsonify({"error": "Prompt required"}), 400

        payload = {
            "model": model,
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}]
        }
        if use_tools:
            payload["tools"] = [{"type": "web_search_20250305", "name": "web_search", "max_uses": max_uses}]

        try:
            if use_tools and not _LLM_PROVIDER_INFO.get(llm_provider, {}).get("supports_web_search", True):
                # DeepSeek cannot browse the web; behave like the "workspace doesn't
                # support tools" case below and continue without live search rather
                # than erroring the whole request.
                no_tools_payload = dict(payload)
                no_tools_payload.pop("tools", None)
                data = call_llm(llm_provider, api_key, no_tools_payload)
                usage = data.get("usage", {})
                out = {
                    "ok": True,
                    "content": data.get("content", []),
                    "usage": usage,
                    "model": model,
                    "provider": llm_provider,
                    "warning": "DeepSeek cannot browse the web itself; continued without live search. Results may be less current.",
                }
                out.update(_llm_response_cost_fields(model, usage, llm_provider))
                return jsonify(out)
            data = call_llm(llm_provider, api_key, payload)
        except urllib.error.HTTPError as e:
            err_body = e.read()
            try:
                err = json.loads(err_body)
                msg = err.get("error", {}).get("message", "API error")
            except Exception:
                msg = err_body.decode("utf-8", errors="replace")[:300]

            # Some accounts/workspaces may not have web_search enabled. Do not break the tab; retry without tools.
            if use_tools and re.search(r"web[_\s-]?search|tool|permission|not enabled|not authorized|unsupported", msg, re.I):
                payload.pop("tools", None)
                data = call_llm(llm_provider, api_key, payload)
                usage = data.get("usage", {})
                out = {
                    "ok": True,
                    "content": data.get("content", []),
                    "usage": usage,
                    "model": model,
                    "provider": llm_provider,
                    "warning": "Web search unavailable; continued without web search."
                }
                out.update(_llm_response_cost_fields(model, usage, llm_provider))
                return jsonify(out)
            out = {"error": _provider_error_message(llm_provider, msg)}
            out.update(_llm_paid_failure_fields(
                e,
                model,
                llm_provider,
                attempted=True,
            ))
            return jsonify(out), 400

        usage = data.get("usage", {})
        out = {"ok": True, "content": data.get("content", []), "usage": usage, "model": model, "provider": llm_provider}
        out.update(_llm_response_cost_fields(model, usage, llm_provider))
        return jsonify(out)

    except Exception as e:
        traceback.print_exc()
        out = {"error": str(e)}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            attempted=bool(
                locals().get("api_key")
                and locals().get("prompt")
                and "payload" in locals()
            ),
        ))
        return jsonify(out), 500


# ── Lead Finder / Market Leads ───────────────────────────────────────────────


_LLM_MODEL_PRICING = _PHASE5B_MODEL_PRICING


def _llm_cost_details(model, usage, provider=None):
    """Return one auditable token/cost breakdown for every AI feature."""
    return _phase5b_cost_details(model, usage, provider)


def _llm_response_cost_fields(model, usage, provider=None):
    details = _llm_cost_details(model, usage, provider)
    return {"cost": float(details.get("usd") or 0.0), "cost_details": details}


def _llm_paid_failure_fields(
    error,
    model,
    provider=None,
    usage=None,
    attempted=None,
):
    """Return additive paid-call ambiguity and billing status fields."""
    return _phase5b_paid_failure_reconciliation(
        error,
        provider=provider,
        model=model,
        usage=usage,
        attempted=attempted,
    )


def _lead_provider_from_model(model, provider=None):
    return _phase5b_normalize_provider(model, provider)


def _lead_pricing_for_model(model, provider=None):
    return _phase5b_pricing_for_model(model, provider)


def _lead_cost(model, usage, provider=None):
    return float(_llm_cost_details(model, usage, provider).get("usd") or 0.0)


def _lead_merge_usage(*usages):
    return _merge_llm_usage(*usages)


def _lead_cost_details(model, usage, provider=None):
    return _llm_cost_details(model, usage, provider)


from cvstudio_lead_cache import LeadCacheService as _LeadCacheService

from cvstudio_lead_match import (
    _lead_blob_has_alias,
    _lead_company_is_actually_role_text,
    _lead_company_matches_title_angles,
    _lead_company_names_match,
    _lead_has_apac,
    _lead_location_allowed,
    _lead_location_instruction,
    _lead_norm_company_name,
    _lead_region_aliases,
    _lead_regions_allow_apac_remote,
    _lead_text_blob,
)

from cvstudio_lead_enrich import (
    _LEAD_EMAIL_BLOCKED_DOMAIN_SUFFIXES,
    _LEAD_EMAIL_BLOCKED_PERSONAL_DOMAINS,
    _LEAD_EMAIL_RE,
    _LEAD_FAMILY_BOOST_TOKENS,
    _LEAD_FAMILY_DEFAULT_TITLE,
    _LEAD_FAMILY_PATTERNS,
    _LEAD_NON_JOB_CONTENT_PATH_MARKERS,
    _LEAD_SENIORITY_TOKENS,
    _LEAD_TERMINAL_ROLE_TOKENS,
    _lead_apply_basic_job_filters,
    _lead_best_selected_source_verification_url,
    _lead_best_verification_url,
    _lead_canonical_direct_job_url,
    _lead_clean_company_guess,
    _lead_clean_csv,
    _lead_clean_job_filters,
    _lead_cleanup_urls_by_selected_sources,
    _lead_contains_any_token,
    _lead_cv_content_signature,
    _lead_cv_evidence_tokens,
    _lead_email_domain,
    _lead_exclude_query_terms,
    _lead_extract_json,
    _lead_fake_anthropic_response,
    _lead_families_from_text,
    _lead_family_scores,
    _lead_filter_by_regions,
    _lead_filter_by_title_angles,
    _lead_filter_people_by_input_companies,
    _lead_filter_query_terms,
    _lead_freshness_label,
    _lead_google_jobs_query,
    _lead_guess_company_from_url,
    _lead_guess_portal_from_url,
    _lead_guess_role_company_from_search_result,
    _lead_guess_title_company_from_context,
    _lead_has_any,
    _lead_int_0_100,
    _lead_is_company_domain_email,
    _lead_is_direct_job_url,
    _lead_is_generated_verification_search,
    _lead_is_generic_portal_category_url,
    _lead_is_non_job_content_url,
    _lead_job_filter_instruction,
    _lead_job_title_angles,
    _lead_keep_only_direct_job_leads,
    _lead_make_search_fallback_leads,
    _lead_normalize_company_fit_freshness,
    _lead_normalize_company_urls,
    _lead_normalize_linkedin_url,
    _lead_parse_posted_date,
    _lead_portal_domain_hint,
    _lead_portal_search_url,
    _lead_primary_role_families,
    _lead_provider_results_to_leads,
    _lead_recover_job_leads_from_text,
    _lead_resolve_search_target_role,
    _lead_reviewable_job_lead,
    _lead_role_specific_title_bank,
    _lead_sanitize_public_business_emails,
    _lead_search_provider_queries,
    _lead_selected_source_site,
    _lead_source_allowed_by_selection,
    _lead_url_portal,
    _lead_verification_company_text,
)


































































def _lead_call_with_optional_web(api_key, payload, warning_prefix="", graceful_json=None, provider="anthropic"):
    return _LEAD_SEARCH_ORCHESTRATOR.call_with_optional_web(
        api_key, payload, warning_prefix=warning_prefix, graceful_json=graceful_json, provider=provider
    )
























def _lead_search_provider_config(raw):
    return _LEAD_SEARCH_ORCHESTRATOR.search_provider_config(raw)






# _lead_ssl_context / _lead_fetch_json_url / _lead_search_tavily /
# _lead_search_serpapi are the search-provider HTTP primitives, extracted
# verbatim into cvstudio_lead_search (Phase 7B). Re-imported here so the
# app-level names, call sites and mock.patch.object seams are unchanged.
from cvstudio_lead_search import (
    LeadSearchOrchestrator as _LeadSearchOrchestrator,
    _lead_fetch_json_url,
    _lead_search_serpapi,
    _lead_search_tavily,
    _lead_ssl_context,
)


# _lead_search_provider_config + _lead_collect_search_provider_results are
# gathered into LeadSearchOrchestrator (Phase 7B). These stay thin app-level
# delegators so the names remain patchable module attributes and every call
# site + mock.patch.object seam is unchanged; the orchestrator reaches the
# cross-called _lead_* names and the secret store back through app-resolving
# callables (see construction below).
def _lead_collect_search_provider_results(search_provider, job_sources, title_angles, regions, target_role="", job_filters=None, depth="standard"):
    return _LEAD_SEARCH_ORCHESTRATOR.collect_search_provider_results(
        search_provider, job_sources, title_angles, regions, target_role, job_filters, depth
    )


_LEAD_SEARCH_ORCHESTRATOR = _LeadSearchOrchestrator(
    # All injected as callables that resolve the CURRENT app-level name on each
    # call: phase5b patches app._lead_search_provider_config / _lead_search_tavily
    # / _lead_search_serpapi and expects the collection loop to honour the patch,
    # and _ai_secret_store can be reassigned by tests.
    secret_store=lambda: _ai_secret_store,
    search_provider_config=lambda raw: _lead_search_provider_config(raw),
    search_tavily=lambda *a, **k: _lead_search_tavily(*a, **k),
    search_serpapi=lambda *a, **k: _lead_search_serpapi(*a, **k),
    search_provider_queries=lambda *a, **k: _lead_search_provider_queries(*a, **k),
    # call_llm and _lead_call_with_optional_web are patched at app level by the
    # phase5b tests, and _ai_secret_store / _LLM_PROVIDER_INFO can be reassigned;
    # resolve each current app global on every call.
    call_llm=lambda *a, **k: call_llm(*a, **k),
    provider_info=lambda: _LLM_PROVIDER_INFO,
    call_with_optional_web=lambda *a, **k: _lead_call_with_optional_web(*a, **k),
)














def _lead_extract_from_search_provider_results(api_key, model, search_provider, provider_results, regions, job_sources, title_angles, target_role, industries, candidate_context, cv_excerpt, company_n=12, job_filters=None, llm_provider="anthropic"):
    return _LEAD_SEARCH_ORCHESTRATOR.extract_from_search_provider_results(
        api_key, model, search_provider, provider_results, regions, job_sources, title_angles,
        target_role, industries, candidate_context, cv_excerpt, company_n=company_n,
        job_filters=job_filters, llm_provider=llm_provider,
    )

def _lead_quick_job_search(api_key, model, regions, job_sources, title_angles, target_role, industries, candidate_context, cv_excerpt, company_n=6, job_filters=None, llm_provider="anthropic"):
    return _LEAD_SEARCH_ORCHESTRATOR.quick_job_search(
        api_key, model, regions, job_sources, title_angles, target_role, industries,
        candidate_context, cv_excerpt, company_n=company_n, job_filters=job_filters,
        llm_provider=llm_provider,
    )



_LEAD_FINDER_LOCK_CODE = "1571"

def _lead_finder_lock_allowed(body=None):
    """Casual local UI lock guard for Lead Finder routes.

    This is not enterprise authentication; it prevents accidental/casual use of
    the Lead Finder tab and direct route calls on the local offline app.
    """
    try:
        body = body if isinstance(body, dict) else {}
        supplied = (request.headers.get("X-Lead-Finder-Code") or body.get("lead_lock_code") or "")
        return str(supplied).strip() == _LEAD_FINDER_LOCK_CODE
    except Exception:
        return False

def _lead_finder_locked_response():
    return jsonify({"error": "Lead Finder is locked. Unlock the Lead Finder tab with the 4-digit code first."}), 423

@app.route("/lead-finder/test-search-provider", methods=["POST"])
def lead_finder_test_search_provider():
    """Small connectivity check for optional job-search providers."""
    try:
        body = request.get_json(force=True, silent=True) or {}
        cfg = _lead_search_provider_config(body)
        if not cfg.get("enabled"):
            return jsonify({"ok": False, "error": "Choose Tavily or SerpAPI and provide an API key."}), 400
        q = "site:linkedin.com/jobs recruiter Malaysia job"
        if cfg["provider"] == "tavily":
            results = _lead_search_tavily(cfg["api_key"], q, max_results=3, timeout=18)
        elif cfg["provider"] == "serpapi":
            results = _lead_search_serpapi(cfg["api_key"], q, max_results=3, timeout=18)
        else:
            results = []
        return jsonify({
            "ok": True,
            "provider": cfg["provider"],
            "count": len(results),
            "sample": results[:2],
            "external_billing": _phase5b_unavailable_external_billing(
                cfg["provider"],
                observed_operations=1,
                reason=(
                    "The connectivity response did not include "
                    "provider-authoritative billing."
                ),
            ),
        })
    except urllib.error.HTTPError as e:
        try:
            err = json.loads(e.read())
            msg = err.get("error") or err.get("message") or str(err)[:300]
        except Exception:
            msg = "Search provider HTTP error"
        out = {"ok": False, "error": str(msg)[:500]}
        if "cfg" in locals() and cfg.get("provider"):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                cfg["provider"],
                observed_operations=1,
                reason=(
                    "The failed connectivity request returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 400
    except Exception as e:
        out = {"ok": False, "error": str(e)[:500]}
        if "cfg" in locals() and cfg.get("provider"):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                cfg["provider"],
                observed_operations=1,
                reason=(
                    "The connectivity request returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 500

@app.route("/lead-finder/title-cache/stats", methods=["GET"])
def lead_finder_title_cache_stats():
    """Report how many CV content-signatures are cached, and how they break down by role family."""
    try:
        data = _lead_title_cache_load()
        entries = data.get("entries", [])
        by_family = {}
        for e in entries:
            fam = e.get("family") or "unknown"
            by_family[fam] = by_family.get(fam, 0) + 1
        return jsonify({"ok": True, "total_entries": len(entries), "by_family": by_family})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)[:500]}), 500

@app.route("/lead-finder/title-cache/clear", methods=["POST"])
def lead_finder_title_cache_clear():
    """Wipe the local title-angle cache (e.g. if a bad AI response got cached)."""
    try:
        _lead_title_cache_save({"entries": []})
        return jsonify({"ok": True, "message": "Title angle cache cleared."})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)[:500]}), 500

@app.route("/lead-finder/search", methods=["POST"])
def lead_finder_search():
    """Find selected-country job leads first. People discovery is a separate lighter step."""
    try:
        ai_request_started = False
        body = request.get_json(force=True, silent=True) or {}
        if not _lead_finder_lock_allowed(body):
            return _lead_finder_locked_response()
        api_key = _resolve_request_api_key(body, "lead")
        model = body.get("model") or "claude-sonnet-4-6"
        llm_provider = (body.get("provider") or "anthropic").strip().lower()
        cv_text = (body.get("cv_text") or "").strip()
        regions = body.get("regions") or []
        if not isinstance(regions, list):
            regions = []
        cleaned_regions = []
        seen_regions = set()
        for r in regions:
            r = str(r or "").strip()
            if not r:
                continue
            k = r.lower()
            if k in seen_regions:
                continue
            seen_regions.add(k)
            cleaned_regions.append(r)
        regions = cleaned_regions

        job_sources = body.get("job_sources") or []
        target_role = (body.get("target_role") or "").strip()
        candidate_context = (body.get("candidate_context") or "").strip()
        industries = (body.get("industries") or "").strip()
        job_filters = _lead_clean_job_filters(body.get("job_filters") or {})
        search_provider = _lead_search_provider_config(body.get("search_provider") or {})
        depth = (body.get("depth") or "standard").strip().lower()

        if not api_key:
            return jsonify({"error": "API key required"}), 400
        if not cv_text:
            return jsonify({"error": "CV text required"}), 400
        if not regions:
            return jsonify({"error": "At least one region is required"}), 400
        if not isinstance(job_sources, list):
            job_sources = []
        job_sources = [str(x).strip() for x in job_sources if str(x).strip()]
        if not job_sources:
            job_sources = ["JobStreet", "LinkedIn Jobs public search results", "Company career pages"]

        depth_cfg = {
            # Proven-small mode: earlier Lead Finder versions worked best when
            # the AI web-search call aimed for a compact 4-6 direct job leads.
            # Keep enough output budget for valid JSON; too little output was
            # causing malformed/non-JSON responses after paid web-search runs.
            "light":    {"company_n": 6,  "max_uses": 2, "tokens": 4200, "cv_chars": 7000,  "primary_timeout": 75,  "fallback_timeout": 18},
            "standard": {"company_n": 6,  "max_uses": 2, "tokens": 4600, "cv_chars": 8500,  "primary_timeout": 90,  "fallback_timeout": 20},
            "deep":     {"company_n": 8,  "max_uses": 3, "tokens": 5600, "cv_chars": 10500, "primary_timeout": 110, "fallback_timeout": 24},
        }.get(depth, {"company_n": 6, "max_uses": 2, "tokens": 4600, "cv_chars": 8500, "primary_timeout": 90, "fallback_timeout": 20})

        # Keep requests bounded. Long CVs were a major timeout driver.
        cv_text = re.sub(r"\n{3,}", "\n\n", cv_text)
        cv_text = re.sub(r"[ \t]{2,}", " ", cv_text).strip()
        cv_excerpt = cv_text[:depth_cfg["cv_chars"]]
        sources_text = ", ".join(job_sources[:8])
        regions_text = ", ".join(regions)
        today_iso = date.today().isoformat()
        role_search_target, role_anchor_warning = _lead_resolve_search_target_role(target_role, cv_excerpt, candidate_context, industries)
        job_title_angles = _lead_job_title_angles(role_search_target, cv_excerpt, candidate_context, industries)
        title_expansion_usage = {}
        title_expansion_warning = ""
        if bool(body.get("allow_ai_title_expansion")) and api_key:
            primary_families = _lead_primary_role_families(role_search_target, cv_excerpt, candidate_context, industries)
            cache_sig = _lead_cv_content_signature(role_search_target, cv_excerpt, candidate_context, industries)
            cached_entry, cache_score = (None, 0.0)
            if cache_sig:
                cached_entry, cache_score = _lead_title_cache_find(cache_sig[0], cache_sig[1])
            if cached_entry:
                ai_titles = cached_entry.get("titles") or []
                _lead_title_cache_touch(cached_entry)
                title_expansion_warning = f"Reused cached AI title angles from a similar previous CV ({int(cache_score * 100)}% content-evidence match) — no extra AI cost."
            else:
                ai_request_started = True
                ai_titles, title_expansion_usage, title_expansion_warning = _lead_ai_expand_title_angles(
                    api_key, model, role_search_target, cv_excerpt, candidate_context, industries,
                    primary_families=primary_families, llm_provider=llm_provider
                )
                if ai_titles and cache_sig:
                    _lead_title_cache_store(cache_sig[0], cache_sig[1], ai_titles)
            if ai_titles:
                existing_lower = {t.lower() for t in job_title_angles}
                added = 0
                for t in ai_titles:
                    if t.lower() not in existing_lower:
                        job_title_angles.append(t)
                        existing_lower.add(t.lower())
                        added += 1
                if added:
                    added_msg = f"Added {added} title angle(s) from cache." if cached_entry else f"AI title expansion added {added} CV-specific title angle(s) to the search."
                    title_expansion_warning = (title_expansion_warning + " " + added_msg).strip() if title_expansion_warning else added_msg
        job_title_angles = job_title_angles[:70]
        job_title_angles_for_search = job_title_angles[:5]
        job_title_angles_text = "\n".join(f"- {t}" for t in job_title_angles_for_search) or "- Infer job titles directly from the CV/context."
        job_filter_instruction = _lead_job_filter_instruction(job_filters)

        graceful_json = {
            "summary": {
                "candidate_title": role_search_target or target_role or "",
                "original_target_role": target_role or "",
                "seniority": "",
                "core_skills": [],
                "target_roles": [role_search_target or target_role] if (role_search_target or target_role) else [],
                "target_job_title_angles": job_title_angles,
                "job_filters_used": job_filters,
                "regions": regions,
                "location_filter": "selected countries are hard filters",
                "job_sources_used": job_sources,
                "source_strategy": "The live job-portal call timed out before returning usable data.",
      "job_fit_method": "Job Fit % estimates role/skill/seniority/location alignment, but lower-fit leads are still shown for recruiter review unless explicitly excluded.",
                "compliance_note": "No private/protected scraping performed."
            },
            "companies": [],
            "people": []
        }

        # Preferred production-style path: use a dedicated search provider API
        # to retrieve job-result URLs/snippets, then ask Claude only to classify
        # those results. This avoids the unreliable one-shot AI web-crawl path.
        if search_provider.get("enabled"):
            provider_results, provider_warnings = _lead_collect_search_provider_results(
                search_provider, job_sources, job_title_angles, regions, role_search_target, job_filters, depth
            )
            if provider_results:
                # First pass: deterministic provider extraction. This avoids spending
                # another AI call when Tavily/SerpAPI already returned direct job URLs
                # or company/role snippets. It also restores the earlier behaviour of
                # returning a small set of real leads instead of falling straight to
                # generic portal search links.
                parsed, provider_warning = _lead_provider_results_to_leads(
                    provider_results, regions, job_sources, job_title_angles, role_search_target,
                    job_filters=job_filters, max_n=depth_cfg["company_n"]
                )
                parsed, direct_only_warning = _lead_keep_only_direct_job_leads(parsed)
                real_count = len(parsed.get("companies") or [])
                warnings = []
                if role_anchor_warning:
                    warnings.append(role_anchor_warning)
                if title_expansion_warning:
                    warnings.append(title_expansion_warning)
                if direct_only_warning:
                    warnings.append(direct_only_warning)
                if provider_warnings:
                    warnings.append("Search provider warning(s): " + " | ".join(provider_warnings[:3]))
                if provider_warning:
                    warnings.append(provider_warning)
                usage = {}
                if real_count == 0 and bool(body.get("allow_provider_ai_refine")):
                    ai_request_started = True
                    parsed, ai_warning, usage = _lead_extract_from_search_provider_results(
                        api_key, model, search_provider, provider_results, regions, job_sources, job_title_angles,
                        role_search_target, industries, candidate_context, cv_excerpt,
                        company_n=min(depth_cfg["company_n"], {"light": 10, "standard": 18, "deep": 30}.get(depth, 18)),
                        job_filters=job_filters, llm_provider=llm_provider
                    )
                    # The refine call has its own prompt/parsing; it does not
                    # automatically inherit the fast-extract path's role-family and
                    # direct-URL/downgrade cleanup. Apply the same guards here so
                    # refine cannot reintroduce mismatched-family leads (e.g. HR
                    # jobs for a non-HR candidate) or raw search-listing links.
                    parsed = _lead_normalize_company_fit_freshness(parsed)
                    parsed, refine_role_family_warning = _lead_filter_by_title_angles(parsed, job_title_angles)
                    parsed = _lead_normalize_company_fit_freshness(parsed)
                    parsed, refine_direct_only_warning = _lead_keep_only_direct_job_leads(parsed)
                    if refine_role_family_warning:
                        warnings.append(refine_role_family_warning)
                    if refine_direct_only_warning:
                        warnings.append(refine_direct_only_warning)
                    if ai_warning:
                        warnings.append(ai_warning)
                elif real_count == 0:
                    # Do not show generic portal/search links as "leads" and do not spend
                    # on an extra Claude refine call by default. The user's feedback was
                    # clear: 0 extracted leads should look like 0 leads, not 8 fallback links.
                    parsed["companies"] = []
                    parsed.setdefault("summary", {})
                    parsed["summary"]["source_strategy"] = "Search provider returned results, but none could be safely converted into direct job leads. Generic portal/search links are hidden and no extra AI refine was run by default to protect cost."
                    warnings.append("Search provider returned results, but no direct job leads could be extracted locally. Generic portal/search links are no longer shown, and AI refine was skipped to avoid extra cost.")
                else:
                    warnings.append("Search provider fast-extract path used; avoided extra AI classification cost. Third-party provider usage/credits are not included in CV Studio cost.")
                parsed.setdefault("summary", {})
                if isinstance(parsed.get("summary"), dict):
                    parsed["summary"].setdefault("target_job_title_angles", job_title_angles)
                    parsed["summary"].setdefault("job_filters_used", job_filters)
                    parsed["summary"].setdefault("regions", regions)
                    parsed["summary"].setdefault("job_sources_used", job_sources)
                    parsed["summary"]["search_provider_used"] = search_provider.get("provider")
                    parsed["summary"]["search_provider_results_seen"] = len(provider_results)
                    if real_count:
                        parsed["summary"]["source_strategy"] = "Search provider results converted directly into provisional extracted job leads without extra AI classification cost."
                parsed["people"] = []
                parsed, source_selection_warning = _lead_cleanup_urls_by_selected_sources(parsed, job_sources)
                if source_selection_warning:
                    warnings.append(source_selection_warning)
                parsed["ok"] = True
                usage = _lead_merge_usage(usage, title_expansion_usage)
                parsed["usage"] = usage or {}
                parsed["model"] = model
                parsed["provider"] = llm_provider
                parsed["cost"] = _lead_cost(model, usage or {}, llm_provider)
                parsed["cost_details"] = _lead_cost_details(model, usage or {}, llm_provider)
                parsed["external_billing"] = _phase5b_unavailable_external_billing(
                    search_provider.get("provider"),
                    reason=(
                        "Search-provider results do not include "
                        "provider-authoritative billing."
                    ),
                )
                parsed["phase"] = "job_search_provider_fast_extract"
                if warnings:
                    parsed["warning"] = " ".join(warnings)
                return jsonify(parsed)
            else:
                # Provider was explicitly configured but returned no results. Do not fall
                # through to the expensive AI web-search route by default; this was causing
                # paid AI usage with no extracted leads.
                out = {
                    "ok": True,
                    "summary": {
                        "candidate_title": role_search_target or target_role or "Candidate",
                        "target_roles": [role_search_target or target_role] if (role_search_target or target_role) else [],
                        "target_job_title_angles": job_title_angles,
                        "job_filters_used": job_filters,
                        "regions": regions,
                        "job_sources_used": job_sources,
                        "search_provider_used": search_provider.get("provider"),
                        "search_provider_results_seen": 0,
                        "source_strategy": f"Search provider {search_provider.get('provider')} returned no results. No generic fallback search links are shown and no AI web-search fallback was run to avoid extra cost.",
                        "job_fit_method": "No extracted job leads were returned.",
                        "compliance_note": "Public search/job sources only."
                    },
                    "companies": [],
                    "people": [],
                    "usage": title_expansion_usage or {},
                    "model": model,
                    "provider": llm_provider,
                    "cost": _lead_cost(model, title_expansion_usage or {}, llm_provider),
                    "cost_details": _lead_cost_details(model, title_expansion_usage or {}, llm_provider),
                    "phase": "job_search_provider_empty",
                    "warning": (f"Search provider {search_provider.get('provider')} returned no usable job results. No fallback search links shown and no AI fallback cost incurred." + (" " + title_expansion_warning if title_expansion_warning else "")).strip()
                }
                out["external_billing"] = _phase5b_unavailable_external_billing(
                    search_provider.get("provider"),
                    reason=(
                        "Search-provider results do not include "
                        "provider-authoritative billing."
                    ),
                )
                return jsonify(out)

        prompt = f"""
You are building a compliant recruiting job-lead list for Hyppies.

IMPORTANT PERFORMANCE RULE
- This is PHASE 1 ONLY: find matching JOB ADS / JOB LEADS.
- Do NOT search for people or emails in this call. Return "people": []. People discovery is done separately from saved job leads.
- You have at most {depth_cfg['max_uses']} web search(es) total for this run. A confirmed DIRECT job-ad URL for one company is worth far more to the recruiter than an extra company whose only evidence is a search-results/listing-page snippet. If your search budget is tight, it is correct to return FEWER companies with a real, confirmed direct URL rather than MORE companies that all fall back to Needs Verification.
- Recommended budget allocation: use your first search broadly to find matching companies/roles, then spend at least one more search specifically CONFIRMING the direct job-ad URL for your strongest 2-4 candidates -- search the portal + the exact company name + the exact role title together (e.g. site:my.jobstreet.com "Acme Corp" "Data Engineer") to try to land on that company's individual job page, not just the listing page you first found them on. Only fall back to Needs Verification for a candidate after that confirmation search doesn't surface a direct permalink.
- Use the earlier proven pattern: search selected portal + country + role-title angle, open/inspect the result enough to extract company/title/direct URL, then return compact JSON.
- Do not search broad CV text. Search role/title angles and filters only, then evaluate results against the CV.

JOB-FIRST TASK
1) Infer the candidate's best-fit job titles, seniority, skills, and industries from the CV/context.
2) Expand the search beyond exact title wording using the role-family title angles below. This tool is for GENERALIST recruiters across HR, operations, finance, marketing, sales, legal, supply chain, oil & gas, consulting, fresh graduate roles, technology, data, SAP/ERP and niche functions. Do not bias toward tech unless the CV is tech.
3) Search ONLY the selected job source(s)/portal(s): {sources_text}
4) Search only for selected country/region leads: {regions_text}
5) Keep only actual public job ads or individual company-career requisition pages that match the candidate.
6) Convert those job ads into job leads. Do not include generic portal search/listing pages.
7) For every job lead, extract job-fit percentage and job freshness/date posted from the job ad or search result snippet where available.

COUNTRY / REMOTE LOCATION RULES
- Selected country/region is a hard filter, not a preference.
- Include on-site/hybrid roles located in the selected country/region.
- Include remote roles ONLY when public evidence clearly says the role accepts candidates based in the selected country/region.
- For APAC/SEA selections, APAC/SEA remote coverage is acceptable.
- If the location is unclear, exclude the lead.

LOCATION INSTRUCTION FOR THIS RUN
{_lead_location_instruction(regions)}

TODAY / DATE BASELINE
{today_iso}
Use this date to convert relative freshness such as "today", "3 days ago", or "2 weeks ago" into date_posted and days_open where possible.

STRICT COMPLIANCE RULES
- Use public business/job information only.
- Do NOT log in, scrape protected pages, bypass access controls, or extract private data.
- Do NOT guess emails. This phase should not return emails.
- job_url must be a DIRECT job-ad URL. Do NOT return portal homepages, generic job-search URLs, company careers landing pages, or search-results pages as leads.
- If you cannot find a direct job-post URL but employer + role evidence is clear, return the row as Needs Verification with source_url only. Do not return generic source-only pages.

REGIONS
{json.dumps(regions, ensure_ascii=False)}

SELECTED JOB SOURCES / PORTALS
{json.dumps(job_sources, ensure_ascii=False)}

ROLE-FAMILY JOB TITLE ANGLES TO SEARCH
{job_title_angles_text}

RELEVANCE FILTERS FOR THIS RUN
{job_filter_instruction}

TARGET ROLE
{role_search_target or target_role or "Infer suitable roles from the CV"}

INDUSTRY FOCUS
{industries or "Infer from the CV; include adjacent sectors where the candidate can fit"}

RECRUITER CONTEXT
{candidate_context[:1200] or "No extra context provided."}

CANDIDATE CV TEXT
{cv_excerpt}

RETURN ONLY VALID JSON. No markdown. Shape exactly:
{{
  "summary": {{
    "candidate_title": "best-fit candidate title",
    "seniority": "junior/mid/senior/lead/head/etc",
    "core_skills": ["skill"],
    "target_roles": ["role"],
    "target_job_title_angles": ["role/title angle searched"],
    "job_filters_used": {{}},
    "regions": ["region"],
    "location_filter": "selected countries are hard filters; include selected-country jobs and remote roles only when selected-country eligibility is evident",
    "job_sources_used": ["portal/source"],
    "source_strategy": "job-portal-first job search; people discovery is separate",
    "job_fit_method": "Job Fit % estimates role/skill/seniority/location alignment, but lower-fit leads are still shown for recruiter review unless explicitly excluded.",
    "freshness_method": "date_posted and days_open are taken from the job ad/search snippet when available; otherwise mark Unknown",
    "compliance_note": "public job/business sources only"
  }},
  "companies": [
    {{
      "id": "co_1",
      "company": "Company name",
      "country": "Country",
      "region": "City/region if known",
      "industry": "Industry",
      "job_portal": "selected portal/source",
      "matched_role": "Relevant open role title",
      "hiring_signal": "Fresh public job ad/hiring signal",
      "date_posted": "YYYY-MM-DD if known, or raw relative text like 3 days ago, or empty string",
      "days_open": 0,
      "job_freshness": "Today / 1-7 days open / 8-14 days open / 15-30 days open / 30+ days open / Unknown",
      "job_url": "DIRECT public job ad URL only; empty string if only a search/listing/portal page is available",
      "job_url_quality": "direct_job_ad / needs_verification / source_or_search_result_only",
      "company_url": "URL or empty string",
      "job_fit_percent": 0,
      "match_score": 0,
      "why_matched": "short reason based on CV and role/skills",
      "source_url": "best evidence URL or empty string",
      "source_note": "job portal/source/evidence summary",
      "date_found": "YYYY-MM-DD or empty string"
    }}
  ],
  "people": []
}}

TARGET COUNT
- companies: up to {depth_cfg['company_n']} usable job leads. Prefer 4-6 solid leads over a broad but noisy list. Direct job URLs are best; Needs Verification leads are acceptable when employer + role evidence is clear.
- Avoid duplicate company+role+portal combinations. Prefer company/title diversity, but do not chase volume if direct job URLs are weak.
- Apply the relevance filters in ranking and selection. Exclude clear mismatches only from explicit exclude/company-exclude keywords and hard location rules.
- Do NOT filter out jobs only because tech stack/skill fit is weak or incomplete; keep them as lower-fit review leads so the recruiter can decide.
- job_fit_percent must be an integer 0-100 representing how well the job fits this candidate.
- match_score must be the same integer as job_fit_percent for backwards compatibility.
- date_posted/days_open/job_freshness should reflect the public job ad freshness. Use empty string/Unknown when not visible.
- job_url should point to the direct job ad when available. If only a portal/search/company-careers source is visible but it clearly shows employer + role evidence, leave job_url empty, set source_url, and set job_url_quality='needs_verification'. Do not return generic source-only pages.
""".strip()

        payload = {
            "model": model,
            "max_tokens": depth_cfg["tokens"],
            "messages": [{"role": "user", "content": prompt}],
            "tools": [{"type": "web_search_20250305", "name": "web_search", "max_uses": depth_cfg["max_uses"]}],
            "_primary_timeout_seconds": depth_cfg["primary_timeout"],
            "_fallback_timeout_seconds": depth_cfg["fallback_timeout"],
            "_skip_no_web_fallback": True,
        }
        ai_request_started = True
        data, warning = _lead_call_with_optional_web(api_key, payload, graceful_json=graceful_json, provider=llm_provider)
        raw_text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text" or "text" in b)
        parse_warning = ""
        json_parse_failed = False
        try:
            parsed = _lead_extract_json(raw_text)
        except Exception:
            json_parse_failed = True
            parsed, recovered_warning = _lead_recover_job_leads_from_text(raw_text, regions, job_sources, job_title_angles, max_n=depth_cfg["company_n"])
            if recovered_warning:
                parse_warning = "API returned non-JSON Lead Finder output, but usable direct job-post URLs were recovered. " + recovered_warning
            else:
                parsed = json.loads(json.dumps(graceful_json, ensure_ascii=False))
                parse_warning = "API returned malformed/non-JSON Lead Finder output and no direct job-post URLs could be recovered; returned a safe empty job-lead result instead of failing."
        # Phase 1 is job-leads only. Some models may ignore the prompt
        # and return people anyway; wipe them here so the separate Find People
        # phase remains the single source for people leads and costs.
        parsed["people"] = []
        parsed = _lead_normalize_company_fit_freshness(parsed)
        parsed, location_warning = _lead_filter_by_regions(parsed, regions)
        parsed["people"] = []
        parsed = _lead_normalize_company_fit_freshness(parsed)
        parsed, filter_warning = _lead_apply_basic_job_filters(parsed, job_filters)
        parsed = _lead_normalize_company_fit_freshness(parsed)
        parsed, role_family_warning = _lead_filter_by_title_angles(parsed, job_title_angles)
        parsed = _lead_normalize_company_fit_freshness(parsed)
        parsed, direct_only_warning = _lead_keep_only_direct_job_leads(parsed)
        parsed = _lead_normalize_company_fit_freshness(parsed)
        # Last safety net: if strict cleanup removed everything but the paid
        # model output contained direct job URLs in prose/source notes, recover
        # those direct job posts instead of wasting the run.
        if json_parse_failed and not parsed.get("companies") and raw_text:
            recovered, recovered_warning = _lead_recover_job_leads_from_text(raw_text, regions, job_sources, job_title_angles, max_n=depth_cfg["company_n"])
            if recovered.get("companies"):
                recovered = _lead_normalize_company_fit_freshness(recovered)
                recovered, location_warning_2 = _lead_filter_by_regions(recovered, regions)
                recovered = _lead_normalize_company_fit_freshness(recovered)
                recovered, filter_warning_2 = _lead_apply_basic_job_filters(recovered, job_filters)
                recovered = _lead_normalize_company_fit_freshness(recovered)
                recovered, role_family_warning_2 = _lead_filter_by_title_angles(recovered, job_title_angles)
                recovered = _lead_normalize_company_fit_freshness(recovered)
                recovered, direct_only_warning_2 = _lead_keep_only_direct_job_leads(recovered)
                if recovered.get("companies"):
                    parsed = recovered
                    if recovered_warning:
                        warning = (warning + " " + recovered_warning).strip() if warning else recovered_warning
                    for _w in (location_warning_2, filter_warning_2, role_family_warning_2, direct_only_warning_2):
                        if _w:
                            warning = (warning + " " + _w).strip() if warning else _w
        usage = data.get("usage", {})
        real_count = len(parsed.get("companies") or [])

        # If the main live crawl came back empty, do not make a second expensive
        # AI call and do not show generic portal/search links as leads. This keeps
        # cost honest: 0 extracted leads should be displayed as 0 extracted leads.
        if not parsed.get("companies"):
            parsed["companies"] = []
            parsed.setdefault("summary", {})
            parsed["summary"].update({
                "candidate_title": parsed.get("summary", {}).get("candidate_title") or role_search_target or target_role,
                "original_target_role": target_role or "",
                "target_job_title_angles": job_title_angles,
                "job_filters_used": job_filters,
                "regions": regions,
                "job_sources_used": job_sources,
                "source_strategy": "AI live job extraction returned no direct job leads after JSON parsing, direct URL cleanup, and non-JSON URL recovery. Generic portal/search fallback links are hidden."
            })
            empty_warning = "No extracted job leads were returned after direct URL cleanup and non-JSON recovery. Generic JobStreet/LinkedIn search links are hidden."
            warning = (warning + " " + empty_warning).strip() if warning else empty_warning

        parsed.setdefault("summary", {})
        if isinstance(parsed.get("summary"), dict):
            parsed["summary"].setdefault("target_job_title_angles", job_title_angles)
            parsed["summary"].setdefault("job_filters_used", job_filters)
            parsed["summary"].setdefault("regions", regions)
            parsed["summary"].setdefault("job_sources_used", job_sources)
        parsed, source_selection_warning = _lead_cleanup_urls_by_selected_sources(parsed, job_sources)
        if source_selection_warning:
            warning = (warning + " " + source_selection_warning).strip() if warning else source_selection_warning
        parsed["ok"] = True
        usage = _lead_merge_usage(usage, title_expansion_usage)
        parsed["usage"] = usage
        if parse_warning:
            warning = (warning + " " + parse_warning).strip() if warning else parse_warning
        if location_warning:
            warning = (warning + " " + location_warning).strip() if warning else location_warning
        if 'filter_warning' in locals() and filter_warning:
            warning = (warning + " " + filter_warning).strip() if warning else filter_warning
        if 'role_family_warning' in locals() and role_family_warning:
            warning = (warning + " " + role_family_warning).strip() if warning else role_family_warning
        if 'role_anchor_warning' in locals() and role_anchor_warning:
            warning = (warning + " " + role_anchor_warning).strip() if warning else role_anchor_warning
        if title_expansion_warning:
            warning = (warning + " " + title_expansion_warning).strip() if warning else title_expansion_warning
        if 'direct_only_warning' in locals() and direct_only_warning:
            warning = (warning + " " + direct_only_warning).strip() if warning else direct_only_warning
        parsed["model"] = model
        parsed["provider"] = llm_provider
        # Always report the API usage/cost returned by the provider, even when
        # 0 leads were extracted. Recruiters need visibility into wasted spend.
        parsed["cost"] = _lead_cost(model, usage, llm_provider)
        parsed["cost_details"] = _lead_cost_details(model, usage, llm_provider)
        if search_provider.get("enabled"):
            parsed["external_billing"] = _phase5b_unavailable_external_billing(
                search_provider.get("provider"),
                reason=(
                    "Search-provider results do not include "
                    "provider-authoritative billing."
                ),
            )
        if real_count == 0:
            spend_msg = "Cost transparency: this run returned 0 extracted direct job leads, but any API usage returned by the provider is still displayed so you can monitor wasted spend."
            warning = (warning + " " + spend_msg).strip() if warning else spend_msg
        parsed["phase"] = "job_search"
        if warning:
            parsed["warning"] = warning
        return jsonify(parsed)

    except urllib.error.HTTPError as e:
        try:
            err = json.loads(e.read())
            msg = err.get("error", {}).get("message", "API error")
        except Exception:
            msg = "API error"
        out = {"error": "API provider error: " + msg}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage") or locals().get("title_expansion_usage"),
            attempted=bool(
                locals().get("ai_request_started")
                or _llm_usage_int(
                    locals().get("usage")
                    or locals().get("title_expansion_usage"),
                    "api_calls",
                )
            ),
        ))
        if locals().get("search_provider", {}).get("enabled"):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                search_provider.get("provider"),
                reason=(
                    "The failed search-provider workflow returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 400
    except (TimeoutError, socket.timeout, urllib.error.URLError) as e:
        out = {"error": "Lead Finder could not reach the API provider. Check internet/API key, then try again with Light depth."}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage") or locals().get("title_expansion_usage"),
            attempted=bool(
                locals().get("ai_request_started")
                or _llm_usage_int(
                    locals().get("usage")
                    or locals().get("title_expansion_usage"),
                    "api_calls",
                )
            ),
        ))
        if locals().get("search_provider", {}).get("enabled"):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                search_provider.get("provider"),
                reason=(
                    "The timed-out search-provider workflow returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 504
    except Exception as e:
        traceback.print_exc()
        out = {"error": str(e)}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage") or locals().get("title_expansion_usage"),
            attempted=bool(
                locals().get("ai_request_started")
                or _llm_usage_int(
                    locals().get("usage")
                    or locals().get("title_expansion_usage"),
                    "api_calls",
                )
            ),
        ))
        if locals().get("search_provider", {}).get("enabled"):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                search_provider.get("provider"),
                reason=(
                    "The search-provider workflow returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 500







# Strong role-family detection for Lead Finder job-search angles.
# This deliberately scores actual role/function evidence instead of every loose
# keyword in the CV. Example: a Data Engineer CV that mentions an HR stakeholder
# should still generate Data Engineer angles, not HR Executive angles.

# Any of these already function as a complete, terminal job-title noun on their
# own. If the typed/inferred role already ends in one of these, do not stack
# another generic suffix (Manager/Executive/Specialist/Consultant/Analyst) on
# top of it — that produced nonsense like "Financial Controller Manager",
# "Legal Counsel Manager", "Company Secretary Manager", "Accountant Executive",
# and "Growth Marketer Manager" in earlier builds.

# Any of these already signal a seniority/career level. If the typed/inferred
# role already carries one, do not layer another level modifier on top —
# earlier builds produced "Junior Senior Accountant" and "Assistant Senior
# Accountant" this way.













_LEAD_TITLE_CACHE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "lead_title_cache.json")
_LEAD_TITLE_CACHE_MAX_ENTRIES = 10000
_LEAD_TITLE_CACHE_SIM_THRESHOLD = 0.45

# Same boost-token lists used in _lead_family_scores, reused here so the cache
# signature reflects the same evidence the family-detector actually keyed on.


def _lead_title_cache_load():
    return _LEAD_CACHE_SERVICE.title_cache_load()


def _lead_title_cache_save(data):
    return _LEAD_CACHE_SERVICE.title_cache_save(data)






def _lead_title_cache_find(family, evidence, threshold=_LEAD_TITLE_CACHE_SIM_THRESHOLD):
    return _LEAD_CACHE_SERVICE.title_cache_find(family, evidence, threshold)


def _lead_title_cache_store(family, evidence, titles):
    return _LEAD_CACHE_SERVICE.title_cache_store(family, evidence, titles)


_LEAD_CONTACT_CACHE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "lead_contact_cache.json")
_LEAD_CONTACT_CACHE_MAX_ENTRIES = 10000

# Repositories and paths are passed as callables because the test-suite
# reassigns these module globals after import (e.g. swapping in a corrupt-storage
# repository) and expects the cache to read the current value on each call.
_LEAD_CACHE_SERVICE = _LeadCacheService(
    title_repository=lambda: _CVSTUDIO_LEAD_TITLE_REPOSITORY,
    contact_repository=lambda: _CVSTUDIO_LEAD_CONTACT_REPOSITORY,
    title_path=lambda: _LEAD_TITLE_CACHE_PATH,
    contact_path=lambda: _LEAD_CONTACT_CACHE_PATH,
    legacy_json_read=_cvstudio_legacy_json_read,
    legacy_json_write=_cvstudio_legacy_json_write,
    storage_error=StorageError,
    title_max_entries=_LEAD_TITLE_CACHE_MAX_ENTRIES,
    contact_max_entries=_LEAD_CONTACT_CACHE_MAX_ENTRIES,
    title_sim_threshold=_LEAD_TITLE_CACHE_SIM_THRESHOLD,
    normalize_linkedin_url=_lead_normalize_linkedin_url,
    is_company_domain_email=_lead_is_company_domain_email,
)


def _lead_contact_cache_load():
    return _LEAD_CACHE_SERVICE.contact_cache_load()


def _lead_contact_cache_save(data):
    return _LEAD_CACHE_SERVICE.contact_cache_save(data)


def _lead_contact_cache_key(person):
    return _LEAD_CACHE_SERVICE.contact_cache_key(person)


def _lead_contact_cache_find(person):
    return _LEAD_CACHE_SERVICE.contact_cache_find(person)


def _lead_contact_cache_store(person, email_data):
    return _LEAD_CACHE_SERVICE.contact_cache_store(person, email_data)


def _lead_contact_cache_touch(person):
    return _LEAD_CACHE_SERVICE.contact_cache_touch(person)


class _LeadApolloRateLimited(Exception):
    """Raised when Apollo returns 429, so a batch enrichment run can stop
    immediately instead of continuing to burn through remaining quota-limited
    calls that would just fail the same way."""
    pass


class _LeadApolloAuthError(Exception):
    """Raised on 401/403 so the person can be told to check their API key,
    rather than every person in the batch silently coming back 'Not found'."""
    pass


def _lead_apollo_enrich_person(api_key, person, timeout=15):
    """Look up one person via Apollo's People Match endpoint.

    Uses name + organization_name + linkedin_url (when available) rather than
    Apollo's internal person IDs, since we never have those — this is the
    documented, reliably-working input shape for this endpoint. Only the
    business email is requested (reveal_personal_emails/reveal_phone_number
    are left False) to stay consistent with this tool's existing
    business-email-only stance.
    """
    name = re.sub(r"\s+", " ", str(person.get("name") or "").strip())
    company = re.sub(r"\s+", " ", str(person.get("company") or "").strip())
    linkedin_url = _lead_normalize_linkedin_url(person.get("profile_url") or person.get("linkedin_url"))
    if not name and not linkedin_url:
        return {"email": "", "email_confidence": "", "email_source": "", "verification_status": "Not found"}

    name_parts = name.split(" ", 1)
    body = {
        "reveal_personal_emails": False,
        "reveal_phone_number": False,
    }
    if name_parts:
        body["first_name"] = name_parts[0]
    if len(name_parts) > 1:
        body["last_name"] = name_parts[1]
    if name:
        body["name"] = name
    if company:
        body["organization_name"] = company
    if linkedin_url:
        body["linkedin_url"] = linkedin_url

    params = urllib.parse.urlencode(
        {k: (str(v).lower() if isinstance(v, bool) else v) for k, v in body.items()},
        doseq=True,
    )
    req_url = "https://api.apollo.io/api/v1/people/match" + ("?" + params if params else "")
    req = urllib.request.Request(req_url, data=b"", method="POST", headers={
        "accept": "application/json",
        "x-api-key": str(api_key or "").strip(),
    })
    context = _lead_ssl_context()
    try:
        if context is not None:
            resp_cm = urllib.request.urlopen(req, timeout=timeout, context=context)
        else:
            resp_cm = urllib.request.urlopen(req, timeout=timeout)
        with resp_cm as resp:
            data = json.loads(resp.read().decode("utf-8", errors="replace"))
    except urllib.error.HTTPError as e:
        if e.code == 429:
            raise _LeadApolloRateLimited("Apollo API rate limit hit (429).")
        if e.code in (401, 403):
            raise _LeadApolloAuthError(f"Apollo API authentication failed ({e.code}). Check the Apollo API key.")
        try:
            body_txt = e.read().decode("utf-8", errors="replace")[:300]
        except Exception:
            body_txt = ""
        raise RuntimeError(f"Apollo API error {e.code}: {body_txt or str(e)}")

    p = data.get("person") if isinstance(data, dict) else None
    if not isinstance(p, dict) or not p.get("email"):
        return {"email": "", "email_confidence": "", "email_source": "", "verification_status": "Not found"}

    email = str(p.get("email") or "").strip()
    if not _lead_is_company_domain_email(email, company):
        return {"email": "", "email_confidence": "", "email_source": "", "verification_status": "Not found"}

    email_status = str(p.get("email_status") or "").strip().lower()
    confidence = "High" if email_status == "verified" else ("Medium" if email_status else "Low")
    return {
        "email": email,
        "email_confidence": confidence,
        "email_source": "Apollo People Match API",
        "verification_status": "Verified business email" if email_status == "verified" else "Public business email",
    }


def _lead_title_cache_touch(matched_entry):
    return _LEAD_CACHE_SERVICE.title_cache_touch(matched_entry)


def _lead_ai_expand_title_angles(api_key, model, target_role, cv_excerpt, candidate_context, industries, primary_families=None, max_titles=18, timeout=20, llm_provider="anthropic"):
    """Ask Claude to generate recruiter title-search angles directly from the CV.

    This SUPPLEMENTS the static family-based title bank rather than replacing
    it. The static bank is free, fast, and deterministic but only covers
    verticals someone has explicitly coded a pattern for (which is why DBA,
    SAP Basis, payroll, etc. kept slipping through). This call is a small,
    cheap, text-only request (no web search tool) that reads the actual CV and
    proposes title angles for whatever role the candidate actually has, so
    unusual/niche titles get reasonable coverage without a code change.
    """
    prompt = f"""
You are helping a recruiter build a list of alternative job-title search angles
for ONE candidate, based only on their actual CV and stated target role.

STRICT RULES
- Stay within the candidate's own job function/department. Do NOT branch into
  an unrelated function (e.g. do not suggest HR titles for a Finance candidate,
  or Sales titles for an Engineering candidate).
- Include the exact target role/title itself, plus realistic seniority variants
  (junior/senior/lead/head/manager as appropriate for that function) and
  equivalent titles used by different companies/countries for the same role.
- Do not invent a title with no basis in the CV or target role.
- Return 10-18 titles, ranked most relevant first.

TARGET ROLE
{target_role or "Infer from the CV"}

INDUSTRY FOCUS
{industries or "Infer from the CV"}

RECRUITER CONTEXT
{candidate_context[:800] or "None provided."}

CANDIDATE CV EXCERPT
{cv_excerpt[:4000]}

RETURN ONLY VALID JSON. No markdown. Shape exactly:
{{"titles": ["title 1", "title 2"]}}
""".strip()
    payload = {
        "model": model,
        "max_tokens": 600,
        "messages": [{"role": "user", "content": prompt}],
        "_timeout_seconds": timeout,
    }
    try:
        data = call_llm(llm_provider, api_key, payload)
    except Exception:
        return [], {}, "AI title expansion call failed; continuing with the built-in title bank only."
    raw_text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text" or "text" in b)
    usage = data.get("usage", {}) or {}
    try:
        parsed = _lead_extract_json(raw_text)
    except Exception:
        return [], usage, "AI title expansion returned malformed output; continuing with the built-in title bank only."
    titles = parsed.get("titles") if isinstance(parsed, dict) else None
    if not isinstance(titles, list):
        return [], usage, "AI title expansion returned no usable titles; continuing with the built-in title bank only."
    clean = []
    seen = set()
    for t in titles:
        t = re.sub(r"\s+", " ", str(t or "")).strip()
        if not t or len(t) > 80:
            continue
        k = t.lower()
        if k in seen:
            continue
        seen.add(k)
        clean.append(t)
    # If we already confidently know the candidate's primary role family, drop
    # any AI-suggested title that clearly maps to a *different* known family.
    # Titles with no detectable family (novel/niche roles) are kept as-is,
    # since the static bank has no opinion on them either.
    if primary_families:
        filtered = []
        for t in clean:
            t_fams = _lead_families_from_text(t, max_families=2)
            if not t_fams or set(t_fams).intersection(primary_families):
                filtered.append(t)
        clean = filtered
    return clean[:max_titles], usage, ""



@app.route("/lead-finder/find-people", methods=["POST"])
def lead_finder_find_people():
    """Second phase: find likely HR/hiring managers from already-found job leads."""
    try:
        body = request.get_json(force=True, silent=True) or {}
        if not _lead_finder_lock_allowed(body):
            return _lead_finder_locked_response()
        api_key = _resolve_request_api_key(body, "lead")
        model = body.get("model") or "claude-sonnet-4-6"
        llm_provider = (body.get("provider") or "anthropic").strip().lower()
        companies = body.get("companies") or []
        regions = body.get("regions") or []
        target_role = (body.get("target_role") or "").strip()
        candidate_context = (body.get("candidate_context") or "").strip()
        cv_text = (body.get("cv_text") or "").strip()

        if not api_key:
            return jsonify({"error": "API key required"}), 400
        if not isinstance(companies, list) or not companies:
            return jsonify({"error": "Job leads required. Run Find Leads first."}), 400
        if not isinstance(regions, list):
            regions = []

        # Keep this phase targeted, but not too narrow. Earlier builds were capped
        # around five companies / twelve people, which made valid HR and hiring
        # manager discovery look empty for common public LinkedIn searches.
        # The UI lets users select exact job leads first; this backend
        # now handles up to eight selected leads in one run.
        companies_lite = companies[:8]
        cv_excerpt = re.sub(r"[ 	]{2,}", " ", cv_text or "").strip()[:4500]
        role_terms = []
        if target_role:
            role_terms.append(target_role)
        for c in companies_lite:
            if isinstance(c, dict):
                for k in ("matched_role", "hiring_signal", "industry"):
                    v = str(c.get(k) or "").strip()
                    if v and v not in role_terms:
                        role_terms.append(v)
        role_terms_text = "; ".join(role_terms[:8]) or "infer functional hiring-manager titles from the matched job roles"
        role_specific_titles = _lead_role_specific_title_bank(target_role, role_terms, cv_text)
        role_specific_titles_text = "\n".join(f"- {t}" for t in role_specific_titles) or "- Infer role-specific hiring manager titles from the matched role and CV"
        # Safe fallback payload used when web/API search times out or returns malformed JSON.
        # v21.49 regression: this was referenced before assignment in the People search route.
        graceful_json = {"people": [], "title_search_angles": role_specific_titles}
        people_title_bank = """
HR / Talent Acquisition title bank:
- Talent Acquisition Specialist, Senior Talent Acquisition Specialist, Talent Acquisition Partner, Talent Acquisition Manager, Regional Talent Acquisition Manager
- Recruiter, Technical Recruiter, Senior Recruiter, Corporate Recruiter, Recruitment Consultant, Sourcing Specialist
- HR Executive, HR Specialist, HR Manager, Senior HR Manager, HR Business Partner, Senior HR Business Partner, People Partner
- People Operations Manager, People & Culture Manager, Head of People, Head of Talent Acquisition, Talent Lead
- Employer Branding / Campus Recruitment / Recruitment Lead / Staffing Lead

Functional hiring-manager title bank:
- Hiring Manager, Reporting Manager, Line Manager, Functional Lead, Team Lead, Principal Lead, Manager, Senior Manager, Associate Director, Director, VP, SVP, Head of, General Manager
- HR roles: HR Manager, HR Director, Head of HR, Senior HRBP, People & Culture Manager, Head of People, Talent Manager, L&D Manager, C&B Manager, Employee Relations Manager
- Operations roles: Operations Manager, Head of Operations, Operations Director, COO, Site Manager, Plant Manager, Production Manager, Supply Chain Manager, Logistics Manager, Procurement Manager, Quality Manager, Customer Service Manager
- Finance roles: Finance Manager, Financial Controller, Head of Finance, Finance Director, CFO, FP&A Manager, Treasury Manager, Tax Manager, Accounting Manager, Shared Services Manager, Finance Operations Manager
- Marketing roles: Marketing Manager, Head of Marketing, Marketing Director, CMO, Brand Manager, Digital Marketing Manager, Performance Marketing Manager, Growth Lead, Communications Manager, CRM Manager, Ecommerce Manager
- Sales/commercial roles: Sales Manager, Head of Sales, Commercial Manager, Commercial Director, Business Development Manager, Key Account Manager, Country Manager, Customer Success Manager, Revenue Operations Manager
- Legal/compliance/risk roles: Head of Legal, Legal Counsel, General Counsel, Compliance Manager, Risk Manager, AML Manager, KYC Manager, Governance Manager, Data Protection Officer
- Admin/corporate services roles: Admin Manager, Office Manager, Business Support Manager, Facilities Manager, Corporate Services Manager
- Product/CX/design roles: Product Manager, Head of Product, Product Director, UX Manager, Design Manager, Customer Experience Manager, Service Design Lead
- Technology/data roles: Head of Engineering, Engineering Manager, Head of Data, Data Engineering Manager, Analytics Manager, Head of Technology, CTO, CIO, IT Director, DevOps Manager, SRE Manager, Application Support Manager, Production Support Manager
- Project/transformation roles: Programme Manager, Project Director, PMO Lead, Transformation Lead, Delivery Director, Business Change Manager
""".strip()

        prompt = f"""
You are performing compliant PUBLIC people-lead discovery for recruiting outreach review.

TASK
- Use the input job leads below.
- For each company, search PUBLIC web/profile results for likely hiring contacts.
- Prioritize public LinkedIn profile/search-result evidence, company team pages, press/interview pages, and public HR/recruiter pages.
- LinkedIn is allowed only through public search result/profile URL evidence. Do not log in, scrape protected pages, bypass access controls, or extract hidden/private data.
- Do NOT search or return emails in this phase. Public email lookup is a separate button.
- Return real people only when a public result/snippet/page supports name + title/company.

VERY IMPORTANT DISCOVERY STRATEGY
For EACH job lead, run focused public searches similar to these patterns:
1) site:linkedin.com/in "<Company>" ("Talent Acquisition" OR Recruiter OR "Technical Recruiter" OR "HR Business Partner" OR "People Partner")
2) site:linkedin.com/in "<Company>" ("Hiring Manager" OR "Head of" OR Manager OR Director OR Lead) plus the matched role/function
3) site:linkedin.com/in "<Company>" ("Head of Data" OR "Engineering Manager" OR "IT Manager" OR "Technology Director" OR "Application Support Manager") adjusted to the matched role
4) "<Company>" "Talent Acquisition" LinkedIn
5) "<Company>" "<matched role/function>" "LinkedIn"
6) "<Company>" "HR Business Partner" OR "People Partner" OR Recruiter
7) "<Company>" "Head of" "<matched function>" OR "Manager" "<matched function>"
8) For niche roles and non-tech roles, use the role-specific title angles. Examples: SAP FICO -> Head of SAP / SAP FICO Manager / SAP Solution Architect / Finance Systems Manager. Data Engineer -> Head of Data Engineering / Data Engineering Manager / Director of Data Analytics. HR Manager -> HR Director / Head of HR / People & Culture Manager. Finance Manager -> Financial Controller / Head of Finance / CFO. Marketing -> Head of Marketing / Brand Manager / Growth Lead. Operations -> Head of Operations / COO / Supply Chain Manager.

COMMON HR / HIRING-MANAGER TITLE BANK
Use these titles as search angles and variations. Do not return fake people; use them to search public evidence.
{people_title_bank}

ROLE-SPECIFIC HIRING-MANAGER TITLE ANGLES FOR THIS CV/JOB
These were generated from the uploaded CV, target role, matched roles and industry signals. Prioritize them when looking for functional hiring managers.
{role_specific_titles_text}

FUNCTIONAL HIRING-MANAGER TARGETS
Use the candidate/job role to infer likely managers across ALL functions, not only technology. Examples:
- HR roles: HR Director, Head of HR, People & Culture Manager, Senior HRBP, Talent Manager.
- Operations roles: Head of Operations, Operations Manager, COO, Plant Manager, Supply Chain Manager.
- Finance roles: Financial Controller, Head of Finance, Finance Director, CFO, FP&A Manager.
- Marketing roles: Marketing Manager, Head of Marketing, Brand Manager, Digital Marketing Lead, Growth Lead.
- Sales/commercial roles: Head of Sales, Commercial Director, Business Development Manager, Country Manager.
- Legal/compliance/risk roles: Head of Legal, General Counsel, Compliance Manager, Risk Manager.
- Data/technology roles: Head of Data, Engineering Manager, IT Director, Application Support Manager.
- Project/transformation roles: Programme Manager, Project Director, PMO Lead, Transformation Lead.

ROLE/FUNCTION TERMS FOR THIS RUN
{role_terms_text}

STRICT RULES
- Do NOT invent or guess names.
- Do NOT invent titles, profile URLs, or emails.
- A public LinkedIn result/profile URL is acceptable evidence if the result/snippet clearly shows the person's name, company and/or relevant title.
- Do not require the person to be based in the selected country. The job lead was already location-filtered; a regional HR or hiring manager can still be valid.
- If a person is relevant but profile_url is unavailable, still return the person if source_url or notes explain the public evidence.
- Prefer 1 HR/TA contact and 1 functional hiring manager per company when available.

REGIONS
{json.dumps(regions, ensure_ascii=False)}

TARGET ROLE
{target_role or "Infer from company matched roles and CV"}

RECRUITER CONTEXT
{candidate_context[:1800] or "No extra context provided."}

CANDIDATE CV EXCERPT
{cv_excerpt}

COMPANY/JOB LEADS
{json.dumps(companies_lite, ensure_ascii=False)}

RETURN ONLY VALID JSON. No markdown. Shape exactly:
{{
  "people": [
    {{
      "id": "pe_1",
      "name": "Person name",
      "title": "Job title",
      "company": "Company name",
      "country": "Country if known, otherwise empty string",
      "contact_type": "Hiring Manager / HR / Talent Acquisition / Department Head",
      "profile_url": "public LinkedIn/profile/company/team URL or empty string",
      "email": "",
      "email_confidence": "",
      "email_source": "",
      "verification_status": "Not found",
      "source_url": "best source URL or empty string",
      "notes": "short public-evidence note explaining why this person is relevant",
      "status": "New"
    }}
  ]
}}

TARGET COUNT
- people: up to 24 relevant contacts across the job leads.
- Aim for 2-4 people per company where public evidence exists: at least one HR/TA/recruiter and one functional hiring manager when possible.
- If no person can be publicly identified for a company, skip that company rather than returning a fake person.
""".strip()

        people_max_uses = 8 if len(companies_lite) >= 6 else (6 if len(companies_lite) >= 3 else 5)
        payload = {
            "model": model,
            "max_tokens": 6200,
            "messages": [{"role": "user", "content": prompt}],
            "tools": [{"type": "web_search_20250305", "name": "web_search", "max_uses": people_max_uses}],
            "_primary_timeout_seconds": 92,
            "_fallback_timeout_seconds": 18,
            # If the selected provider cannot browse (DeepSeek), return the safe
            # empty people result instead of asking a no-web model to invent
            # public contacts.
            "_skip_no_web_fallback": True,
        }
        data, warning = _lead_call_with_optional_web(api_key, payload, graceful_json=graceful_json, provider=llm_provider)
        raw_text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text" or "text" in b)
        parse_warning = ""
        try:
            parsed = _lead_extract_json(raw_text)
        except Exception:
            parsed = json.loads(json.dumps(graceful_json, ensure_ascii=False))
            parse_warning = "API returned malformed/non-JSON people-search output; returned a safe empty people result instead of failing."
        people_filtered, location_warning = _lead_filter_people_by_input_companies(parsed.get("people") or [], companies_lite)
        usage = data.get("usage", {})
        out = {
            "ok": True,
            "people": people_filtered,
            "title_search_angles": role_specific_titles,
            "usage": usage,
            "model": model,
            "provider": llm_provider,
            "cost": _lead_cost(model, usage, llm_provider),
            "cost_details": _lead_cost_details(model, usage, llm_provider),
            "phase": "people_search",
        }
        if not people_filtered and not parse_warning:
            zero_hint = "No public people contacts were returned. Try selecting 1-3 job leads with clear company names/job URLs, or run again because public profile search can be inconsistent."
            warning = (warning + " " + zero_hint).strip() if warning else zero_hint
        if parse_warning:
            warning = (warning + " " + parse_warning).strip() if warning else parse_warning
        if location_warning:
            warning = (warning + " " + location_warning).strip() if warning else location_warning
        if warning:
            out["warning"] = warning
        return jsonify(out)

    except urllib.error.HTTPError as e:
        try:
            err = json.loads(e.read())
            msg = err.get("error", {}).get("message", "API error")
        except Exception:
            msg = "API error"
        out = {"error": "API provider error: " + msg}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage"),
            attempted=True,
        ))
        return jsonify(out), 400
    except (TimeoutError, socket.timeout, urllib.error.URLError) as e:
        out = {"error": "People search could not reach the API provider. Try again after saving job leads."}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage"),
            attempted=True,
        ))
        return jsonify(out), 504
    except Exception as e:
        traceback.print_exc()
        out = {"error": str(e)}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage"),
            attempted=bool(locals().get("api_key") and locals().get("companies")),
        ))
        return jsonify(out), 500


@app.route("/lead-finder/find-emails", methods=["POST"])
def lead_finder_find_emails():
    """Second-pass business email enrichment for people leads.

    Order of operations, each one skipping the next when it already has an answer:
    1) Contact cache (free, instant) — skip anyone already looked up before.
    2) Apollo People Match API, if configured — no Anthropic call needed at all.
    3) Fallback: Claude + web search for explicitly public/published emails only.
    Only people in `selected_person_ids` are processed when that list is
    non-empty, so quota (Apollo credits or Anthropic tokens) is not spent on
    contacts the recruiter has not actually chosen to enrich.
    """
    try:
        ai_request_started = False
        apollo_request_started = False
        body = request.get_json(force=True, silent=True) or {}
        if not _lead_finder_lock_allowed(body):
            return _lead_finder_locked_response()
        api_key = _resolve_request_api_key(body, "lead")
        model = body.get("model") or "claude-sonnet-4-6"
        llm_provider = (body.get("provider") or "anthropic").strip().lower()
        people = body.get("people") or []
        companies = body.get("companies") or []
        regions = body.get("regions") or []
        enrichment_provider = body.get("enrichment_provider") or {}
        provider_name = str(enrichment_provider.get("provider") or "none").strip().lower()
        provider_api_key = str(enrichment_provider.get("api_key") or "").strip()
        if not provider_api_key or provider_api_key == "__BACKEND_SECURE__":
            provider_api_key = str(_ai_secret_store.get("enrichment_apollo") or "").strip()
        selected_ids = {str(x).strip() for x in (body.get("selected_person_ids") or []) if str(x).strip()}

        if not isinstance(people, list) or not people:
            return jsonify({"error": "People leads required"}), 400

        target_people = [p for p in people if isinstance(p, dict) and (not selected_ids or str(p.get("id") or "").strip() in selected_ids)]
        if not target_people:
            return jsonify({"error": "No matching selected people found. Select at least one person, or clear the selection to enrich everyone."}), 400

        # --- Step 1: cache lookup (applies regardless of which method fills misses) ---
        cache_hits = 0
        already_resolved = 0
        need_lookup = []
        for p in target_people:
            # Never re-attempt someone who already has a real email on file (from
            # an earlier run, whichever method found it). A failed lookup this
            # time must not overwrite previously-found good data with "Not found".
            if str(p.get("email") or "").strip() and str(p.get("verification_status") or "").strip().lower() not in ("", "not found"):
                already_resolved += 1
                continue
            cached = _lead_contact_cache_find(p)
            if cached and _lead_is_company_domain_email(cached.get("email", ""), p.get("company") or ""):
                p["email"] = cached.get("email", "")
                p["email_confidence"] = cached.get("email_confidence", "")
                p["email_source"] = cached.get("email_source", "")
                p["verification_status"] = cached.get("verification_status", "")
                p["email_lookup_method"] = "cache"
                _lead_contact_cache_touch(p)
                cache_hits += 1
            else:
                need_lookup.append(p)

        apollo_warning = ""
        apollo_looked_up = 0
        apollo_found = 0

        # --- Step 2: Apollo, if configured, for whatever the cache did not answer ---
        if provider_name == "apollo" and provider_api_key and need_lookup:
            stop_early = ""
            for p in need_lookup:
                try:
                    apollo_request_started = True
                    apollo_looked_up += 1
                    result = _lead_apollo_enrich_person(provider_api_key, p)
                except _LeadApolloAuthError as e:
                    stop_early = str(e)
                    break
                except _LeadApolloRateLimited as e:
                    stop_early = str(e) + " Remaining contacts in this batch were not looked up; try again later or select fewer people per run."
                    break
                except Exception as e:
                    # One person's lookup failing should not abort the whole batch.
                    result = {"email": "", "email_confidence": "", "email_source": "", "verification_status": "Not found"}
                    apollo_warning = (apollo_warning + f" Apollo lookup failed for {p.get('name') or 'one contact'}: {str(e)[:150]}.").strip()
                if result.get("email"):
                    apollo_found += 1
                p["email"] = result.get("email", "")
                p["email_confidence"] = result.get("email_confidence", "")
                p["email_source"] = result.get("email_source", "")
                p["verification_status"] = result.get("verification_status", "Not found")
                p["email_lookup_method"] = "apollo"
                _lead_contact_cache_store(p, result)
            if stop_early:
                apollo_warning = (apollo_warning + " " + stop_early).strip()

            usage = {"input_tokens": 0, "output_tokens": 0}
            summary_bits = [f"{already_resolved} already had an email"] if already_resolved else []
            if cache_hits:
                summary_bits.append(f"{cache_hits} from cache")
            summary_bits.append(f"{apollo_looked_up} looked up via Apollo ({apollo_found} found)")
            warning = ("Email enrichment: " + ", ".join(summary_bits) + ".").strip()
            if apollo_warning:
                warning = (warning + " " + apollo_warning).strip()
            people = _lead_sanitize_public_business_emails(people)
            return jsonify({
                "ok": True,
                "people": people,
                "usage": usage,
                "model": model,
                "provider": llm_provider,
                "cost": 0,
                "cost_details": _lead_cost_details(model, usage, llm_provider),
                "apollo_looked_up": apollo_looked_up,
                "apollo_found": apollo_found,
                "cache_hits": cache_hits,
                "already_resolved": already_resolved,
                "external_billing": _phase5b_unavailable_external_billing(
                    "apollo",
                    observed_operations=apollo_looked_up,
                    reason=(
                        "Apollo enrichment responses do not include "
                        "provider-authoritative billing."
                    ),
                ),
                "warning": warning,
            })

        # --- Step 3: fall back to Claude + public web search for anything not resolved above ---
        if not need_lookup:
            # Everything was already answered by existing data/cache; no Anthropic call needed.
            summary_bits = [f"{already_resolved} already had an email"] if already_resolved else []
            if cache_hits:
                summary_bits.append(f"{cache_hits} from cache")
            summary = ("Email enrichment: " + ", ".join(summary_bits) + ". No new lookups were needed.") if summary_bits else "Email enrichment: nothing new to look up."
            people = _lead_sanitize_public_business_emails(people)
            return jsonify({
                "ok": True, "people": people, "usage": {"input_tokens": 0, "output_tokens": 0},
                "model": model, "provider": llm_provider, "cost": 0, "cost_details": _lead_cost_details(model, {}, llm_provider),
                "cache_hits": cache_hits, "already_resolved": already_resolved, "warning": summary,
            })

        if not api_key:
            return jsonify({"error": "API key required for AI-based email search (or configure an Apollo API key instead)."}), 400

        people_lite = need_lookup[:15]
        companies_lite = companies[:15] if isinstance(companies, list) else []
        # If email enrichment times out, preserve the full existing people list.
        # The prompt still sends a capped batch, but graceful fallback must not
        # silently drop contacts beyond the first batch.
        graceful_json = {"people": people if isinstance(people, list) else people_lite}

        prompt = f"""
You are performing compliant public business email enrichment for recruiting outreach review.

STRICT RULES
- Do NOT scrape LinkedIn, logged-in LinkedIn pages, LinkedIn profile pages, or any site that requires login/CAPTCHA/rate-limit bypass.
- Treat LinkedIn/profile URLs as identity context only: name, title, company. Never extract emails from LinkedIn pages.
- Do NOT guess, infer, pattern-generate, or invent email addresses.
- Only return an email when it is explicitly public/business-sourced from non-LinkedIn business sources such as company websites, official team/staff pages, conference/speaker pages, press pages, or business directories.
- Return company-domain/business-domain emails only. Do NOT return personal/free-mail addresses such as Gmail, Yahoo, Hotmail, Outlook, iCloud, Proton, or similar domains.
- Keep every input person, preserving name/title/company/contact_type/profile_url/source_url/status.
- If no public business email is found, set email to empty string, email_confidence to empty string, email_source to empty string, verification_status to "Not found".
- If an email is found, set verification_status to "Public business email" or "Needs review" and provide email_source.

INPUT PEOPLE
{json.dumps(people_lite, ensure_ascii=False)}

INPUT COMPANIES
{json.dumps(companies_lite, ensure_ascii=False)}

RETURN ONLY VALID JSON. No markdown. Shape exactly:
{{
  "people": [
    {{
      "id": "same id",
      "name": "same name",
      "title": "same title",
      "company": "same company",
      "country": "same country",
      "contact_type": "same contact type",
      "profile_url": "same or improved public profile URL",
      "email": "public business email only or empty",
      "email_confidence": "High/Medium/Low or empty",
      "email_source": "URL/evidence or empty",
      "verification_status": "Public business email / Needs review / Not found",
      "source_url": "source URL or empty",
      "notes": "brief note",
      "status": "same status or New"
    }}
  ]
}}
""".strip()

        payload = {
            "model": model,
            "max_tokens": 3400,
            "messages": [{"role": "user", "content": prompt}],
            "tools": [{"type": "web_search_20250305", "name": "web_search", "max_uses": 1}],
            "_primary_timeout_seconds": 45,
            "_fallback_timeout_seconds": 18,
            # If the selected provider cannot browse (DeepSeek), preserve existing
            # people and ask the user to use Apollo/Claude/GPT instead of letting
            # a no-web model guess public business emails.
            "_skip_no_web_fallback": True,
        }
        ai_request_started = True
        data, warning = _lead_call_with_optional_web(api_key, payload, graceful_json=graceful_json, provider=llm_provider)
        raw_text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text" or "text" in b)
        parse_warning = ""
        try:
            parsed = _lead_extract_json(raw_text)
        except Exception:
            parsed = json.loads(json.dumps(graceful_json, ensure_ascii=False))
            parse_warning = "API returned malformed/non-JSON email-search output; preserved existing people leads instead of failing."
        # Email enrichment only sends a capped batch to control cost/time, so a
        # successful model response may contain fewer people than the full table.
        # Merge returned enrichment fields back into the full original list instead
        # of replacing/dropping contacts beyond the capped batch.
        if isinstance(parsed.get("people"), list):
            returned_people = parsed.get("people") or []
            by_id = {str(p.get("id") or "").strip(): p for p in returned_people if isinstance(p, dict) and str(p.get("id") or "").strip()}
            by_key = {
                (str(p.get("name") or "").strip().lower(), str(p.get("company") or "").strip().lower()): p
                for p in returned_people if isinstance(p, dict)
            }
            merged_people = []
            for original in people if isinstance(people, list) else []:
                if not isinstance(original, dict):
                    continue
                rid = str(original.get("id") or "").strip()
                key = (str(original.get("name") or "").strip().lower(), str(original.get("company") or "").strip().lower())
                enriched = by_id.get(rid) or by_key.get(key)
                item = dict(original)
                if isinstance(enriched, dict):
                    for k, v in enriched.items():
                        # Preserve original identity fields unless the enrichment improves blank values.
                        if k in ("id", "name", "title", "company", "country", "contact_type") and item.get(k):
                            continue
                        item[k] = v
                    item["email_lookup_method"] = "ai_public_search"
                if not item.get("verification_status"):
                    item["verification_status"] = "Public business email" if item.get("email") else "Not found"
                item = _lead_sanitize_public_business_emails([item])[0]
                if item.get("email") and item.get("email_lookup_method") == "ai_public_search":
                    _lead_contact_cache_store(item, item)
                merged_people.append(item)
            parsed["people"] = merged_people
        else:
            parsed["people"] = people if isinstance(people, list) else []
        parsed["people"] = _lead_sanitize_public_business_emails(parsed.get("people") or [])
        # Email enrichment should preserve the current people list; do not drop
        # contacts here just because an email-source response omitted location.
        location_warning = ""
        usage = data.get("usage", {})
        parsed["ok"] = True
        parsed["usage"] = usage
        cache_bits = ([f"{already_resolved} already had an email"] if already_resolved else []) + ([f"{cache_hits} from cache"] if cache_hits else [])
        cache_summary = (", ".join(cache_bits) + f", {len(people_lite)} looked up via AI public search.") if cache_bits else ""
        if cache_summary:
            warning = (cache_summary + " " + warning).strip() if warning else cache_summary
        if parse_warning:
            warning = (warning + " " + parse_warning).strip() if warning else parse_warning
        if location_warning:
            warning = (warning + " " + location_warning).strip() if warning else location_warning
        parsed["cache_hits"] = cache_hits
        parsed["model"] = model
        parsed["provider"] = llm_provider
        parsed["cost"] = _lead_cost(model, usage, llm_provider)
        parsed["cost_details"] = _lead_cost_details(model, usage, llm_provider)
        if warning:
            parsed["warning"] = warning
        return jsonify(parsed)

    except urllib.error.HTTPError as e:
        try:
            err = json.loads(e.read())
            msg = err.get("error", {}).get("message", "API error")
        except Exception:
            msg = "API error"
        out = {"error": "API provider error: " + msg}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage"),
            attempted=bool(
                locals().get("ai_request_started")
                or _llm_usage_int(locals().get("usage"), "api_calls")
            ),
        ))
        if (
            locals().get("provider_name") == "apollo"
            and locals().get("apollo_request_started")
        ):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                "apollo",
                observed_operations=locals().get("apollo_looked_up"),
                reason=(
                    "The failed Apollo workflow returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 400
    except (TimeoutError, socket.timeout, urllib.error.URLError) as e:
        safe_people = _lead_sanitize_public_business_emails(people if isinstance(people, list) else []) if 'people' in locals() else []
        out = {"ok": True, "people": safe_people, "warning": "Email search timed out before enrichment finished; preserved existing people leads without adding emails.", "usage": {"input_tokens": 0, "output_tokens": 0}, "model": model if 'model' in locals() else "claude-sonnet-4-6", "provider": llm_provider if 'llm_provider' in locals() else "anthropic", "cost": 0, "cost_details": _lead_cost_details(model if 'model' in locals() else "claude-sonnet-4-6", {}, llm_provider if 'llm_provider' in locals() else "anthropic")}
        out.update(_llm_paid_failure_fields(
            e,
            out["model"],
            out["provider"],
            attempted=bool(
                locals().get("ai_request_started")
                or _llm_usage_int(locals().get("usage"), "api_calls")
            ),
        ))
        if (
            locals().get("provider_name") == "apollo"
            and locals().get("apollo_request_started")
        ):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                "apollo",
                observed_operations=locals().get("apollo_looked_up"),
                reason=(
                    "The timed-out Apollo workflow returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out)
    except Exception as e:
        traceback.print_exc()
        out = {"error": str(e)}
        out.update(_llm_paid_failure_fields(
            e,
            locals().get("model", ""),
            locals().get("llm_provider", "anthropic"),
            locals().get("usage"),
            attempted=bool(
                locals().get("ai_request_started")
                or _llm_usage_int(locals().get("usage"), "api_calls")
            ),
        ))
        if (
            locals().get("provider_name") == "apollo"
            and locals().get("apollo_request_started")
        ):
            out["external_billing"] = _phase5b_unavailable_external_billing(
                "apollo",
                observed_operations=locals().get("apollo_looked_up"),
                reason=(
                    "The Apollo workflow returned no "
                    "provider-authoritative billing."
                ),
            )
        return jsonify(out), 500


@app.route("/lead-finder/contact-cache/stats", methods=["GET"])
def lead_finder_contact_cache_stats():
    try:
        data = _lead_contact_cache_load()
        entries = data.get("entries", {})
        found = sum(1 for e in entries.values() if e.get("email"))
        return jsonify({"ok": True, "total_entries": len(entries), "with_email": found})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)[:500]}), 500

@app.route("/lead-finder/contact-cache/clear", methods=["POST"])
def lead_finder_contact_cache_clear():
    try:
        _lead_contact_cache_save({"entries": {}})
        return jsonify({"ok": True, "message": "Contact cache cleared."})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)[:500]}), 500


def _normalize_cv_text_alignment(value):
    """Return the only supported Word paragraph alignment values.

    Alignment is a deterministic export preference, not AI-generated content.
    Any missing, stale or malformed browser value safely falls back to left.
    """
    return "justify" if str(value or "").strip().lower() == "justify" else "left"


@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    data_path = None
    out_path = None
    try:
        body = request.get_json(force=True, silent=True)
        if not body:
            return jsonify({"error": "Invalid JSON body"}), 400

        cv_data = body.get("data")
        if not cv_data:
            return jsonify({"error": "No CV data provided"}), 400
        # Outline-label nesting is inferred from the raw bullet text, so compute
        # it BEFORE _normalize_cv_structured_content strips the labels.
        _label_levels = _infer_label_bullet_levels(cv_data)
        cv_data = _normalize_cv_structured_content(cv_data)
        cv_data = _normalize_cv_data_for_output(cv_data)
        cv_data["_document_alignment"] = _normalize_cv_text_alignment(body.get("alignment"))
        # Nested-list indent map: the source-derived map (docx w:ilvl / pdf
        # geometry, sent by the browser) merged with the outline-label inference,
        # deepest level wins. Unmatched bullets stay at level 0.
        _levels = body.get("bullet_levels")
        if not isinstance(_levels, list):
            _levels = cv_data.get("_bullet_levels")
        cv_data["_bullet_levels"] = _merge_bullet_level_lists(
            _levels if isinstance(_levels, list) else [],
            _label_levels,
        )

        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
            json.dump(cv_data, f)
            data_path = f.name

        out_path = data_path.replace(".json", ".docx")
        script_dir = os.path.dirname(os.path.abspath(__file__))
        script = os.path.join(script_dir, "generate.js")

        result = subprocess.run(
            ["node", script, data_path, out_path],
            capture_output=True, text=True, timeout=30
        )

        if result.returncode != 0:
            err = _summarize_node_failure(result.stdout, result.stderr)
            return jsonify({"error": err}), 500

        # Read the completed document into memory before returning it. Flask may
        # stream a path after the route's finally block; using BytesIO lets us
        # delete both temporary files deterministically without breaking the
        # download response.
        with open(out_path, "rb") as handle:
            document_bytes = handle.read()

        return send_file(
            io.BytesIO(document_bytes),
            as_attachment=True,
            download_name="Hyppies_CV_Formatted.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        for temporary_path in (data_path, out_path):
            if not temporary_path:
                continue
            try:
                os.unlink(temporary_path)
            except FileNotFoundError:
                pass
            except Exception:
                pass


def _cv_bullet_match_key(text):
    """Normalize a bullet for source<->parsed matching, mirrored in generate.js.

    Strips the same leading marker/enumerator the formatter strips for display
    (_CV_LEADING_BULLET_MARKER_RE) and lowercases, so a level looked up from the
    raw source bullet matches the rendered (already-stripped) bullet text.
    """
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    for _ in range(5):
        stripped = _CV_LEADING_BULLET_MARKER_RE.sub("", value, count=1).strip()
        if stripped == value:
            break
        value = stripped
    return value.lower()


_ROMAN_RE = re.compile(r"^[ivxlcdm]+$", re.I)
_LABEL_PAREN_RE = re.compile(r"^\((?P<body>[ivxlcdm]{1,7}|[a-z]|\d{1,3})\)", re.I)
_LABEL_NUM_RE = re.compile(
    r"^(?P<body>\d{1,3})(?:[.)](?=\s|[^\W\d_])|-(?=\s))"
)
_LABEL_BARE_RE = re.compile(
    r"^(?P<body>(?:[a-z]|[ivxlcdm]{2,7}))(?P<punct>[.)-])(?=\s)"
)


def _label_body_kind(body):
    if body.isdigit():
        return "digit"
    if len(body) > 1 and _ROMAN_RE.match(body):
        return "roman"
    return "alpha"


def _bullet_label_style(text):
    """Classify a leading outline label into a style key, or None.

    Restricted to the label forms the formatter actually strips
    (_CV_LEADING_BULLET_MARKER_RE). Keeping classification and stripping in
    lockstep means a nested item's label is both removed AND counted -- never one
    without the other. A plain glyph or unlabeled line returns None (base level).
    """
    s = str(text or "").lstrip()
    m = _LABEL_PAREN_RE.match(s)
    if m:
        return ("paren", _label_body_kind(m.group("body")))
    m = _LABEL_NUM_RE.match(s)
    if m:
        return ("num", "digit")
    m = _LABEL_BARE_RE.match(s)
    if m:
        return ("bare", _label_body_kind(m.group("body")))
    return None


def _infer_outline_levels(texts):
    """Infer a nesting level per bullet from the sequence of outline labels.

    Appearance-order stack: an unlabeled line is level 0 and resets the outline;
    the first labelled style below it is level 1; a new deeper style nests one
    level further; returning to a seen style pops back to that level. Mirrors
    cvInferLevelsForSequence in index.html.
    """
    levels = []
    stack = []
    for text in texts:
        style = _bullet_label_style(text)
        if style is None:
            stack = []
            levels.append(0)
        elif style in stack:
            idx = stack.index(style)
            del stack[idx + 1:]
            levels.append(idx + 1)
        else:
            stack.append(style)
            levels.append(len(stack))
    return levels


def _infer_label_bullet_levels(cv_data):
    """Outline-label nesting levels for a parsed CV's role bullets.

    Runs _infer_outline_levels over each role's bullet sequence (raw text, before
    labels are stripped) and returns [{text: match_key, level}] for levels >= 1.
    """
    out = []
    for exp in (cv_data or {}).get("work_experiences") or []:
        if not isinstance(exp, dict):
            continue
        for role in exp.get("roles") or []:
            if not isinstance(role, dict):
                continue
            texts = []
            for item in role.get("bullets") or []:
                if isinstance(item, str):
                    texts.append(item)
                elif isinstance(item, dict):
                    for sub in item.get("bullets") or []:
                        if isinstance(sub, str):
                            texts.append(sub)
            for text, level in zip(texts, _infer_outline_levels(texts)):
                if level >= 1:
                    key = _cv_bullet_match_key(text)
                    if key:
                        out.append({"text": key, "level": min(level, 8)})
    return out


def _merge_bullet_level_lists(*lists):
    """Merge {text, level} lists by key, keeping the deepest level (>= 1)."""
    best = {}
    for entries in lists:
        for entry in entries or []:
            if not isinstance(entry, dict):
                continue
            key = str(entry.get("text") or "")
            if not key:
                continue
            try:
                level = int(entry.get("level") or 0)
            except (TypeError, ValueError):
                continue
            level = max(0, min(level, 8))
            if key not in best or level > best[key]:
                best[key] = level
    return [{"text": k, "level": v} for k, v in best.items() if v >= 1]


def _extract_docx_bullet_levels(file_bytes):
    """Map nested (w:ilvl >= 1) list paragraphs to their indent level.

    Word stores list nesting in w:numPr/w:ilvl, which the plain-text extraction
    discards. This side-channel lets the renderer restore the indent
    deterministically by matching parsed bullet text back to its source level.
    Only levels >= 1 are recorded; every unmatched bullet renders at level 0.
    """
    import io as _io
    import docx as _docx
    from docx.oxml.ns import qn
    out = []
    seen = set()
    try:
        doc = _docx.Document(_io.BytesIO(file_bytes))
    except Exception:
        return out
    for p in doc.element.body.iter(qn("w:p")):
        ilvl_vals = p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val")
        if not ilvl_vals:
            continue
        try:
            level = int(ilvl_vals[0])
        except (TypeError, ValueError):
            continue
        if level < 1:
            continue
        text = "".join(node.text or "" for node in p.iter(qn("w:t"))).strip()
        key = _cv_bullet_match_key(text)
        if not key:
            continue
        dedup = (key, level)
        if dedup in seen:
            continue
        seen.add(dedup)
        out.append({"text": key, "level": min(level, 8)})
    return out


_PDF_BULLET_LEADIN_RE = re.compile(
    r"^(?:[•●▪◦‣⁃∙·‧*]|[-–—]\s|\(?[0-9]{1,3}[.)\-]\s|\(?[a-z]{1,3}[.)\-]\s)",
    re.I,
)


def _pdf_bullet_levels_from_lines(lines):
    """Assign indent levels to marker-led lines from their x-positions.

    ``lines`` is ``[(x0, text)]`` in document order. The dominant left column of
    marker-led lines is level 0; deeper columns become levels by their indent
    step. Conservative: needs >= 4 marker-led lines and a clear indent step,
    otherwise returns [] (everything flat). Pure -- unit-testable without a PDF.
    """
    from collections import Counter
    threshold = 14
    tol = 6
    min_step = 12
    bullet_lines = [(x0, t) for x0, t in lines if _PDF_BULLET_LEADIN_RE.match(str(t or ""))]
    if len(bullet_lines) < 4:
        return []

    def bucket(x):
        return round(x / tol) * tol

    counts = Counter(bucket(x0) for x0, _ in bullet_lines)
    top_count = max(counts.values())
    base = min(b for b, c in counts.items() if c == top_count)
    offsets = sorted({bucket(x0) - base for x0, _ in bullet_lines if bucket(x0) > base + threshold})
    if not offsets or offsets[0] < min_step:
        return []
    step = offsets[0]
    out = []
    seen = set()
    for x0, t in bullet_lines:
        off = bucket(x0) - base
        if off <= threshold:
            continue
        level = min(max(1, round(off / step)), 8)
        key = _cv_bullet_match_key(t)
        if not key:
            continue
        dedup = (key, level)
        if dedup in seen:
            continue
        seen.add(dedup)
        out.append({"text": key, "level": level})
    return out


def _extract_pdf_bullet_levels(file_bytes):
    """Recover nested-list indent from a PDF using text x-positions."""
    import io as _io
    try:
        import pdfplumber
    except Exception:
        return []
    lines = []
    try:
        with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                by_top = {}
                for word in page.extract_words(use_text_flow=False):
                    by_top.setdefault(round(word["top"] / 3.0), []).append(word)
                for key in sorted(by_top):
                    row = sorted(by_top[key], key=lambda w: w["x0"])
                    text = " ".join(w["text"] for w in row).strip()
                    if text:
                        lines.append((row[0]["x0"], text))
    except Exception:
        return []
    return _pdf_bullet_levels_from_lines(lines)


def _extract_docx_text_preserve_tables(file_bytes):
    _validate_zip_payload(file_bytes, "DOCX")
    """Extract DOCX text, preserving table-cell paragraph/bullet boundaries.

    Corporate mandates and many CVs place multiple responsibilities inside a
    single table cell.  Joining those paragraphs with spaces destroys bullet
    boundaries and lets the JD term extractor keep only the first few concepts.
    Rows containing only one paragraph per cell remain compact `` | `` lines;
    multi-paragraph cells are emitted as separate lines in document order.
    """
    import io as _io
    import docx as _docx
    from docx.table import Table as _DocxTable
    from docx.text.paragraph import Paragraph as _DocxParagraph
    doc = _docx.Document(_io.BytesIO(file_bytes))
    lines = []

    def clean(text):
        return re.sub(r"\s+", " ", (text or "")).strip()

    def cell_lines(cell):
        out = []
        for paragraph in cell.paragraphs:
            value = clean(paragraph.text)
            if value and (not out or value != out[-1]):
                out.append(value)
        return out

    def row_lines(row):
        groups = []
        seen_groups = set()
        for cell in row.cells:
            group = cell_lines(cell)
            key = tuple(group)
            # Merged cells are repeated by python-docx; keep the first copy only.
            if group and key not in seen_groups:
                seen_groups.add(key)
                groups.append(group)
        if not groups:
            return []
        if all(len(group) == 1 for group in groups):
            return [" | ".join(group[0] for group in groups)]
        out = []
        for group in groups:
            for value in group:
                if value and (not out or value != out[-1]):
                    out.append(value)
        return out

    seen_headers = set()
    def append_header(value):
        value = clean(value)
        if value and value not in seen_headers:
            seen_headers.add(value)
            lines.append(value)

    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            try:
                for paragraph in hdr.paragraphs:
                    append_header(paragraph.text)
                for table in getattr(hdr, "tables", []) or []:
                    for row in table.rows:
                        for value in row_lines(row):
                            append_header(value)
            except Exception:
                pass

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            value = clean(_DocxParagraph(child, doc).text)
            if value:
                lines.append(value)
        elif tag == 'tbl':
            table = _DocxTable(child, doc)
            for row in table.rows:
                lines.extend(row_lines(row))
    return "\n".join(lines)

def _normalize_extracted_text_artifacts(text):
    """Clean common PDF/OCR font-encoding artefacts from extracted JD/CV text.

    Some PDFs embed fonts where common ligatures are exposed by pdfplumber as
    placeholders such as ``(cid:415)`` instead of real text. Replacing them with
    spaces gives broken JD text like ``Accoun ng`` / ``Loca on``. Keep this
    cleanup conservative and focused on known extraction artefacts so genuine
    wording and tech stack terms are preserved.
    """
    if not text:
        return text
    import re as _re
    t = str(text)

    # Known glyph-code mappings seen in PDFs generated from Word/Canva-like
    # layouts. These map to common ligature fragments, not company names.
    _cid_map = {
        '332': 'ft',   # so(cid:332)ware -> software, shi(cid:332) -> shift
        '414': 'tf',   # pla(cid:414)orm -> platform
        '415': 'ti',   # Accoun(cid:415)ng -> Accounting, Loca(cid:415)on -> Location
        '425': 'tt',   # Commi(cid:425)ed -> Committed, wri(cid:425)en -> written
        '427': 'tti',  # a(cid:427)tude -> attitude
    }
    def _cid_repl(m):
        return _cid_map.get(m.group(1), ' ')
    t = _re.sub(r'\(?cid:(\d+)\)?', _cid_repl, t, flags=_re.I)

    # Same artefacts sometimes arrive as odd Unicode glyphs instead of cid tags.
    t = t.translate(str.maketrans({
        '\u019f': 'ti',  # Ɵ
        '\u014c': 'ft',  # Ō
        '\u019e': 'tf',  # ƞ
        '\u01a9': 'tt',  # Ʃ
        '\u01ab': 'tti', # ƫ
        '\uf0b7': '•',   # private-use bullet
        '\ufffe': '-',   # malformed soft-hyphen separator seen as buy￾ins
    }))

    # If a previous extraction already replaced cid tags with spaces, recover
    # the most common JD words. Only exact broken-word patterns are touched.
    _word_fixes = {
        r'\bAccoun\s+ng\b': 'Accounting', r'\baccoun\s+ng\b': 'accounting',
        r'\bLoca\s+on\b': 'Location', r'\bloca\s+on\b': 'location',
        r'\bFull\s+me\b': 'Full time', r'\bfull\s+me\b': 'full time',
        r'\bso\s+ware\b': 'software', r'\bSo\s+ware\b': 'Software',
        r'\bpla\s+orm\b': 'platform', r'\bPla\s+orm\b': 'Platform',
        r'\bwri\s+en\b': 'written', r'\bWri\s+en\b': 'Written',
        r'\bCommi\s+ed\b': 'Committed', r'\bcommi\s+ed\b': 'committed',
        r'\ba\s+tude\b': 'attitude', r'\bA\s+tude\b': 'Attitude',
        r'\bcri\s+cal\b': 'critical', r'\bCri\s+cal\b': 'Critical',
        r'\bpoten\s+al\b': 'potential', r'\bPoten\s+al\b': 'Potential',
        r'\bop\s+mize\b': 'optimize', r'\bOp\s+mize\b': 'Optimize',
        r'\bcollec\s+on\b': 'collection', r'\bCollec\s+on\b': 'Collection',
        r'\bsolu\s+on\b': 'solution', r'\bSolu\s+on\b': 'Solution',
        r'\bsolu\s+ons\b': 'solutions', r'\bSolu\s+ons\b': 'Solutions',
        r'\bopera\s+on\b': 'operation', r'\bOpera\s+on\b': 'Operation',
        r'\bopera\s+ons\b': 'operations', r'\bOpera\s+ons\b': 'Operations',
        r'\bini\s+a\s+ves\b': 'initiatives', r'\bIni\s+a\s+ves\b': 'Initiatives',
        r'\biden\s+fy\b': 'identify', r'\bIden\s+fy\b': 'Identify',
        r'\beffec\s+vely\b': 'effectively', r'\bEffec\s+vely\b': 'Effectively',
        r'\bCollabora\s+ve\b': 'Collaborative', r'\bcollabora\s+ve\b': 'collaborative',
        r'\bac\s+vely\b': 'actively', r'\bAc\s+vely\b': 'Actively',
        r'\bconstruc\s+ve\b': 'constructive', r'\bConstruc\s+ve\b': 'Constructive',
        r'\bcon\s+nuous\b': 'continuous', r'\bCon\s+nuous\b': 'Continuous',
        r'\bsuppor\s+ve\b': 'supportive', r'\bSuppor\s+ve\b': 'Supportive',
        r'\bmo\s+vates\b': 'motivates', r'\bMo\s+vates\b': 'Motivates',
        r'\bCoordina\s+on\b': 'Coordination', r'\bcoordina\s+on\b': 'coordination',
        r'\bme\s+culous\b': 'meticulous', r'\bMe\s+culous\b': 'Meticulous',
        r'\bproac\s+ve\b': 'proactive', r'\bProac\s+ve\b': 'Proactive',
        r'\bmul\s+ple\b': 'multiple', r'\bMul\s+ple\b': 'Multiple',
        r'\bpriori\s+es\b': 'priorities', r'\bPriori\s+es\b': 'Priorities',
        r'\bshi\s+ing\b': 'shifting', r'\bShi\s+ing\b': 'Shifting',
        r'\bnight\s+shi\s+\b': 'night shift ', r'\bNight\s+shi\s+\b': 'Night shift ',
        r'\b\s+me\b': ' time',
        r'\bcke\s+ng\b': 'ticketing', r'\bCke\s+ng\b': 'Ticketing',
        r'\bma\s+ers\b': 'matters', r'\bMa\s+ers\b': 'Matters',
        r'\bsa\s+sfac\s+on\b': 'satisfaction', r'\bSa\s+sfac\s+on\b': 'Satisfaction',
    }
    for _pat, _rep in _word_fixes.items():
        t = _re.sub(_pat, _rep, t)

    # Clean spacing without flattening paragraphs.
    t = _re.sub(r'[ \t]{2,}', ' ', t)
    t = _re.sub(r' *\n *', '\n', t)
    t = _re.sub(r'\n{3,}', '\n\n', t).strip()
    return t



def _find_soffice_binary():
    """Return a LibreOffice/soffice binary path when available, else None."""
    import shutil as _shutil
    candidates = [
        _shutil.which('soffice'),
        _shutil.which('libreoffice'),
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        '/usr/local/bin/soffice',
        '/opt/homebrew/bin/soffice',
        '/usr/bin/soffice',
        '/usr/bin/libreoffice',
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return None


def _legacy_doc_word_powershell_candidates():
    """Return Windows PowerShell hosts that may match installed Word's bitness."""
    if os.name != "nt":
        return []
    import shutil as _shutil

    system_root = os.environ.get("SystemRoot") or r"C:\Windows"
    candidates = [
        os.path.join(
            system_root,
            "SysWOW64",
            "WindowsPowerShell",
            "v1.0",
            "powershell.exe",
        ),
        os.path.join(
            system_root,
            "System32",
            "WindowsPowerShell",
            "v1.0",
            "powershell.exe",
        ),
        _shutil.which("powershell.exe"),
        _shutil.which("pwsh.exe"),
    ]
    found = []
    seen = set()
    for candidate in candidates:
        if not candidate or not os.path.isfile(candidate):
            continue
        normalized = os.path.normcase(os.path.realpath(candidate))
        if normalized in seen:
            continue
        seen.add(normalized)
        found.append(candidate)
    return found


def _converter_process_creation_options():
    """Return process-group options used by optional document converters."""
    import subprocess as _subprocess

    if os.name == "nt":
        return {
            "creationflags": int(
                getattr(_subprocess, "CREATE_NO_WINDOW", 0) or 0
            ) | int(
                getattr(_subprocess, "CREATE_NEW_PROCESS_GROUP", 0) or 0
            )
        }
    return {"start_new_session": True}


def _terminate_converter_process_tree(process, *, extra_pids=()):
    """Best-effort termination for a failed converter and its children."""
    import signal as _signal
    import shutil as _shutil
    import subprocess as _subprocess

    try:
        process_running = process.poll() is None
    except Exception:
        process_running = True
    if os.name == "nt":
        system_root = os.environ.get("SystemRoot") or r"C:\Windows"
        taskkill = os.path.join(system_root, "System32", "taskkill.exe")
        if not os.path.isfile(taskkill):
            taskkill = _shutil.which("taskkill.exe") or ""
        pids = []
        process_pids = [getattr(process, "pid", None)] if process_running else []
        for value in list(extra_pids or ()) + process_pids:
            try:
                pid = int(value)
            except (TypeError, ValueError):
                continue
            if pid > 0 and pid not in pids:
                pids.append(pid)
        if taskkill:
            for pid in pids:
                try:
                    _subprocess.run(
                        [taskkill, "/PID", str(pid), "/T", "/F"],
                        stdin=_subprocess.DEVNULL,
                        stdout=_subprocess.DEVNULL,
                        stderr=_subprocess.DEVNULL,
                        timeout=5,
                        check=False,
                        creationflags=int(
                            getattr(_subprocess, "CREATE_NO_WINDOW", 0) or 0
                        ),
                    )
                except Exception:
                    # cleanup-only: the direct process kill below remains.
                    pass
    else:
        try:
            os.killpg(int(process.pid), _signal.SIGKILL)
        except Exception:
            # cleanup-only: fall back to killing the direct process below.
            pass
    if process_running:
        try:
            process.kill()
        except Exception:
            # cleanup-only: the process may already have exited.
            pass
        try:
            process.communicate(timeout=5)
        except Exception:
            # cleanup-only: termination is already best-effort and bounded.
            pass


def _recorded_converter_process_ids(extra_pid_path):
    """Return the positive native PID recorded by a converter wrapper."""
    if not extra_pid_path:
        return []
    try:
        with open(extra_pid_path, "r", encoding="ascii") as handle:
            pid = int(handle.read().strip())
        return [pid] if pid > 0 else []
    except Exception:
        return []


def _run_converter_subprocess(command, *, timeout, extra_pid_path=""):
    """Run one converter and tear down its process tree on failure/timeout."""
    import subprocess as _subprocess

    process = _subprocess.Popen(
        command,
        stdin=_subprocess.DEVNULL,
        stdout=_subprocess.PIPE,
        stderr=_subprocess.PIPE,
        **_converter_process_creation_options(),
    )
    try:
        stdout, stderr = process.communicate(timeout=max(1.0, float(timeout)))
    except Exception:
        _terminate_converter_process_tree(
            process,
            extra_pids=_recorded_converter_process_ids(extra_pid_path),
        )
        raise
    if process.returncode:
        recorded_pids = _recorded_converter_process_ids(extra_pid_path)
        if recorded_pids:
            _terminate_converter_process_tree(
                process,
                extra_pids=recorded_pids,
            )
    return _subprocess.CompletedProcess(
        command,
        process.returncode,
        stdout,
        stderr,
    )


def _convert_legacy_doc_with_microsoft_word(file_bytes, *, timeout=45):
    """Convert one legacy DOC to macro-free DOCX through installed Word.

    Word is an optional local desktop dependency, not part of the protected
    package. Both 32-bit and 64-bit PowerShell hosts are tried because COM
    registration follows the installed Office bitness.
    """
    if os.name != "nt" or not file_bytes:
        return b""
    import tempfile as _tempfile

    script = r'''param([string]$InputPath, [string]$OutputPath, [string]$ProcessIdPath)
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
$existingWordProcessIds = @(
    Get-Process -Name WINWORD -ErrorAction SilentlyContinue |
        ForEach-Object { [uint32]$_.Id }
)
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $word.AutomationSecurity = 3
    $word.Options.UpdateLinksAtOpen = $false
    try {
        Add-Type @"
using System;
using System.Runtime.InteropServices;
public static class CVStudioWordNativeMethods {
    [DllImport("user32.dll")]
    public static extern uint GetWindowThreadProcessId(IntPtr hWnd, out uint processId);
}
"@
        [uint32]$wordProcessId = 0
        [void][CVStudioWordNativeMethods]::GetWindowThreadProcessId(
            [IntPtr]$word.Hwnd,
            [ref]$wordProcessId
        )
        if (
            $wordProcessId -gt 0 -and
            $existingWordProcessIds -notcontains $wordProcessId
        ) {
            [IO.File]::WriteAllText($ProcessIdPath, [string]$wordProcessId)
        }
    } catch {}
    $document = $word.Documents.Open($InputPath, $false, $true, $false)
    $document.SaveAs2($OutputPath, 16)
    if (-not (Test-Path -LiteralPath $OutputPath)) {
        throw 'Word did not create the converted document.'
    }
}
finally {
    if ($null -ne $document) {
        $document.Close(0)
        [void][Runtime.InteropServices.Marshal]::FinalReleaseComObject($document)
    }
    if ($null -ne $word) {
        $word.Quit()
        [void][Runtime.InteropServices.Marshal]::FinalReleaseComObject($word)
    }
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}
'''
    deadline = time.monotonic() + max(1.0, float(timeout))
    for powershell in _legacy_doc_word_powershell_candidates():
        remaining = deadline - time.monotonic()
        if remaining <= 0:
            break
        try:
            with _tempfile.TemporaryDirectory(
                prefix="cvstudio-legacy-doc-word-"
            ) as td:
                input_path = os.path.join(td, "input.doc")
                output_path = os.path.join(td, "input.docx")
                script_path = os.path.join(td, "convert.ps1")
                process_id_path = os.path.join(td, "word-process-id.txt")
                with open(input_path, "wb") as handle:
                    handle.write(file_bytes)
                with open(script_path, "w", encoding="utf-8", newline="\n") as handle:
                    handle.write(script)
                result = _run_converter_subprocess(
                    [
                        powershell,
                        "-NoLogo",
                        "-NoProfile",
                        "-NonInteractive",
                        "-ExecutionPolicy",
                        "Bypass",
                        "-File",
                        script_path,
                        "-InputPath",
                        input_path,
                        "-OutputPath",
                        output_path,
                        "-ProcessIdPath",
                        process_id_path,
                    ],
                    timeout=max(1.0, remaining),
                    extra_pid_path=process_id_path,
                )
                if result.returncode == 0 and os.path.isfile(output_path):
                    with open(output_path, "rb") as handle:
                        return handle.read()
        except Exception:
            continue
    return b""


def _convert_legacy_doc_with_libreoffice(file_bytes, *, timeout=45):
    """Convert one legacy DOC to DOCX through an optional LibreOffice install."""
    if not file_bytes:
        return b""
    import pathlib as _pathlib
    import tempfile as _tempfile

    soffice = _find_soffice_binary()
    if not soffice:
        return b""
    try:
        with _tempfile.TemporaryDirectory(
            prefix="cvstudio-legacy-doc-libreoffice-"
        ) as td:
            input_path = os.path.join(td, "input.doc")
            output_path = os.path.join(td, "input.docx")
            profile_path = _pathlib.Path(td, "profile")
            profile_path.mkdir()
            with open(input_path, "wb") as handle:
                handle.write(file_bytes)
            result = _run_converter_subprocess(
                [
                    soffice,
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nolockcheck",
                    "--nofirststartwizard",
                    "-env:UserInstallation=" + profile_path.resolve().as_uri(),
                    "--convert-to",
                    "docx:Office Open XML Text",
                    "--outdir",
                    td,
                    input_path,
                ],
                timeout=max(1.0, float(timeout)),
            )
            if result.returncode == 0 and os.path.isfile(output_path):
                with open(output_path, "rb") as handle:
                    return handle.read()
    except Exception:  # best-effort optional converter; caller retains the 424 path
        pass
    return b""


def _convert_legacy_doc_to_docx_bytes(file_bytes):
    """Return a validated temporary DOCX conversion and its converter label.

    This helper is called only after the mandatory verified Antiword runtime
    has run and rejected the exact document. Conversion cannot make a missing
    or untrusted Antiword installation pass the legacy-DOC gate.
    """
    if not (
        file_bytes
        and file_bytes.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
    ):
        return b"", ""
    converters = (
        ("microsoft-word", _convert_legacy_doc_with_microsoft_word),
        ("libreoffice", _convert_legacy_doc_with_libreoffice),
    )
    deadline = time.monotonic() + 35.0
    for label, converter in converters:
        remaining = deadline - time.monotonic()
        if remaining <= 0:
            break
        try:
            converted = converter(file_bytes, timeout=remaining) or b""
            if not converted.startswith(b"PK"):
                continue
            _validate_zip_payload(converted, "DOCX")
            from docx import Document as _ConvertedDocument

            _ConvertedDocument(io.BytesIO(converted))
            return converted, label
        except Exception:
            continue
    return b"", ""


def _iter_antiword_binaries():
    """Yield only the exact hash- and function-verified managed runtime."""
    candidate = _verified_antiword_finder(
        _CVSTUDIO_ROOT,
        os.path.dirname(os.path.abspath(__file__)),
    )
    if candidate:
        yield candidate

def _find_antiword_binary():
    """Return the verified Antiword binary, kept for compatibility callers."""
    return next(_iter_antiword_binaries(), None)


def _antiword_env_for_binary(binary_path):
    """Return the fixed child environment for the pinned Antiword runtime."""
    return _verified_antiword_subprocess_env(binary_path)


def _require_verified_antiword():
    return _require_verified_antiword_runtime(
        _CVSTUDIO_ROOT,
        os.path.dirname(os.path.abspath(__file__)),
    )


def _run_verified_antiword(arguments, *, timeout=20):
    return _run_verified_antiword_runtime(
        _CVSTUDIO_ROOT,
        arguments,
        os.path.dirname(os.path.abspath(__file__)),
        timeout=timeout,
    )


def _antiword_dependency_response(error):
    return jsonify(_antiword_dependency_payload(error)), 424


def _antiword_dependency_payload(error):
    reason = str(getattr(error, "reason", "unavailable") or "unavailable")
    document_failure = reason == "document-extraction-failed"
    message = _cvstudio_safe_error_message(
        str(
            getattr(
                error,
                "public_message",
                "Verified Antiword is required for legacy .doc files. Re-run "
                "the CV Studio installer, then retry this file.",
            )
        )
    )
    return {
        "ok": False,
        "error": message,
        "message": message,
        "code": (
            "LEGACY_DOC_EXTRACTION_FAILED"
            if document_failure
            else "ANTIWORD_DEPENDENCY_UNAVAILABLE"
        ),
        "retryable": False,
        "request_id": _cvstudio_current_request_id(),
        "severity": "warning",
        "action": (
            "convert_to_docx_or_pdf"
            if document_failure
            else "run_installer"
        ),
        "details": {
            "dependency": "antiword",
            "reason": reason,
            "required_for": "legacy_doc",
        },
    }


def _office_bytes_to_pdf_preview(file_bytes, ext, cancel_check=None):
    """Best-effort Office/ODT/RTF visual preview through LibreOffice PDF conversion.

    Optional cancellation is used only by AI Crawler background prefetch. A
    foreground open or Clear terminates the disposable headless LibreOffice
    process instead of letting it compete with the active preview.
    """
    ext = (ext or '.docx').lower()
    if not ext.startswith('.'):
        ext = '.' + ext
    is_legacy_doc = _document_is_legacy_doc(
        file_bytes,
        filename=ext,
    )
    if is_legacy_doc:
        ext = ".doc"
    if is_legacy_doc:
        # A verified installation alone is insufficient: Antiword must also
        # decode this exact document before LibreOffice may render its layout.
        # The rendered PDF is visual-only and is never treated as extracted
        # Antiword text.
        _spider_extract_legacy_doc_text_for_preview(file_bytes)
    soffice = _find_soffice_binary()
    if not soffice:
        return None, 'Exact layout preview for DOC/DOCX/RTF/ODT needs LibreOffice installed; showing text-style fallback.'
    import tempfile as _tempfile, subprocess as _subprocess, os as _os
    try:
        with _tempfile.TemporaryDirectory() as td:
            if callable(cancel_check):
                cancel_check()
            in_path = _os.path.join(td, 'input' + ext)
            with open(in_path, 'wb') as fh:
                fh.write(file_bytes)
            cmd = [soffice, '--headless', '--convert-to', 'pdf', '--outdir', td, in_path]
            proc = _subprocess.Popen(cmd, stdout=_subprocess.PIPE, stderr=_subprocess.PIPE)
            deadline = time.time() + 45
            try:
                while proc.poll() is None:
                    if callable(cancel_check):
                        try:
                            cancel_check()
                        except Exception:
                            try:
                                proc.terminate()
                                proc.wait(timeout=3)
                            except Exception:
                                try:
                                    proc.kill()
                                except Exception:
                                    pass
                            raise
                    if time.time() >= deadline:
                        try:
                            proc.terminate()
                            proc.wait(timeout=3)
                        except Exception:
                            try:
                                proc.kill()
                            except Exception:
                                pass
                        return None, 'Office visual conversion timed out; showing text-style fallback.'
                    time.sleep(0.15)
                stdout, stderr = proc.communicate(timeout=3)
            finally:
                if proc.poll() is None:
                    try:
                        proc.kill()
                    except Exception:
                        pass
            if callable(cancel_check):
                cancel_check()
            out_path = _os.path.join(td, 'input.pdf')
            if not _os.path.exists(out_path):
                pdfs = [name for name in _os.listdir(td) if name.lower().endswith('.pdf')]
                if pdfs:
                    out_path = _os.path.join(td, pdfs[0])
            if _os.path.exists(out_path):
                if _os.path.getsize(out_path) > 24 * 1024 * 1024:
                    return None, 'Office visual conversion produced an oversized preview; showing text-style fallback.'
                with open(out_path, 'rb') as fh:
                    return fh.read(), 'Visual preview converted to PDF via LibreOffice.'
            err = (stderr or b'').decode('utf-8', errors='ignore').strip()
            return None, ('Office visual conversion failed; showing text-style fallback.' + ((' ' + err[:180]) if err else ''))
    except _SpiderPreviewCancelled:
        raise
    except Exception as exc:
        return None, 'Office visual conversion failed; showing text-style fallback. ' + str(exc)[:160]


def _docx_bytes_to_preview_html(file_bytes):
    """Approximate DOCX preview preserving document order, bullets, bold/italic and tables.

    This is a fallback only when LibreOffice visual conversion is unavailable.
    It should still feel sane: paragraphs and tables are rendered in the same
    order they appear in the DOCX body instead of showing all paragraphs first
    and tables later.
    """
    import html as _html
    import re as _re
    import io as _io
    import docx as _docx
    from docx.table import Table as _DocxTable
    from docx.text.paragraph import Paragraph as _DocxParagraph
    doc = _docx.Document(_io.BytesIO(file_bytes))

    def run_html(run):
        txt = _html.escape(run.text or '')
        if not txt:
            return ''
        if run.bold:
            txt = '<strong>' + txt + '</strong>'
        if run.italic:
            txt = '<em>' + txt + '</em>'
        if run.underline:
            txt = '<u>' + txt + '</u>'
        return txt

    def para_html(p):
        text = ''.join(run_html(r) for r in p.runs) or _html.escape(p.text or '')
        if not text.strip():
            return ''
        style_name = (getattr(p.style, 'name', '') or '').lower()
        raw_text = (p.text or '').strip()
        if 'heading 1' in style_name or style_name == 'title':
            return '<h2 style="font-size:18px;margin:12px 0 6px;">' + text + '</h2>'
        if 'heading 2' in style_name or 'subtitle' in style_name:
            return '<h3 style="font-size:15px;margin:10px 0 5px;">' + text + '</h3>'
        is_list = ('list' in style_name) or bool(_re.match(r'^[\u2022\-\u2013\*]\s+', raw_text))
        if is_list:
            clean = _re.sub(r'^[\u2022\-\u2013\*]\s*', '', text).strip()
            return '<div style="margin:4px 0 4px 18px;text-indent:-12px;">&bull; ' + clean + '</div>'
        css = 'margin:4px 0;'
        try:
            align_val = getattr(p.paragraph_format, 'alignment', None)
            # python-docx alignment repr varies by version; numeric 1/2 are center/right.
            if align_val == 1 or str(align_val).upper().endswith('CENTER'):
                css += 'text-align:center;'
            elif align_val == 2 or str(align_val).upper().endswith('RIGHT'):
                css += 'text-align:right;'
        except Exception:
            pass
        return '<p style="' + css + '">' + text + '</p>'

    def table_html(table):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                bits = []
                for cp in cell.paragraphs:
                    h = para_html(cp)
                    if h:
                        bits.append(h)
                cell_html = ''.join(bits) or _html.escape((cell.text or '').strip())
                cells.append('<td style="border:1px solid #ddd;padding:5px;vertical-align:top;">' + cell_html + '</td>')
            rows.append('<tr>' + ''.join(cells) + '</tr>')
        return '<table style="border-collapse:collapse;width:100%;margin:8px 0;font-size:11px;">' + ''.join(rows) + '</table>' if rows else ''

    parts = []
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            h = para_html(_DocxParagraph(child, doc))
            if h:
                parts.append(h)
        elif tag == 'tbl':
            h = table_html(_DocxTable(child, doc))
            if h:
                parts.append(h)
    return ''.join(parts) or '<div style="color:#666;">Could not render DOCX layout preview. Use Text mode.</div>'


@app.route("/preview-file", methods=["POST"])
def preview_file():
    """Return best-effort visual preview data for JD original comparison."""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        f = request.files['file']
        raw_name = f.filename or ''
        filename = raw_name.lower()
        file_bytes = f.read()
        import base64 as _base64, html as _html, os as _os
        ext = _os.path.splitext(filename)[1].lower()
        content_kind = _document_content_kind(file_bytes)
        if content_kind == "legacy_doc":
            ext = ".doc"
        elif content_kind == "pdf":
            ext = ".pdf"
        elif content_kind == "image":
            ext = _document_image_extension(file_bytes)
        elif content_kind == "zip" and ext not in (".docx", ".odt"):
            ext = ".docx"
        # Avoid freezing the browser by returning huge base64 data URLs. Text
        # extraction can still run separately; only the visual preview is skipped.
        max_visual_bytes = 12 * 1024 * 1024
        if len(file_bytes) > max_visual_bytes:
            if ext == ".doc":
                _spider_extract_legacy_doc_text_for_preview(file_bytes)
            return jsonify({"ok": False, "error": "Visual preview skipped for files larger than 12 MB; use Text mode."})

        if ext == '.pdf':
            _pdf_page_count(file_bytes)
            data = _base64.b64encode(file_bytes).decode('ascii')
            return jsonify({"ok": True, "mode": "pdf", "mime": "application/pdf", "data_url": "data:application/pdf;base64," + data, "note": "Visual preview uses the original uploaded PDF."})

        if ext in ('.png', '.jpg', '.jpeg'):
            # Compressed image byte size alone does not prevent a decompression
            # bomb from freezing the browser preview. Validate dimensions and
            # decoding on the backend before returning the data URL.
            checked_image = _safe_image_open(file_bytes)
            try:
                pass  # _safe_image_open already decodes the image via load().
            finally:
                try: checked_image.close()
                except Exception: pass
            mime = 'image/png' if ext == '.png' else 'image/jpeg'
            data = _base64.b64encode(file_bytes).decode('ascii')
            return jsonify({"ok": True, "mode": "image", "mime": mime, "data_url": "data:" + mime + ";base64," + data, "note": "Visual preview uses the original uploaded image."})

        if ext in ('.docx', '.doc', '.rtf', '.odt'):
            if ext in ('.docx', '.odt'):
                _validate_zip_payload(file_bytes, ext[1:].upper())
            pdf_bytes, note = _office_bytes_to_pdf_preview(file_bytes, ext)
            if pdf_bytes:
                _pdf_page_count(pdf_bytes)
                data = _base64.b64encode(pdf_bytes).decode('ascii')
                return jsonify({"ok": True, "mode": "pdf", "mime": "application/pdf", "data_url": "data:application/pdf;base64," + data, "note": note})
            if ext == '.docx':
                try:
                    html_preview = _docx_bytes_to_preview_html(file_bytes)
                    return jsonify({"ok": True, "mode": "html", "html": html_preview, "note": note or "Approximate DOCX layout preview; Text mode is used for anonymising."})
                except Exception:
                    pass
            return jsonify({"ok": False, "error": note or "Visual preview unavailable for this file type; use Text mode."})

        if ext == '.txt':
            return jsonify({"ok": False, "error": "TXT has no document layout. Use Text mode."})

        return jsonify({"ok": False, "error": "Visual preview unavailable for this file type."})
    except AntiwordDependencyError as e:
        return _antiword_dependency_response(e)
    except Exception as e:
        status = _document_validation_status(e)
        if status == 400 or status == 413:
            return jsonify({"error": str(e)}), status
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/extract-text", methods=["POST"])
def extract_text():
    """Extract plain text from uploaded PDF or DOCX file"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        f = request.files['file']
        filename = str(f.filename or '').lower()
        file_bytes = f.read()
        content_kind = _document_content_kind(file_bytes)
        if content_kind == "legacy_doc":
            filename = "verified-upload.doc"
        elif content_kind == "pdf":
            filename = "verified-upload.pdf"
        elif content_kind == "image":
            filename = "verified-upload" + _document_image_extension(file_bytes)
        elif content_kind == "zip" and not filename.endswith((".docx", ".odt")):
            filename = "verified-upload.docx"

        text = ""
        bullet_levels = []
        legacy_doc_converter = ""

        if filename.endswith('.pdf'):
            import pdfplumber
            _pdf_page_count(file_bytes)
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                text = "\n\n".join(pages)

            if text.strip():
                try:
                    bullet_levels = _extract_pdf_bullet_levels(file_bytes)
                except Exception:
                    bullet_levels = []

            # Fallback: if pdfplumber got nothing, try OCR via pytesseract
            if not text.strip():
                try:
                    import shutil

                    # Never mutate the runtime environment from a request. In a
                    # native Nuitka build sys.executable is CVStudio.exe, not a
                    # Python interpreter, and a network-stalled pip process could
                    # hang this request indefinitely. The installer owns mandatory
                    # OCR dependencies and gives a deterministic repair path.
                    try:
                        import pytesseract
                    except ImportError as dependency_error:
                        raise RuntimeError(
                            "OCR Python components are missing. Run INSTALL.bat on Windows "
                            "or install.sh on macOS, then restart CV Studio. Missing component: {}"
                            .format(getattr(dependency_error, "name", "pytesseract"))
                        )

                    # Check Tesseract binary — also look in local tesseract/ folder
                    tess_path = _find_mandatory_tesseract(_CVSTUDIO_ROOT)
                    if not tess_path:
                        raise RuntimeError(
                            "Tesseract is not installed. On Mac run: brew install tesseract. On Windows run INSTALL.bat."
                        )
                    pytesseract.pytesseract.tesseract_cmd = tess_path

                    # Check poppler — local folder first, then system-wide, then PATH
                    _script_dir = os.path.dirname(os.path.abspath(__file__))
                    local_poppler = os.path.join(_script_dir, 'poppler', 'Library', 'bin')
                    local_poppler_bin = os.path.join(_script_dir, 'poppler', 'bin')
                    _poppler_candidates = [
                        local_poppler,
                        local_poppler_bin,
                        r'C:\Program Files\poppler\Library\bin',
                        r'C:\Program Files\poppler\bin',
                        r'C:\Program Files (x86)\poppler\Library\bin',
                        r'C:\Program Files (x86)\poppler\bin',
                        '/usr/local/bin',
                        '/opt/homebrew/bin',
                        '/usr/bin',
                    ]
                    poppler_path = None
                    for _p in _poppler_candidates:
                        if os.path.exists(os.path.join(_p, 'pdfinfo.exe')) or os.path.exists(os.path.join(_p, 'pdftoppm.exe')):
                            poppler_path = _p
                            break
                    if poppler_path is None and (shutil.which('pdfinfo') or shutil.which('pdftoppm')):
                        poppler_path = None  # available through PATH for fallback

                    # Convert and OCR one bounded page at a time. PDFium is the
                    # packaged renderer; Poppler is only an optional fallback.
                    # many-page PDF as a giant in-memory image list.
                    text = _ocr_pdf_pagewise(
                        file_bytes,
                        pytesseract,
                        poppler_path=poppler_path if 'poppler_path' in dir() else None,
                        dpi=220,
                    )

                except Exception as ocr_err:
                    traceback.print_exc()
                    return jsonify({"error": f"OCR failed: {str(ocr_err)}"}), 400

        elif filename.endswith('.docx'):
            _validate_zip_payload(file_bytes, 'DOCX')
            from docx import Document
            import zipfile, re
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = []

            # Extract from headers (name is often stored here)
            for section in doc.sections:
                for hdr in [section.header, section.first_page_header, section.even_page_header]:
                    try:
                        for para in hdr.paragraphs:
                            if para.text.strip():
                                paragraphs.append(para.text)
                    except:
                        pass

            # Also extract headers from raw XML (catches edge cases)
            try:
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                    header_files = [n for n in z.namelist() if re.match(r'word/header\d+\.xml', n)]
                    for hf in header_files:
                        raw = z.read(hf).decode('utf-8', errors='ignore')
                        hdr_text = re.sub(r'<[^>]+>', ' ', raw)
                        hdr_text = ' '.join(hdr_text.split()).strip()
                        if hdr_text and hdr_text not in paragraphs:
                            paragraphs.insert(0, hdr_text)
                    # Also extract textboxes from document body
                    doc_xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
                    for txbx in re.findall(r'<w:txbxContent>(.*?)</w:txbxContent>', doc_xml, re.DOTALL):
                        txbx_text = re.sub(r'<[^>]+>', ' ', txbx)
                        txbx_text = ' '.join(txbx_text.split()).strip()
                        if txbx_text and txbx_text not in paragraphs:
                            paragraphs.append(txbx_text)
            except:
                pass

            # Extract body paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            text = "\n".join(paragraphs)
            try:
                bullet_levels = _extract_docx_bullet_levels(file_bytes)
            except Exception:
                bullet_levels = []

        elif filename.endswith('.doc'):
            # Legacy .doc — try multiple strategies. Important: python-docx only reads .docx, not genuine OLE .doc.
            # Guard corrupt/non-Word payloads early so blank/gibberish bytes do not flow into downstream CV/profile generation.
            _legacy_doc_error = "Could not extract readable text from this legacy .doc. Try converting it to DOCX or PDF."
            _is_ole_doc = file_bytes.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1')
            _is_zip_doc = file_bytes[:2] == b'PK'
            if not (_is_ole_doc or _is_zip_doc):
                return jsonify({"error": _legacy_doc_error}), 400
            text = ""
            _verified_antiword_text = False
            _antiword_run_completed = False
            _require_verified_antiword()

            def _clean_legacy_doc_text(_raw_text):
                import re as _re_doc
                _raw_text = _re_doc.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]+', ' ', _raw_text)
                _raw_text = _re_doc.sub(r'[ \t]{2,}', ' ', _raw_text)
                _raw_text = _re_doc.sub(r'\n{3,}', '\n\n', _raw_text).strip()
                _lines = []
                for _line in _raw_text.split('\n'):
                    _line = _line.strip()
                    if len(_line) <= 3:
                        continue
                    _printable = sum(1 for _c in _line if 0x20 <= ord(_c) <= 0x7E) / max(len(_line), 1)
                    _has_resume_signal = any(_sig in _line.lower() for _sig in ['email', '@', 'phone', 'mobile', 'experience', 'education', 'linkedin', 'skills', 'summary'])
                    if _printable > 0.55 or _has_resume_signal:
                        _lines.append(_line)
                return '\n'.join(_lines).strip()

            def _looks_like_good_doc_text(_txt):
                import re as _re_score
                if not _txt or len(_txt.strip()) < 80:
                    return False
                _sample = _txt[:4000]
                _bad = sum(1 for _c in _sample if (ord(_c) > 0xFFFF) or (ord(_c) >= 0x3000 and not ('\u4e00' <= _c <= '\u9fff')))
                _visible = sum(1 for _c in _sample if not _c.isspace()) or 1
                _signals = len(_re_score.findall(r'(?i)\b(name|email|e-mail|phone|mobile|experience|education|skills|work|project|position)\b|@', _txt[:6000]))
                return (_bad / _visible) < 0.20 and _signals >= 2

            # Strategy 1: the exact verified managed/bundled Antiword runtime.
            # PATH and other arbitrary executable locations are never searched.
            try:
                import tempfile as _tmp_aw, os as _os_aw
                with _tmp_aw.TemporaryDirectory() as _td_aw:
                    _doc_aw = _os_aw.path.join(_td_aw, 'input.doc')
                    with open(_doc_aw, 'wb') as _fh_aw:
                        _fh_aw.write(file_bytes)
                    r2 = _run_verified_antiword(
                        (_doc_aw,),
                        timeout=20,
                    )
                    _antiword_run_completed = True
                    if r2.returncode == 0:
                        _aw_text = r2.stdout.decode('utf-8', errors='ignore').strip()
                        if not _aw_text:
                            _aw_text = r2.stdout.decode('cp1252', errors='ignore').strip()
                        if _looks_like_good_doc_text(_aw_text):
                            text = _aw_text
                            _verified_antiword_text = True
            except AntiwordDependencyError:
                raise
            except Exception:
                pass

            # Strategy 2: after the verified runtime has rejected this exact
            # document, convert it to a validated macro-free DOCX. Microsoft
            # Word is preferred on Windows for fidelity; LibreOffice is the
            # cross-platform optional fallback. Neither converter may bypass a
            # missing or untrusted Antiword runtime because the mandatory gate
            # above runs first.
            if _antiword_run_completed and not _verified_antiword_text:
                try:
                    _converted_docx, _converter = (
                        _convert_legacy_doc_to_docx_bytes(file_bytes)
                    )
                    if _converted_docx:
                        _converted_text = _extract_docx_text_preserve_tables(
                            _converted_docx
                        )
                        if _looks_like_good_doc_text(_converted_text):
                            text = _converted_text
                            legacy_doc_converter = _converter
                            bullet_levels = _extract_docx_bullet_levels(
                                _converted_docx
                            )
                except Exception:  # best-effort converter; diagnostic probes remain available
                    pass

            # Strategy 3: LibreOffice/soffice conversion to txt (if installed).
            # This remains a diagnostic-only probe when DOCX conversion failed.
            if not text:
                try:
                    import subprocess as _sp_lo, tempfile as _tmp_lo, os as _os_lo, shutil as _sh_lo
                    _soffice = _sh_lo.which('soffice') or _sh_lo.which('libreoffice')
                    if _soffice:
                        with _tmp_lo.TemporaryDirectory() as _td:
                            _in = _os_lo.path.join(_td, 'input.doc')
                            with open(_in, 'wb') as _fh:
                                _fh.write(file_bytes)
                            _sp_lo.run([_soffice, '--headless', '--convert-to', 'txt:Text', '--outdir', _td, _in], capture_output=True, timeout=45)
                            _out = _os_lo.path.join(_td, 'input.txt')
                            if _os_lo.path.exists(_out):
                                with open(_out, 'rb') as _fh:
                                    _converted = _fh.read()
                                for _enc in ['utf-8', 'utf-16', 'cp1252', 'latin-1']:
                                    try:
                                        _lo_text = _converted.decode(_enc, errors='ignore').strip()
                                        if _looks_like_good_doc_text(_lo_text):
                                            text = _lo_text
                                            break
                                    except Exception:
                                        pass
                except Exception:
                    pass

            def _extract_word_binary_piece_table(_doc_bytes):
                """Minimal native extractor for legacy binary Word .doc files.
                This avoids blank/gibberish extraction on installs without antiword or LibreOffice.
                It reads the WordDocument piece table and extracts ANSI/Unicode text pieces.
                """
                import struct as _struct, re as _re_wb
                import olefile as _ole_wb

                def _clean_piece_text(_s):
                    # Keep displayed hyperlink text, drop hidden field instruction.
                    _s = _re_wb.sub(r'\x13[^\x14]{0,500}\x14([^\x15]{0,500})\x15', r'\1', _s)
                    _s = _s.replace('\r', '\n').replace('\x07', ' | ')
                    _s = _re_wb.sub(r'[\x00-\x06\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]+', ' ', _s)
                    _s = _re_wb.sub(r'[ \t]{2,}', ' ', _s)
                    _s = _re_wb.sub(r' *\| *\| *', ' | ', _s)
                    _s = _re_wb.sub(r'\n[ \t]+', '\n', _s)
                    _s = _re_wb.sub(r'\n{3,}', '\n\n', _s)
                    _lines = []
                    for _line in _s.split('\n'):
                        _line = _line.strip()
                        if not _line:
                            if _lines and _lines[-1] != '':
                                _lines.append('')
                            continue
                        _lines.append(_line)
                    return '\n'.join(_lines).strip()

                _ofile = _ole_wb.OleFileIO(io.BytesIO(_doc_bytes))
                try:
                    _word = _ofile.openstream('WordDocument').read()
                    _flags = _struct.unpack_from('<H', _word, 0x0A)[0]
                    _table_name = '1Table' if (_flags & 0x0200) else '0Table'
                    if not _ofile.exists(_table_name):
                        return ''
                    _table = _ofile.openstream(_table_name).read()
                finally:
                    _ofile.close()

                if len(_word) < 0x01AA:
                    return ''
                _fc_clx = _struct.unpack_from('<I', _word, 0x01A2)[0]
                _lcb_clx = _struct.unpack_from('<I', _word, 0x01A6)[0]
                if not _lcb_clx or _fc_clx >= len(_table):
                    return ''
                _clx = _table[_fc_clx:_fc_clx + _lcb_clx]
                _pos = 0
                _out = []
                while _pos < len(_clx):
                    _b = _clx[_pos]
                    if _b == 0x01 and _pos + 3 <= len(_clx):
                        _cb_grpprl = _struct.unpack_from('<H', _clx, _pos + 1)[0]
                        _pos += 3 + _cb_grpprl
                        continue
                    if _b != 0x02 or _pos + 5 > len(_clx):
                        _pos += 1
                        continue
                    _lcb_pcdt = _struct.unpack_from('<I', _clx, _pos + 1)[0]
                    _pcdt = _clx[_pos + 5:_pos + 5 + _lcb_pcdt]
                    if len(_pcdt) < 16:
                        break
                    _n = (len(_pcdt) - 4) // 12
                    if _n <= 0:
                        break
                    _cps = [_struct.unpack_from('<I', _pcdt, _i * 4)[0] for _i in range(_n + 1)]
                    _pcd_base = 4 * (_n + 1)
                    for _i in range(_n):
                        _cp0, _cp1 = _cps[_i], _cps[_i + 1]
                        if _cp1 <= _cp0:
                            continue
                        _pcd_off = _pcd_base + (_i * 8)
                        if _pcd_off + 6 > len(_pcdt):
                            continue
                        _fc_comp = _struct.unpack_from('<I', _pcdt, _pcd_off + 2)[0]
                        _compressed = bool(_fc_comp & 0x40000000)
                        _fc = _fc_comp & 0x3FFFFFFF
                        _chars = _cp1 - _cp0
                        if _compressed:
                            # In compressed Word pieces the stored fc is doubled.
                            _start = _fc // 2
                            _raw = _word[_start:_start + _chars]
                            _piece = _raw.decode('cp1252', errors='ignore')
                        else:
                            _start = _fc
                            _raw = _word[_start:_start + (_chars * 2)]
                            _piece = _raw.decode('utf-16-le', errors='ignore')
                        if _piece:
                            _out.append(_piece)
                    break
                return _clean_piece_text(''.join(_out))

            # Strategy 4: native OLE/WordDocument piece-table extraction.
            # This is a defense-in-depth probe and cannot satisfy success.
            if not text:
                try:
                    _piece_text = _extract_word_binary_piece_table(file_bytes)
                    if _looks_like_good_doc_text(_piece_text):
                        text = _piece_text
                except Exception:
                    pass

            # Strategy 5: last-resort raw scans, but only accept them if they look like real text.
            if not text:
                try:
                    _raw_try = _clean_legacy_doc_text(file_bytes.decode('utf-16-le', errors='ignore'))
                    if _looks_like_good_doc_text(_raw_try):
                        text = _raw_try
                except Exception:
                    pass

            # Strategy 6: python-docx (works only if .doc is actually .docx renamed)
            if not text:
                try:
                    from docx import Document as _Doc
                    _doc = _Doc(io.BytesIO(file_bytes))
                    _docx_renamed_text = "\n".join(p.text for p in _doc.paragraphs if p.text.strip())
                    if _looks_like_good_doc_text(_docx_renamed_text):
                        text = _docx_renamed_text
                except Exception:
                    pass
            if not (_verified_antiword_text or legacy_doc_converter):
                raise AntiwordDependencyError(
                    "document-extraction-failed",
                    "Verified Antiword could not decode this legacy .doc. "
                    "Convert it to DOCX or PDF and retry.",
                )
            if not text:
                return jsonify({"error": _legacy_doc_error}), 400
            if not _looks_like_good_doc_text(text):
                # Keep the final .doc gate specific to legacy Word so normal PDF/DOCX/TXT extraction is untouched.
                return jsonify({"error": _legacy_doc_error}), 400

        elif filename.endswith('.txt'):
            # Try common encodings
            for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
                try:
                    text = file_bytes.decode(enc)
                    break
                except Exception:
                    continue

        elif filename.endswith('.rtf'):
            # Strip RTF control codes to extract plain text
            import re as _re
            raw = file_bytes.decode('latin-1', errors='ignore')
            # Remove RTF groups and control words
            text = _re.sub(r'\\\w+', ' ', raw)
            text = _re.sub(r'[{}\\]', '', text)
            text = _re.sub(r'\s+', ' ', text).strip()

        elif filename.endswith('.odt'):
            _validate_zip_payload(file_bytes, 'ODT')
            # ODT is a zip with content.xml inside
            import zipfile as _zf, re as _re
            with _zf.ZipFile(io.BytesIO(file_bytes)) as z:
                if 'content.xml' in z.namelist():
                    xml = z.read('content.xml').decode('utf-8', errors='ignore')
                    text = _re.sub(r'<[^>]+>', ' ', xml)
                    text = ' '.join(text.split())
                else:
                    return jsonify({"error": "Could not read ODT file."}), 400

        elif filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            # Image file — run OCR directly
            try:
                from PIL import Image
                import pytesseract

                tess_path = _find_mandatory_tesseract(_CVSTUDIO_ROOT)
                if not tess_path:
                    raise RuntimeError(
                        "Mandatory Tesseract OCR is unavailable. Re-run the CV Studio installer."
                    )
                pytesseract.pytesseract.tesseract_cmd = tess_path

                img = _safe_image_open(file_bytes)
                try:
                    text = _ocr_image_text(pytesseract, img, lang='eng', timeout=35)
                finally:
                    img.close()
            except Exception as e:
                return jsonify({"error": f"Image OCR failed: {str(e)}. Make sure Tesseract is installed."}), 400

        else:
            return jsonify({"error": "Unsupported file type. Supported: PDF, DOCX, DOC, TXT, RTF, ODT, PNG, JPG, JPEG."}), 400

        if not text.strip():
            return jsonify({"error": "Could not extract text from file. It may be scanned or image-based."}), 400

        # Clean PDF/font/OCR artefacts without deleting ligature fragments.
        text = _normalize_extracted_text_artifacts(text)
        response_payload = {
            "ok": True,
            "text": text,
            "bullet_levels": bullet_levels,
        }
        if legacy_doc_converter:
            response_payload["legacy_doc_converter"] = legacy_doc_converter
        return jsonify(response_payload)

    except AntiwordDependencyError as e:
        return _antiword_dependency_response(e)
    except Exception as e:
        _extract_name = str(locals().get('filename', '') or '').lower()
        _extract_msg = str(e)
        _extract_status = _document_validation_status(e)
        if _extract_status == 413:
            return jsonify({"error": _extract_msg}), 413
        import re as _re_extract
        if _extract_name.endswith('.pdf') and _re_extract.search(r'pdf|/root object|xref|trailer|syntax', _extract_msg, _re_extract.I):
            return jsonify({"error": "Could not read PDF file. It may be corrupt, password-protected, or not a valid PDF."}), 400
        if _extract_name.endswith('.docx') and _re_extract.search(r'zip|package|opc|word/document|not a zip', _extract_msg, _re_extract.I):
            return jsonify({"error": "Could not read DOCX file. It may be corrupt or not a valid DOCX."}), 400
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ── Blind CV deterministic organisation masking guard ─────────────────────
# AI providers, especially cheaper models, can miss client/competitor/company
# names embedded inside project descriptions or achievement bullets.  The blind
# route therefore does a final deterministic sweep using names seen in the
# original parsed CV plus a small curated brand list.  This is intentionally
# conservative: it targets organisation/client names only and avoids common
# technologies/tools so that bullets keep their technical meaning.
from cvstudio_blind_mask import (
    _BLIND_CURATED_ORG_NAMES,
    _BLIND_ORG_SUFFIX_RE,
    _BLIND_ORG_TECH_ALLOWLIST,
    _blind_add_mask_term,
    _blind_collect_org_mask_terms,
    _blind_mask_org_terms_recursive,
    _blind_postprocess_company_mentions,
    _blind_replace_org_terms_in_text,
    _blind_walk_strings,
)

# Common Malaysia/SEA/global employers, clients and competitors that often
# appear as single-token names in project bullets without legal suffixes.  This
# list is deliberately used only for exact phrase matching; it does not infer
# unknown names.















BLIND_SYSTEM_PROMPT = """You are a CV anonymisation specialist. You receive a parsed CV as a JSON object and must return the SAME JSON structure with all identifying information redacted. Output ONLY valid JSON. No markdown fences. No explanation.

IDENTITY REDACTION — remove or replace these completely:
- candidate.name → replace with "Candidate"
- candidate.current_company → apply company masking rules below
- Any URL, LinkedIn link, GitHub link, personal website, portfolio URL, Behance, Dribbble → replace the full URL with "[Link Redacted]". If the category is "Portfolio & Links", replace the items value with "[Link Redacted]"
- Email addresses → "[Email Redacted]"
- Phone numbers → "[Phone Redacted]"

COMPANY NAME MASKING — replace each company name with a descriptor that conveys its scale, prestige, and sector without revealing its identity. This applies not only to employer fields, but also to every company/client/competitor name mentioned anywhere in work descriptions, project descriptions, achievements, highlights, summaries, notes, and additional information. For example, if a bullet says the candidate built a platform for Maxis, Maybank, AIA, IHH, Grab, Shopee, or any other client/competitor/employer, mask that name too. Use your knowledge of the actual company to pick the most accurate and specific label for employer fields. In free-text bullets, project descriptions and achievements, replace incidental company/client names with a neutral placeholder such as "[Company]" unless a natural descriptor is clearly safer. Vary your phrasing naturally — do not always use the same template. Choose from the examples below or craft a fitting variation:

BANKING & FINANCE:
- Bulge-bracket / global investment bank (e.g. Goldman Sachs, JP Morgan, Morgan Stanley) → "Top Bulge-Bracket Investment Bank" / "One of the World's Largest Investment Banks" / "Leading Global Investment Bank"
- Top-3 largest bank by assets globally → "One of the World's Three Largest Banks" / "Globally Systemically Important Bank"
- Major regional bank in SEA / Asia → "One of Southeast Asia's Largest Banks" / "Leading Regional Bank in Asia" / "Top-Tier ASEAN Commercial Bank"
- National / domestic bank (well-known in home country) → "One of the Country's Most Prominent Banks" / "Leading Domestic Commercial Bank"
- Islamic bank → "Leading Islamic Financial Institution" / "Prominent Shariah-Compliant Bank"
- Major fintech / digital bank / neobank → "One of the World's Largest Fintech Giants" / "Leading Digital Bank" / "Prominent Fintech Unicorn"
- Payment network / card scheme (Visa, Mastercard, etc.) → "Global Payment Network" / "World's Leading Card Scheme"
- Insurance company (global) → "Fortune 500 Global Insurance Group" / "One of the World's Largest Insurers"
- Insurance company (regional) → "Leading Regional Insurance Company" / "Prominent Insurance Provider in Asia"

CONSULTING & PROFESSIONAL SERVICES:
- MBB (McKinsey / BCG / Bain) → "Top-3 Global Management Consultancy" / "One of the Most Prominent Strategy Consulting Firms" / "Elite Global Strategy Consultancy"
- Big 4 (Deloitte / PwC / EY / KPMG) → "Big 4 Consulting & Accounting Firm" / "One of the Big 4 Professional Services Firms"
- Tier-2 consulting (Accenture, Oliver Wyman, Roland Berger, etc.) → "Leading Global Management Consultancy" / "Top-Tier Strategy & Operations Consultancy"
- Boutique / specialist consultancy → "Boutique Management Consultancy" / "Specialist Strategy Consulting Firm"

TECHNOLOGY:
- FAANG / MAANG (Meta, Apple, Amazon, Netflix, Google, Microsoft) → "Leading Global Technology Giant" / "One of the World's Largest Technology Companies" / "Top-5 Global Tech Company"
- Major enterprise software company (SAP, Oracle, Salesforce, etc.) → "Leading Enterprise Software Giant" / "Prominent Global Enterprise Technology Company"
- Major regional tech company in Asia (Grab, Sea, Goto, Tencent, Alibaba, etc.) → "One of Southeast Asia's Largest Technology Companies" / "Prominent Asian Super-App" / "Leading Asian Technology Conglomerate"
- Large IT services firm (Infosys, TCS, Wipro, etc.) → "Top Global IT Services Company" / "Leading Technology Services Provider"
- Semiconductor / hardware company → "Major Global Semiconductor Company" / "Leading Chip & Hardware Manufacturer"
- Cybersecurity firm → "Leading Cybersecurity Company" / "Prominent Information Security Firm"
- Tech startup / scaleup → "[Sector] Tech Startup" / "Fast-Growing Technology Scaleup"

TELECOMMUNICATIONS:
- Global telecom operator → "Major Global Telecommunications Operator" / "One of the World's Largest Telcos"
- Regional / national telecom → "Leading National Telecommunications Provider" / "Prominent Telco in Southeast Asia"

E-COMMERCE & CONSUMER INTERNET:
- Global e-commerce giant (Amazon, Alibaba, etc.) → "One of the World's Largest E-Commerce Platforms" / "Global E-Commerce & Cloud Giant"
- Regional marketplace (Shopee, Lazada, Tokopedia, etc.) → "Leading Southeast Asian E-Commerce Marketplace" / "Top Regional Online Marketplace"
- Ride-hailing / super-app → "Leading Ride-Hailing & Super-App Platform" / "Prominent Consumer Tech Super-App"
- Food delivery platform → "Major Food Delivery Platform" / "Leading On-Demand Delivery Company"

MULTINATIONAL CORPORATIONS (NON-TECH/FINANCE):
- Fortune 500 global MNC (US-headquartered) → "Fortune 500 American MNC" / "Fortune 500 Global Conglomerate"
- Fortune 500 European MNC → "Fortune 500 European MNC" / "Leading European Multinational Corporation"
- Fortune 500 Asian MNC → "Fortune 500 Asian Multinational" / "Leading Asian Industrial Conglomerate"
- Large FMCG company (Unilever, P&G, Nestle, etc.) → "One of the World's Largest FMCG Companies" / "Global Consumer Goods Giant"
- Automotive manufacturer (global top-10) → "Top-10 Global Automotive Manufacturer" / "Leading Global Car Maker"
- Oil & gas supermajor → "Global Oil & Gas Supermajor" / "One of the World's Largest Energy Companies"
- Oil & gas national company / regional → "National Oil Corporation" / "Leading Regional Energy Company"
- Pharmaceutical / biotech (global) → "Top-10 Global Pharmaceutical Company" / "Leading Global Biopharmaceutical Firm"
- Healthcare system / hospital group → "Leading Private Hospital Group" / "Prominent Healthcare Provider"
- Logistics & supply chain (global) → "One of the World's Largest Logistics Companies" / "Global Supply Chain & Freight Giant"
- Aviation / airline → "Major National Carrier" / "Leading Low-Cost Airline in Asia"
- Hospitality / hotel chain → "Global Luxury Hotel Group" / "Leading International Hospitality Chain"
- Media & entertainment (global) → "Major Global Media & Entertainment Conglomerate" / "Leading Streaming & Entertainment Company"
- Retail (global) → "Fortune 500 Global Retailer" / "One of the World's Largest Retail Chains"

REGIONAL / LOCAL:
- Large regional conglomerate → "Leading Regional Conglomerate" / "One of the Largest Conglomerates in Southeast Asia"
- Government-linked company (GLC) → "Government-Linked Company (GLC)" / "Prominent Government-Linked Corporation"
- Government agency / statutory body / ministry → "Government Agency" / "Statutory Body" / "Public Sector Institution"
- SME or lesser-known local company → "[Industry] Company" / "Local [Sector] Firm"
- Early-stage startup → "[Sector] Startup" / "Early-Stage Technology Startup"

IMPORTANT: Use judgment — if a company falls neatly into one tier, use the most precise label. If multiple candidates fit, pick the one that best preserves prestige without revealing identity. Vary phrasing across entries so the document does not read repetitively (e.g. do not use "Leading" for every single entry). After masking employer fields, scan ALL remaining text fields again for residual company/client/competitor names and remove them. Do not leave client names inside bullets, project descriptions, achievements, highlights, summaries, or additional information.

PRODUCT NAME REDACTION — scan ALL text fields (bullets, items, certifications) for branded product names that would reveal the employer company. Redact those product-brand prefixes only. Rules:
- If a product name starts with the company name or brand (e.g. "Digi Prepaid", "Maybank2u", "GrabFood", "AirAsia Go"), replace ONLY the company/brand prefix with "[Company]" — e.g. → "[Company] Prepaid", "[Company]2u", "[Company] Food", "[Company] Go"
- Do NOT redact generic product-type words (Prepaid, App, Portal, Platform, Dashboard, Wallet, Pay, etc.) — keep those
- Do NOT redact technology names, frameworks, tools (AWS, Salesforce, Python, etc.)
- Do NOT redact certification or course names from external bodies

UNIVERSITY / EDUCATION MASKING — replace each institution name with a descriptor that reflects its true prestige tier. Vary phrasing naturally:

IVY LEAGUE & OXBRIDGE:
- Harvard, Yale, Princeton, Columbia, Penn, Cornell, Dartmouth, Brown, Oxford, Cambridge → "Ivy League University" / "Top Ivy League University" / "World-Renowned Ivy League Institution" / "University of Oxford / Cambridge" (Oxbridge can be named by city only if not obvious)

GLOBAL TOP-10 (non-Ivy):
- MIT, Stanford, Caltech, Imperial, ETH Zurich, NUS (top ranked), etc. → "Top-10 QS Global University" / "World-Leading Research University" / "Elite Global STEM University"

GLOBAL TOP-30:
- Top-30 QS but outside top-10 → "Top 30 QS World University" / "Globally Ranked Research University" / "Prominent World-Ranked University"

GLOBAL TOP-50:
- Top-50 QS but outside top-30 → "Top 50 QS World University" / "World-Renowned University" / "Internationally Recognised Research University"

TOP ASIAN UNIVERSITY (outside global top-50):
- Well-ranked in Asia but not global top-50 (e.g. Monash Malaysia, UTAR, APU, MMU, UTP, XMUM) → "Leading University in Asia" / "Top-Ranked University in Southeast Asia" / "Prominent University in the Region"

NATIONAL PUBLIC UNIVERSITY (Malaysia / SEA):
- UM, UKM, UPM, UTM, UiTM, USM, UNIMAS, UMS, and equivalents → "Leading National University" / "One of the Country's Premier Public Universities" / "Top Public Research University"

WELL-KNOWN PRIVATE UNIVERSITY:
- Taylor's, Sunway, HELP, UCSI, SEGi, Multimedia University, UTAR, etc. → "Established Private University" / "Prominent Private University in Malaysia" / "Reputable Private Institution"

INTERNATIONAL BRANCH CAMPUS:
- Monash University Malaysia, University of Nottingham Malaysia, etc. → "Branch Campus of a Global Top-50 University" / "International Branch Campus in Malaysia"

COMMUNITY COLLEGE / POLYTECHNIC / FOUNDATION:
- Politeknik, community college, A-levels centre, matriculation → "Polytechnic / College" / "Pre-University Institution"

PROFESSIONAL / EXECUTIVE EDUCATION:
- INSEAD, London Business School, IESE, etc. → "Top Global Business School" / "Elite Executive Education Institution"
- Local executive programmes → "Professional Development Institution"

If unsure of tier, use "Established University" or "Recognised Higher Learning Institution".

PRESERVE EVERYTHING ELSE exactly as-is:
- All job titles (do not mask these)
- All date ranges
- All bullet point content (except product-name redactions above)
- All skills, certifications (except URLs), languages
- The entire JSON structure and all field names
- candidate.notice_period, candidate.current_position, candidate.languages, candidate.is_employed
- summary_bullets array (keep as empty array)

Output starts with { and ends with }. No markdown fences. No backticks."""


@app.route("/blind", methods=["POST"])
def blind_cv():
    """Take already-parsed CV JSON and return a blinded version of it."""
    try:
        body = request.get_json(force=True, silent=True)
        if not body:
            return jsonify({"error": "Invalid JSON body"}), 400

        api_key = _resolve_request_api_key(body, "main")
        cv_data = body.get("cv_data")
        model   = body.get("model") or "claude-sonnet-4-6"
        llm_provider = (body.get("provider") or "anthropic").strip().lower()
        usage = _merge_llm_usage()

        if not api_key:
            return jsonify({"error": "API key required"}), 400
        if not cv_data:
            return jsonify({"error": "cv_data required"}), 400

        data = call_llm(llm_provider, api_key, {
            "model": model,
            "max_tokens": 64000,
            "system": BLIND_SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": f"Anonymise this parsed CV JSON:\n\n{json.dumps(cv_data)}"}]
        })
        usage = _merge_llm_usage(usage, data.get("usage"))

        raw_text = "".join(b.get("text", "") for b in data.get("content", []))
        raw_text = raw_text.strip()
        if raw_text.startswith("```json"):
            raw_text = raw_text[7:]
        if raw_text.startswith("```"):
            raw_text = raw_text[3:]
        if raw_text.endswith("```"):
            raw_text = raw_text[:-3]
        raw_text = raw_text.strip()

        try:
            blinded = json.loads(raw_text)
        except json.JSONDecodeError as e:
            out = {"error": f"AI provider returned invalid JSON. Error: {str(e)}\n\nRaw response (first 1000 chars):\n{raw_text[:1000]}",
                   "paid_ai_failure": True, "usage": usage, "model": model, "provider": llm_provider}
            out.update(_llm_response_cost_fields(model, usage, llm_provider))
            out.update(_llm_paid_failure_fields(
                e,
                model,
                llm_provider,
                usage,
                attempted=True,
            ))
            return jsonify(out), 500

        # Deterministic last line of defence: mask residual employer/client/competitor
        # names that the AI may have left inside bullets, project descriptions,
        # achievements, highlights, summaries, or additional information.
        blinded = _blind_postprocess_company_mentions(blinded, cv_data)

        out = {"ok": True, "data": blinded, "usage": usage, "model": model, "provider": llm_provider}
        out.update(_llm_response_cost_fields(model, usage, llm_provider))
        return jsonify(out)

    except urllib.error.HTTPError as e:
        err_body = e.read()
        try:
            err = json.loads(err_body)
            msg = err.get("error", {}).get("message", "API error")
        except:
            msg = err_body.decode("utf-8", errors="replace")[:300]
        out = {"error": _provider_error_message(llm_provider, msg)}
        out.update(_llm_paid_failure_fields(
            e,
            model,
            llm_provider,
            locals().get("usage"),
            attempted=True,
        ))
        return jsonify(out), 400

    except Exception as e:
        traceback.print_exc()
        usage = locals().get("usage") or _merge_llm_usage()
        out = {"error": str(e), "paid_ai_failure": bool(_llm_usage_int(usage, "api_calls")), "usage": usage,
               "model": locals().get("model", ""), "provider": locals().get("llm_provider", "anthropic")}
        out.update(_llm_response_cost_fields(out["model"], usage, out["provider"]))
        out.update(_llm_paid_failure_fields(
            e,
            out["model"],
            out["provider"],
            usage,
            attempted=bool(locals().get("api_key") and locals().get("cv_data")),
        ))
        return jsonify(out), 500


def _summarize_node_failure(stdout='', stderr=''):
    """Return the most useful Node/generate.js failure line for users."""
    text = "\n".join([str(stderr or ""), str(stdout or "")])
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    preferred = []
    for ln in lines:
        low = ln.lower()
        if low.startswith("node.js v") or low.startswith("at ") or low.startswith("{") or low.startswith("}"):
            continue
        if "cannot find module" in low or "error:" in low or "enoent" in low or "timed out" in low:
            preferred.append(ln)
    if preferred:
        return preferred[0][:300]
    for ln in lines:
        if not ln.lower().startswith("node.js v"):
            return ln[:300]
    return "DOCX generation failed"


# The browser Batch workflow deliberately calls /generate-docx once per CV so
# each item keeps independent retry, Dashboard-record and JobAdder-upload state.
# The unused /generate-docx-batch ZIP endpoint was removed in v24.6.182.

# ── PPC / Post Placement Care ───────────────────────────────────────────────
# The pure PPC placement helpers, the per-run detail cache and the two
# JobAdder-client network helpers live in ``cvstudio_ppc`` (Phase 7B). The web
# shell keeps the Flask route and OAuth token lifecycle here and binds the
# shared ``_JOBADDER_CLIENT`` into the network helpers at call time so the
# module never imports ``app``.
import cvstudio_ppc as _cvstudio_ppc
from cvstudio_ppc import (
    _PPC_PLACEMENT_TYPES,
    _ppc_clean_date,
    _ppc_detail_cache,
    _ppc_detail_cache_clear,
    _ppc_detail_cache_lock,
    _ppc_int,
    _ppc_normalize_placement,
    _ppc_person_name,
)


def _ppc_get_json(token, path, params=None, timeout=30, retry_429=True):
    return _cvstudio_ppc.ppc_get_json(
        _JOBADDER_CLIENT, token, path, params=params, timeout=timeout, retry_429=retry_429
    )


def _ppc_fetch_type_collection(token, placement_type, base_params, max_records):
    return _cvstudio_ppc.ppc_fetch_type_collection(
        _JOBADDER_CLIENT, _ppc_get_json, token, placement_type, base_params, max_records
    )


@app.route("/jobadder/ppc/placements", methods=["POST"])
def jobadder_ppc_placements():
    """Read JobAdder placements for PPC. This route never writes to JobAdder."""
    data = request.get_json(silent=True) or {}
    with _ja_store_lock:
        token = str(_ja_creds_store.get("access_token") or "").strip()
        account_cache_namespace = str(
            _ja_creds_store.get("cache_namespace") or ""
        ).strip()
    if not token:
        return jsonify({"error": "Connect to JobAdder first"}), 401
    start_date = _ppc_clean_date(data.get("start_date"))
    end_date = _ppc_clean_date(data.get("end_date"))
    if start_date and end_date and start_date > end_date:
        return jsonify({"error": "Start date cannot be after end date"}), 400
    approved_only = bool(data.get("approved_only", True))
    force_details = bool(data.get("force"))
    try:
        max_records = max(1, min(20000, int(data.get("max_records") or 20000)))
    except Exception:
        max_records = 20000
    try:
        detail_raw = data.get("detail_limit")
        detail_limit = max(0, min(20000, int(max_records if detail_raw is None else detail_raw)))
    except Exception:
        detail_limit = min(max_records, 20000)

    base_params = []
    if approved_only:
        base_params.append(("Approved", "true"))
    if start_date:
        base_params.append(("StartDate", ">" + start_date))
    if end_date:
        base_params.append(("StartDate", "<" + end_date))

    raw_items = []
    type_diagnostics = []
    total_expected = 0
    try:
        for placement_type in _PPC_PLACEMENT_TYPES:
            remaining = max_records - len(raw_items)
            if remaining <= 0:
                type_diagnostics.append({
                    "type": placement_type, "expected": None, "returned": 0,
                    "pages": 0, "complete": False, "truncated": True,
                    "skipped": "overall record cap reached",
                })
                continue
            rows, diag = _ppc_fetch_type_collection(token, placement_type, base_params, remaining)
            raw_items.extend(rows)
            type_diagnostics.append(diag)
            total_expected += int(diag.get("expected") or 0)
    except RuntimeError as exc:
        code = int(getattr(exc, "http_code", 502) or 502)
        payload = {"error": str(exc)}
        retry_after = getattr(exc, "retry_after", None)
        if retry_after:
            payload["retry_after"] = retry_after
        return jsonify(payload), code
    except Exception as exc:
        return jsonify({"error": "Could not load JobAdder placements", "detail": str(exc)}), 500

    # De-duplicate by Placement ID after the independent type queries. Preserve
    # the first row and merge any later non-empty summary values defensively.
    by_placement = {}
    for raw in raw_items:
        if not isinstance(raw, dict):
            continue
        pid = raw.get("placementId")
        key = str(pid or "").strip()
        if not key:
            continue
        if key not in by_placement:
            by_placement[key] = raw
        else:
            base = by_placement[key]
            for field, value in raw.items():
                if value not in (None, "", [], {}):
                    base[field] = value
    items = [_ppc_normalize_placement(raw) for raw in by_placement.values()]

    # Treat the selected calendar range as authoritative even if an upstream
    # response includes a boundary/out-of-range row unexpectedly. JobAdder's
    # > and < date prefixes are inclusive according to the packaged OpenAPI.
    if start_date or end_date:
        ranged_items = []
        for row in items:
            row_date = _ppc_clean_date(row.get("start_date"))
            if start_date and (not row_date or row_date < start_date):
                continue
            if end_date and (not row_date or row_date > end_date):
                continue
            ranged_items.append(row)
        items = ranged_items

    items.sort(key=lambda row: str(row.get("start_date") or ""), reverse=True)

    # Full placement representations contain recruiters and invoice fields.
    # Enrich every retrieved row by default (up to max_records), while caching
    # details by Placement ID + updatedAt for fast range changes in the same run.
    detail_failures = 0
    detail_enriched = 0
    detail_cached = 0
    detail_targets = min(len(items), detail_limit)
    if detail_targets:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        by_id = {str(row.get("placement_id")): row for row in items if row.get("placement_id") not in (None, "")}
        target_ids = list(by_id.keys())[:detail_targets]
        network_ids = []
        for pid in target_ids:
            base = by_id.get(pid) or {}
            cache_key = (
                account_cache_namespace,
                pid,
                str(base.get("updated_at") or ""),
            )
            cached = None
            if account_cache_namespace and not force_details:
                with _ppc_detail_cache_lock:
                    cached = _ppc_detail_cache.get(cache_key)
            if isinstance(cached, dict):
                for key, value in cached.items():
                    if value not in (None, "", [], {}):
                        base[key] = value
                detail_cached += 1
            else:
                network_ids.append(pid)

        def _load_detail(pid):
            detail = _ppc_get_json(token, "placements/{}".format(urllib.parse.quote(pid, safe="")), timeout=35)
            return pid, detail

        if network_ids:
            with ThreadPoolExecutor(max_workers=min(8, max(1, len(network_ids)))) as pool:
                futures = [pool.submit(_load_detail, pid) for pid in network_ids]
                for future in as_completed(futures):
                    try:
                        pid, detail = future.result()
                        if isinstance(detail, dict) and detail.get("placementId") not in (None, ""):
                            enriched = _ppc_normalize_placement(detail)
                            base = by_id.get(str(pid))
                            if base is not None:
                                for key, value in enriched.items():
                                    if value not in (None, "", [], {}):
                                        base[key] = value
                                cache_key = (
                                    account_cache_namespace,
                                    str(pid),
                                    str(base.get("updated_at") or ""),
                                )
                                with _ja_store_lock:
                                    namespace_is_current = bool(
                                        account_cache_namespace
                                        and str(
                                            _ja_creds_store.get(
                                                "cache_namespace"
                                            )
                                            or ""
                                        ).strip()
                                        == account_cache_namespace
                                        and str(
                                            _ja_creds_store.get("access_token")
                                            or ""
                                        ).strip()
                                    )
                                    if namespace_is_current:
                                        with _ppc_detail_cache_lock:
                                            _ppc_detail_cache[cache_key] = dict(
                                                enriched
                                            )
                                detail_enriched += 1
                        else:
                            detail_failures += 1
                    except Exception:
                        detail_failures += 1

    items.sort(key=lambda row: str(row.get("start_date") or ""), reverse=True)
    type_complete = all(bool(x.get("complete")) for x in type_diagnostics)
    truncated = any(bool(x.get("truncated")) for x in type_diagnostics) or total_expected > max_records
    complete = bool(type_complete and not truncated and len(items) >= min(total_expected, max_records))
    details_complete = bool(detail_targets >= len(items) and detail_failures == 0)
    warning_parts = []
    if truncated:
        warning_parts.append("Placement retrieval reached the {}-record safety cap".format(max_records))
    incomplete_types = [str(x.get("type")) for x in type_diagnostics if not x.get("complete")]
    if incomplete_types:
        warning_parts.append("Incomplete JobAdder pagination for: {}".format(", ".join(incomplete_types)))
    if detail_limit < len(items):
        warning_parts.append("Full recruiter/invoice detail loaded for {} of {} placements".format(detail_limit, len(items)))
    if detail_failures:
        warning_parts.append("{} placement detail request(s) used summary fallbacks".format(detail_failures))
    return jsonify({
        "ok": True,
        "items": items,
        "returned": len(items),
        "total_count": total_expected,
        "truncated": truncated,
        "complete": complete,
        "details_complete": details_complete,
        "max_records": max_records,
        "detail_limit": detail_limit,
        "detail_enriched": detail_enriched,
        "detail_cached": detail_cached,
        "detail_failures": detail_failures,
        "type_diagnostics": type_diagnostics,
        "warning": "; ".join(warning_parts),
        "fetched_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "filters": {"start_date": start_date, "end_date": end_date, "approved_only": approved_only},
        "source": "JobAdder /v2/placements (type-complete pagination)",
        "read_only": True,
    })


@app.errorhandler(Exception)
def handle_exception(e):
    # Preserve normal HTTP status codes and emit the same structured contract
    # used by route-level failures. The legacy ``error`` field remains present.
    try:
        from werkzeug.exceptions import HTTPException
        if isinstance(e, HTTPException):
            status = int(getattr(e, "code", 500) or 500)
            message = "Not found" if status == 404 else (getattr(e, "description", None) or str(e))
            code, retryable, action = _cvstudio_classify_error(status, message, request.path)
            return _cvstudio_error_payload(code, message, status, retryable=retryable, action=action)
    except Exception:
        pass
    try:
        logging.getLogger("cvstudio").exception("Unhandled request failure request_id=%s path=%s", _cvstudio_current_request_id(), request.path)
    except Exception:
        traceback.print_exc()
    message = _cvstudio_safe_error_message(str(e) or type(e).__name__, "Unexpected local server error")
    return _cvstudio_error_payload("INTERNAL_SERVER_ERROR", message, 500, retryable=True, action="retry", details={"exception_type": type(e).__name__})


# --- Salary Comparison feature module ---------------------------------------
# Self-contained Flask blueprint. Its deterministic tax/FX engine, exporters and
# approved-rule store live entirely in the salary_comparison package, so the
# legacy web shell gains only this registration. Persistent data is kept beside
# the other CV Studio state so version upgrades never reset approved tax rules.
import salary_comparison as _salary_comparison

app.config.setdefault(
    "SALARY_COMPARISON_DATA_DIR",
    os.path.join(_RUNTIME_STATE_DIR, "salary_comparison"),
)


def _salary_comparison_resolve_key(provider):
    """Resolve a saved DeepSeek/Claude key from CV Studio's shared AI secret
    store for the Salary Comparison rule-updater.  The stored key never leaves
    the server and a browser-supplied key is never persisted."""
    return _resolve_request_api_key(
        {"provider": provider, "api_key": "__BACKEND_SECURE__"}, scope="main"
    )


app.config["SALARY_COMPARISON_KEY_RESOLVER"] = _salary_comparison_resolve_key
_salary_comparison.init_app(app, url_prefix="/salary-comparison")


_CVSTUDIO_ARCHITECTURE = _finalize_modular_monolith_app(
    app,
    expected_route_count=116,
    expected_route_contract_sha256=(
        "855e04d56c550c35739c70d2dc8d35fc9d2b37d35f76453b7f3d472cf702d18e"
    ),
    expected_before_request_handlers=(
        "_assign_cvstudio_request_id",
        "_reject_declared_oversize_request",
        "_reject_non_local_host_header",
        "_require_ai_spend_browser_session",
        "_reject_cross_site_unsafe_request",
    ),
)


def _cvstudio_http_channel_timeout():
    """Waitress ``channel_timeout`` that safely covers the worst-case /parse hold.

    ``channel_timeout`` is an *inactivity* timeout: while a request is blocked
    waiting on an upstream AI call, no bytes flow to the browser, so the channel
    looks idle for the whole duration. A single ``/parse`` can chain up to three
    sequential DeepSeek calls (initial parse + retry Strategy 2 re-send +
    Strategy 3 continue), each capped by ``_cv_parse_backend_timeout_seconds``
    (currently 300s for long CVs), plus deterministic post-processing. The
    ceiling is derived from that function so the two can never drift -- if the
    per-call parse budget is raised, this follows it automatically.
    """
    per_call_ceiling = _cv_parse_backend_timeout_seconds("x" * 20000)  # long-CV branch
    max_sequential_calls = 3  # initial + Strategy 2 re-send + Strategy 3 continue
    post_processing_margin_seconds = 120
    return per_call_ceiling * max_sequential_calls + post_processing_margin_seconds


def _run_cvstudio_server():
    """Serve the app, preferring Waitress with the built-in Werkzeug server as a
    fallback.

    Waitress is a real, cross-platform, pure-Python WSGI server that freezes
    cleanly into the protected build; Werkzeug's ``app.run`` is the development
    server. Waitress is used when it is importable and not explicitly disabled;
    if the import fails (or ``CVSTUDIO_SERVER`` selects the legacy server), the
    behaviour is exactly the previous ``app.run`` call, so nothing regresses.

    Loopback only -- never bind a non-local interface; the ``_reject_non_local_
    host_header`` request guard and the whole security posture assume 127.0.0.1.
    Set ``CVSTUDIO_SERVER=werkzeug`` to force the legacy server.
    """
    prefer = str(os.environ.get("CVSTUDIO_SERVER") or "").strip().lower()
    if prefer not in ("werkzeug", "flask", "dev"):
        try:
            from waitress import serve as _waitress_serve
        except Exception:
            _waitress_serve = None
        if _waitress_serve is not None:
            _waitress_serve(
                app,
                host="127.0.0.1",
                port=_CVSTUDIO_PORT,
                # Requests routinely block for minutes (AI parse, OCR, legacy
                # .doc extraction); a generous fixed pool keeps the single-user
                # UI responsive while a long parse is in flight.
                threads=16,
                channel_timeout=_cvstudio_http_channel_timeout(),
                ident="CV Studio",
            )
            return
    app.run(port=_CVSTUDIO_PORT, debug=False, threaded=True)


if __name__ == "__main__":
    try:
        _boot_elapsed = _boot_time.time() - _BOOT_T0
        _log_path = os.path.join(_CVSTUDIO_ROOT, "startup_timing.log")
        if os.path.isfile(_log_path) and os.path.getsize(_log_path) >= 1024 * 1024:
            _old = _log_path + ".1"
            try:
                if os.path.exists(_old): os.remove(_old)
                os.replace(_log_path, _old)
            except Exception:
                pass
        with open(_log_path, "a", encoding="utf-8") as _f:
            _f.write(f"{datetime.now().isoformat(timespec='seconds')} | Python cold start (imports + route registration): {_boot_elapsed:.2f}s, about to bind port {_CVSTUDIO_PORT}\n")
    except Exception:
        pass
    # Native Windows builds may run without an attached console. Never let
    # informational startup output prevent the Flask server from binding.
    try:
        if getattr(sys, "stdout", None) is not None:
            print("\n✅  Hyppies CV Formatter is running!")
            print("👉  Open this in your browser: {}\n".format(_CVSTUDIO_BASE_URL))
    except Exception:
        pass
    _run_cvstudio_server()
